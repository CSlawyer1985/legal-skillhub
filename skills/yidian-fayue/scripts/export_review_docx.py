# -*- coding: utf-8 -*-
"""
WordBaby 合同导出模块
适配WordBaby环境 | 生成原生修订痕迹合同 + 标准化审查报告
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --------------------------
# 字体统一设置（法务标准格式）
# --------------------------
def set_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

# --------------------------
# Word 原生修订痕迹（关键！）
# --------------------------
def add_revision(paragraph, original_text, revised_text, changes):
    """添加Word原生修订：删除原文字+插入新文字，支持审阅"""
    p = paragraph
    # 删除标记（原内容）
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:author'), '合同审查工具')
    del_elem.set(qn('w:date'), '2026-03-23T00:00:00Z')
    del_text = OxmlElement('w:delText')
    del_text.text = original_text
    del_elem.append(del_text)
    p._element.append(del_elem)

    # 插入标记（修订内容）
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:author'), '合同审查工具')
    ins_elem.set(qn('w:date'), '2026-03-23T00:00:00Z')
    run = p.add_run(revised_text)
    ins_elem.append(run._element)
    p._element.append(ins_elem)

# --------------------------
# 导出 1：审查报告（标准版）
# --------------------------
def export_report(contract_name, review_result, save_dir="./"):
    doc = Document()
    # 标题
    title = doc.add_heading(0)
    t_run = title.add_run(f"{contract_name} 审查报告")
    set_font(t_run, "黑体", 18, True)

    # 基础信息
    doc.add_paragraph(f"📅 审查时间：{review_result['time']}")
    doc.add_paragraph(f"📊 资源来源：{review_result['source']}")
    doc.add_paragraph(f"⚠️ 发现问题：{len(review_result['issues'])} 项")
    doc.add_paragraph("-" * 60)

    # 一、问题清单
    doc.add_heading("一、问题清单", level=1)
    for i, item in enumerate(review_result['issues'], 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(f"【{item['type']}】{item['desc']}")
        set_font(p.runs[0])

    # 二、法律风险分析
    doc.add_heading("二、法律风险分析", level=1)
    doc.add_paragraph(review_result['legal_analysis'])

    # 三、修改建议
    doc.add_heading("三、修改建议", level=1)
    for i, sug in enumerate(review_result['suggestions'], 1):
        doc.add_paragraph(f"{i}. {sug}")

    # 四、法规依据
    doc.add_heading("四、法规依据", level=1)
    doc.add_paragraph(review_result['legal_basis'])

    # 保存
    filename = f"{contract_name}_审查报告.docx"
    path = os.path.join(save_dir, filename)
    doc.save(path)
    return filename

# --------------------------
# 导出 2：修订版合同（原生痕迹）
# --------------------------
def export_revised(contract_name, original_text, revised_text, changes, save_dir="./"):
    doc = Document()
    # 标题
    title = doc.add_heading(0)
    t_run = title.add_run(f"{contract_name}（修订版）")
    set_font(t_run, "黑体", 16, True)

    # 修订说明
    doc.add_paragraph("🔍 修订说明：")
    doc.add_paragraph("1. 本文档含 Word 原生修订痕迹，可直接【接受/拒绝】修改")
    doc.add_paragraph("2. 修订依据：民法典 + 本地审查规则")
    doc.add_paragraph("-" * 60)

    # 合同正文（带修订标记）
    doc.add_heading("合同正文", level=1)
    content_p = doc.add_paragraph()
    set_font(content_p.runs[0] if content_p.runs else content_p.add_run(), "宋体", 12)
    
    # 写入带修订痕迹的文本
    if changes:
        add_revision(content_p, original_text, revised_text, changes)
    else:
        content_p.add_run(revised_text)

    # 保存
    filename = f"{contract_name}_修订版.docx"
    path = os.path.join(save_dir, filename)
    doc.save(path)
    return filename
