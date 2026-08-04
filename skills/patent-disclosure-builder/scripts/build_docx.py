#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py — 按官方《发明专利交底书模板 2026VT1.0》结构生成 .docx 交底书。

用法:
    python build_docx.py content.json output.docx

content.json 结构（所有字段均为字符串，除 terms/meta 外）:
{
  "meta": {                       # 抬头信息表
    "交底书名称": "一种...方法",
    "发明人": "张三、李四",
    "撰写人": "张三",
    "所在部门": "研发中心",
    "涉及产品和技术": "...",
    "竞争对手产品": "...",         # 可空
    "紧急联络方式": "..."          # 可空
  },
  "terms": [                      # 关键术语表，可空数组
    {"术语": "LLM", "译名": "大语言模型", "说明": "..."}
  ],
  "发明点概述": "……（<=300字）",
  "背景_现有技术方案": "……",       # 2.1
  "背景_缺点及要解决问题": "……",   # 2.2
  "发明内容_产品侧": "……",         # 3.1 ，纯后台可写"不涉及产品侧"
  "发明内容_技术侧": "……",         # 3.2
  "有益效果": "……",
  "替代方案": "……",               # 非强制，空则填"无"
  "参考文献": "……",              # 非强制，空则填"无"

  "附加联系方式": true,           # 可省略，默认 true；置 false 则不加文末反馈区块
  "联系方式": {                   # 可省略；填了则覆盖脚本顶部 CONTACT_DEFAULT 对应项
    "署名": "IP老张",
    "微信": "IPlaozhang",
    "邮箱": "45752733@qq.com",
    "在线反馈链接": "https://...",# 可选
    "二维码列表": [               # 可选：多张码并排，默认已内置加好友/打赏两张
      {"图片": "wechat_friend.jpg", "标题": "微信（加好友）"},
      {"图片": "reward.jpg",        "标题": "打赏码"}
    ],
    "二维码目录": ""              # 可选：二维码所在目录；空则用 skill 自带 assets/qr
  }
}
# 二维码默认放在 skill 自带的 assets/qr/ 下（wechat_friend.jpg/wechat_pay.jpg/reward.jpg），
# 脚本会自动定位，无需在 content.json 里指定路径。

设计说明:
- 段落内用换行符 '\n' 分隔的文本会被拆成多个段落。
- 以 "图" 开头且形如 "图X. 说明" 的行会居中，作为图占位说明。
- 不依赖原 .doc 二进制（老式 OLE 无法作为 python-docx 模板），改为按官方章节结构重建。

插图（图文混排，符合官方模板）:
- 在任意正文字段里用标记 [[IMG:图片路径|图注文字]] 单占一行，即可在该位置插入图片。
  例: 发明内容_技术侧 里写
      "本系统整体架构如下：\n[[IMG:figs/arch.png|系统整体架构图]]\n各模块说明如下……"
- 图片路径相对于 content.json 所在目录解析（也支持绝对路径）。
- 图片自动等比缩放到正文宽度内、居中；下方自动加"图N  图注"（黑体五号居中，N 全文自动累加）。
- 图片文件缺失时，插入红色占位提示"【缺图：路径】"，不中断生成。

字体/字号体系（统一协调）:
- 文档标题：黑体 二号(22pt) 加粗 居中
- 一级标题(1、2、…)：黑体 四号(14pt) 加粗
- 二级标题(2.1/3.1…)：黑体 小四(12pt) 加粗
- 正文：宋体 小四(12pt)，行距固定 1.5 倍
- 表格文字：宋体 五号(10.5pt)；表头 黑体 五号 加粗
- 页眉右侧："保护创新，积累资产" 黑体 小五(9pt)
"""
import os, sys, json, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_HEI = "黑体"
FONT_SONG = "宋体"

# skill 根目录（本脚本在 scripts/ 下，上一级即 skill 根），用于定位自带资源如 assets/qr
SKILL_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# 默认二维码目录：skill 自带的 assets/qr
QR_DEFAULT_DIR = os.path.join(SKILL_ROOT, "assets", "qr")

# ============ 反馈与联系（文末区块）默认配置 ============
# 每次生成交底书会在正文末尾自动附上此区块，方便发明人反馈疑问/建议。
# 可在 content.json 里传 "联系方式" 对象覆盖任意字段（含 二维码图片/在线反馈链接）。
CONTACT_DEFAULT = {
    "署名": "IP老张",
    "微信": "IPlaozhang",
    "邮箱": "45752733@qq.com",
    "在线反馈链接": "",     # 可选：腾讯文档/问卷/表单等 URL
    "引导语": "本交底书由「专利交底书生成助手」协助生成，仅是交底书初稿。"
              "如对生成结果有任何疑问、补充或建议，以及如果希望IP老张帮忙进行"
              "专业把关和完善建议，欢迎随时与我联系：",
    # 二维码：支持多张并排展示，每项 {"图片": 路径, "标题": 说明文字}
    # 路径相对 SKILL 根目录（脚本会自动定位）或绝对路径均可。
    "二维码列表": [
        {"图片": "wechat_friend.jpg", "标题": "微信（加好友）"},
        {"图片": "reward.jpg", "标题": "打赏码"},
    ],
    # 二维码文件所在目录（默认指向 skill 的 assets/qr）；空则以 content.json 目录为基准
    "二维码目录": "",
}
# 是否默认附加反馈区块（content.json 里 "附加联系方式": false 可关闭）
CONTACT_ENABLED_DEFAULT = True

# 正文可用宽度（A4 21cm - 左右各 3cm = 15cm），图片最大宽度略留边
BODY_WIDTH_CM = 15.0
IMG_MAX_WIDTH_CM = 14.0

# 插图标记：[[IMG:路径|图注]]  （图注可省略）
IMG_PATTERN = re.compile(r'\[\[IMG:([^|\]]+)(?:\|([^\]]*))?\]\]')

# 全文图片计数器（模块级，main 里重置）
_FIG_COUNTER = {"n": 0}
# content.json 所在目录，用于解析相对图片路径
_BASE_DIR = {"path": "."}

# 统一字号体系（pt）
SZ_TITLE = 22      # 文档大标题 二号
SZ_H1 = 14         # 一级标题 四号
SZ_H2 = 12         # 二级标题 小四
SZ_BODY = 12       # 正文 小四
SZ_TABLE = 10.5    # 表格 五号
SZ_HEADER = 9      # 页眉 小五
SZ_NOTE = 9        # 说明文字 小五


def set_run_font(run, name=FONT_SONG, size=SZ_BODY, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体需单独设置 eastAsia
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    """level=1 一级标题(14pt)，level=2 二级标题(12pt)。"""
    size = SZ_H1 if level == 1 else SZ_H2
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 6)
    pf.space_after = Pt(6 if level == 1 else 3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_run_font(r, name=FONT_HEI, size=size, bold=True)
    return p


def _resolve_img_path(path):
    path = path.strip()
    if os.path.isabs(path):
        return path
    return os.path.normpath(os.path.join(_BASE_DIR["path"], path))


def add_image(doc, img_path, caption=None):
    """插入图片：等比缩放到正文宽内、居中；下方加"图N 图注"。图缺失则红字占位。"""
    real = _resolve_img_path(img_path)
    if not os.path.isfile(real):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"【缺图：{img_path}】")
        set_run_font(r, name=FONT_SONG, size=SZ_BODY, color=(0xC0, 0x00, 0x00))
        return
    # 图片段落（居中）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    try:
        from PIL import Image  # 若可用，用真实像素判断是否超宽
        with Image.open(real) as im:
            w_px, _ = im.size
            dpi = im.info.get("dpi", (96, 96))[0] or 96
        w_cm = w_px / dpi * 2.54
        if w_cm > IMG_MAX_WIDTH_CM:
            run.add_picture(real, width=Cm(IMG_MAX_WIDTH_CM))
        else:
            run.add_picture(real)
    except Exception:
        # 无 PIL 或读取失败：保守地按最大宽度插入
        try:
            run.add_picture(real, width=Cm(IMG_MAX_WIDTH_CM))
        except Exception as e:
            r = p.add_run(f"【图插入失败：{img_path} ({e})】")
            set_run_font(r, name=FONT_SONG, size=SZ_BODY, color=(0xC0, 0x00, 0x00))
            return
    # 图注
    _FIG_COUNTER["n"] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cap_text = f"图{_FIG_COUNTER['n']}" + (f"  {caption.strip()}" if caption and caption.strip() else "")
    cr = cap.add_run(cap_text)
    set_run_font(cr, name=FONT_HEI, size=SZ_NOTE)


def add_body(doc, text, size=SZ_BODY, empty_placeholder="【待补充】", placeholder_color=(0xC0, 0x00, 0x00)):
    """按 \n 拆段；识别 [[IMG:路径|图注]] 插图；'图X.' 文字行居中。空文本填占位。
    empty_placeholder / placeholder_color 可自定义（非强制章节传 '无' + 黑色）。"""
    if text is None:
        text = ""
    text = str(text).strip()
    if not text:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r = p.add_run(empty_placeholder)
        set_run_font(r, name=FONT_SONG, size=size, color=placeholder_color)
        return
    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            continue
        # 整行是插图标记 → 插图
        m = IMG_PATTERN.fullmatch(line.strip())
        if m:
            add_image(doc, m.group(1), m.group(2))
            continue
        # 行内混有插图标记 → 拆分处理（标记单独成图，其余文字成段）
        if IMG_PATTERN.search(line):
            parts = IMG_PATTERN.split(line)
            # split 结果: [前文, 路径, 图注, 中间文, 路径, 图注, ...]
            idx = 0
            while idx < len(parts):
                seg = parts[idx]
                if seg and seg.strip():
                    _emit_text_para(doc, seg.strip(), size)
                # 后面紧跟一组 (路径, 图注)
                if idx + 2 < len(parts) and parts[idx + 1] is not None:
                    add_image(doc, parts[idx + 1], parts[idx + 2])
                    idx += 3
                else:
                    idx += 1
            continue
        _emit_text_para(doc, line, size)


def _emit_text_para(doc, line, size=SZ_BODY):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(3)
    r = p.add_run(line)
    set_run_font(r, name=FONT_SONG, size=size)
    if re.match(r'^图\s*\d+', line):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_text(cell, text, name=FONT_SONG, size=SZ_TABLE, bold=False, align_left=True):
    cell.text = ""
    p = cell.paragraphs[0]
    # 清掉 python-docx 默认段落自带的空 run，避免单元格顶部多余空白
    for r0 in list(p.runs):
        r0._element.getparent().remove(r0._element)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(str(text) if text is not None else "")
    set_run_font(r, name=name, size=size, bold=bold)
    # 垂直居中
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def set_col_widths(table, widths_cm):
    """稳定设置列宽：需同时关闭自动布局并对每个单元格设宽。"""
    table.autofit = False
    table.allow_autofit = False
    # 关闭 autofit：设置 tblLayout=fixed
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            row.cells[idx].width = Cm(w)


def build_meta_table(doc, meta):
    rows = [
        ("交底书名称 *", meta.get("交底书名称", "")),
        ("发明人 *", meta.get("发明人", "")),
        ("撰写人 *", meta.get("撰写人", "")),
        ("所在部门 *", meta.get("所在部门", "")),
        ("涉及产品和技术 *", meta.get("涉及产品和技术", "")),
        ("竞争对手产品", meta.get("竞争对手产品", "")),
        ("紧急联络方式", meta.get("紧急联络方式", "")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_text(c0, k, name=FONT_HEI, size=SZ_TABLE, bold=True)
        set_cell_text(c1, v, name=FONT_SONG, size=SZ_TABLE)
    # A4 正文宽约 16cm：标签列 4cm，内容列 12cm
    set_col_widths(table, [4.0, 12.0])


def build_terms_table(doc, terms):
    if not terms:
        add_body(doc, "（无特别需要定义的术语）")
        return
    table = doc.add_table(rows=1 + len(terms), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for j, h in enumerate(["术语", "中文译名", "说明"]):
        set_cell_text(hdr[j], h, name=FONT_HEI, size=SZ_TABLE, bold=True)
    for i, t in enumerate(terms, start=1):
        cells = table.rows[i].cells
        vals = [t.get("术语", ""), t.get("译名", ""), t.get("说明", "")]
        for j, v in enumerate(vals):
            set_cell_text(cells[j], v, name=FONT_SONG, size=SZ_TABLE)
    # 术语 3.5cm / 译名 3.5cm / 说明 9cm
    set_col_widths(table, [3.5, 3.5, 9.0])


def setup_header(doc):
    """页眉右侧显示 '保护创新，积累资产'。"""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("保护创新，积累资产")
    set_run_font(r, name=FONT_HEI, size=SZ_HEADER, bold=True, color=(0x59, 0x59, 0x59))


def add_contact_block(doc, contact):
    """文末追加"反馈与联系"区块：引导语 + 联系方式清单 + 可选二维码。"""
    # 分隔：一条居中细横线效果（用短破折线段落近似）
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(14)
    sr = sep.add_run("— " * 12)
    set_run_font(sr, name=FONT_SONG, size=SZ_NOTE, color=(0xA0, 0xA0, 0xA0))

    add_heading(doc, "反馈与联系", level=1)

    # 引导语
    guide = str(contact.get("引导语", "") or "").strip()
    if guide:
        add_body(doc, guide)

    # 联系方式逐项（有值才显示）
    items = []
    if contact.get("微信"):
        items.append(("微信", str(contact["微信"]).strip()))
    if contact.get("邮箱"):
        items.append(("邮箱", str(contact["邮箱"]).strip()))
    if contact.get("在线反馈链接"):
        items.append(("在线反馈", str(contact["在线反馈链接"]).strip()))
    for label, val in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_after = Pt(2)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, name=FONT_HEI, size=SZ_BODY, bold=True)
        r2 = p.add_run(val)
        set_run_font(r2, name=FONT_SONG, size=SZ_BODY)

    # 二维码：支持多张并排（收款码/打赏码/加好友码等）
    qr_dir = str(contact.get("二维码目录", "") or "").strip()
    # 收集有效二维码：兼容旧字段 "二维码图片"（单张）与新字段 "二维码列表"（多张）
    qr_items = list(contact.get("二维码列表", []) or [])
    if contact.get("二维码图片"):
        qr_items.append({"图片": contact["二维码图片"], "标题": "扫码添加 / 反馈"})
    # 逐项解析真实路径（优先 二维码目录 → skill 自带 QR_DEFAULT_DIR → content.json 目录）
    resolved = []
    for it in qr_items:
        img = str(it.get("图片", "") or "").strip()
        if not img:
            continue
        cands = []
        if os.path.isabs(img):
            cands.append(img)
        else:
            if qr_dir:
                cands.append(os.path.join(qr_dir, img))
            cands.append(os.path.join(QR_DEFAULT_DIR, img))
            cands.append(_resolve_img_path(img))
        real = next((c for c in cands if os.path.isfile(c)), None)
        if real:
            resolved.append((real, str(it.get("标题", "") or "").strip()))

    if resolved:
        n = len(resolved)
        qtable = doc.add_table(rows=2, cols=n)
        qtable.alignment = WD_TABLE_ALIGNMENT.CENTER
        qtable.autofit = True
        for j, (real, title) in enumerate(resolved):
            # 上行：图片居中
            cimg = qtable.rows[0].cells[j]
            cimg.text = ""
            pimg = cimg.paragraphs[0]
            for r0 in list(pimg.runs):
                r0._element.getparent().remove(r0._element)
            pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pimg.add_run()
            try:
                run.add_picture(real, width=Cm(3.6))
            except Exception:
                pass
            # 下行：标题居中
            ccap = qtable.rows[1].cells[j]
            ccap.text = ""
            pcap = ccap.paragraphs[0]
            for r0 in list(pcap.runs):
                r0._element.getparent().remove(r0._element)
            pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = pcap.add_run(title)
            set_run_font(cr, name=FONT_HEI, size=SZ_NOTE, color=(0x59, 0x59, 0x59))

    # 署名
    sign = str(contact.get("署名", "") or "").strip()
    if sign:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(f"—— {sign}")
        set_run_font(r, name=FONT_SONG, size=SZ_BODY, color=(0x59, 0x59, 0x59))


def main():
    if len(sys.argv) < 3:
        print("用法: python build_docx.py content.json output.docx")
        sys.exit(1)
    content_path, out_path = sys.argv[1], sys.argv[2]
    with open(content_path, "r", encoding="utf-8") as f:
        c = json.load(f)

    # 相对图片路径以 content.json 所在目录为基准；重置图号
    _BASE_DIR["path"] = os.path.dirname(os.path.abspath(content_path))
    _FIG_COUNTER["n"] = 0

    doc = Document()

    # 页边距（A4 常规）
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(3.0)

    # 默认正文样式：宋体 小四 1.5 倍行距
    normal = doc.styles["Normal"]
    normal.font.name = FONT_SONG
    normal.font.size = Pt(SZ_BODY)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 页眉
    setup_header(doc)

    # 大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    tr = title_p.add_run("专利技术交底书")
    set_run_font(tr, name=FONT_HEI, size=SZ_TITLE, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("（* 标注为必填项）")
    set_run_font(nr, name=FONT_SONG, size=SZ_NOTE, color=(0x80, 0x80, 0x80))

    # 抬头信息表
    build_meta_table(doc, c.get("meta", {}))

    add_heading(doc, "【关键术语】本发明涉及到的关键术语和定义", level=1)
    build_terms_table(doc, c.get("terms", []))

    add_heading(doc, "1、*发明点概述", level=1)
    add_body(doc, c.get("发明点概述", ""))

    add_heading(doc, "2、*【背景技术】与本方案最相近的现有技术", level=1)
    add_heading(doc, "2.1 现有技术的技术方案", level=2)
    add_body(doc, c.get("背景_现有技术方案", ""))
    add_heading(doc, "2.2 现有技术的缺点及本技术方案会解决的问题", level=2)
    add_body(doc, c.get("背景_缺点及要解决问题", ""))

    add_heading(doc, "3、*【发明内容】本技术方案的详细阐述", level=1)
    add_heading(doc, "3.1 产品侧", level=2)
    add_body(doc, c.get("发明内容_产品侧", ""))
    add_heading(doc, "3.2 技术侧", level=2)
    add_body(doc, c.get("发明内容_技术侧", ""))

    add_heading(doc, "4、技术方案所产生的有益效果", level=1)
    add_body(doc, c.get("有益效果", ""))

    # 5、发散思维（替代方案）——非强制，空则填"无"
    add_heading(doc, "5、发散思维：针对 3 中的技术方案，是否还有其他替代方案（非必填，无则填“无”）", level=1)
    add_body(doc, c.get("替代方案", ""), empty_placeholder="无", placeholder_color=(0x00, 0x00, 0x00))

    # 6、参考文献——非强制，空则填"无"
    add_heading(doc, "6、参考文献（如：专利/论文/网页/期刊；非必填，无则填“无”）", level=1)
    add_body(doc, c.get("参考文献", ""), empty_placeholder="无", placeholder_color=(0x00, 0x00, 0x00))

    # 文末：反馈与联系（默认附加；content.json 里 "附加联系方式": false 可关闭）
    if c.get("附加联系方式", CONTACT_ENABLED_DEFAULT):
        contact = dict(CONTACT_DEFAULT)
        contact.update(c.get("联系方式", {}) or {})  # content.json 可覆盖
        add_contact_block(doc, contact)

    doc.save(out_path)
    print("OK ->", out_path)


if __name__ == "__main__":
    main()
