#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import re
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

DEFAULT_LINE_SPACING = 1.3
DEFAULT_FONT_NAME = "仿宋"
DEFAULT_FONT_SIZE = 12


def _set_style_east_asia_font(style, font_name: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _set_run_font(run, font_name: str, size: Optional[int] = None, bold: Optional[bool] = None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _apply_line_spacing(para, line_spacing: float) -> None:
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = line_spacing


def _set_document_font(doc: Document, font_name: str, base_font_size: int, line_spacing: float) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(base_font_size)
        _set_style_east_asia_font(style, font_name)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = line_spacing


def _add_heading(doc: Document, text: str, level: int, font_name: str, line_spacing: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    size = 16 if level == 1 else 14
    _set_run_font(run, font_name, size=size, bold=True)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    _apply_line_spacing(para, line_spacing)
    if level == 1:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _add_paragraph(doc: Document, text: str, font_name: str, line_spacing: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, font_name)
    para.paragraph_format.space_after = Pt(2)
    _apply_line_spacing(para, line_spacing)


def _add_table(doc: Document, headers: List[str], rows: List[List[str]],
               font_name: str, line_spacing: float, first_col_ratio: float = 0.28) -> None:
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.add_row().cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = ""
        para = header_cells[idx].paragraphs[0]
        run = para.add_run(header)
        _set_run_font(run, font_name, bold=True)
        _apply_line_spacing(para, line_spacing)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, cell_text in enumerate(row_data):
            cells[idx].text = ""
            para = cells[idx].paragraphs[0]
            run = para.add_run(cell_text)
            _set_run_font(run, font_name)
            _apply_line_spacing(para, line_spacing)


def _add_page_numbers(doc: Document, font_name: str, base_font_size: int) -> None:
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = para.add_run()
        _set_run_font(run, font_name, size=base_font_size)

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        run._r.append(instr_text)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)


def render_review_report(
    reviews_path: str,
    output_path: str,
    input_doc_name: str = "",
    annotation_stats: Optional[Dict] = None,
    font_name: str = DEFAULT_FONT_NAME,
    base_font_size: int = DEFAULT_FONT_SIZE,
    line_spacing: float = DEFAULT_LINE_SPACING,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(reviews_path, 'r', encoding='utf-8') as f:
        reviews = json.load(f)

    doc = Document()
    _set_document_font(doc, font_name, base_font_size, line_spacing)

    _add_heading(doc, "专利申请文件审查报告", 1, font_name, line_spacing)

    _add_heading(doc, "一、基本信息", 2, font_name, line_spacing)
    info_rows = [
        ["审查文件", input_doc_name or "未指定"],
        ["发现问题", str(len(reviews))],
    ]
    if annotation_stats:
        info_rows.append(["批注成功", str(annotation_stats.get('successful', 0))])
        info_rows.append(["批注跳过", str(annotation_stats.get('skipped', 0))])
        info_rows.append(["批注失败", str(annotation_stats.get('failed', 0))])
        rate = annotation_stats.get('success_rate', 0)
        info_rows.append(["批注成功率", f"{rate:.1f}%"])
    _add_table(doc, ["项目", "内容"], info_rows, font_name, line_spacing, first_col_ratio=0.3)

    section_groups: Dict[str, List] = {}
    for review in reviews:
        section = review.get('section', '未知')
        section_groups.setdefault(section, []).append(review)

    _add_heading(doc, "二、各章节问题统计", 2, font_name, line_spacing)
    stat_rows = []
    for section, items in sorted(section_groups.items()):
        action_counts = {}
        for item in items:
            action = item.get('action_type', 'comment')
            action_counts[action] = action_counts.get(action, 0) + 1
        action_str = "、".join(f"{k}:{v}" for k, v in sorted(action_counts.items()))
        stat_rows.append([section, str(len(items)), action_str])
    _add_table(doc, ["章节", "问题数", "操作类型分布"], stat_rows, font_name, line_spacing, first_col_ratio=0.25)

    _add_heading(doc, "三、问题详情", 2, font_name, line_spacing)
    for section, items in sorted(section_groups.items()):
        _add_heading(doc, f"{section}（{len(items)} 处）", 2, font_name, line_spacing)
        for idx, item in enumerate(items, 1):
            claim_num = item.get('claim_number')
            claim_prefix = f"权利要求{claim_num}: " if claim_num else ""
            issue_text = item.get('issue', '')
            suggestion_text = item.get('suggestion', '')
            action_type = item.get('action_type', 'comment')

            para = doc.add_paragraph()
            run = para.add_run(f"{idx}. {claim_prefix}{issue_text}")
            _set_run_font(run, font_name, bold=True)
            _apply_line_spacing(para, line_spacing)

            if suggestion_text:
                para2 = doc.add_paragraph()
                run2 = para2.add_run(f"   修改建议：{suggestion_text}")
                _set_run_font(run2, font_name)
                _apply_line_spacing(para2, line_spacing)

            if action_type in ('replace', 'delete'):
                old_text = item.get('old_text', '')
                new_text = item.get('new_text', '')
                if old_text:
                    para3 = doc.add_paragraph()
                    run3 = para3.add_run(f"   原文：{old_text[:80]}{'...' if len(old_text) > 80 else ''}")
                    _set_run_font(run3, font_name)
                    _apply_line_spacing(para3, line_spacing)
                if new_text and action_type == 'replace':
                    para4 = doc.add_paragraph()
                    run4 = para4.add_run(f"   改为：{new_text[:80]}{'...' if len(new_text) > 80 else ''}")
                    _set_run_font(run4, font_name)
                    _apply_line_spacing(para4, line_spacing)

    _add_page_numbers(doc, font_name, base_font_size)
    doc.save(output_path)
    return output_path


def render_review_summary(
    reviews_path: str,
    output_path: str,
    input_doc_name: str = "",
    patent_type: str = "",
    font_name: str = DEFAULT_FONT_NAME,
    base_font_size: int = DEFAULT_FONT_SIZE,
    line_spacing: float = DEFAULT_LINE_SPACING,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(reviews_path, 'r', encoding='utf-8') as f:
        reviews = json.load(f)

    doc = Document()
    _set_document_font(doc, font_name, base_font_size, line_spacing)

    _add_heading(doc, "专利申请文件审查摘要", 1, font_name, line_spacing)

    _add_heading(doc, "一、基本信息", 2, font_name, line_spacing)
    info_rows = [
        ["文件名称", input_doc_name or "未指定"],
        ["专利类型", patent_type or "未识别"],
        ["发现问题总数", str(len(reviews))],
    ]
    _add_table(doc, ["项目", "内容"], info_rows, font_name, line_spacing, first_col_ratio=0.35)

    section_groups: Dict[str, List] = {}
    for review in reviews:
        section = review.get('section', '未知')
        section_groups.setdefault(section, []).append(review)

    _add_heading(doc, "二、各章节问题统计", 2, font_name, line_spacing)
    stat_rows = []
    for section, items in sorted(section_groups.items()):
        stat_rows.append([section, str(len(items))])
    _add_table(doc, ["章节", "问题数"], stat_rows, font_name, line_spacing, first_col_ratio=0.5)

    _add_heading(doc, "三、关键问题摘要", 2, font_name, line_spacing)
    replace_items = [r for r in reviews if r.get('action_type') in ('replace', 'delete')]
    if replace_items:
        for idx, item in enumerate(replace_items[:10], 1):
            section = item.get('section', '')
            issue = item.get('issue', '')
            para = doc.add_paragraph()
            run = para.add_run(f"{idx}. [{section}] {issue}")
            _set_run_font(run, font_name)
            _apply_line_spacing(para, line_spacing)
    else:
        _add_paragraph(doc, "无需要修改的关键问题", font_name, line_spacing)

    _add_page_numbers(doc, font_name, base_font_size)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    print("专利审查报告生成器")
    print("=" * 60)
    print()
    print("使用示例:")
    print()
    print("from scripts.report_renderer import render_review_report, render_review_summary")
    print()
    print("render_review_report('reviews.json', '审查报告.docx', '专利文件.docx')")
    print("render_review_summary('reviews.json', '审查摘要.docx', '专利文件.docx')")
