# -*- coding: utf-8 -*-
"""
合同审查报告生成器 v2.0
东润律师事务所 · 合同审查技能辅助工具

功能：从结构化审查数据生成格式化的DOCX审查报告。
适用场景：合同审查技能在完成分析后调用本脚本输出最终报告。

使用方式1（命令行）：
  python gen_report.py <contract_name> <output_path> <analysis_json_path>

使用方式2（Python调用）：
  from gen_report import generate_report
  generate_report(contract_name, output_path, analysis_dict)
"""

import sys, io, os, json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class ReportGenerator:
    """合同审查报告 DOCX 生成器"""

    # 风险配色
    COLORS = {
        'red':    RGBColor(200, 0, 0),
        'yellow': RGBColor(180, 140, 0),
        'blue':   RGBColor(0, 80, 180),
        'gray':   RGBColor(140, 140, 140),
        'black':  RGBColor(0, 0, 0),
        'dark':   RGBColor(50, 50, 50),
    }

    ICONS = {'red': '🔴', 'yellow': '🟡', 'blue': '🔵'}

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """设置文档默认样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(11)
        font.color.rgb = self.COLORS['dark']
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(4)

    def _set_east_asian_font(self, run, font_name='宋体'):
        """设置东亚字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def cover_page(self, contract_name, standpoint='甲方', date=None):
        """生成封面"""
        date_str = date or datetime.now().strftime('%Y年%m月%d日')

        # 空白区域
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(100)

        # 主标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('合 同 审 查 报 告')
        self._set_east_asian_font(r, '黑体')
        r.font.size = Pt(26)
        r.bold = True

        # 合同名称
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        r = p.add_run(contract_name)
        self._set_east_asian_font(r, '宋体')
        r.font.size = Pt(16)

        # 信息区
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        info_lines = [
            f'审查立场：{standpoint}',
            f'审查日期：{date_str}',
            '审查依据：合同起草与审查SOP（通用版）',
            '出具机构：山东东润律师事务所',
        ]
        for line in info_lines:
            r = p.add_run(line + '\n')
            self._set_east_asian_font(r, '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = self.COLORS['dark']

        self.doc.add_page_break()

    def h1(self, text):
        """一级标题"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        self._set_east_asian_font(r, '黑体')
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = self.COLORS['black']
        return p

    def h2(self, text):
        """二级标题"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        self._set_east_asian_font(r, '黑体')
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = self.COLORS['dark']
        return p

    def para(self, text, bold=False, indent=True):
        """正文段落"""
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        self._set_east_asian_font(r, '宋体')
        r.font.size = Pt(11)
        if bold:
            r.bold = True
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        return p

    def risk(self, level, text, source=''):
        """
        风险标注段落
        level: 'red' (必须修改), 'yellow' (建议修改), 'blue' (提示注意)
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)

        # 图标
        r = p.add_run(f'{self.ICONS[level]} ')
        self._set_east_asian_font(r, '宋体')
        r.bold = True

        # 正文
        r = p.add_run(text)
        self._set_east_asian_font(r, '宋体')
        r.font.size = Pt(11)
        r.font.color.rgb = self.COLORS[level]

        # 来源标注
        if source:
            r = p.add_run(f'  [{source}]')
            self._set_east_asian_font(r, '宋体')
            r.font.size = Pt(9)
            r.font.color.rgb = self.COLORS['gray']

        return p

    def separator(self):
        """分隔线"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run('─' * 50)
        r.font.size = Pt(8)
        r.font.color.rgb = self.COLORS['gray']

    def info_table(self, rows):
        """信息表格（键值对）"""
        table = self.doc.add_table(rows=len(rows), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, (k, v) in enumerate(rows):
            table.cell(i, 0).text = k
            table.cell(i, 1).text = v
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        self._set_east_asian_font(r, '宋体')
                        r.font.size = Pt(10.5)
        self.doc.add_paragraph()  # spacing after table
        return table

    def generate(self, contract_name, analysis, standpoint='甲方'):
        """
        主生成方法

        analysis 结构（字典）：
        {
            "risk_level": "🟡 中风险",
            "signing_advice": "修改后签署",
            "overall_assessment": "总体评价文字...",
            "commercial_purpose": {"q1": "...", "q2": "...", "q3": "..."},
            "modules": {
                "M1": {"title": "主体与效力", "status": "✅/⚠️/🔴", "findings": [{"level":"red","text":"...","source":"..."}, ...]},
                ...
            },
            "special_clauses": [{"level":"red","text":"...", "source":"..."}, ...],
            "logic_check": [{"level":"red","text":"...", "source":"..."}, ...],
            "language_purification": ["建议1", "建议2", ...],
            "professional_clauses": [{"level":"yellow","text":"...", "source":"..."}, ...],
            "summary": {
                "必须修改": ["项1", "项2", ...],
                "建议修改": ["项1", "项2", ...],
                "提示注意": ["项1", "项2", ...]
            },
            "conclusion": "审查结论..."
        }
        """

        # === 封面 ===
        self.cover_page(contract_name, standpoint)

        # === 一、合同基本信息 ===
        self.h1('一、合同基本信息')
        basic = analysis.get('basic_info', {})
        for k, v in basic.items():
            self.para(f'{k}：{v}', indent=False)

        # === 二、总体评价 ===
        self.h1('二、总体评价')
        rows = [
            ('审查立场', standpoint),
            ('风险等级', analysis.get('risk_level', '待评估')),
            ('签约建议', analysis.get('signing_advice', '待评估')),
        ]
        self.info_table(rows)
        self.para(analysis.get('overall_assessment', ''))

        # === 三、商业目的穿透 ===
        self.h1('三、商业目的穿透分析')
        cp = analysis.get('commercial_purpose', {})
        for key, label in [('q1', '第一问：客户想通过合同得到什么？'),
                           ('q2', '第二问：谁强势，谁弱势？'),
                           ('q3', '第三问：警惕合同名实不符')]:
            if key in cp:
                self.para(label, bold=True)
                self.para(cp[key])

        if 'validity_check' in cp:
            self.para('效力预判：', bold=True)
            for item in cp['validity_check']:
                self.para(item)

        # === 四、逐模块审查 ===
        self.h1('四、逐模块审查结果（M1-M10）')
        for m_num in range(1, 11):
            key = f'M{m_num}'
            if key in analysis.get('modules', {}):
                mod = analysis['modules'][key]
                self.h2(f'{key} {mod.get("title", "")}  {mod.get("status", "")}')
                for f in mod.get('findings', []):
                    self.risk(f.get('level', 'blue'), f.get('text', ''), f.get('source', ''))
                if mod.get('positive_notes'):
                    for note in mod['positive_notes']:
                        self.para(f'✅ {note}')

        # === 五、特殊条款审查 ===
        self.h1('五、特殊条款审查')
        for item in analysis.get('special_clauses', []):
            self.risk(item.get('level', 'blue'), item.get('text', ''), item.get('source', ''))

        # === 六、逻辑一致性 ===
        self.h1('六、逻辑一致性检验')
        for item in analysis.get('logic_check', []):
            self.risk(item.get('level', 'blue'), item.get('text', ''), item.get('source', ''))

        # === 七、语言净化 ===
        self.h1('七、语言净化建议')
        for item in analysis.get('language_purification', []):
            self.para(f'• {item}')

        # === 八、专业性条款补充 ===
        self.h1('八、专业性条款补充建议')
        for item in analysis.get('professional_clauses', []):
            self.risk(item.get('level', 'blue'), item.get('text', ''), item.get('source', ''))

        # === 九、修改建议汇总 ===
        self.h1('九、修改建议汇总')
        summary = analysis.get('summary', {})
        cat_colors = {
            '必须修改': ('🔴 必须修改（签约前必须完成）', self.COLORS['red']),
            '建议修改': ('🟡 建议修改（建议签约前调整）', self.COLORS['yellow']),
            '提示注意': ('🔵 提示注意', self.COLORS['blue']),
        }
        for cat, (label, color) in cat_colors.items():
            items = summary.get(cat, [])
            if items:
                self.para(f'\n{label}（共{len(items)}项）：', bold=True)
                for i, item in enumerate(items, 1):
                    self.para(f'{i}. {item}')

        # === 十、审查结论 ===
        self.h1('十、审查结论')
        self.para(analysis.get('signing_advice', '待评估'), bold=True)
        self.para(analysis.get('conclusion', ''))

        # === 页脚 ===
        self.doc.add_paragraph()
        self.separator()
        self.para('本报告由山东东润律师事务所出具，仅供客户内部参考使用。', indent=False)
        self.para(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}', indent=False)

        return self.doc

    def save(self, output_path):
        """保存文档"""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path


# ============================================================
# 客户友好版报告生成器 v2.1
# ============================================================

class ClientReportGenerator:
    """面向客户的合同审查报告 — 通俗、简洁、可操作"""

    # 品牌配色（客户友好：更柔和的色调）
    CLIENT_COLORS = {
        'red':    RGBColor(180, 30, 30),
        'yellow': RGBColor(160, 120, 0),
        'blue':   RGBColor(30, 80, 160),
        'green':  RGBColor(30, 130, 76),
        'gray':   RGBColor(130, 130, 130),
        'dark':   RGBColor(60, 60, 60),
        'title':  RGBColor(0, 51, 102),
    }

    ICONS = {'red': '\u25cf', 'yellow': '\u25cf', 'blue': '\u25cf'}  # 改用实心圆点，兼容性更好

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        font.color.rgb = self.CLIENT_COLORS['dark']
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.paragraph_format.line_spacing = 1.4
        style.paragraph_format.space_after = Pt(3)

    def _set_font(self, run, name='微软雅黑'):
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

    def _risk_dot(self, level):
        colors = {'red': self.CLIENT_COLORS['red'],
                  'yellow': self.CLIENT_COLORS['yellow'],
                  'blue': self.CLIENT_COLORS['blue']}
        labels = {'red': '必须改', 'yellow': '建议改', 'blue': '需注意'}
        return f'[{labels[level]}]', colors[level]

    # ---------- 页面结构 ----------

    def cover(self, contract_name, standpoint, date=None):
        """客户版封面"""
        date_str = date or datetime.now().strftime('%Y年%m月%d日')

        p = self.doc.add_paragraph(); p.paragraph_format.space_before = Pt(80)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('合同审查意见书')
        self._set_font(r, '微软雅黑'); r.font.size = Pt(28); r.bold = True
        r.font.color.rgb = self.CLIENT_COLORS['title']

        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        r = p.add_run(contract_name)
        self._set_font(r, '微软雅黑'); r.font.size = Pt(14)

        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(50)
        for line in [f'审查方：{standpoint}', f'日期：{date_str}', '山东东润律师事务所']:
            r = p.add_run(line + '\n')
            self._set_font(r, '微软雅黑'); r.font.size = Pt(11)
            r.font.color.rgb = self.CLIENT_COLORS['gray']

        self.doc.add_page_break()

    def executive_summary(self, analysis):
        """一页纸概览"""
        self._section_title('审查结论')

        # 核心结论卡片（模拟卡片效果）
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(f"签约建议：{analysis.get('signing_advice', '')}")
        self._set_font(r, '微软雅黑'); r.font.size = Pt(16); r.bold = True
        r.font.color.rgb = self.CLIENT_COLORS['yellow']

        # 关键指标
        summary = analysis.get('summary', {})
        must = len(summary.get('必须修改', []))
        suggest = len(summary.get('建议修改', []))
        notice = len(summary.get('提示注意', []))

        p = self.doc.add_paragraph()
        r = p.add_run(f'共发现 {must} 项必须修改、{suggest} 项建议修改、{notice} 项需要注意的问题。')
        self._set_font(r, '微软雅黑'); r.font.size = Pt(11)

        # 一句话总结
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(self._plain_summary(analysis))
        self._set_font(r, '微软雅黑'); r.font.size = Pt(11)

        self.doc.add_page_break()

    def _plain_summary(self, analysis):
        """生成一句话通俗总结"""
        summary = analysis.get('summary', {})
        must = summary.get('必须修改', [])
        risk = analysis.get('risk_level', '')

        parts = []
        if '中' in risk:
            parts.append('整体风险可控，')
        elif '高' in risk:
            parts.append('存在较高风险，')

        parts.append(f'主要问题集中在{self._top_issues(must)}方面。')
        parts.append(f'完成必须修改的{must.__len__()}个问题后可以签署。')
        return ''.join(parts)

    def _top_issues(self, must_items):
        if not must_items:
            return '个别细节'
        keywords = ['信息', '日期', '金额', '罚款', '违约金', '发票', '附件']
        found = []
        for kw in keywords:
            for item in must_items:
                if kw in item:
                    found.append(kw)
                    break
        if not found or len(found) > 3:
            return '合同信息和违约条款'
        return '、'.join(found[:3])

    # ---------- 核心内容 ----------

    def key_issues(self, analysis):
        """核心问题（客户最关心的）"""
        self._section_title('需要您关注的核心问题')

        summary = analysis.get('summary', {})
        must_items = summary.get('必须修改', [])

        if must_items:
            self._subsection('签约前必须解决')
            self._para('以下问题需要在签署合同前完成修改。不改好不能签。', indent=True)
            for i, item in enumerate(must_items, 1):
                self._action_item(item, 'red')

        suggest_items = summary.get('建议修改', [])
        if suggest_items:
            self._subsection('建议一并调整')
            self._para('以下问题建议在本次修改中一并处理，但不影响签约。', indent=True)
            for i, item in enumerate(suggest_items, 1):
                self._action_item(item, 'yellow')

        self.doc.add_page_break()

    def negotiation_guide(self, analysis):
        """谈判指引"""
        self._section_title('谈判策略建议')

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run('以下是关键条款的谈判建议，帮助您在修改合同条款时有的放矢。')
        self._set_font(r, '微软雅黑'); r.font.size = Pt(10)
        r.font.color.rgb = self.CLIENT_COLORS['gray']

        # 谈判优先级表
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        headers = ['条款', '现状问题', '修改目标', '谈判策略']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    self._set_font(r, '微软雅黑'); r.bold = True; r.font.size = Pt(9)

        ng_items = self._extract_negotiation_items(analysis)
        for item in ng_items:
            row = table.add_row()
            for i, val in enumerate(item):
                row.cells[i].text = val
                for p in row.cells[i].paragraphs:
                    for r in p.runs:
                        self._set_font(r, '微软雅黑'); r.font.size = Pt(9)

        self.doc.add_page_break()

    def _extract_negotiation_items(self, analysis):
        """从分析中提取谈判条目"""
        items = []
        modules = analysis.get('modules', {})
        m6 = modules.get('M6', {})
        for f in m6.get('findings', []):
            if f['level'] == 'red':
                text = f['text']
                # 简化为谈判条目
                if '十倍' in text:
                    items.append(['缺件赔偿', '物料成本价十倍，超出法律允许上限', '按物料成本价赔偿，其他损失另行计算', '核心条款，必须修改。态度坚决。'])
                elif '所有货款' in text or '壹佰万' in text:
                    items.append(['违约赔偿', '全部货款/不低于100万元，面临无效风险', '按合同总额10%-20%设违约金', '核心条款，必须修改。可接受较高比例。'])
                elif '300元/分钟' in text:
                    items.append(['延迟罚款', '按300元/分钟计算，日罚432,000元', '按货值每日千分之五计算', '核心条款，必须修改。对方易接受比例式。'])
                elif '不以实际损失' in text:
                    items.append(['重复质量罚款', '写明不以实际损失为基础，自认无效', '删除该表述，保留不低于5万元', '必须修改。删除即可，底线不变。'])
        return items

    def bottom_line(self, analysis):
        """最终建议"""
        self._section_title('我们的建议')

        advice = analysis.get('signing_advice', '')
        if '修改后签署' in advice:
            msg = '我们建议您要求对方修改上述必须修改的问题后签署本合同。当前版本的合同框架有利于贵司，主要问题是几处惩罚条款写得太重，在法庭上反而站不住脚。修改后，这份合同能较好地保护贵司权益。'
        elif '不建议签' in advice:
            msg = '鉴于合同存在重大风险，我们建议暂不签署。请联系我们进一步讨论。'
        else:
            msg = '合同整体风险可控，可以签署。请注意上述需关注的问题。'

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(msg)
        self._set_font(r, '微软雅黑'); r.font.size = Pt(12)

        # 联系信息
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        r = p.add_run('如有疑问，请随时联系我们进行详细讨论。\n山东东润律师事务所')
        self._set_font(r, '微软雅黑'); r.font.size = Pt(10)
        r.font.color.rgb = self.CLIENT_COLORS['gray']

    # ---------- 基础排版 ----------

    def _section_title(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        self._set_font(r, '微软雅黑'); r.font.size = Pt(18); r.bold = True
        r.font.color.rgb = self.CLIENT_COLORS['title']

        # 下划线
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run('\u2500' * 60)
        r.font.size = Pt(6); r.font.color.rgb = self.CLIENT_COLORS['title']

    def _subsection(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        self._set_font(r, '微软雅黑'); r.font.size = Pt(13); r.bold = True
        r.font.color.rgb = self.CLIENT_COLORS['dark']

    def _para(self, text, bold=False, indent=False):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        self._set_font(r, '微软雅黑'); r.font.size = Pt(11)
        if bold: r.bold = True
        if indent: p.paragraph_format.first_line_indent = Cm(0.74)

    def _action_item(self, text, level):
        """操作项：标签 + 内容"""
        label, color = self._risk_dot(level)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f'{label} ')
        self._set_font(r, '微软雅黑'); r.font.size = Pt(10); r.bold = True
        r.font.color.rgb = color
        r = p.add_run(text)
        self._set_font(r, '微软雅黑'); r.font.size = Pt(10)

    def separator(self):
        p = self.doc.add_paragraph()
        r = p.add_run('\u2500' * 50)
        r.font.size = Pt(6); r.font.color.rgb = self.CLIENT_COLORS['gray']

    def generate(self, contract_name, analysis, standpoint='甲方'):
        """
        生成客户友好版报告
        结构：封面 → 审查结论（一页纸）→ 核心问题 → 谈判策略 → 我们的建议
        """
        self.cover(contract_name, standpoint)
        self.executive_summary(analysis)
        self.key_issues(analysis)
        self.negotiation_guide(analysis)
        self.bottom_line(analysis)

        # 页脚
        self.doc.add_paragraph()
        self.separator()
        self._para('本意见书由山东东润律师事务所出具，仅供客户内部参考。', indent=False)
        self._para(f'生成时间：{datetime.now().strftime("%Y-%m-%d")}', indent=False)

        return self.doc

    def save(self, output_path):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path


def generate_client_report(contract_name, output_path, analysis, standpoint='甲方'):
    """便捷函数：生成客户友好版报告"""
    gen = ClientReportGenerator()
    gen.generate(contract_name, analysis, standpoint)
    gen.save(output_path)
    return output_path


def generate_report(contract_name, output_path, analysis, standpoint='甲方'):
    """便捷函数：生成并保存报告"""
    gen = ReportGenerator()
    gen.generate(contract_name, analysis, standpoint)
    gen.save(output_path)
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 4:
        print('用法: python gen_report.py <合同名称> <输出路径.docx> <分析数据.json> [审查立场]')
        print('示例: python gen_report.py "采购合同" "report.docx" "analysis.json" "甲方"')
        sys.exit(1)

    contract_name = sys.argv[1]
    output_path = sys.argv[2]
    json_path = sys.argv[3]
    standpoint = sys.argv[4] if len(sys.argv) > 4 else '甲方'

    with open(json_path, 'r', encoding='utf-8') as f:
        analysis = json.load(f)

    path = generate_report(contract_name, output_path, analysis, standpoint)
    print(f'报告已生成: {path}')
