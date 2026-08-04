#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
律师费报价方案Word文档生成器
根据案件参数自动生成专业格式的律师费报价方案
"""

import argparse
import json
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 各省律师费收费标准（财产案件分段累进）
FEE_STANDARDS = {
    "浙江": {
        "10万以下": {"min": 0.08, "max": 0.10, "mid": 0.09},
        "10-50万": {"min": 0.05, "max": 0.07, "mid": 0.06},
        "50-500万": {"min": 0.03, "max": 0.05, "mid": 0.04},
        "500万以上": {"min": 0.005, "max": 0.01, "mid": 0.008},
    },
    "北京": {
        "10万以下": {"min": 0.08, "max": 0.10, "mid": 0.09},
        "10-50万": {"min": 0.05, "max": 0.07, "mid": 0.06},
        "50-500万": {"min": 0.03, "max": 0.05, "mid": 0.04},
        "500万以上": {"min": 0.005, "max": 0.01, "mid": 0.008},
    },
    "上海": {
        "10万以下": {"min": 0.08, "max": 0.10, "mid": 0.09},
        "10-50万": {"min": 0.05, "max": 0.07, "mid": 0.06},
        "50-500万": {"min": 0.03, "max": 0.05, "mid": 0.04},
        "500万以上": {"min": 0.005, "max": 0.01, "mid": 0.008},
    },
    "广东": {
        "10万以下": {"min": 0.08, "max": 0.10, "mid": 0.09},
        "10-50万": {"min": 0.05, "max": 0.07, "mid": 0.06},
        "50-500万": {"min": 0.03, "max": 0.05, "mid": 0.04},
        "500万以上": {"min": 0.005, "max": 0.01, "mid": 0.008},
    },
}

# 默认使用浙江省标准
DEFAULT_PROVINCE = "浙江"

# 默认难点描述模板
DEFAULT_DIFFICULTIES = {
    1: ["案件事实相对清晰", "法律关系较为简单", "证据材料较为完备"],
    2: ["涉及多方利益协调", "部分事实存在争议", "需要专业法律分析"],
    3: ["案件事实复杂", "证据材料较多", "涉及多个法律关系", "需要专项调查"],
    4: ["案件难度较大", "涉及重大利益", "证据收集困难", "法律适用存在争议", "需要多方协调"],
    5: ["案件极为复杂", "涉及跨境/跨地区法律问题", "证据收集极为困难", "法律适用争议大", "需要专家论证", "多方利益冲突激烈"],
}

# 默认工作量估算
DEFAULT_WORK_HOURS = {
    1: {"证据筹备": "10-15", "法律策略": "8-10", "庭审谈判": "5-8", "成果落地": "3-5"},
    2: {"证据筹备": "15-20", "法律策略": "12-15", "庭审谈判": "10-12", "成果落地": "5-8"},
    3: {"证据筹备": "25-35", "法律策略": "20-25", "庭审谈判": "15-20", "成果落地": "10-15"},
    4: {"证据筹备": "35-45", "法律策略": "25-30", "庭审谈判": "20-25", "成果落地": "10-15"},
    5: {"证据筹备": "45-60", "法律策略": "35-45", "庭审谈判": "30-40", "成果落地": "20-25"},
}


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_run_with_format(paragraph, text, font_name="宋体", font_size=12, bold=False):
    """添加格式化文本"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def set_paragraph_format(paragraph, first_line_indent=True, line_spacing=1.5, space_before=0, space_after=10):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)  # 2字符
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return paragraph


def calculate_standard_fee(property_amount, province):
    """计算标准收费"""
    standards = FEE_STANDARDS.get(province, FEE_STANDARDS[DEFAULT_PROVINCE])
    
    total = 0
    details = []
    
    # 10万以下
    if property_amount <= 10:
        fee = property_amount * standards["10万以下"]["mid"]
        details.append(f"10万元以下部分：按{standards['10万以下']['min']*100:.0f}%-{standards['10万以下']['max']*100:.0f}%计费，取中间值{standards['10万以下']['mid']*100:.0f}%，计{property_amount}万×{standards['10万以下']['mid']*100:.0f}%={fee:.2f}万元")
        total += fee
    else:
        fee = 10 * standards["10万以下"]["mid"]
        details.append(f"10万元以下部分：按{standards['10万以下']['min']*100:.0f}%-{standards['10万以下']['max']*100:.0f}%计费，取中间值{standards['10万以下']['mid']*100:.0f}%，计10万×{standards['10万以下']['mid']*100:.0f}%={fee:.2f}万元")
        total += fee
        
        # 10-50万
        if property_amount <= 50:
            amount = property_amount - 10
            fee = amount * standards["10-50万"]["mid"]
            details.append(f"10-50万元部分（共计{amount}万元）：按{standards['10-50万']['min']*100:.0f}%-{standards['10-50万']['max']*100:.0f}%计费，取中间值{standards['10-50万']['mid']*100:.0f}%，计{amount}万×{standards['10-50万']['mid']*100:.0f}%={fee:.2f}万元")
            total += fee
        else:
            fee = 40 * standards["10-50万"]["mid"]
            details.append(f"10-50万元部分（共计40万元）：按{standards['10-50万']['min']*100:.0f}%-{standards['10-50万']['max']*100:.0f}%计费，取中间值{standards['10-50万']['mid']*100:.0f}%，计40万×{standards['10-50万']['mid']*100:.0f}%={fee:.2f}万元")
            total += fee
            
            # 50-500万
            if property_amount <= 500:
                amount = property_amount - 50
                fee = amount * standards["50-500万"]["mid"]
                details.append(f"50-500万元部分（共计{amount}万元）：按{standards['50-500万']['min']*100:.0f}%-{standards['50-500万']['max']*100:.0f}%计费，取中间值{standards['50-500万']['mid']*100:.0f}%，计{amount}万×{standards['50-500万']['mid']*100:.0f}%={fee:.2f}万元")
                total += fee
            else:
                fee = 450 * standards["50-500万"]["mid"]
                details.append(f"50-500万元部分（共计450万元）：按{standards['50-500万']['min']*100:.0f}%-{standards['50-500万']['max']*100:.0f}%计费，取中间值{standards['50-500万']['mid']*100:.0f}%，计450万×{standards['50-500万']['mid']*100:.0f}%={fee:.2f}万元")
                total += fee
                
                # 500万以上
                amount = property_amount - 500
                fee = amount * standards["500万以上"]["mid"]
                details.append(f"500万元以上部分（{property_amount-500}万元）：按{standards['500万以上']['min']*100:.1f}%-{standards['500万以上']['max']*100:.1f}%计费，取中间值{standards['500万以上']['mid']*100:.1f}%，计{amount}万×{standards['500万以上']['mid']*100:.1f}%={fee:.2f}万元")
                total += fee
    
    return total, details


def calculate_advanced_fee(property_amount, difficulty_level):
    """计算进阶收费（基础固定+激励）"""
    base_fee = 5 + property_amount * 0.02  # 基础费：5万 + 标的额的2%
    
    # 激励费计算
    perfect_bonus = property_amount * 0.006  # 完美达标0.6%
    core_bonus = property_amount * 0.003  # 核心达标0.3%
    
    return base_fee, perfect_bonus, core_bonus


def calculate_stage_fee(property_amount, difficulty_level):
    """计算分阶段收费"""
    stage1 = 3 + property_amount * 0.02  # 立案阶段
    stage2 = 3 + property_amount * 0.02  # 庭前阶段
    stage3 = 3 + property_amount * 0.01  # 执行阶段
    
    return stage1, stage2, stage3


def generate_difficulty_text(case_type, difficulty_level, custom_difficulties=None):
    """生成难度描述文本"""
    if custom_difficulties:
        difficulties = custom_difficulties
    else:
        difficulties = DEFAULT_DIFFICULTIES.get(difficulty_level, DEFAULT_DIFFICULTIES[3])
    
    case_specific_intro = {
        "离婚纠纷": "本案涉及婚姻关系解除、子女抚养、财产分割、债务承担等多维度法律问题，",
        "合同纠纷": "本案涉及合同订立、履行、变更、解除等多环节法律问题，",
        "劳动争议": "本案涉及劳动关系建立、履行、解除等多方面法律问题，",
        "交通事故": "本案涉及交通事故责任认定、人身损害赔偿、保险理赔等法律问题，",
        "刑事辩护": "本案涉及刑事指控、证据审查、法律适用等核心法律问题，",
        "房产纠纷": "本案涉及房产产权、交易安全、租赁关系等法律问题，",
        "债务纠纷": "本案涉及债权债务关系、担保责任、追偿权等法律问题，",
    }
    
    intro = case_specific_intro.get(case_type, "本案涉及多维度法律问题，")
    
    difficulty_desc = {
        1: "案件整体难度较低，事实较为清晰，法律关系相对简单，预计工作量较小。",
        2: "案件存在一定复杂性，部分事实需要进一步核实，预计工作量适中。",
        3: "案件复杂度较高，涉及多个法律关系，需要进行专项调查和论证，预计工作量较大。",
        4: "案件难度较大，涉及重大利益争议，证据收集困难，法律适用存在争议，预计工作量很大。",
        5: "案件极为复杂，涉及跨境/跨地区法律问题、重大利益冲突或多方博弈，证据收集极为困难，需要专家论证，预计工作量巨大。",
    }
    
    result = intro + difficulty_desc.get(difficulty_level, difficulty_desc[3])
    
    if custom_difficulties:
        result += "\n\n本案核心难点包括：\n" + "\n".join([f"{i+1}. {d}" for i, d in enumerate(custom_difficulties)])
    
    return result


def generate_work_hours_text(work_hours_dict=None, difficulty_level=3):
    """生成工作量描述文本"""
    if work_hours_dict is None:
        work_hours_dict = DEFAULT_WORK_HOURS.get(difficulty_level, DEFAULT_WORK_HOURS[3])
    
    total_min = 0
    total_max = 0
    
    for phase, hours in work_hours_dict.items():
        if isinstance(hours, str):
            if '-' in hours:
                parts = hours.split('-')
                total_min += int(parts[0])
                total_max += int(parts[1])
    
    result = "本案全流程专属工作量如下：\n\n"
    
    phase_mapping = {
        "证据筹备": "证据专项筹备阶段",
        "法律策略": "法律策略与文书阶段",
        "庭审谈判": "庭审与专项谈判阶段",
        "成果落地": "成果落地跟进阶段",
    }
    
    for phase, hours in work_hours_dict.items():
        phase_name = phase_mapping.get(phase, phase)
        result += f"{phase_name}：预计 {hours} 小时。\n"
    
    result += f"\n总专属工作量：{total_min}-{total_max} 小时。"
    
    return result


def create_document(args):
    """创建Word文档"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title, first_line_indent=False, space_before=0, space_after=20)
    run = title.add_run(f"{args.case_type}律师费报价方案")
    run.font.name = "黑体"
    run.font.size = Pt(22)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    # ===== 一、案件核心难度与专属工作量分析 =====
    h1 = doc.add_paragraph()
    set_paragraph_format(h1, first_line_indent=False, space_before=15, space_after=10)
    run = h1.add_run("一、案件核心难度与专属工作量分析")
    run.font.name = "黑体"
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    # （一）案件核心难度
    h2 = doc.add_paragraph()
    set_paragraph_format(h2, first_line_indent=False, space_before=10, space_after=8)
    run = h2.add_run("（一）案件核心难度")
    run.font.name = "黑体"
    run.font.size = Pt(12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    # 难度描述
    difficulty_text = generate_difficulty_text(args.case_type, args.difficulty_level, args.core_difficulties)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=True, space_before=5, space_after=10)
    add_run_with_format(p, difficulty_text)
    
    # （二）律师专属工作量
    h2_2 = doc.add_paragraph()
    set_paragraph_format(h2_2, first_line_indent=False, space_before=10, space_after=8)
    run = h2_2.add_run("（二）律师专属工作量")
    run.font.name = "黑体"
    run.font.size = Pt(12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    work_hours_text = generate_work_hours_text(args.work_hours, args.difficulty_level)
    p2 = doc.add_paragraph()
    set_paragraph_format(p2, first_line_indent=True, space_before=5, space_after=10)
    add_run_with_format(p2, work_hours_text)
    
    # ===== 二、优化后律师费报价方案 =====
    h1_2 = doc.add_paragraph()
    set_paragraph_format(h1_2, first_line_indent=False, space_before=15, space_after=10)
    run = h1_2.add_run("二、优化后律师费报价方案")
    run.font.name = "黑体"
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    # 方案一：正常收费标准
    h2_plan1 = doc.add_paragraph()
    set_paragraph_format(h2_plan1, first_line_indent=False, space_before=10, space_after=8)
    run = h2_plan1.add_run("方案一：正常收费标准（参照" + args.province + "省律师费收费标准）")
    run.font.name = "黑体"
    run.font.size = Pt(12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    # 收费依据
    p_basis = doc.add_paragraph()
    set_paragraph_format(p_basis, first_line_indent=True, space_before=5, space_after=5)
    add_run_with_format(p_basis, f"收费依据：严格参照《{args.province}省律师事务所律师服务费标准制定指引（试行）》《{args.province}省律师协会计时收费规则》及财产案件分段累进收费相关规定，结合本案实际情况，兼顾收费合规性与案件实际工作量，为本案正常收费标准。")
    
    # 基础收费计算
    p_calc_intro = doc.add_paragraph()
    set_paragraph_format(p_calc_intro, first_line_indent=True, space_before=5, space_after=5)
    add_run_with_format(p_calc_intro, f"基础收费计算（按财产标的额分段累进+复杂案件上浮）：\n结合本案明确财产标的总价值{args.property_amount}万元，严格参照{args.province}省财产案件政府指导价分段累进计算基础收费；同时因本案属重大疑难复杂案件，按{args.province}省相关规定上浮50%（符合重大疑难复杂案件上浮上限要求），具体计算如下：")
    
    # 计算明细
    base_fee, fee_details = calculate_standard_fee(args.property_amount, args.province)
    advanced_fee = base_fee * 1.5  # 上浮50%
    
    for detail in fee_details:
        p_detail = doc.add_paragraph()
        set_paragraph_format(p_detail, first_line_indent=True, space_before=3, space_after=3)
        add_run_with_format(p_detail, detail)
    
    # 上浮计算
    p_advanced = doc.add_paragraph()
    set_paragraph_format(p_advanced, first_line_indent=True, space_before=5, space_after=5)
    add_run_with_format(p_advanced, f"复杂案件上浮：上述累计基础收费（{base_fee:.2f}万元）×150%={advanced_fee:.2f}万元；\n综上，基础收费合计：{advanced_fee:.2f}万元。")
    
    # 涵盖服务
    p_service1 = doc.add_paragraph()
    set_paragraph_format(p_service1, first_line_indent=True, space_before=5, space_after=10)
    add_run_with_format(p_service1, "涵盖服务：含案情梳理、证据筹备、法律文书起草、庭审代理、调解谈判、成果落地跟进、后续法律咨询，全程遵循专业诉讼策略，保障核心诉求推进。")
    
    # 方案二：基础固定收费 + 核心成果激励收费
    h2_plan2 = doc.add_paragraph()
    set_paragraph_format(h2_plan2, first_line_indent=False, space_before=10, space_after=8)
    run = h2_plan2.add_run("方案二：基础固定收费 + 核心成果激励收费（首选·权责清晰）")
    run.font.name = "黑体"
    run.font.size = Pt(12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    base_fee2, perfect_bonus, core_bonus = calculate_advanced_fee(args.property_amount, args.difficulty_level)
    
    p_plan2_1 = doc.add_paragraph()
    set_paragraph_format(p_plan2_1, first_line_indent=True, space_before=5, space_after=5)
    add_run_with_format(p_plan2_1, f"基础固定服务费：人民币 {base_fee2:.2f} 万元\n涵盖全案：案情深度梳理、证据专项筹备、法律文书起草、诉讼策略制定、全程庭审代理、常规调解谈判、基础法律意见出具。")
    
    p_plan2_2 = doc.add_paragraph()
    set_paragraph_format(p_plan2_2, first_line_indent=True, space_before=5, space_after=5)
    add_run_with_format(p_plan2_2, f"核心成果激励费（达成诉求后支付）\n完美达标（核心诉求全部实现）：支付激励费 = 财产分割总标的额 ×0.6% = {perfect_bonus:.2f} 万元；\n核心达标（实现单一核心诉求）：支付激励费 = 财产分割总标的额 ×0.3% = {core_bonus:.2f} 万元。")
    
    # 方案三：分阶段递进收费
    h2_plan3 = doc.add_paragraph()
    set_paragraph_format(h2_plan3, first_line_indent=False, space_before=10, space_after=8)
    run = h2_plan3.add_run("方案三：分阶段递进收费 + 核心目标激励费（灵活·风险共担）")
    run.font.name = "黑体"
    run.font.size = Pt(12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    stage1, stage2, stage3 = calculate_stage_fee(args.property_amount, args.difficulty_level)
    
    p_plan3 = doc.add_paragraph()
    set_paragraph_format(p_plan3, first_line_indent=True, space_before=5, space_after=5)
    add_run_with_format(p_plan3, f"第一阶段·立案与证据筹备（签订委托即付）：{stage1:.2f} 万元\n服务：案情分析、证据收集/公证/保全、起诉状起草、法院立案。\n\n第二阶段·庭前准备与庭审代理（开庭前7日付）：{stage2:.2f} 万元\n服务：庭审策略定制、专项文书撰写、全程庭审、举证质证、专项辩论、多轮调解。\n\n第三阶段·成果落地与执行跟进（生效后付）：{stage3:.2f} 万元\n服务：财产过户、股权变更、债务厘清、执行跟进、后续咨询。")
    
    # ===== 三、其他重要说明 =====
    h1_3 = doc.add_paragraph()
    set_paragraph_format(h1_3, first_line_indent=False, space_before=15, space_after=10)
    run = h1_3.add_run("三、其他重要说明")
    run.font.name = "黑体"
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    other_notes = [
        "第三方费用：案件受理费、保全费、财产评估费、证据公证费等，由委托人自行承担，实报实销。",
        "付费规则：签订委托合同后支付首期费用，后续费用按方案约定节点支付；激励费在核心诉求达成后3日内支付。",
        "服务保障：全程由专业律师主办，全力保障委托人核心诉求实现。",
        "保密义务：双方应对合作内容严格保密，未经同意不得向第三方披露。",
    ]
    
    for note in other_notes:
        p_note = doc.add_paragraph()
        set_paragraph_format(p_note, first_line_indent=True, space_before=5, space_after=5)
        add_run_with_format(p_note, note)
    
    # ===== 落款 =====
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 律所名称
    p_firm = doc.add_paragraph()
    p_firm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(p_firm, first_line_indent=False, space_before=10, space_after=10)
    add_run_with_format(p_firm, args.law_firm, font_size=14, bold=True)
    
    # 律师姓名
    p_lawyer = doc.add_paragraph()
    p_lawyer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(p_lawyer, first_line_indent=False, space_before=5, space_after=5)
    add_run_with_format(p_lawyer, args.lawyer_name + " 律师", font_size=14, bold=True)
    
    # 日期
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(p_date, first_line_indent=False, space_before=5, space_after=5)
    date_str = datetime.now().strftime("%Y年%m月%d日")
    add_run_with_format(p_date, date_str, font_size=12)
    
    return doc


def main():
    parser = argparse.ArgumentParser(description="律师费报价方案Word文档生成器")
    
    # 必需参数
    parser.add_argument("--case-type", required=True, help="案件类型（如：离婚纠纷、合同纠纷、劳动争议等）")
    parser.add_argument("--property-amount", type=float, required=True, help="财产标的额（万元）")
    parser.add_argument("--law-firm", required=True, help="律所名称")
    parser.add_argument("--lawyer-name", required=True, help="律师姓名")
    
    # 可选参数
    parser.add_argument("--difficulty-level", type=int, default=3, choices=[1, 2, 3, 4, 5], help="难度等级（1-5），默认3")
    parser.add_argument("--province", default="浙江", help="省份（用于参照收费标准），默认浙江")
    parser.add_argument("--core-difficulties", help="核心难点（JSON数组格式）")
    parser.add_argument("--work-hours", help="工作量估算（JSON格式）")
    parser.add_argument("--output-path", default="./律师费报价方案.docx", help="输出文件路径")
    
    args = parser.parse_args()
    
    # 解析JSON参数
    if args.core_difficulties:
        try:
            args.core_difficulties = json.loads(args.core_difficulties)
        except json.JSONDecodeError:
            print("错误：core-difficulties 参数格式错误，请使用JSON数组格式")
            sys.exit(1)
    
    if args.work_hours:
        try:
            args.work_hours = json.loads(args.work_hours)
        except json.JSONDecodeError:
            print("错误：work-hours 参数格式错误，请使用JSON格式")
            sys.exit(1)
    
    # 创建文档
    doc = create_document(args)
    
    # 保存文档
    output_path = args.output_path
    doc.save(output_path)
    
    # 输出结果
    result = {
        "status": "success",
        "file_path": os.path.abspath(output_path),
        "message": f"律师费报价方案已生成：{os.path.abspath(output_path)}"
    }
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
