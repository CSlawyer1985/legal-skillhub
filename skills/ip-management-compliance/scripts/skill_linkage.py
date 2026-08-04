# -*- coding: utf-8 -*-
"""
技能联动模块 (Skill Linkage Module)
功能：ip-management-compliance 内部独立技能间及与 patent-infringement-guide 联动
支持数据传递、联动任务生成、深度分析触发
"""

import json
import os
from dataclasses import dataclass, field, asdict
from datetime import datetime
from typing import List, Dict, Optional, Any
from enum import Enum
from pathlib import Path


class RiskLevel(Enum):
    """风险等级枚举"""
    LOW = "低风险"
    MEDIUM = "中风险"
    HIGH = "高风险"
    CRITICAL = "严重风险"


class LinkageTrigger(Enum):
    """联动触发类型"""
    FTO_MEDIUM_RISK = "fto_medium_risk"        # FTO分析发现中风险
    FTO_HIGH_RISK = "fto_high_risk"           # FTO分析发现高风险
    STABILITY_LOW = "stability_low"           # 稳定性评估较低
    INVALID_ANALYSIS = "invalid_analysis"     # 无效分析需要深度侵权比对
    USER_REQUEST = "user_request"             # 用户主动请求联动


@dataclass
class PatentInfo:
    """专利信息（用于联动传递）"""
    patent_no: str = ""           # 专利号
    application_no: str = ""       # 申请号
    title: str = ""               # 专利名称
    patent_type: str = ""         # 专利类型（发明/实用新型/外观设计）
    applicant: str = ""           # 申请人
    owner: str = ""               # 权利人
    application_date: str = ""     # 申请日
    publication_date: str = ""    # 公开/公告日
    ipc_code: str = ""            # IPC分类号
    legal_status: str = ""        # 法律状态
    source_url: str = ""           # 原文链接
    abstract: str = ""            # 摘要
    claims: str = ""              # 权利要求（可选）


@dataclass
class ComparisonFile:
    """对比文件（用于侵权分析）"""
    seq_no: int = 0               # 序号
    patent_info: PatentInfo = field(default_factory=PatentInfo)
    similarity: float = 0.0       # 相似度
    risk_level: str = ""          # 风险等级
    matched_features: List[str] = field(default_factory=list)  # 匹配的技术特征
    missing_features: List[str] = field(default_factory=list)  # 缺失的技术特征
    equivalent_features: List[str] = field(default_factory=list)  # 等同特征
    analysis_notes: str = ""      # 分析备注


@dataclass
class FTOResult:
    """FTO分析结果（用于联动）"""
    target_patent: PatentInfo = field(default_factory=PatentInfo)  # 被分析专利
    target_technical_scheme: str = ""  # 被控技术方案描述
    comparison_files: List[ComparisonFile] = field(default_factory=list)  # 对比文件列表
    overall_risk_level: str = ""   # 总体风险等级
    risk_summary: str = ""         # 风险摘要
    key_findings: List[str] = field(default_factory=list)  # 关键发现
    linkage_recommended: bool = False  # 是否建议联动
    linkage_trigger: str = ""      # 联动触发类型


@dataclass
class LinkageTask:
    """联动任务数据结构"""
    task_id: str = ""             # 任务ID（时间戳+随机数）
    source_skill: str = "ip-management-compliance"  # 源技能
    target_skill: str = "patent-infringement-guide"  # 目标技能
    trigger_type: str = ""        # 触发类型
    created_at: str = ""           # 创建时间
    fto_result: Optional[FTOResult] = None  # FTO分析结果
    status: str = "pending"       # 任务状态：pending/processing/completed
    notes: str = ""               # 备注


class SkillLinkage:
    """技能联动管理器"""

    def __init__(self, workspace_dir: str = None):
        """
        初始化联动管理器

        Args:
            workspace_dir: 工作区目录，用于存储联动数据
        """
        if workspace_dir is None:
            workspace_dir = str(Path(__file__).resolve().parent / "linkage_data")
        self.workspace_dir = workspace_dir
        os.makedirs(self.workspace_dir, exist_ok=True)

    def generate_task_id(self) -> str:
        """生成唯一任务ID"""
        import hashlib
        import time
        timestamp = str(time.time()).encode()
        return hashlib.md5(timestamp).hexdigest()[:12]

    def create_linkage_task(
        self,
        fto_result: FTOResult,
        trigger_type: LinkageTrigger = None
    ) -> LinkageTask:
        """
        创建联动任务

        Args:
            fto_result: FTO分析结果
            trigger_type: 联动触发类型

        Returns:
            LinkageTask: 联动任务对象
        """
        if trigger_type is None:
            # 根据风险等级自动判断触发类型
            risk = fto_result.overall_risk_level
            if risk in ["高风险", "严重风险"]:
                trigger_type = LinkageTrigger.FTO_HIGH_RISK
            elif risk == "中风险":
                trigger_type = LinkageTrigger.FTO_MEDIUM_RISK
            else:
                # 低风险不自动触发联动
                return None

        task = LinkageTask(
            task_id=self.generate_task_id(),
            source_skill="ip-management-compliance",
            target_skill="patent-infringement-guide",
            trigger_type=trigger_type.value,
            created_at=datetime.now().isoformat(),
            fto_result=fto_result,
            status="pending"
        )

        # 判断是否需要联动
        if fto_result.overall_risk_level in ["中风险", "高风险", "严重风险"]:
            task.linkage_recommended = True

        return task

    def save_linkage_task(self, task: LinkageTask) -> str:
        """
        保存联动任务到文件

        Args:
            task: 联动任务对象

        Returns:
            str: 任务文件路径
        """
        task_file = os.path.join(self.workspace_dir, f"linkage_task_{task.task_id}.json")

        # 转换为可序列化格式
        task_dict = {
            "task_id": task.task_id,
            "source_skill": task.source_skill,
            "target_skill": task.target_skill,
            "trigger_type": task.trigger_type,
            "created_at": task.created_at,
            "status": task.status,
            "notes": task.notes,
            "linkage_recommended": task.linkage_recommended,
            "fto_result": None
        }

        if task.fto_result:
            fto_dict = self._fto_to_dict(task.fto_result)
            task_dict["fto_result"] = fto_dict

        with open(task_file, 'w', encoding='utf-8') as f:
            json.dump(task_dict, f, ensure_ascii=False, indent=2)

        return task_file

    def _fto_to_dict(self, fto_result: FTOResult) -> Dict:
        """将FTO结果转换为字典"""
        result = {
            "target_patent": asdict(fto_result.target_patent),
            "target_technical_scheme": fto_result.target_technical_scheme,
            "overall_risk_level": fto_result.overall_risk_level,
            "risk_summary": fto_result.risk_summary,
            "key_findings": fto_result.key_findings,
            "linkage_recommended": fto_result.linkage_recommended,
            "comparison_files": []
        }

        for cf in fto_result.comparison_files:
            cf_dict = {
                "seq_no": cf.seq_no,
                "patent_info": asdict(cf.patent_info),
                "similarity": cf.similarity,
                "risk_level": cf.risk_level,
                "matched_features": cf.matched_features,
                "missing_features": cf.missing_features,
                "equivalent_features": cf.equivalent_features,
                "analysis_notes": cf.analysis_notes
            }
            result["comparison_files"].append(cf_dict)

        return result

    def load_linkage_task(self, task_id: str) -> Optional[LinkageTask]:
        """加载联动任务"""
        task_file = os.path.join(self.workspace_dir, f"linkage_task_{task_id}.json")
        if not os.path.exists(task_file):
            return None

        with open(task_file, 'r', encoding='utf-8') as f:
            task_dict = json.load(f)

        return self._dict_to_task(task_dict)

    def _dict_to_task(self, task_dict: Dict) -> LinkageTask:
        """将字典转换为联动任务"""
        task = LinkageTask(
            task_id=task_dict["task_id"],
            source_skill=task_dict["source_skill"],
            target_skill=task_dict["target_skill"],
            trigger_type=task_dict["trigger_type"],
            created_at=task_dict["created_at"],
            status=task_dict["status"],
            notes=task_dict.get("notes", ""),
            linkage_recommended=task_dict.get("linkage_recommended", False)
        )

        if "fto_result" in task_dict and task_dict["fto_result"]:
            fto_dict = task_dict["fto_result"]
            fto_result = FTOResult(
                target_technical_scheme=fto_dict.get("target_technical_scheme", ""),
                overall_risk_level=fto_dict.get("overall_risk_level", ""),
                risk_summary=fto_dict.get("risk_summary", ""),
                key_findings=fto_dict.get("key_findings", []),
                linkage_recommended=fto_dict.get("linkage_recommended", False)
            )

            # 解析目标专利
            if "target_patent" in fto_dict:
                tp_dict = fto_dict["target_patent"]
                fto_result.target_patent = PatentInfo(**tp_dict)

            # 解析对比文件
            if "comparison_files" in fto_dict:
                for cf_dict in fto_dict["comparison_files"]:
                    cf = ComparisonFile(
                        seq_no=cf_dict.get("seq_no", 0),
                        similarity=cf_dict.get("similarity", 0.0),
                        risk_level=cf_dict.get("risk_level", ""),
                        matched_features=cf_dict.get("matched_features", []),
                        missing_features=cf_dict.get("missing_features", []),
                        equivalent_features=cf_dict.get("equivalent_features", []),
                        analysis_notes=cf_dict.get("analysis_notes", "")
                    )
                    if "patent_info" in cf_dict:
                        cf.patent_info = PatentInfo(**cf_dict["patent_info"])
                    fto_result.comparison_files.append(cf)

            task.fto_result = fto_result

        return task

    def list_pending_tasks(self) -> List[Dict]:
        """列出所有待处理的联动任务"""
        tasks = []
        for fname in os.listdir(self.workspace_dir):
            if fname.startswith("linkage_task_") and fname.endswith(".json"):
                task_path = os.path.join(self.workspace_dir, fname)
                with open(task_path, 'r', encoding='utf-8') as f:
                    task_dict = json.load(f)
                    if task_dict["status"] == "pending":
                        tasks.append({
                            "task_id": task_dict["task_id"],
                            "trigger_type": task_dict["trigger_type"],
                            "created_at": task_dict["created_at"],
                            "risk_level": task_dict.get("fto_result", {}).get("overall_risk_level", "未知"),
                            "task_file": task_path
                        })
        return tasks

    def generate_xml_for_infringement(self, task: LinkageTask) -> str:
        """
        生成 patent-infringement-guide 格式的XML数据

        Args:
            task: 联动任务

        Returns:
            str: XML格式字符串
        """
        if not task.fto_result:
            return ""

        xml_lines = ['<?xml version="1.0" encoding="UTF-8"?>']
        xml_lines.append('<linkage-data>')
        xml_lines.append('  <source-skill>ip-management-compliance</source-skill>')
        xml_lines.append(f'  <task-id>{task.task_id}</task-id>')
        xml_lines.append(f'  <trigger-type>{task.trigger_type}</trigger-type>')
        xml_lines.append(f'  <created-at>{task.created_at}</created-at>')

        # 被控技术方案
        xml_lines.append('  <target-patent>')
        tp = task.fto_result.target_patent
        xml_lines.append(f'    <patent-no>{tp.patent_no}</patent-no>')
        xml_lines.append(f'    <title>{tp.title}</title>')
        xml_lines.append(f'    <patent-type>{tp.patent_type}</patent-type>')
        xml_lines.append(f'    <applicant>{tp.applicant}</applicant>')
        xml_lines.append(f'    <ipc-code>{tp.ipc_code}</ipc-code>')
        xml_lines.append(f'    <abstract>{tp.abstract}</abstract>')
        xml_lines.append('  </target-patent>')

        xml_lines.append(f'  <technical-scheme>{task.fto_result.target_technical_scheme}</technical-scheme>')

        # 风险摘要
        xml_lines.append('  <risk-summary>')
        xml_lines.append(f'    <overall-level>{task.fto_result.overall_risk_level}</overall-level>')
        xml_lines.append(f'    <summary>{task.fto_result.risk_summary}</summary>')
        for finding in task.fto_result.key_findings:
            xml_lines.append(f'    <finding>{finding}</finding>')
        xml_lines.append('  </risk-summary>')

        # 对比文件
        xml_lines.append('  <comparison-files>')
        for cf in task.fto_result.comparison_files:
            xml_lines.append('    <file>')
            xml_lines.append(f'      <seq-no>{cf.seq_no}</seq-no>')
            pi = cf.patent_info
            xml_lines.append('      <patent-info>')
            xml_lines.append(f'        <patent-no>{pi.patent_no}</patent-no>')
            xml_lines.append(f'        <title>{pi.title}</title>')
            xml_lines.append(f'        <patent-type>{pi.patent_type}</patent-type>')
            xml_lines.append(f'        <applicant>{pi.applicant}</applicant>')
            xml_lines.append(f'        <application-date>{pi.application_date}</application-date>')
            xml_lines.append(f'        <publication-date>{pi.publication_date}</publication-date>')
            xml_lines.append(f'        <ipc-code>{pi.ipc_code}</ipc-code>')
            xml_lines.append(f'        <legal-status>{pi.legal_status}</legal-status>')
            xml_lines.append(f'        <source-url>{pi.source_url}</source-url>')
            xml_lines.append('      </patent-info>')
            xml_lines.append(f'      <similarity>{cf.similarity}</similarity>')
            xml_lines.append(f'      <risk-level>{cf.risk_level}</risk-level>')
            for mf in cf.matched_features:
                xml_lines.append(f'      <matched-feature>{mf}</matched-feature>')
            for msf in cf.missing_features:
                xml_lines.append(f'      <missing-feature>{msf}</missing-feature>')
            for ef in cf.equivalent_features:
                xml_lines.append(f'      <equivalent-feature>{ef}</equivalent-feature>')
            xml_lines.append(f'      <analysis-notes>{cf.analysis_notes}</analysis-notes>')
            xml_lines.append('    </file>')
        xml_lines.append('  </comparison-files>')

        xml_lines.append('</linkage-data>')

        return '\n'.join(xml_lines)

    def export_linkage_report(self, task: LinkageTask, output_path: str = None) -> str:
        """
        导出联动任务报告（Word格式）

        Args:
            task: 联动任务
            output_path: 输出路径

        Returns:
            str: 报告文件路径
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if output_path is None:
            output_path = os.path.join(
                self.workspace_dir,
                f"linkage_report_{task.task_id}.docx"
            )

        doc = Document()

        # 标题
        title = doc.add_heading('', level=0)
        run = title.add_run('专利分析与侵权联动报告')
        run.font.size = Pt(22)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'任务ID: {task.task_id}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # 基本信息
        doc.add_heading('一、联动任务信息', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        info = [
            ('源技能', task.source_skill),
            ('目标技能', task.target_skill),
            ('触发类型', task.trigger_type),
            ('创建时间', task.created_at),
            ('任务状态', task.status)
        ]
        for i, (k, v) in enumerate(info):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v

        if task.fto_result:
            # 风险摘要
            doc.add_heading('二、风险摘要', level=1)
            p = doc.add_paragraph()
            p.add_run(f"总体风险等级: ").bold = True
            p.add_run(task.fto_result.overall_risk_level)

            p = doc.add_paragraph()
            p.add_run(f"风险摘要: {task.fto_result.risk_summary}")

            if task.fto_result.key_findings:
                doc.add_heading('关键发现', level=2)
                for finding in task.fto_result.key_findings:
                    doc.add_paragraph(finding, style='List Bullet')

            # 对比文件列表
            if task.fto_result.comparison_files:
                doc.add_heading('三、对比文件列表', level=1)
                table = doc.add_table(rows=len(task.fto_result.comparison_files) + 1, cols=5)
                table.style = 'Table Grid'
                headers = ['序号', '对比专利号', '专利名称', '相似度', '风险等级']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

                for i, cf in enumerate(task.fto_result.comparison_files):
                    row = table.rows[i + 1]
                    row.cells[0].text = str(cf.seq_no)
                    row.cells[1].text = cf.patent_info.patent_no
                    row.cells[2].text = cf.patent_info.title
                    row.cells[3].text = f"{cf.similarity:.2%}"
                    row.cells[4].text = cf.risk_level

        # 联动建议
        doc.add_heading('四、联动建议', level=1)
        if task.linkage_recommended:
            doc.add_paragraph(
                '根据FTO分析结果，建议启动 patent-infringement-guide 进行深度侵权分析。',
                style='List Bullet'
            )
            doc.add_paragraph(
                '可使用以下方式启动联动分析：',
                style='List Bullet'
            )
            doc.add_paragraph(
                f'1. 读取联动数据: 任务ID = {task.task_id}',
                style='List Number'
            )
            doc.add_paragraph(
                '2. 使用XML格式导入被控技术方案进行深度等同判断',
                style='List Number'
            )
            doc.add_paragraph(
                '3. 获取规避设计建议和抗辩策略',
                style='List Number'
            )
        else:
            doc.add_paragraph('当前风险等级较低，无需深度联动分析。')

        doc.save(output_path)
        return output_path


# 全局联动管理器实例
_linkage_manager = None

def get_linkage_manager(workspace_dir: str = None) -> SkillLinkage:
    """获取全局联动管理器实例"""
    global _linkage_manager
    if _linkage_manager is None:
        _linkage_manager = SkillLinkage(workspace_dir)
    return _linkage_manager


def auto_trigger_linkage(fto_result: FTOResult) -> Optional[LinkageTask]:
    """
    自动触发联动（根据FTO结果）

    Args:
        fto_result: FTO分析结果

    Returns:
        Optional[LinkageTask]: 联动任务（如果需要联动），否则返回None
    """
    manager = get_linkage_manager()

    # 判断是否需要联动
    if fto_result.overall_risk_level not in ["中风险", "高风险", "严重风险"]:
        return None

    # 确定触发类型
    if fto_result.overall_risk_level in ["高风险", "严重风险"]:
        trigger_type = LinkageTrigger.FTO_HIGH_RISK
    else:
        trigger_type = LinkageTrigger.FTO_MEDIUM_RISK

    # 创建并保存任务
    task = manager.create_linkage_task(fto_result, trigger_type)
    if task:
        task_file = manager.save_linkage_task(task)
        task.notes = f"任务已保存至: {task_file}"
        return task

    return None
