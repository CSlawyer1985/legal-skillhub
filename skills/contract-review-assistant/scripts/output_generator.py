#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多格式输出生成器
支持Word、PDF、HTML、Markdown等多种输出格式
"""

import os
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime


class OutputGenerator:
    """多格式输出生成器基类"""
    
    def __init__(self, output_dir: str = '.'):
        """
        初始化输出生成器
        
        Args:
            output_dir: 输出目录
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate(self, review_result: Dict, output_path: str, format: str = 'docx'):
        """
        生成指定格式的输出文件
        
        Args:
            review_result: 审查结果字典
            output_path: 输出文件路径
            format: 输出格式 ('docx', 'pdf', 'html', 'md')
        """
        format = format.lower()
        
        if format == 'docx':
            return self._generate_word(review_result, output_path)
        elif format == 'pdf':
            return self._generate_pdf(review_result, output_path)
        elif format == 'html':
            return self._generate_html(review_result, output_path)
        elif format == 'md':
            return self._generate_markdown(review_result, output_path)
        else:
            raise ValueError(f"不支持的输出格式: {format}")
    
    def _generate_word(self, result: Dict, output_path: str) -> str:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 标题
            title = doc.add_heading('合同审查报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 基本信息
            doc.add_heading('一、基本信息', 1)
            basic_info = [
                ('合同类型', result.get('contract_type', '未指定')),
                ('审查日期', datetime.now().strftime('%Y年%m月%d日')),
                ('审查人员', 'AI辅助审查系统')
            ]
            
            for label, value in basic_info:
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)
            
            # 当事人信息
            doc.add_heading('二、当事人信息', 1)
            parties = result.get('parties', [])
            if parties:
                for i, party in enumerate(parties, 1):
                    doc.add_paragraph(f'{i}. {party}')
            else:
                doc.add_paragraph('未识别到当事人信息')
            
            # 条款审查详情
            doc.add_heading('三、条款审查详情', 1)
            clauses = result.get('clauses', [])
            
            for clause in clauses:
                # 条款标题
                heading = doc.add_heading(level=2)
                clause_type = clause.get('clause_type', '未知条款')
                risk_level = clause.get('risk_level', '低风险')
                
                heading_text = f'{clause_type} 【{risk_level}】'
                heading.text = heading_text
                
                # 条款内容
                doc.add_paragraph(f"条款内容：{clause.get('content', '')}")
                
                # 风险分析
                risk_analysis = clause.get('risk_analysis', '无')
                p = doc.add_paragraph(f"风险分析：{risk_analysis}")
                p.paragraph_format.space_after = Pt(6)
                
                # 修改建议
                suggestion = clause.get('suggestion', '无')
                p = doc.add_paragraph(f"修改建议：{suggestion}")
                p.paragraph_format.space_after = Pt(12)
            
            # 风险摘要
            doc.add_heading('四、风险摘要', 1)
            risks = result.get('risks', [])
            if risks:
                for risk in risks:
                    doc.add_paragraph(f'• {risk}', style='List Bullet')
            else:
                doc.add_paragraph('未发现明显风险')
            
            # 审查建议
            doc.add_heading('五、审查建议', 1)
            suggestions = result.get('suggestions', [])
            if suggestions:
                for i, suggestion in enumerate(suggestions, 1):
                    doc.add_paragraph(f'{i}. {suggestion}')
            else:
                doc.add_paragraph('无特别建议')
            
            # 法律依据
            legal_refs = result.get('legal_references', {})
            if legal_refs:
                doc.add_heading('六、相关法律依据', 1)
                for keyword, ref_info in legal_refs.items():
                    heading = doc.add_heading(level=2)
                    heading.text = f'涉及"{keyword}"的相关规定'
                    
                    p = doc.add_paragraph()
                    p.add_run(f'法律来源：').bold = True
                    p.add_run(ref_info.get('law', '未知'))
                    
                    p = doc.add_paragraph()
                    p.add_run(f'相关条文：').bold = True
                    p.add_run(ref_info.get('snippet', ''))
                    p.paragraph_format.space_after = Pt(12)
            
            # 页脚
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('—' * 30).bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer = doc.add_paragraph()
            footer.add_run('本报告由AI辅助审查系统生成，仅供参考。').italic = True
            footer.add_run('\n最终解释权归审查人员所有。').italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存文档
            doc.save(output_path)
            return output_path
            
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
    
    def _generate_pdf(self, result: Dict, output_path: str) -> str:
        """生成PDF文档"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import cm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # 注册中文字体（使用系统字体）
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                font_name = 'SimSun'
            except:
                font_name = 'Helvetica'  # 降级方案
            
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            # 标题
            c.setFont(font_name, 16)
            c.drawString(2*cm, height-2*cm, '合同审查报告')
            
            # 基本信息
            y = height - 3*cm
            c.setFont(font_name, 12)
            c.drawString(2*cm, y, f"合同类型：{result.get('contract_type', '未指定')}")
            y -= 0.5*cm
            c.drawString(2*cm, y, f"审查日期：{datetime.now().strftime('%Y年%m月%d日')}")
            
            # 当事人信息
            y -= 1*cm
            c.setFont(font_name, 14)
            c.drawString(2*cm, y, '当事人信息：')
            c.setFont(font_name, 10)
            for i, party in enumerate(result.get('parties', []), 1):
                y -= 0.5*cm
                c.drawString(2.5*cm, y, f'{i}. {party}')
            
            # 保存PDF
            c.save()
            return output_path
            
        except ImportError:
            raise ImportError("需要安装reportlab库: pip install reportlab")
    
    def _generate_html(self, result: Dict, output_path: str) -> str:
        """生成HTML文档"""
        html_content = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>合同审查报告</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            background-color: white;
            padding: 30px;
            border-radius: 5px;
        }}
        h1 {{
            color: #003366;
            border-bottom: 2px solid #003366;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #0066cc;
            margin-top: 20px;
        }}
        .clause {{
            border-left: 3px solid #0066cc;
            padding-left: 15px;
            margin: 10px 0;
        }}
        .high-risk {{ color: #ff0000; }}
        .medium-risk {{ color: #ff9900; }}
        .low-risk {{ color: #009900; }}
        .footer {{
            text-align: center;
            font-style: italic;
            color: #666;
            margin-top: 30px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>合同审查报告</h1>
        
        <h2>一、基本信息</h2>
        <p><strong>合同类型：</strong>{result.get('contract_type', '未指定')}</p>
        <p><strong>审查日期：</strong>{datetime.now().strftime('%Y年%m月%d日')}</p>
        <p><strong>审查人员：</strong>AI辅助审查系统</p>
        
        <h2>二、当事人信息</h2>
        <ul>
"""
        
        # 添加当事人
        for i, party in enumerate(result.get('parties', []), 1):
            html_content += f'            <li>{party}</li>\n'
        
        html_content += """
        </ul>
        
        <h2>三、条款审查详情</h2>
"""
        
        # 添加条款审查
        for clause in result.get('clauses', []):
            risk_level = clause.get('risk_level', '低风险')
            risk_class = {
                '高风险': 'high-risk',
                '中风险': 'medium-risk',
                '低风险': 'low-risk'
            }.get(risk_level, 'low-risk')
            
            html_content += f"""
        <div class="clause">
            <h3>{clause.get('clause_type', '未知条款')} - <span class="{risk_class}">{risk_level}</span></h3>
            <p><strong>条款内容：</strong>{clause.get('content', '')}</p>
            <p><strong>风险分析：</strong>{clause.get('risk_analysis', '无')}</p>
            <p><strong>修改建议：</strong>{clause.get('suggestion', '无')}</p>
        </div>
"""
        
        html_content += """
        <h2>四、风险摘要</h2>
        <ul>
"""
        
        # 添加风险摘要
        for risk in result.get('risks', []):
            html_content += f'            <li>{risk}</li>\n'
        
        html_content += """
        </ul>
        
        <h2>五、审查建议</h2>
        <ol>
"""
        
        # 添加审查建议
        for suggestion in result.get('suggestions', []):
            html_content += f'            <li>{suggestion}</li>\n'
        
        # 添加法律依据
        legal_refs = result.get('legal_references', {})
        if legal_refs:
            html_content += """
        </ol>
        
        <h2>六、相关法律依据</h2>
"""
            for keyword, ref_info in legal_refs.items():
                html_content += f"""
        <div class="clause">
            <h3>涉及"{keyword}"的相关规定</h3>
            <p><strong>法律来源：</strong>{ref_info.get('law', '未知')}</p>
            <p><strong>相关条文：</strong>{ref_info.get('snippet', '')}</p>
        </div>
"""
        
        html_content += """
        <div class="footer">
            <p>本报告由AI辅助审查系统生成，仅供参考。</p>
            <p>最终解释权归审查人员所有。</p>
        </div>
    </div>
</body>
</html>
"""
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    def _generate_markdown(self, result: Dict, output_path: str) -> str:
        """生成Markdown文档"""
        md_content = f"""# 合同审查报告

## 一、基本信息

- **合同类型：** {result.get('contract_type', '未指定')}
- **审查日期：** {datetime.now().strftime('%Y年%m月%d日')}
- **审查人员：** AI辅助审查系统

## 二、当事人信息

"""
        
        # 添加当事人
        for i, party in enumerate(result.get('parties', []), 1):
            md_content += f'{i}. {party}\n'
        
        md_content += '\n## 三、条款审查详情\n\n'
        
        # 添加条款审查
        for clause in result.get('clauses', []):
            risk_level = clause.get('risk_level', '低风险')
            md_content += f"### {clause.get('clause_type', '未知条款')} - **{risk_level}**\n\n"
            md_content += f"**条款内容：** {clause.get('content', '')}\n\n"
            md_content += f"**风险分析：** {clause.get('risk_analysis', '无')}\n\n"
            md_content += f"**修改建议：** {clause.get('suggestion', '无')}\n\n"
        
        md_content += '## 四、风险摘要\n\n'
        
        # 添加风险摘要
        for risk in result.get('risks', []):
            md_content += f'- {risk}\n'
        
        md_content += '\n## 五、审查建议\n\n'
        
        # 添加审查建议
        for i, suggestion in enumerate(result.get('suggestions', []), 1):
            md_content += f'{i}. {suggestion}\n'
        
        # 添加法律依据
        legal_refs = result.get('legal_references', {})
        if legal_refs:
            md_content += '\n## 六、相关法律依据\n\n'
            for keyword, ref_info in legal_refs.items():
                md_content += f"### 涉及\"{keyword}\"的相关规定\n\n"
                md_content += f"**法律来源：** {ref_info.get('law', '未知')}\n\n"
                md_content += f"**相关条文：** {ref_info.get('snippet', '')}\n\n"
        
        md_content += f"""

---

*本报告由AI辅助审查系统生成，仅供参考。*
*最终解释权归审查人员所有。*

生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return output_path


def create_output_generator(output_dir: str = '.') -> OutputGenerator:
    """
    创建输出生成器
    
    Args:
        output_dir: 输出目录
        
    Returns:
        输出生成器实例
    """
    return OutputGenerator(output_dir)


if __name__ == '__main__':
    # 测试代码
    print('=' * 60)
    print('多格式输出生成器测试')
    print('=' * 60)
    
    # 模拟审查结果
    test_result = {
        'contract_type': '软件开发合同',
        'parties': ['甲方：深圳某某科技有限公司', '乙方：北京某某技术公司'],
        'clauses': [
            {
                'clause_type': '违约责任',
                'content': '任何一方违约，需支付合同总价30%的违约金',
                'risk_level': '高风险',
                'risk_analysis': '违约金比例过高，可能不被法院支持',
                'suggestion': '建议调整为实际损失的1.3倍或不超过合同总价的20%'
            }
        ],
        'risks': ['违约责任条款存在高风险'],
        'suggestions': ['建议修改违约金条款', '建议增加不可抗力条款']
    }
    
    # 创建输出生成器
    generator = create_output_generator()
    
    # 测试各种格式
    formats = ['docx', 'html', 'md']
    for fmt in formats:
        output_file = f'test_report.{fmt}'
        try:
            generator.generate(test_result, output_file, fmt)
            print(f'✓ 生成{fmt.upper()}格式成功: {output_file}')
        except Exception as e:
            print(f'✗ 生成{fmt.upper()}格式失败: {e}')
    
    print('\n' + '=' * 60)
    print('提示：生成PDF需要安装reportlab库')
    print('pip install reportlab')
    print('=' * 60)
