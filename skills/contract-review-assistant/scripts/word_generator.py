#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档生成模块
生成专业的合同审查报告
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional
import os


class WordGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10.5)
        
        # 标题样式
        heading1 = self.doc.styles['Heading 1']
        heading1.font.name = '微软雅黑'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 51, 102)
        
        heading2 = self.doc.styles['Heading 2']
        heading2.font.name = '微软雅黑'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 102, 204)
    
    def generate_review_report(self, review_result: dict, output_path: str) -> str:
        """
        生成审查报告
        
        Args:
            review_result: 审查结果字典，包含：
                - contract_type: 合同类型
                - parties: 当事人列表
                - clauses: 条款列表
                - risks: 风险列表
                - suggestions: 建议列表
            output_path: 输出文件路径
        
        Returns:
            生成的文件路径
        """
        # 添加标题
        self.doc.add_heading('合同审查报告', 0)
        
        # 添加基本信息
        self._add_basic_info(review_result)
        
        # 添加当事人信息
        self._add_parties_info(review_result.get('parties', []))
        
        # 添加条款审查详情
        self._add_clauses_review(review_result.get('clauses', []))
        
        # 添加风险摘要
        self._add_risks_summary(review_result.get('risks', []))
        
        # 添加审查建议
        self._add_suggestions(review_result.get('suggestions', []))
        
        # 添加页脚
        self._add_footer()
        
        # 保存文档
        self.doc.save(output_path)
        return output_path
    
    def _add_basic_info(self, result: dict):
        """添加基本信息"""
        self.doc.add_heading('一、基本信息', 1)
        
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 合同类型
        table.rows[0].cells[0].text = '合同类型'
        table.rows[0].cells[1].text = result.get('contract_type', '未指定')
        
        # 审查日期
        from datetime import datetime
        table.rows[1].cells[0].text = '审查日期'
        table.rows[1].cells[1].text = datetime.now().strftime('%Y年%m月%d日')
        
        # 审查人员
        table.rows[2].cells[0].text = '审查人员'
        table.rows[2].cells[1].text = 'AI辅助审查系统'
        
        self.doc.add_paragraph()
    
    def _add_parties_info(self, parties: List[str]):
        """添加当事人信息"""
        self.doc.add_heading('二、当事人信息', 1)
        
        if not parties:
            self.doc.add_paragraph('未识别到当事人信息')
            return
        
        for i, party in enumerate(parties, 1):
            p = self.doc.add_paragraph(f'{i}. {party}')
            p.paragraph_format.space_after = Pt(6)
        
        self.doc.add_paragraph()
    
    def _add_clauses_review(self, clauses: List[dict]):
        """添加条款审查详情"""
        self.doc.add_heading('三、条款审查详情', 1)
        
        if not clauses:
            self.doc.add_paragraph('未发现需要审查的条款')
            return
        
        for clause in clauses:
            # 条款标题
            heading = self.doc.add_heading(level=2)
            heading_text = f"{clause.get('clause_type', '未知条款')} "
            
            # 根据风险等级添加标记
            risk_level = clause.get('risk_level', '低风险')
            if '高' in risk_level:
                heading_text += '【高风险】'
            elif '中' in risk_level:
                heading_text += '【中风险】'
            else:
                heading_text += '【低风险】'
            
            heading.text = heading_text
            
            # 条款内容
            self.doc.add_paragraph(f"条款内容：{clause.get('content', '')}")
            
            # 风险分析
            p = self.doc.add_paragraph(f"风险分析：{clause.get('risk_analysis', '无')}")
            p.paragraph_format.space_after = Pt(6)
            
            # 修改建议
            p = self.doc.add_paragraph(f"修改建议：{clause.get('suggestion', '无')}")
            p.paragraph_format.space_after = Pt(12)
        
        self.doc.add_paragraph()
    
    def _add_risks_summary(self, risks: List[str]):
        """添加风险摘要"""
        self.doc.add_heading('四、风险摘要', 1)
        
        if not risks:
            self.doc.add_paragraph('未发现明显风险')
            return
        
        for risk in risks:
            p = self.doc.add_paragraph(f'• {risk}')
            p.paragraph_format.space_after = Pt(6)
        
        self.doc.add_paragraph()
    
    def _add_suggestions(self, suggestions: List[str]):
        """添加审查建议"""
        self.doc.add_heading('五、审查建议', 1)
        
        if not suggestions:
            self.doc.add_paragraph('无特别建议')
            return
        
        for i, suggestion in enumerate(suggestions, 1):
            p = self.doc.add_paragraph(f'{i}. {suggestion}')
            p.paragraph_format.space_after = Pt(6)
        
        self.doc.add_paragraph()
    
    def _add_footer(self):
        """添加页脚"""
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run('—' * 30).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer = self.doc.add_paragraph()
        footer.add_run('本报告由AI辅助审查系统生成，仅供参考。').italic = True
        footer.add_run('\n最终解释权归审查人员所有。').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_diff_comparison(self, diff_text: str):
        """添加版本差异对比"""
        self.doc.add_heading('附录：版本差异对比', 1)
        
        # 添加差异内容
        for line in diff_text.split('\n'):
            if line.startswith('+'):
                p = self.doc.add_paragraph(line)
                p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif line.startswith('-'):
                p = self.doc.add_paragraph(line)
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色
            else:
                self.doc.add_paragraph(line)


if __name__ == '__main__':
    # 测试代码
    generator = WordGenerator()
    
    # 模拟审查结果
    test_result = {
        'contract_type': '软件开发合同',
        'parties': ['甲方：深圳某某科技有限公司', '乙方：北京某某技术公司'],
        'clauses': [
            {
                'clause_type': '违约责任',
                'content': '任何一方违约，需支付合同总价30%的违约金',
                'risk_level': '高风险',
                'risk_analysis': '违约金比例过高，可能不被法院支持',
                'suggestion': '建议调整为实际损失的1.3倍或不超过合同总价的20%'
            }
        ],
        'risks': ['违约责任条款存在高风险'],
        'suggestions': ['建议修改违约金条款', '建议增加不可抗力条款']
    }
    
    output_file = 'test_review_report.docx'
    generator.generate_review_report(test_result, output_file)
    print(f'测试报告已生成：{output_file}')
