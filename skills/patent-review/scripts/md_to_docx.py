#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_to_docx.py - 将 Markdown 核稿意见通知书转换为 .docx 文件

依赖：python-docx（pip install python-docx）
用法：
    python md_to_docx.py <input.md> [-o <output.docx>]

省略 -o 时输出为 <input>.docx。
输出的 .docx 使用中文字体（宋体/黑体），A4 页面，符合正式核稿意见通知书排版要求。
"""
import argparse
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── 常量 ──────────────────────────────────────────────
FONT_BODY = '宋体'
FONT_HEI = '黑体'
FONT_KAI = '楷体'
SIZE_TITLE = Pt(16)
SIZE_H1 = Pt(14)
SIZE_H2 = Pt(12)
SIZE_H3 = Pt(11)
SIZE_BODY = Pt(10.5)
SIZE_TABLE = Pt(9)

# ── 辅助函数 ──────────────────────────────────────────

def set_font(run, name=FONT_BODY, size=SIZE_BODY, bold=False, color=None):
    """设置 run 的字体、字号、粗细、颜色"""
    run.font.size = size
    run.font.name = name
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph(doc, text, font=FONT_BODY, size=SIZE_BODY, bold=False,
                  alignment=None, space_after=Pt(6), space_before=Pt(0),
                  first_line_indent=None, color=None):
    """添加段落并设置格式"""
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    pf = para.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    if text:
        run = para.add_run(text)
        set_font(run, font, size, bold, color)
    return para


def add_styled_paragraph(doc, segments, alignment=None, space_after=Pt(6),
                         space_before=Pt(0), first_line_indent=None):
    """添加带内联样式的段落（segments 为 [(text, bold, italic, color), ...]）"""
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    pf = para.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    for text, bold, italic, color in segments:
        if not text:
            continue
        run = para.add_run(text)
        set_font(run, FONT_BODY if not bold else FONT_HEI, SIZE_BODY, bold, color)
        if italic:
            run.italic = True
    return para


def set_table_cell(cell, text, bold=False, font=FONT_BODY, size=SIZE_TABLE,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """设置表格单元格文本和格式"""
    # 清空默认段落
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    if text:
        run = p.add_run(text)
        set_font(run, font, size, bold, color)


# ── Markdown 解析 ──────────────────────────────────────

def parse_inline(text):
    """解析行内格式：**粗体**、*斜体*，返回 segments 列表"""
    segments = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|☐|★★★|★★|★|—)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False, False, None))
        raw = m.group(0)
        if raw.startswith('**') and raw.endswith('**'):
            segments.append((m.group(2), True, False, None))
        elif raw.startswith('*') and raw.endswith('*'):
            segments.append((m.group(3), False, True, None))
        else:
            # 特殊符号
            if '★★★' in raw:
                segments.append((raw, True, False, RGBColor(0xCC, 0x00, 0x00)))
            elif '★★' in raw:
                segments.append((raw, True, False, RGBColor(0xFF, 0x66, 0x00)))
            elif '★' in raw:
                segments.append((raw, True, False, RGBColor(0x00, 0x80, 0x00)))
            else:
                segments.append((raw, False, False, None))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False, False, None))
    if not segments:
        segments.append((text, False, False, None))
    return segments


def is_table_separator(line):
    """判断是否是 Markdown 表格分隔行 (| --- | --- |)"""
    return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))


def parse_table_header(line):
    """解析表头行，返回列名列表"""
    return [c.strip() for c in line.strip().split('|')[1:-1]]


def parse_table_row(line):
    """解析表格数据行，返回单元格列表"""
    return [c.strip() for c in line.strip().split('|')[1:-1]]


def is_horizontal_rule(line):
    """判断是否水平分隔线"""
    return re.match(r'^(\-{3,}|\*{3,}|\_{3,})$', line.strip()) is not None


# ── 主转换函数 ────────────────────────────────────────

def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ── 页面设置 A4 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── 设置默认样式 ──
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    i = 0
    in_table = False
    in_code_block = False
    table_rows = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip('\n\r')

        # 代码块
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            add_paragraph(doc, raw.rstrip(), font=FONT_KAI, size=Pt(8),
                          space_after=Pt(0), first_line_indent=Cm(0.75))
            i += 1
            continue

        # 空行
        if not line.strip():
            if in_table:
                in_table = False
                _flush_table(doc, table_rows)
                table_rows = []
            i += 1
            continue

        # 水平分隔线
        if is_horizontal_rule(line):
            if in_table:
                in_table = False
                _flush_table(doc, table_rows)
                table_rows = []
            add_paragraph(doc, '─' * 40, font=FONT_BODY, size=Pt(6),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
            i += 1
            continue

        # 表格 - 分隔行
        if is_table_separator(line):
            i += 1
            continue

        # 表格 - 数据行
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            row = parse_table_row(line)
            if row and any(c.strip() for c in row):
                table_rows.append(row)
            i += 1
            continue

        # 表格结束（当前行不是表格行但之前是）
        if in_table:
            in_table = False
            _flush_table(doc, table_rows)
            table_rows = []

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            _add_heading(doc, title_text, level)
            i += 1
            continue

        # 无序列表
        ul_match = re.match(r'^(\s*)[\-*]\s+(.*)', line)
        if ul_match:
            text = ul_match.group(2)
            add_paragraph(doc, '• ' + text, font=FONT_BODY, size=SIZE_BODY,
                          first_line_indent=Cm(0.75), space_after=Pt(3))
            i += 1
            continue

        # 有序列表（基本支持）
        ol_match = re.match(r'^\s*(\d+)[\.)]\s+(.*)', line)
        if ol_match:
            num, text = ol_match.groups()
            add_paragraph(doc, f'{num}. {text}', font=FONT_BODY, size=SIZE_BODY,
                          first_line_indent=Cm(0.75), space_after=Pt(3))
            i += 1
            continue

        # 普通段落（带行内样式）
        segments = parse_inline(line.strip())
        # 判断是否以 checkbox 开头
        is_checkbox = line.strip().startswith('☐')
        add_styled_paragraph(doc, segments,
                             first_line_indent=None if is_checkbox else Cm(0.75),
                             space_after=Pt(4))
        i += 1

    # 处理末尾未 flush 的表格
    if in_table and table_rows:
        _flush_table(doc, table_rows)

    doc.save(docx_path)
    return docx_path


# ── 内部辅助 ──────────────────────────────────────────

def _add_heading(doc, text, level):
    """添加标题（黑体，按层级设字号）"""
    sizes = {1: SIZE_TITLE, 2: SIZE_H1, 3: SIZE_H2, 4: SIZE_H3}
    sz = sizes.get(level, SIZE_BODY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(18) if level <= 2 else Pt(12)
    pf.space_after = Pt(12) if level <= 2 else Pt(8)
    run = para.add_run(text)
    set_font(run, FONT_HEI, sz, bold=True)
    return para


def _flush_table(doc, rows):
    """输出累积的表格到文档"""
    if not rows:
        return
    # 移除纯空行
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    # 补齐列数
    for r in rows:
        while len(r) < max_cols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=max_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头加灰底
    header_row = table.rows[0]
    for cell in header_row.cells:
        set_cell_shading(cell, 'D9E2F3')
        set_table_cell(cell, '', bold=True, font=FONT_HEI)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            is_header = (row_idx == 0)
            font = FONT_HEI if is_header else FONT_BODY
            set_table_cell(cell, cell_text, bold=is_header, font=font)

    # 表格后空行
    add_paragraph(doc, '', font=FONT_BODY, size=Pt(6), space_after=Pt(6))


# ── 入口 ───────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(description='将 Markdown 核稿意见通知书转换为 .docx')
    ap.add_argument('input', help='输入 .md 文件路径')
    ap.add_argument('-o', '--output', help='输出 .docx 文件路径（默认为输入名 .docx）')
    args = ap.parse_args()

    out_path = args.output or re.sub(r'\.md$', '.docx', args.input)
    result = convert(args.input, out_path)
    print(f'已写出: {result}')


if __name__ == '__main__':
    main()
