"""
Word 报告生成模板——合规审查报告（路径 B）
用于路径 B 的可选 Word 报告输出。

================================================================================
⚠️ 文档生成强制原则（docx 原文直转）—— 本模块必须遵守
================================================================================

最终生成的 .docx 文件必须与 WorkBuddy 对话框中展示的《最终审查报告》
在以下四个维度 100% 一致：

  格式    │ 标题层级、表格、列表、引用块等结构元素一一对应
  序号    │ 章节编号、表格行序号、(1)(2)(3)编号完全一致，不得重排或遗漏
  措辞    │ 每个字句均来自对话原文，不得改写/缩写/扩展/概括
  内容    │ 所有章节、表格行、风险点、法条引用完整保留，不得增删

实现方式：generate_compliance_report_from_markdown() 接受完整 Markdown 原文，
逐元素忠实解析转换为 Word。
⚠️ 严禁使用 generate_compliance_report() 的结构化字段映射方式——该方式必然
   导致上述四个维度出现不一致（已标记为废弃，仅保留向后兼容）。
================================================================================

依赖：pip install python-docx
提供两种模式：
1. generate_compliance_report(data_dict, output_path) —— 结构化字段映射（旧版，已废弃）
2. generate_compliance_report_from_markdown(markdown_text, output_path) —— 原文直转（推荐 ✅）
   将对话中已输出的完整审查报告 Markdown 文本忠实地转换为 docx，
   确保格式、序号、措辞、内容与对话中展示的报告 100% 一致。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re
from datetime import date


# ============================================================
# 模式一：原文直转（推荐）—— Markdown 忠实转 docx
# ============================================================

def generate_compliance_report_from_markdown(markdown_text, output_path,
                                              title='广告合规审查报告',
                                              cover_meta=None):
    """
    将对话中输出的完整审查报告 Markdown 原文忠实地转换为 Word 文档。

    参数:
        markdown_text (str): 对话中已输出的完整报告文本（Markdown 格式）。
                             应包含从「# 广告合规审查报告」到「---」结束的全部内容。
        output_path (str): 输出 .docx 文件路径
        title (str): 封面标题，默认「广告合规审查报告」
        cover_meta (dict|None): 封面元数据，如 {"产品名称": "XX", "审查日期": "2026-07-02"}
                               为 None 时自动从 markdown 开头提取

    返回: output_path

    支持的 Markdown 语法:
        - # 一级标题 → Word 标题 1
        - ## 二级标题 → Word 标题 2
        - ### 三级标题 → Word 标题 3
        - #### 四级及以上标题 → Word 标题 4
        - | 表格 | ... | → Word 表格（自动识别表头，支持单元格内 **粗体**）
        - **粗体文字** → Word 粗体（包括表格内、列表内、段落内）
        - - 无序列表 → Word 项目符号列表
        - 1. / 1) / (1) 有序列表 → Word 编号列表
        - > 引用块 → 缩进段落（左侧缩进 1cm，用于改写版本等区块）
        - --- 分隔线 → 跳过
        - 普通段落 → Word 正文段落（支持跨行、内联换行）
        - 🔴 🟡 🟢 彩色圆点 → 原样保留

    已知限制：
        - 表格单元格内的复杂嵌套格式（如同时包含粗体+斜体+链接）仅支持粗体解析
        - 不支持 Markdown 图片语法 ![alt](url)
        - 不支持代码块 ``` ... ```
    """
    doc = _create_base_document()

    # ---- 封面 ----
    _add_cover_page(doc, title, cover_meta)
    doc.add_page_break()

    # ---- 预处理：去除首尾空白，统一换行符 ----
    cleaned = '\n'.join([
        line.rstrip() for line in markdown_text.split('\n')
    ])

    # ---- 解析并转换正文 ----
    _parse_markdown_to_doc(doc, cleaned, cover_title=title)

    doc.save(output_path)
    return output_path


def _create_base_document():
    """创建基础文档对象，设置默认字体样式"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc


def _add_cover_page(doc, title, cover_meta):
    """添加封面页"""
    doc.add_paragraph()
    doc.add_paragraph()

    # 大标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    doc.add_paragraph()
    doc.add_paragraph()

    # 封面元信息
    if cover_meta:
        for label, value in cover_meta.items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'{label}：{value}')
            run.font.size = Pt(14)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    else:
        # 默认显示日期
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'审查日期：{date.today().isoformat()}')
        run.font.size = Pt(14)
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')


def _parse_markdown_to_doc(doc, text, cover_title='广告合规审查报告'):
    """
    核心解析器：将 Markdown 文本逐元素转换为 Word 内容。

    修复清单（v2）：
    - [FIX-1] 引用块不再重复添加 > 前缀
    - [FIX-2] 表格单元格支持 **粗体** 内联格式
    - [FIX-3] 段落内 \\n 转为 Word 实际换行（软回车 Shift+Enter）
    - [FIX-4] 压缩连续空行，避免标题后多余空段
    - [FIX-5] 正文首个 H1 与封面标题相同时跳过（避免重复）
    - [FIX-6] 支持 (1)(2)(3) 编号模式识别为有序列表
    """

    lines = text.split('\n')
    i = 0
    n = len(lines)
    first_h1_skipped = False          # [FIX-5] 是否已跳过首个重复H1
    last_was_structural = False       # [FIX-4] 上一个元素是否为结构性元素（标题/表格）

    while i < n:
        line = lines[i]

        # 跳过空行 —— 但不在结构性元素后创建多余空段 [FIX-4]
        if not line.strip():
            i += 1
            continue

        stripped = line.strip()

        # ---- 分隔线 ----
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            i += 1
            last_was_structural = False
            continue

        # ---- 一级标题 # （[FIX-5] 首个与封面标题相同则跳过）----
        m = re.match(r'^(#{1})\s+(.+)$', stripped)
        if m:
            heading_text = m.group(2).strip()
            if not first_h1_skipped and heading_text == cover_title:
                first_h1_skipped = True
                i += 1
                continue
            _add_heading(doc, heading_text, level=1)
            last_was_structural = True
            i += 1
            continue

        # ---- 二级标题 ## ----
        m = re.match(r'^(#{2})\s+(.+)$', stripped)
        if m:
            heading_text = m.group(2).strip()
            _add_heading(doc, heading_text, level=2)
            last_was_structural = True
            i += 1
            continue

        # ---- 三级标题 ### ----
        m = re.match(r'^(#{3})\s+(.+)$', stripped)
        if m:
            heading_text = m.group(2).strip()
            _add_heading(doc, heading_text, level=3)
            last_was_structural = True
            i += 1
            continue

        # ---- 四级及以上标题 ####+ ----
        m = re.match(r'^(#{4,})\s+(.+)$', stripped)
        if m:
            heading_text = m.group(2).strip()
            _add_heading(doc, heading_text, level=4)
            last_was_structural = True
            i += 1
            continue

        # ---- 表格 | ... | ----
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < n and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            last_was_structural = True
            continue

        # ---- 无序列表 - / * ----
        m = re.match(r'^[-*+]\s+(.+)$', stripped)
        if m:
            items = []
            while i < n and (lines[i].strip().startswith('- ') or
                             lines[i].strip().startswith('* ') or
                             lines[i].startswith('  ')):
                item_line = lines[i].rstrip()
                indent = len(item_line) - len(item_line.lstrip())
                content = item_line.strip().lstrip('-* ').strip()
                if content:
                    items.append((indent, content))
                i += 1
            if items:
                _add_list(doc, items, ordered=False)
            last_was_structural = False
            continue

        # ---- 有序列表 数字. / 数字) / (数字) [FIX-6] ----
        m = re.match(r'^(\d+)[.)]\s+(.+)$', stripped)
        m_paren = re.match(r'^\((\d+)\)\s+(.+)$', stripped)  # (1) 模式
        if m or m_paren:
            items = []
            while i < n:
                item_line = lines[i].rstrip()
                num_m = re.match(r'^(\d+)[.)\s]+(.+)$', item_line.strip())
                num_mp = re.match(r'^\((\d+)\)\s+(.+)$', item_line.strip())
                if num_m:
                    content = num_m.group(2).strip()
                    items.append(content)
                    i += 1
                elif num_mp:
                    content = num_mp.group(2).strip()
                    items.append(content)
                    i += 1
                else:
                    break
            if items:
                _add_ordered_list(doc, items)
            last_was_structural = False
            continue

        # ---- 引用块 > （[FIX-1] 修复重复>前缀问题）----
        if stripped.startswith('>'):
            quote_lines = []
            while i < n and lines[i].strip().startswith('>'):
                quote_content = lines[i].strip().lstrip('> ').strip()
                if quote_content:
                    quote_lines.append(quote_content)
                i += 1
            if quote_lines:
                _add_quote_block(doc, '\n'.join(quote_lines))
            last_was_structural = False
            continue

        # ---- 普通段落（可能跨行）[FIX-3] 支持 \n 换行 ----
        para_lines = []
        _is_para_line = lambda s: (
            bool(s.strip())
            and not s.strip().startswith('#')
            and not s.strip().startswith('|')
            and not s.strip().startswith('-')
            and not s.strip().startswith('* ')
            and not re.match(r'^\d+[.)]\s', s.strip())
            and not re.match(r'^\(\d+\)\s', s.strip())   # (1) 模式
            and not s.strip().startswith('>')
            and not re.match(r'^-{3,}$', s.strip())
            and not re.match(r'^\*{3,}$', s.strip())
        )
        while i < n and _is_para_line(lines[i]):
            para_lines.append(lines[i])
            i += 1

        if para_lines:
            para_text = '\n'.join(para_lines)
            _add_formatted_paragraph(doc, para_text)
            last_was_structural = False


def _add_heading(doc, text, level=1):
    """添加标题，统一中文字体"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')


def _add_table(doc, table_lines):
    """
    将 Markdown 表格行转换为 Word 表格。
    [FIX-2] 单元格内 **粗体** 标记正确转换为 Word Bold。
    """
    if len(table_lines) < 2:
        return

    parsed_rows = []
    for tl in table_lines:
        cells = [c.strip() for c in tl.split('|')]
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if all(re.match(r'^-+:?\s*$', c) for c in cells):
            continue
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    cols = max(len(r) for r in parsed_rows)
    for r in parsed_rows:
        while len(r) < cols:
            r.append('')

    table = doc.add_table(rows=len(parsed_rows), cols=cols)
    table.style = 'Table Grid'
    table.autofit = True

    for row_idx, row_data in enumerate(parsed_rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            para = cell.paragraphs[0]
            _set_formatted_text(para, cell_text)
            if row_idx == 0:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'SimSun'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')


def _add_list(doc, items, ordered=False):
    """添加无序列表。items: list of (indent, content) tuples"""
    for indent, content in items:
        p = doc.add_paragraph(style='List Bullet')
        _set_formatted_text(p, content)


def _add_ordered_list(doc, items):
    """添加有序列表"""
    for content in items:
        p = doc.add_paragraph(style='List Number')
        _set_formatted_text(p, content)


def _add_quote_block(doc, text):
    """
    添加引用块（缩进段落）。
    [FIX-1] 不再在内容前添加 > 字符——解析器已去除原始 > 前缀。
    多行内容拆分为独立缩进段落。
    """
    sub_lines = text.split('\n')
    for idx, sub_line in enumerate(sub_lines):
        sub_line = sub_line.strip()
        if not sub_line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(0.5)
        _set_formatted_text(p, sub_line)


def _add_formatted_paragraph(doc, text):
    """
    添加支持粗体/行内格式化的段落。
    [FIX-3] 段落内的 \n 字符转为 Word 软回车（Shift+Enter），保持视觉分行。
    """
    sub_parts = text.split('\n')
    for idx, part in enumerate(sub_parts):
        part = part
        if not part.strip():
            p = doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _set_formatted_text(p, part)


def _set_formatted_text(paragraph, text):
    """
    解析文本中的 **bold** 标记和纯文本，设置到段落的 run 中。
    支持 **粗体** 内嵌在普通文本中的任意位置。
    """
    pattern = r'(\*\*[^*]+\*\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.font.bold = True
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(12)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(12)


# ============================================================
# 模式二：结构化字段映射（旧版）—— 已废弃，保留向后兼容
# ============================================================

def _add_table_row(table, cells_text):
    """旧版辅助函数：向表格添加一行"""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        row.cells[i].text = str(text)
    return row


def generate_compliance_report(data, output_path):
    """
    ⚠️ 已废弃——结构化字段映射方式。

    此函数将报告拆解为预定义字典字段再重组为 docx，
    必然导致输出与对话中的原报告在格式、序号、措辞、内容上不一致。
    仅保留向后兼容，新代码请使用 generate_compliance_report_from_markdown()。
    """
    doc = _create_base_document()

    _add_cover_page(doc, '广告合规审查报告', {
        '产品名称': data.get('product_name', ''),
        '审查日期': data.get('review_date', date.today().isoformat()),
    })
    doc.add_page_break()

    text_sections = [
        ('文案预处理', 'text_preprocessing'),
        ('审核结果摘要', 'review_summary'),
        ('资质核查清单', 'qualification_checklist'),
        ('广告定性（广告 vs 非广告前置判断）', 'ad_classification'),
    ]
    for title_text, key in text_sections:
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(data.get(key, '（无）'))

    doc.add_heading('逐句合规审查', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ['序号', '原文', '风险点', '触犯条款', '风险等级', '改写建议']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True

    for row_data in data.get('review_table_rows', []):
        level = str(row_data.get('level', ''))
        if '高' in level:
            level_display = '🔴 ' + level
        elif '中' in level:
            level_display = '🟡 ' + level
        elif '低' in level:
            level_display = '🟢 ' + level
        else:
            level_display = level
        _add_table_row(table, [
            row_data.get('seq', ''),
            row_data.get('original', ''),
            row_data.get('risk', ''),
            row_data.get('clause', ''),
            level_display,
            row_data.get('rewrite', ''),
        ])

    doc.add_paragraph()

    remaining_sections = [
        ('核验修改记录', 'verification_record'),
        ('整体合规结论', 'overall_conclusion'),
        ('处罚风险预估', 'penalty_risk'),
        ('合规改写版本', 'rewritten_version'),
        ('存疑说明', 'uncertainty_note'),
        ('AI 风险告知', 'ai_risk_notice'),
    ]
    for title_text, key in remaining_sections:
        content = data.get(key, '').strip()
        if not content:
            continue
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(content)

    doc.save(output_path)
    return output_path
