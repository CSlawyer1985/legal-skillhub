"""
Word 报告生成模板——法律咨询意见书
用于路径 A 的可选 Word 报告输出。

依赖：pip install python-docx
用法：调用 generate_legal_opinion(data_dict, output_path)
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date


def generate_legal_opinion(data, output_path):
    """
    data 字典结构：
    {
        "consultation_item": "咨询事项",
        "answer_date": "YYYY-MM-DD",
        "question_classification": "问题定性",
        "core_conclusion": "核心结论",
        "verification_record": "核验修改记录（仅当执行外部核验时传入；传入空字符串则跳过此节）",
        "legal_basis": "法律依据（可换行）",
        "judicial_views": "司法/监管观点",
        "practical_analysis": "实务分析",
        "risk_warnings": "风险提示与建议（文本中已含 🔴🟡🟢 彩色圆点标注风险等级）",
        "uncertainty_note": "存疑说明",
        "ai_risk_notice": "AI 风险告知"
    }
    注意：正式报告不包含「检索覆盖说明」「时效性说明」「引用清单」独立章节，
    相关信息融入法律依据的逐条引用标注中。
    """
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)

    # ---- 封面 ----
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph('法律咨询意见书')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_paragraph()
    cover_items = [
        f"咨询事项：{data.get('consultation_item', '')}",
        f"答疑日期：{data.get('answer_date', date.today().isoformat())}",
    ]
    for item in cover_items:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- 正文各节 ----
    sections = [
        ('用户咨询事项', 'consultation_item'),
        ('问题定性', 'question_classification'),
        ('核心结论', 'core_conclusion'),
        ('核验修改记录', 'verification_record'),
        ('法律依据', 'legal_basis'),
        ('司法/监管观点', 'judicial_views'),
        ('实务分析', 'practical_analysis'),
        ('风险提示与建议', 'risk_warnings'),
        ('存疑说明', 'uncertainty_note'),
        ('AI 风险告知', 'ai_risk_notice'),
    ]

    for title_text, key in sections:
        content = data.get(key, '').strip()
        # 跳过空内容章节（如未执行外部核验时 verification_record 为空）
        if not content:
            continue
        heading = doc.add_heading(title_text, level=1)
        doc.add_paragraph(content)

    doc.save(output_path)
    return output_path
