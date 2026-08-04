# -*- coding: utf-8 -*-
"""
要素式起诉状填充脚本
====================

以最高人民法院《要素式起诉状示范文本》（金融借款合同纠纷）模板为基底，
将结构化数据（JSON）填充进模板，**完整保留原文档的表格线、合并单元格、
列宽与字体等布局**；当原告 / 被告 / 第三人数量大于 1 时，按模板既有格式
克隆对应的当事人行（重复同名标签行，不做重新编号）。

用法：
    python fill_element_complaint.py --template <模板.docx> --data <数据.json> --out <输出.docx>

依赖：python-docx（仅用于读写 docx，所有表格线/合并信息都来自模板原始 XML，
      仅替换单元格内的文字内容，不改 tcPr / tblPr，因此边框零丢失）。
"""

import argparse
import copy
import json
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------------------------------------------------------- #
# 基础工具
# --------------------------------------------------------------------------- #
def add_run_to_para(p, text, rpr=None):
    """向段落追加一个 run。rpr 为模板原有 <w:rPr> 深拷贝（保留字体/字号等）；
    为 None 时不写 rPr，交给样式/段落继承，与模板空白单元格行为完全一致。"""
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)


def set_cell_text(cell, text):
    """整单元格写入文字（多行用 \\n 拆分为独立段落）。
    - 保留模板原段落的 pPr（段落格式）；
    - 新 run 继承模板首个 run 的 rPr（字体/字号），使输出与模板逐字一致；
    仅替换 <w:p> 内文字，不触碰 tcPr（边框/合并/列宽）与 tblPr。"""
    tc = cell._tc
    existing = tc.findall(qn("w:p"))
    pPr_tpl = None
    rPr_tpl = None
    if existing:
        p0 = existing[0].find(qn("w:pPr"))
        if p0 is not None:
            pPr_tpl = copy.deepcopy(p0)
        # 抓首个带 rPr 的 run 作为新 run 的格式模板（无则继承样式）
        for p in existing:
            for r in p.findall(qn("w:r")):
                rpr = r.find(qn("w:rPr"))
                if rpr is not None:
                    rPr_tpl = copy.deepcopy(rpr)
                    break
            if rPr_tpl is not None:
                break
    for p in existing:
        tc.remove(p)
    for line in text.split("\n"):
        p = OxmlElement("w:p")
        if pPr_tpl is not None:
            p.append(copy.deepcopy(pPr_tpl))
        if line:
            add_run_to_para(p, line, rPr_tpl)
        tc.append(p)


def _replace_para_text(p, new_text):
    """把段落内所有 run 的文字合并替换为 new_text，保留首个 run 原有 rPr
    （字体/字号等完全沿用模板，不做任何强制覆盖）。"""
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    elif new_text:
        add_run_to_para(p, new_text)


# --------------------------------------------------------------------------- #
# 当事人单元格填充
# --------------------------------------------------------------------------- #
def fill_natural_cell(cell, f):
    """自然人单元格：按每行前缀定位并替换占位内容（保留段落数 = 保留行布局）。"""
    for p in cell.paragraphs:
        txt = "".join(r.text for r in p.runs)
        if not txt.strip():
            continue
        new = None
        if txt.startswith("姓名："):
            new = "姓名：" + (f.get("name") or "")
        elif txt.startswith("性别："):
            g = (f.get("gender") or "").strip()
            line = txt
            if g == "男":
                line = line.replace("男□", "男☑", 1)
            elif g == "女":
                line = line.replace("女□", "女☑", 1)
            new = line
        elif txt.startswith("出生日期："):
            new = "出生日期：" + (f.get("birth") or "")
        elif txt.startswith("民族："):
            new = "民族：" + (f.get("nation") or "汉族")
        elif txt.startswith("工作单位："):
            new = "工作单位：" + (f.get("work_unit") or "") + "    职务：" + (f.get("title") or "")
        elif txt.startswith("联系电话："):
            new = "联系电话：" + (f.get("phone") or "")
        elif txt.startswith("住所地（户籍所在地）："):
            new = "住所地（户籍所在地）：" + (f.get("address") or "")
        elif txt.startswith("经常居住地："):
            new = "经常居住地：" + (f.get("residence") or "")
        elif txt.startswith("证件类型："):
            new = "证件类型：" + (f.get("id_type") or "居民身份证") + "    证件号码：" + (f.get("id_number") or "")
        if new is not None and new != txt:
            _replace_para_text(p, new)


def fill_legal_cell(cell, f):
    """法人 / 非法人组织单元格：按前缀定位替换；org_type / ownership 命中即勾选。"""
    for p in cell.paragraphs:
        txt = "".join(r.text for r in p.runs)
        if not txt.strip():
            continue
        new = None
        if txt.startswith("名称："):
            new = "名称：" + (f.get("name") or "")
        elif txt.startswith("住所地（主要办事机构所在地）："):
            new = "住所地（主要办事机构所在地）：" + (f.get("address") or "")
        elif txt.startswith("注册地 / 登记地："):
            new = "注册地 / 登记地：" + (f.get("reg_address") or "")
        elif txt.startswith("法定代表人 / 负责人："):
            new = "法定代表人 / 负责人：" + (f.get("legal_rep") or "")
        elif txt.startswith("统一社会信用代码："):
            new = "统一社会信用代码：" + (f.get("credit_code") or "")
        elif txt.startswith("类型："):
            line = txt
            ot = (f.get("org_type") or "").strip()
            if ot:
                line = line.replace(ot + "□", ot + "☑", 1)
            new = line
        elif txt.startswith("职务："):
            new = "职务：" + (f.get("rep_title") or "")
        elif txt.startswith("联系电话："):
            new = "联系电话：" + (f.get("phone") or "")
        elif "所有制性质" in txt:
            line = txt
            ow = (f.get("ownership") or "").strip()
            if ow:
                line = line.replace(ow + "□", ow + "☑", 1)
            new = line
        if new is not None and new != txt:
            _replace_para_text(p, new)


# --------------------------------------------------------------------------- #
# 当事人区块定位与克隆
# --------------------------------------------------------------------------- #
def _raw_text(tc):
    return "".join(t.text or "" for t in tc.iter(qn("w:t")))


def _cell_vmerge(tc):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        return None
    return vm.get(qn("w:val"))


def find_party_rows(table, keyword):
    """返回 (natural_row_idx, legal_row_indices)。
    natural_row_idx：关键字+自然人 标签所在行；
    legal_row_indices：法人区块所有行（含 vMerge 续行）。
    直接用行内首个 <w:tc> 原始 XML 判空，避免 python-docx 把纵向合并单元格
    的文本回填导致误判（续行本身文本为空）。"""
    rows = table.rows
    natural_idx = None
    legal_start = None
    legal_idxs = []
    for i, row in enumerate(rows):
        tc0 = row._tr.findall(qn("w:tc"))[0]
        label = _raw_text(tc0).strip()
        if keyword in label and "自然人" in label and natural_idx is None:
            natural_idx = i
        if keyword in label and "法人、非法人" in label and legal_start is None:
            legal_start = i
    if legal_start is not None:
        legal_idxs = [legal_start]
        j = legal_start + 1
        while j < len(rows):
            tc0 = rows[j]._tr.findall(qn("w:tc"))[0]
            if _raw_text(tc0) == "":
                legal_idxs.append(j)
                j += 1
            else:
                break
    return natural_idx, legal_idxs


def clone_rows(table, src_indices):
    """深拷贝源行，返回新的 <w:tr> 元素列表。"""
    return [copy.deepcopy(table.rows[i]._tr) for i in src_indices]


def fill_block(table, new_trs, party, is_legal):
    """填充（尚未插入的）克隆行内容。"""
    from docx.table import _Cell
    fn = fill_legal_cell if is_legal else fill_natural_cell
    for tr in new_trs:
        tcs = tr.findall(qn("w:tc"))
        for ci, tc in enumerate(tcs):
            if ci == 0:
                continue  # 标签列跳过
            cell = _Cell(tc, table)
            fn(cell, party.get("fields", {}))


# --------------------------------------------------------------------------- #
# 章节（诉讼请求 / 管辖 / 事实 / 意愿）按标签填充
# --------------------------------------------------------------------------- #
def norm(s):
    return (s or "").replace(" ", "").replace("　", "")


def _tc_gridspan(tc):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 1
    g = tcPr.find(qn("w:gridSpan"))
    return int(g.get(qn("w:val"))) if g is not None else 1


def fill_sections(table, sections):
    """sections: {标签子串(可含空格): 填充文本}。
    对所有行，若其首列标签（去空格后）包含子串（去空格后），则写入内容单元格。
    内容单元格 = 该行第 2 个真实 <w:tc> 元素（标签格之后的第一格），不受标签列
    gridSpan 影响；跳过整行仅 1 个真实 <w:tc> 的标题/说明行（写入会覆盖标题）。"""
    from docx.table import _Cell
    if not sections:
        return
    for row in table.rows:
        tcs = row._tr.findall(qn("w:tc"))
        if len(tcs) <= 1:
            continue  # 整行合并的标题/说明行，不填
        label = _raw_text(tcs[0]).strip()
        if not label:
            continue
        nlabel = norm(label)
        for sub, text in sections.items():
            if norm(sub) and norm(sub) in nlabel:
                content_tc = tcs[1]
                set_cell_text(_Cell(content_tc, table), text)
                break


# --------------------------------------------------------------------------- #
# 主流程
# --------------------------------------------------------------------------- #
def process_party_group(table, keyword, parties, insert_before_rowidx):
    """处理一个当事人分组（原告/被告/第三人）。
    insert_before_rowidx: 该组在表中结束的边界行索引（额外法人区块插入其前）；
                          若为 None 表示追加到表尾（原告）。"""
    naturals = [p for p in parties if p.get("type") == "自然人"]
    legals = [p for p in parties if p.get("type") == "法人"]

    natural_idx, legal_idxs = find_party_rows(table, keyword)
    legal_start = legal_idxs[0] if legal_idxs else None

    # 插入锚点：在插入任何行之前先抓取元素引用，避免索引漂移
    legal_ref_tr = table.rows[legal_start]._tr if legal_start is not None else None
    boundary_tr = table.rows[insert_before_rowidx]._tr if insert_before_rowidx is not None else None

    # 先从「原样模板行」克隆额外当事人行（此时尚未填充，保证克隆源是纯净模板）
    natural_clones = [clone_rows(table, [natural_idx]) for _ in naturals[1:]] if natural_idx is not None else []
    legal_clones = [clone_rows(table, list(legal_idxs)) for _ in legals[1:]] if legal_idxs else []

    # 1) 填充既有首名当事人
    if natural_idx is not None and naturals:
        fill_block(table, [table.rows[natural_idx]._tr], naturals[0], False)
    if legal_idxs and legals:
        fill_block(table, [table.rows[i]._tr for i in legal_idxs], legals[0], True)

    # 2) 额外自然人：克隆行填充后插入到法人区块之前
    for party, cl in zip(naturals[1:], natural_clones):
        fill_block(table, cl, party, False)
        if legal_ref_tr is not None:
            legal_ref_tr.addprevious(cl[0])
        else:
            table._tbl.append(cl[0])

    # 3) 额外法人：克隆整套法人区块，插入到组边界之前
    #    按原顺序逐行 addprevious(boundary)：A1,A2,A3 依次插入 boundary 之前，
    #    最终在文档中即为 [A1, A2, A3, boundary]，块内顺序与 vMerge 完整保留。
    for party, cl in zip(legals[1:], legal_clones):
        fill_block(table, cl, party, True)
        if boundary_tr is not None:
            for tr in cl:
                boundary_tr.addprevious(tr)
        else:
            for tr in cl:
                table._tbl.append(tr)


def find_row_index_by_label(table, label_prefix, gridspan_only=False):
    for i, row in enumerate(table.rows):
        if row.cells[0].text.strip().startswith(label_prefix):
            return i
    return None


def fill_agent(table, agent):
    if not agent or not agent.get("has"):
        return
    # 委托诉讼代理人行（TABLE1 首行）
    for row in table.rows:
        if "委托诉讼代理人" in row.cells[0].text:
            lines = []
            lines.append("有□" if agent.get("has") else "无□")
            if agent.get("name"):
                lines.append("姓名：" + agent["name"])
            if agent.get("unit"):
                lines.append("单位：" + agent["unit"])
            if agent.get("title"):
                lines.append("职务：" + agent["title"])
            if agent.get("phone"):
                lines.append("联系电话：" + agent["phone"])
            lines.append("代理权限：" + (agent.get("authority") or "一般授权") + "□")
            set_cell_text(row.cells[1], "\n".join(lines))
            break


def update_title(doc, case_type):
    if not case_type:
        return
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("（") and t.endswith("）") and "纠纷" in t:
            if p.runs:
                p.runs[0].text = "（" + case_type + "）"
                for r in p.runs[1:]:
                    r.text = ""
                # 不调用 set_run_font：保留模板原有 rPr（方正小标宋_GBK / sz=36）
            break


def fill_signoff(doc, signer, date):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("具状人"):
            if signer and p.runs:
                p.runs[0].text = "具状人（签字、盖章）：" + signer
                for r in p.runs[1:]:
                    r.text = ""
                # 保留模板原有 rPr（方正小标宋_GBK / sz=30）
        elif t.startswith("日期"):
            if date and p.runs:
                p.runs[0].text = "日期：" + date
                for r in p.runs[1:]:
                    r.text = ""
                # 保留模板原有 rPr（方正小标宋_GBK / sz=30）


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True, help="要素式模板 docx")
    ap.add_argument("--data", required=True, help="结构化数据 JSON")
    ap.add_argument("--out", required=True, help="输出 docx 路径")
    args = ap.parse_args()

    with open(args.data, "r", encoding="utf-8") as fh:
        data = json.load(fh)

    doc = Document(args.template)
    tables = doc.tables
    t0 = tables[0]  # 说明 + 原告
    t1 = tables[1]  # 被告 / 第三人 / 诉讼请求 / 事实 / 意愿

    # 边界行索引（用于插入额外被告/第三人之前）
    third_natural_idx, _ = find_party_rows(t1, "第三人")
    claim_idx = find_row_index_by_label(t1, "诉讼请求")

    # 原告：表尾追加
    process_party_group(t0, "原告", data.get("plaintiffs", []), None)
    # 被告：插入到第三人之前
    process_party_group(t1, "被告", data.get("defendants", []), third_natural_idx)
    # 第三人：插入到诉讼请求之前
    process_party_group(t1, "第三人", data.get("third_parties", []), claim_idx)

    # 委托诉讼代理人
    fill_agent(t1, data.get("agent"))

    # 章节要素
    fill_sections(t1, data.get("sections", {}))

    # 标题案由 + 落款
    update_title(doc, data.get("case_type"))
    fill_signoff(doc, data.get("signer"), data.get("date"))

    doc.save(args.out)
    print("OK ->", args.out)


if __name__ == "__main__":
    main()
