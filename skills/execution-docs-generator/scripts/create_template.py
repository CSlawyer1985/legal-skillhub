#!/usr/bin/env python3
"""
创建 .docx 模板：将文档中的真实值替换为 [变量] 占位符
用法：python3 create_template.py input.docx output.docx --map '{"赵建华":"[申请人姓名]","张茂凤":"[被申请人姓名]",...}'
"""
import re, json, argparse
from docx import Document

def create_template(input_path: str, output_path: str, mapping: dict):
    doc = Document(input_path)
    replaced = 0
    
    for para in doc.paragraphs:
        for run in para.runs:
            for old_val, placeholder in mapping.items():
                if old_val in run.text:
                    run.text = run.text.replace(old_val, placeholder)
                    replaced += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for old_val, placeholder in mapping.items():
                            if old_val in run.text:
                                run.text = run.text.replace(old_val, placeholder)
                                replaced += 1
    
    doc.save(output_path)
    print(f'✅ 模板已创建：{output_path}（{replaced} 处替换）')
    return output_path

if __name__ == '__main__':
    p = argparse.ArgumentParser(description='创建模板：真实值 → [变量]')
    p.add_argument('input', help='原始 .docx 文件')
    p.add_argument('output', help='输出模板 .docx')
    p.add_argument('--map', default='{}', help='替换映射 JSON {"真实值":"[变量名]",...}')
    args = p.parse_args()
    create_template(args.input, args.output, json.loads(args.map))
