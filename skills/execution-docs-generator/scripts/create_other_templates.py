#!/usr/bin/env python3
"""创建失信被执行人名单申请书 + 限制消费申请书 + 授权委托书模板 .docx"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys

def setup_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(3); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(2.5)
    return doc

F = '宋体'

def add_p(doc, text, size=12, bold=False, align=None, indent=0):
    p = doc.add_paragraph()
    if align: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Pt(indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = F; run._element.rPr.rFonts.set(qn('w:eastAsia'), F)
    run.font.size = Pt(size); run.bold = bold
    return p

def add_mix(doc, segments, indent=0):
    """segments: [(text, bold), ...]"""
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Pt(indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for text, bold in segments:
        run = p.add_run(text)
        run.font.name = F; run._element.rPr.rFonts.set(qn('w:eastAsia'), F)
        run.font.size = Pt(12); run.bold = bold
    return p

def make_dishonest():
    doc = setup_doc()
    add_p(doc, '请求将被执行人纳入失信\n被执行人名单申请书', 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_mix(doc, [('申请人：', True), ('[申请人姓名]，[申请人身份信息]', False)])
    add_mix(doc, [('被申请人：', True), ('[被申请人姓名]，[被申请人身份信息]', False)])
    add_p(doc, '请求事项：', 12, True)
    add_p(doc, '请求贵院依法将被申请人纳入失信被执行人名单，并在征信系统予以记录。', 12, False, indent=24)
    add_p(doc, '事实及理由：', 12, True)
    add_p(doc, '[申请人姓名]申请强制执行[被申请人姓名]一案，贵院已依法受理。', 12, False, indent=24)
    add_p(doc, '[被申请人姓名]在执行阶段规避执行，不履行生效法律文书所确定的义务，其行为已严重侵犯了申请人的合法权益，应依法对其进行信用惩戒。', 12, False, indent=24)
    add_p(doc, '申请人为维护自身合法权益，依据《最高人民法院关于公布失信被执行人名单信息的若干规定》第一条第二项、第六项和第二条之规定，特提出上述申请，望予以支持。', 12, False, indent=24)
    add_p(doc, '此致', 12, False, indent=24)
    add_p(doc, '[管辖法院全称]', 12, False)
    add_p(doc, '申请人：', 12, False, indent=200)
    add_p(doc, '[year]年  月  日', 12, False, WD_ALIGN_PARAGRAPH.RIGHT)
    return doc

def make_consumption():
    doc = setup_doc()
    add_p(doc, '限制消费申请书', 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_mix(doc, [('申请人：', True), ('[申请人姓名]，[申请人身份信息]', False)])
    add_mix(doc, [('被申请人：', True), ('[被申请人姓名]，[被申请人身份信息]', False)])
    add_p(doc, '请求事项：', 12, True)
    add_p(doc, '请求贵院对被申请人作出限制消费的措施并执行。', 12, False, indent=24)
    add_p(doc, '事实及理由：', 12, True)
    add_p(doc, '[申请人姓名]申请强制执行[被申请人姓名]一案，贵院已依法受理。', 12, False, indent=24)
    add_p(doc, '由于[被申请人姓名]拒不履行付款义务，现申请人依据《最高人民法院关于限制被执行人高消费及有关消费的若干规定》第一条、第三条之规定，请求贵院对[被申请人姓名]采取限制消费的措施。', 12, False, indent=24)
    add_p(doc, '申请人为维护自身合法权益，特提出上述申请，望予以支持。', 12, False, indent=24)
    add_p(doc, '此致', 12, False, indent=24)
    add_p(doc, '[管辖法院全称]', 12, False)
    add_p(doc, '申请人：', 12, False, indent=200)
    add_p(doc, '[year]年  月  日', 12, False, WD_ALIGN_PARAGRAPH.RIGHT)
    return doc

def make_poa():
    doc = setup_doc()
    add_p(doc, '授权委托书', 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_mix(doc, [('委托人：', True), ('[委托人姓名]', False)])
    add_mix(doc, [('受托人：', True), ('[律师姓名]，[律师事务所全称]律师[律师联系方式行]', False)])
    add_p(doc, '现委托[律师姓名]在委托人申请强制执行[被申请人姓名]一案中，作为委托人申请执行的委托代理人，委托权限如下：', 12, False, indent=24)
    add_p(doc, '申请执行；参与执行程序；放弃或变更执行请求；承认或者反驳对方请求；提出异议；收发法律文书；执行和解；代领本案执行款及其他款项（特别授权）。', 12, False, indent=24)
    add_p(doc, '委托人：', 12, False, indent=200)
    add_p(doc, '[year]年  月  日', 12, False, WD_ALIGN_PARAGRAPH.RIGHT)
    return doc

if __name__ == '__main__':
    base = sys.argv[1] if len(sys.argv) > 1 else '/tmp'
    for name, maker in [('失信被执行人名单申请书_模板', make_dishonest),
                         ('限制消费申请书_模板', make_consumption),
                         ('授权委托书_模板', make_poa)]:
        doc = maker()
        path = f'{base}/{name}.docx'
        doc.save(path)
        print(f'✅ {path}')
