#!/usr/bin/env python3
"""创建强制执行申请书模板 .docx — 格式与用户原始 .doc 完全一致"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys, os

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.top_margin = Cm(3); s.bottom_margin = Cm(2.5)
s.left_margin = Cm(3); s.right_margin = Cm(2.5)

F = '宋体'
def r(p, text, size=12, bold=False, align=None, indent=0):
    if align: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Pt(indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = F; run._element.rPr.rFonts.set(qn('w:eastAsia'), F)
    run.font.size = Pt(size); run.bold = bold
    return run

def mr(p, segments):
    """mixed run: list of (text, bold) tuples"""
    for text, bold in segments:
        r(p, text, 12, bold)

# P0: Title
p = doc.add_paragraph(); r(p, '强制执行申请书', 14, True, WD_ALIGN_PARAGRAPH.CENTER)

# P1: 申请人
p = doc.add_paragraph(); mr(p, [('申请人：', True), ('[申请人姓名]，[申请人身份信息]', False)])

# P2: 被申请人
p = doc.add_paragraph(); mr(p, [('被申请人：', True), ('[被申请人姓名]，[被申请人身份信息]', False)])

# P3: 请求事项
p = doc.add_paragraph(); r(p, '请求事项：', 12, True)

# P4-7: 请求事项列表
for item in [
    '1、强制被申请人向申请人支付人民币[本金金额]元；',
    '2、强制被申请人向申请人支付利息，以[本金金额]元为基数，自[利息起算日]起至实际付清之日止按年利率[年利率]计算的利息，暂算至[利息暂算日]为[利息暂算金额]元；',
    '3、强制被申请人加倍支付迟延履行期间的债务利息，按日万分之一点七五计算，自[加倍利息起算日]暂算至[加倍利息暂算日]为[加倍利息暂算金额]元；（以上三项共计为[合计金额]元）',
    '4、本案执行费用由被申请人承担。',
]:
    p = doc.add_paragraph(); r(p, item, 12, False, indent=24)

# P8: 事实与理由
p = doc.add_paragraph(); r(p, '事实与理由：', 12, True)

# P9: 案件信息
p = doc.add_paragraph(); r(p, '申请人与被申请人[案由]一案，经[作出机构]出具了[案号][文书类型]。', 12, False, indent=24)

# P10: 判决摘要
p = doc.add_paragraph(); r(p, '根据[文书类型]，被申请人应于本判决发生法律效力之日起[履行期限]内向申请人归还本金[本金金额]元，支付自[利息起算日]起至实际付清之日止按年利率[年利率]计算的利息，并负担诉讼费用[诉讼费金额]元。', 12, False, indent=24)

# P11: 迟延履行
p = doc.add_paragraph(); r(p, '如果被申请人未按照判决指定的期间履行给付义务，应当依照《中华人民共和国民事诉讼法》第二百六十四条之规定，加倍支付迟延履行期间的债务利息。', 12, False, indent=24)

# P12: 执行理由
p = doc.add_paragraph(); r(p, '现被申请人未按判决书支付任何款项，为维护申请人的合法权益，特提出上述强制执行请求，请予以支持。', 12, False, indent=24)

# P13: 此致
p = doc.add_paragraph(); r(p, '此致', 12, False, indent=24)

# P14: 法院
p = doc.add_paragraph(); r(p, '[管辖法院全称]', 12, False)

# P15: 签名行
p = doc.add_paragraph(); r(p, '申请人：', 12, False, indent=200)

# P16: 日期
p = doc.add_paragraph(); r(p, '[year]年  月  日', 12, False, WD_ALIGN_PARAGRAPH.RIGHT)

out = sys.argv[1] if len(sys.argv) > 1 else '/tmp/强制执行申请书_模板.docx'
doc.save(out); print(f'✅ {out}')
