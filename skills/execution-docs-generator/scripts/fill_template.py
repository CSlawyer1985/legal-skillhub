#!/usr/bin/env python3
"""
精确模板复制引擎 — 格式 1:1 保留
读取 .docx 模板文件 → 逐段逐 run 替换 [变量] 占位符 → 保存新 .docx
完全不动任何格式：字体、字号、加粗、对齐、缩进、行距全部保持原样。
"""

import re, json, argparse, copy
from docx import Document
from docx.oxml.ns import qn


def fill_template(template_path: str, variables: dict, output_path: str):
    """
    打开模板 .docx，找到所有 [变量名] 占位符，替换为实际值。
    每个 run 独立处理，不合并不拆分，保持原始格式。
    """
    doc = Document(template_path)
    
    replacements = 0
    
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if '[' not in text:
                continue
            
            # 查找所有 [变量] 并进行替换
            new_text = text
            for key, val in variables.items():
                placeholder = f'[{key}]'
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(val))
                    replacements += 1
            
            if new_text != text:
                run.text = new_text
    
    # 也处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text
                        if '[' not in text:
                            continue
                        new_text = text
                        for key, val in variables.items():
                            placeholder = f'[{key}]'
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(val))
                                replacements += 1
                        if new_text != text:
                            run.text = new_text
    
    doc.save(output_path)
    print(f'✅ {output_path}（替换 {replacements} 处变量）')
    return output_path


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='精确模板复制 — 格式完全不动，只换文字')
    parser.add_argument('template', help='模板 .docx 文件路径（含 [变量] 占位符）')
    parser.add_argument('output', help='输出 .docx 文件路径')
    parser.add_argument('--vars', default='{}', help='变量 JSON')
    args = parser.parse_args()
    
    variables = json.loads(args.vars)
    fill_template(args.template, variables, args.output)
