#!/usr/bin/env python3
"""
将 accurLex 知法返回的审查意见转为 Word 文档。

核心原则：**审查意见原文不做任何内容修改，只做展示形式处理**。

用法:
  python3 gen_review_docx.py <input.json> [--output <output.docx>]
  python3 gen_review_docx.py <input.md>   [--output <output.docx>]

input.json: task_contract_review.js 输出的结构化 JSON
input.md:   task_contract_review.js 输出的 Markdown 原文

如果不指定 --output，默认输出到同目录下同名 .docx 文件。
"""

import sys
import json
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 工具函数 ──

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name='宋体', font_size=12, bold=False, color=None):
    """设置 run 的字体属性"""
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, '黑体', 22, bold=True)
    elif level == 1:
        run = p.add_run(text)
        set_run_font(run, '黑体', 16, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, '黑体', 14, bold=True)
    elif level == 3:
        run = p.add_run(text)
        set_run_font(run, '黑体', 13, bold=True)
    else:
        run = p.add_run(text)
        set_run_font(run, '黑体', 12, bold=True)
    return p


def add_risk_badge(paragraph, risk_level):
    """在段落中添加风险等级标记"""
    badges = {
        '高风险': ('🔴 高风险', (255, 0, 0)),
        '中风险': ('🟡 中风险', (255, 165, 0)),
        '低风险': ('🟢 低风险', (0, 128, 0)),
    }
    text, color = badges.get(risk_level, ('', None))
    if text:
        run = paragraph.add_run(f'【{text}】')
        set_run_font(run, '黑体', 12, bold=True, color=color)


def parse_markdown_to_docx(doc, markdown_text):
    """
    将 Markdown 格式的审查意见转为 Word 段落。
    **只做展示形式转换，不修改原文内容。**
    
    处理规则：
    - ### / #### / ##### → 对应级别标题
    - **bold** → 加粗
    - > 引用 → 缩进灰色斜体
    - - 列表 → Word 列表
    - 1. 2. 有序列表 → Word 有序列表
    - --- → 分隔线
    - 纯文本 → 普通段落
    - 风险等级标记（高风险/中风险/低风险）→ 彩色加粗
    """
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 空行跳过
        if not stripped:
            i += 1
            continue
        
        # 水平分隔线
        if stripped == '---':
            # 添加一条细线分隔
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # 标题（### 及以上）
        heading_match = re.match(r'^(#{3,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1)) - 2  # ### → 1, #### → 2, etc.
            title_text = heading_match.group(2)
            add_heading_styled(doc, title_text, level)
            i += 1
            continue
        
        # 引用块（> ）
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_text = re.sub(r'^>\s?', '', lines[i].strip())
                quote_lines.append(quote_text)
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            for qi, ql in enumerate(quote_lines):
                run = p.add_run(ql)
                set_run_font(run, '宋体', 11, color=(128, 128, 128))
                if qi < len(quote_lines) - 1:
                    run.add_break()
            continue
        
        # 无序列表（- 或 * 开头）
        if re.match(r'^[-*]\s+', stripped):
            list_items = []
            while i < len(lines):
                s = lines[i].strip()
                if re.match(r'^[-*]\s+', s):
                    item_text = re.sub(r'^[-*]\s+', '', s)
                    list_items.append(item_text)
                    i += 1
                else:
                    break
            for item in list_items:
                p = doc.add_paragraph(style='List Bullet')
                parse_inline_formatting(p, item)
            continue
        
        # 有序列表（1. 2. 等开头）
        if re.match(r'^\d+[.)]\s+', stripped):
            list_items = []
            while i < len(lines):
                s = lines[i].strip()
                if re.match(r'^\d+[.)]\s+', s):
                    item_text = re.sub(r'^\d+[.)]\s+', '', s)
                    list_items.append(item_text)
                    i += 1
                else:
                    break
            for idx, item in enumerate(list_items, 1):
                p = doc.add_paragraph(style='List Number')
                parse_inline_formatting(p, item)
            continue
        
        # 普通段落
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if (not next_stripped or 
                next_stripped.startswith('#') or 
                next_stripped.startswith('>') or 
                next_stripped == '---' or
                re.match(r'^[-*]\s+', next_stripped) or
                re.match(r'^\d+[.)]\s+', next_stripped)):
                break
            para_lines.append(next_stripped)
            i += 1
        
        full_text = '\n'.join(para_lines)
        p = doc.add_paragraph()
        parse_inline_formatting(p, full_text)


def parse_inline_formatting(paragraph, text):
    """
    解析行内格式（加粗、风险标记等），添加到段落。
    **不修改文本内容，只处理展示格式。**
    """
    # 匹配 **bold** 或 【高风险】【中风险】【低风险】标记
    pattern = r'(\*\*(.+?)\*\*)|(【高风险】)|(【中风险】)|(【低风险】)'
    
    last_end = 0
    for m in re.finditer(pattern, text):
        # 匹配前的普通文本
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            run = paragraph.add_run(plain)
            set_run_font(run, '宋体', 12)
        
        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            set_run_font(run, '宋体', 12, bold=True)
        elif m.group(3):  # 【高风险】
            run = paragraph.add_run('【🔴 高风险】')
            set_run_font(run, '黑体', 12, bold=True, color=(255, 0, 0))
        elif m.group(4):  # 【中风险】
            run = paragraph.add_run('【🟡 中风险】')
            set_run_font(run, '黑体', 12, bold=True, color=(255, 165, 0))
        elif m.group(5):  # 【低风险】
            run = paragraph.add_run('【🟢 低风险】')
            set_run_font(run, '黑体', 12, bold=True, color=(0, 128, 0))
        
        last_end = m.end()
    
    # 剩余文本
    if last_end < len(text):
        plain = text[last_end:]
        run = paragraph.add_run(plain)
        set_run_font(run, '宋体', 12)


# ── 主流程 ──

def generate_review_docx(input_path, output_path=None):
    """
    从 JSON 或 Markdown 文件生成审查意见书 Word 文档。
    **审查意见内容原文不修改，只做展示形式处理。**
    """
    # 确定输出路径
    if not output_path:
        base = os.path.splitext(input_path)[0]
        output_path = base + '.docx'
    
    # 读取输入
    ext = os.path.splitext(input_path)[1].lower()
    
    if ext == '.json':
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        review_text = data.get('reviewText', '')
        citations = data.get('citations', [])
        error = data.get('error')
        
        if error:
            print(f'[error] 审查失败: {error}', file=sys.stderr)
            sys.exit(1)
        
        # 组装完整 Markdown（与 JS 脚本 formatRawOutput 一致）
        full_markdown = review_text
        if citations:
            unique = list(dict.fromkeys(citations))  # 去重保序
            full_markdown += '\n\n---\n\n### 引用法条\n\n'
            full_markdown += '\n\n'.join(unique)
        full_markdown += '\n\n---\n\n> ⚠️ 以上审查意见由 AI 辅助生成，仅供参考。重大决策请咨询专业律师。\n>\n> 🔗 [accurLex知法](https://accurlex.com) 提供专业AI合同审查服务'
    
    elif ext in ('.md', '.markdown', '.txt'):
        with open(input_path, 'r', encoding='utf-8') as f:
            full_markdown = f.read()
    
    else:
        print(f'[error] 不支持的文件格式: {ext}', file=sys.stderr)
        sys.exit(1)
    
    # 创建 Word 文档
    doc = Document()
    
    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 文档标题
    add_heading_styled(doc, '合同审查意见书', level=0)
    
    # 分隔线
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # 核心内容：解析 Markdown 原文 → Word（不改内容）
    parse_markdown_to_docx(doc, full_markdown)
    
    # 保存
    doc.save(output_path)
    print(f'[ok] 审查意见书 Word 已生成: {output_path}', file=sys.stderr)
    return output_path


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='将 accurLex 审查意见转为 Word 文档（原文不改内容）')
    parser.add_argument('input', help='输入文件路径（.json 或 .md）')
    parser.add_argument('--output', '-o', help='输出 .docx 文件路径（默认同目录同名）')
    args = parser.parse_args()
    
    generate_review_docx(args.input, args.output)
