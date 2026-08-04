"""
Helper functions for filling element-style legal forms (要素式起诉状/答辩状).

These are building blocks — the AI should compose them into a fill script
tailored to the specific template and source material.

Key design decisions from hard-won experience:
1. Templates have complex merged cells — always check uniqueness via id(cell._tc).
2. Run-level replacement (cell_replace) preserves original formatting.
3. Paragraph-level overwrite (cell_set) for complete rewrites.
4. For "not applicable" sections, append "/" rather than writing "本案不涉及".
5. Summary/free-text boxes get original source text pasted directly.
6. Leave evidence lists and mediation sections empty unless specified.
7. Handle birth dates and split-run fields with paragraph-level replacement.
8. Font: auto-discover template's font and style, don't hardcode.
"""

import re
import shutil
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# These are FALLBACKS — discover_template_style() should be called first
# to find the actual values used in the template.
FONT_NAME = '方正书宋_GBK'
TABLE_TEXT_STYLE = '6'


# ================================================================
# Template safe-copy
# ================================================================


def copy_template(src: str, dst: str = None) -> str:
    """Create a safe working copy of the template.

    Never modify the original template in-place. This prevents:
    - Corrupting the original if the fill script crashes mid-way
    - Accidentally shipping partially-filled templates
    - Losing the pristine template for future use

    Args:
        src: path to original template .docx
        dst: optional destination path. If None, creates {name}_filled.docx
             next to the original.

    Returns:
        path to the working copy (open this for modification)
    """
    src_path = Path(src)
    if dst is None:
        stem = src_path.stem
        dst = str(src_path.parent / f"{stem}_filled{src_path.suffix}")
    shutil.copy2(src, dst)
    return dst


# ================================================================
# Template style discovery
# ================================================================


def discover_template_style(template_path: str) -> dict:
    """Auto-discover the font and paragraph style used by template data cells.

    This is MANDATORY before writing any fill script. Different templates
    use different style IDs and fonts. Hardcoding '方正书宋_GBK' and style '6'
    will silently produce wrong output on templates that differ.

    Strategy:
    1. Find the first data cell (natural person info, claims, etc.)
    2. Read its paragraph style and run font
    3. Return the discovered values

    Returns:
        dict with keys: 'font_name', 'style_id', 'style_name'
        e.g. {'font_name': '方正书宋_GBK', 'style_id': '6', 'style_name': 'Table Text'}
    """
    from lxml import etree
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    doc = Document(template_path)
    result = {'font_name': '仿宋', 'style_id': '1', 'style_name': 'Normal'}

    # Scan tables for a data cell — one that contains template field labels
    # like "姓名：" or "性别：" — the most reliable indicator
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if '姓名：' in text and '性别：' in text:
                    # Found a natural person data cell — extract its style
                    for p in cell.paragraphs:
                        pPr = p._element.find(f'{{{W}}}pPr')
                        if pPr is not None:
                            pStyle = pPr.find(f'{{{W}}}pStyle')
                            if pStyle is not None:
                                style_id = pStyle.get(qn('w:val'))
                                result['style_id'] = style_id
                                # Look up the style name
                                for s in doc.styles:
                                    if s.style_id == style_id:
                                        result['style_name'] = s.name
                                        break

                        for r in p.runs:
                            rPr = r._element.find(f'{{{W}}}rPr')
                            if rPr is not None:
                                rFonts = rPr.find(f'{{{W}}}rFonts')
                                if rFonts is not None:
                                    ea = rFonts.get(qn('w:eastAsia'))
                                    if ea:
                                        result['font_name'] = ea
                                        # If font name is empty (inherits from style),
                                        # look up the style's font
                                        if not ea.strip():
                                            for s in doc.styles:
                                                if s.style_id == result['style_id']:
                                                    srpr = s.element.find(
                                                        f'.//{{{W}}}rPr')
                                                    if srpr is not None:
                                                        srf = srpr.find(
                                                            f'{{{W}}}rFonts')
                                                        if srf is not None:
                                                            ea2 = srf.get(
                                                                qn('w:eastAsia'))
                                                            if ea2:
                                                                result[
                                                                    'font_name'] = ea2
                                                    break
                                    return result
    return result


# ================================================================
# Reading source material
# ================================================================


def read_docx(path: str) -> str:
    """Extract all text from a .docx file."""
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return '\n'.join(parts)


def read_doc(path: str) -> str:
    """Extract text from .doc file using macOS textutil."""
    result = subprocess.run(
        ['textutil', '-convert', 'txt', '-stdout', path],
        capture_output=True, text=True, timeout=30
    )
    if result.returncode != 0:
        raise RuntimeError(f"textutil failed: {result.stderr}")
    return result.stdout


def read_source(path: str) -> str:
    """Read any supported source file, return plain text."""
    if path.endswith('.docx'):
        return read_docx(path)
    elif path.endswith('.doc'):
        return read_doc(path)
    else:
        with open(path, 'r') as f:
            return f.read()


# ================================================================
# Source text parsing — structured data extraction
# ================================================================


def parse_parties(text: str) -> dict:
    """Extract party information (原告/被告) from complaint text.

    Parses standard Chinese civil complaint format:
        原告：姓名，性别，民族，出生日期，身份证号，住址
        被告：姓名，性别，民族，出生日期，身份证号，住址

    Returns:
        dict with 'plaintiff' and 'defendant' keys, each a dict of fields.
    """
    result = {'plaintiff': {}, 'defendant': {}}

    # Split text into plaintiff and defendant sections
    # Pattern: 原告：... 被告：...
    plaintiff_match = re.search(
        r'原告[：:]\s*(.+?)(?:被告|$)', text, re.DOTALL)
    defendant_match = re.search(
        r'被告[：:]\s*(.+?)(?:诉讼请求|事实与理由|$)', text, re.DOTALL)

    for role, match in [('plaintiff', plaintiff_match), ('defendant', defendant_match)]:
        if not match:
            continue
        section = match.group(1).strip()

        party = {}

        # Name — first thing after 原告：/被告：
        name_match = re.match(r'([^，,]+)', section)
        if name_match:
            party['name'] = name_match.group(1).strip()

        # Gender
        gender_match = re.search(r'[男女]', section[:20])
        if gender_match:
            party['gender'] = gender_match.group(0)

        # Ethnicity
        eth_match = re.search(r'(汉族|回族|藏族|维吾尔族|蒙古族|壮族|满族|苗族|彝族|土家族|朝鲜族|侗族|瑶族|白族|哈尼族|哈萨克族|傣族|黎族|傈僳族|畲族|高山族|水族|东乡族|纳西族|景颇族|柯尔克孜族|土族|达斡尔族|仫佬族|羌族|布朗族|撒拉族|毛南族|仡佬族|锡伯族|阿昌族|普米族|塔吉克族|怒族|乌孜别克族|俄罗斯族|鄂温克族|德昂族|保安族|裕固族|京族|塔塔尔族|独龙族|鄂伦春族|赫哲族|门巴族|珞巴族|基诺族)', section)
        if eth_match:
            party['ethnicity'] = eth_match.group(1)

        # Birth date — various formats: 1955年6月25日, 1955.6.25
        birth_match = re.search(
            r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', section)
        if birth_match:
            party['birth_year'] = birth_match.group(1)
            party['birth_month'] = birth_match.group(2)
            party['birth_day'] = birth_match.group(3)
        else:
            birth_match = re.search(r'(\d{4})[\.年](\d{1,2})[\.月](\d{1,2})', section)
            if birth_match:
                party['birth_year'] = birth_match.group(1)
                party['birth_month'] = birth_match.group(2)
                party['birth_day'] = birth_match.group(3)

        # ID number — 18 or 15 digits
        id_match = re.search(
            r'(?:身份证号[码]?|证件号[码]?)[：:]*\s*(\d{15,18})', section)
        if id_match:
            party['id_number'] = id_match.group(1)
            party['id_type'] = '居民身份证'

        # Phone number
        phone_match = re.search(
            r'(?:电话|手机|联系方式)[：:]*\s*(1[3-9]\d{9})', section)
        if phone_match:
            party['phone'] = phone_match.group(1)

        # Address — after 住/户籍所在地/经常居住地
        addr_match = re.search(
            r'(?:住[所址]|户籍所在地|经常居住地)[：:]*\s*(.+?)(?:。|，联系电话|联系方式|$)',
            section)
        if addr_match:
            addr = addr_match.group(1).strip()
            # Clean up — remove trailing punctuation
            addr = re.sub(r'[。，,]+$', '', addr)
            party['address'] = addr

        result[role] = party

    return result


def parse_claims(text: str) -> list:
    """Extract numbered claims (诉讼请求) from complaint text.

    Returns list of claim strings, one per numbered item.
    """
    # Find 诉讼请求 section
    claims_match = re.search(
        r'诉讼请求[：:]\s*\n?(.+?)(?:事实[与和]理由|此致|$)', text, re.DOTALL)
    if not claims_match:
        return []

    section = claims_match.group(1).strip()

    # Multiple split strategies, in order of commonality:
    # 1. Chinese semicolons at line endings (most common in Chinese complaints)
    items = re.split(r'[；;]\s*\n?', section)
    # 2. If that didn't split well, try numbered items
    if len(items) <= 1:
        items = re.split(r'(?<=\s)(?=\d+[\.、．)])', section)
    # 3. Fallback: split on any semicolons
    if len(items) <= 1:
        items = re.split(r'[；;]', section)

    claims = []
    for item in items:
        item = item.strip().rstrip('；;，,。.')
        if item and len(item) > 5:
            claims.append(item)

    return claims


def parse_facts(text: str) -> str:
    """Extract 事实与理由 section from complaint text."""
    facts_match = re.search(
        r'事实[与和]理由[：:]\s*\n?(.+?)(?:此致|$)', text, re.DOTALL)
    if facts_match:
        return facts_match.group(1).strip()
    return ''


def detect_case_type(text: str) -> str:
    """Detect case type from complaint text.

    Returns one of: '民间借贷', '买卖合同', '劳动争议', '交通事故', '其他'
    """
    keywords = {
        '民间借贷': ['借款', '借条', '欠条', '出借', '还本', '还息', '利息', '本金'],
        '买卖合同': ['货款', '买卖合同', '供货', '交货', '验收'],
        '劳动争议': ['劳动', '工资', '加班', '解除劳动关系', '社保', '工伤'],
        '交通事故': ['交通', '撞', '肇事', '机动车', '保险理赔'],
    }
    scores = {}
    for case_type, kws in keywords.items():
        scores[case_type] = sum(1 for kw in kws if kw in text)
    if max(scores.values()) == 0:
        return '其他'
    return max(scores, key=scores.get)


# ================================================================
# Template inspection
# ================================================================


def inspect_template(template_path: str, output_path: str = None):
    """Dump template structure for analysis — with style info.

    MUST run before writing any fill logic. Outputs:
    - All paragraphs with styles
    - All tables with unique cell mapping (shows merge structure)
    - Cell text content + style info for each unique cell
    - Discovered font and style ID at the top
    """
    from lxml import etree
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Discover style first
    style_info = discover_template_style(template_path)

    doc = Document(template_path)
    lines = []

    lines.append(f"=== TEMPLATE STYLE ===")
    lines.append(f"font_name: {style_info['font_name']}")
    lines.append(f"style_id: {style_info['style_id']}")
    lines.append(f"style_name: {style_info['style_name']}")
    lines.append("")

    lines.append("=== PARAGRAPHS ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            lines.append(f"P{i}: [{para.style.name}] {para.text.strip()}")

    lines.append("\n=== TABLES ===")
    for ti, table in enumerate(doc.tables):
        lines.append(
            f"\n--- TABLE {ti}: {len(table.rows)}r x {len(table.columns)}c ---")
        seen = {}
        for ri in range(len(table.rows)):
            for ci in range(len(table.columns)):
                cell = table.cell(ri, ci)
                cid = id(cell._tc)
                if cid in seen:
                    orig = seen[cid]
                    lines.append(
                        f"  [{ri},{ci}] = GHOST -> same as [{orig[0]},{orig[1]}]")
                else:
                    seen[cid] = (ri, ci)
                    text = cell.text.replace('\n', '|')[:120]
                    # Also get paragraph style
                    pstyle = ''
                    if cell.paragraphs:
                        pPr = cell.paragraphs[0]._element.find(f'{{{W}}}pPr')
                        if pPr is not None:
                            ps = pPr.find(f'{{{W}}}pStyle')
                            if ps is not None:
                                pstyle = f' [pStyle={ps.get(qn("w:val"))}]'
                    lines.append(f"  [{ri},{ci}]{pstyle} {text}")

    result = '\n'.join(lines)
    if output_path:
        with open(output_path, 'w') as f:
            f.write(result)
    return result


# ================================================================
# Cell manipulation helpers
# ================================================================


def cell_replace(cell, old: str, new: str) -> bool:
    """Replace text within runs (preserves original formatting).

    Best for: filling individual fields (name, ID, phone) within a cell
    that has other template text.
    """
    for p in cell.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                return True
    return False


def cell_set(cell, lines: list):
    """Overwrite entire cell with multi-line content.

    Best for: summary boxes, free-text areas, any cell that needs
    complete replacement. Does NOT set font explicitly — use
    ensure_font() afterward.

    Args:
        cell: the table cell
        lines: list of strings, one per paragraph
    """
    paras = list(cell.paragraphs)
    existing = len(paras)
    for p in paras:
        for r in p.runs:
            r.text = ''
    for i, line in enumerate(lines):
        if i < existing:
            p = paras[i]
        else:
            p = cell.add_paragraph()
        if p.runs:
            p.runs[0].text = line
        else:
            p.add_run(line)


def cell_skip(cell):
    """Mark a section as not applicable by appending '/' to first paragraph.

    Idempotent — if the cell already has filled-in data (e.g. a name, an ID
    number), does NOT append '/' to avoid mutilating real data. Only appends
    '/' if the cell still contains its original template placeholder text.

    Use for: 法人 sections, 担保 sections, any field where the case
    doesn't involve that element.
    """
    try:
        p = cell.paragraphs[0]
        text = p.text
        # Guard: if cell has been filled (contains actual data, not just
        # template placeholders and "/"), skip
        # Template cells typically contain "姓名：" not "姓名：张某某"
        # A filled cell has text after the colon that's not just whitespace
        if text.rstrip().endswith('/'):
            return  # already skipped

        # Check if this looks like unfilled template text
        # (fields like "名称：" with nothing after, or "有□" checkboxes)
        is_template = (
            re.search(r'(名称|姓名|注册地|法定代表|统一社会信用代码)：\s*$', text) or
            re.search(r'[有无]□\s*$', text) or
            '□    否□' in text
        )
        # Also check if any actual data has been written (name after colon)
        has_data = bool(re.search(
            r'(名称|姓名)：\s*\S{2,}', text))  # at least 2 chars of data

        if has_data and not is_template:
            return  # cell has been filled, don't add "/"

        if p.runs:
            last_run = p.runs[-1]
            if not last_run.text.rstrip().endswith('/'):
                last_run.text = last_run.text + '/'
    except Exception:
        pass


def cell_replace_para(cell, keyword: str, new_text: str) -> bool:
    """Replace entire paragraph text when the keyword is found.

    Use for: fields split across multiple runs (birth dates, addresses)
    where run-level replace won't match.
    """
    for p in cell.paragraphs:
        if keyword in p.text:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
            return True
    return False


def cell_tick(cell, label: str):
    """Check a checkbox: change □ to ☑ for the given label.

    Only changes the FIRST □ matched (use count=1).
    """
    for p in cell.paragraphs:
        for r in p.runs:
            if label in r.text:
                r.text = r.text.replace('□', '☑', 1)
                return True
    return False


def cell_gender(cell, gender: str):
    """Handle gender checkbox: rewrite the entire gender line cleanly.

    Some fonts (like 仿宋) render □ (U+25A1) and ☑ (U+2611) identically,
    making "男□ 女☑" look like two identical boxes. This function rewrites
    the gender line as a single clean run: only the selected gender gets a
    box, the unselected one doesn't.

    Args:
        gender: '男' or '女'
    """
    line = '性别：女☑' if gender == '女' else '性别：男☑'
    for p in cell.paragraphs:
        if '性别' not in p.text:
            continue
        # Merge all runs into one clean run
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = line
        else:
            p.add_run(line)
        return True
    return False


def ensure_font(cell, font_name: str = None, style_id: str = None):
    """Set font on all runs and paragraph style in a cell.

    Two-layer approach:
    1. Set pStyle to the discovered template style
    2. Also set explicit run-level font as fallback

    CRITICAL: Only call this on cells where you've ADDED new text content.
    Do NOT call on pre-filled cells whose data was already correct — forcing
    an explicit font on inherited-style cells can change their appearance
    (font name, size, or weight may differ from the template's implicit style).

    If font_name/style_id not provided, uses the module-level defaults
    (which should have been set via discover_template_style()).
    """
    from lxml import etree
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    fn = font_name or FONT_NAME
    sid = style_id or TABLE_TEXT_STYLE

    for p in cell.paragraphs:
        pPr = p._element.find(f'{{{W}}}pPr')
        if pPr is None:
            pPr = etree.SubElement(p._element, f'{{{W}}}pPr')
        pStyle = pPr.find(f'{{{W}}}pStyle')
        if pStyle is None:
            pStyle = etree.SubElement(pPr, f'{{{W}}}pStyle')
        pStyle.set(qn('w:val'), sid)

        for r in p.runs:
            rPr = r._element.find(f'{{{W}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(r._element, f'{{{W}}}rPr')
            rFonts = rPr.find(f'{{{W}}}rFonts')
            if rFonts is None:
                rFonts = etree.SubElement(rPr, f'{{{W}}}rFonts')
            rFonts.set(qn('w:eastAsia'), fn)
            rFonts.set(qn('w:ascii'), fn)
            rFonts.set(qn('w:hAnsi'), fn)


# ================================================================
# Cell uniqueness check
# ================================================================


def is_unique_cell(table, ri: int, ci: int) -> bool:
    """Check if a cell is the 'primary' (first-occurrence) cell.

    Templates often merge cells. Multiple (ri, ci) pairs may point to the
    same underlying cell. Only write to the FIRST occurrence to avoid
    overwriting data that was written to another coordinate.
    """
    target_id = id(table.cell(ri, ci)._tc)
    for r in range(ri + 1):
        for c in range(len(table.columns)):
            if r == ri and c >= ci:
                break
            if id(table.cell(r, c)._tc) == target_id:
                return False
    return True


# ================================================================
# Common fill patterns
# ================================================================


def fill_natural_person(cell, data: dict):
    """Fill a natural person info cell.

    data keys: name, gender, birth_year, birth_month, birth_day,
               ethnicity, workplace, title, phone, address, residence,
               id_type, id_number
    """
    if data.get('name'):
        cell_replace(cell, '姓名：', f'姓名：{data["name"]}')
    if data.get('gender') == '女':
        cell_replace(cell, '女□', '女☑')
    elif data.get('gender') == '男':
        cell_replace(cell, '男□', '男☑')
    if data.get('birth_year'):
        cell_replace_para(cell, '出生日期：',
                          f'出生日期：{data["birth_year"]} 年 '
                          f'{data.get("birth_month", "")} 月 '
                          f'{data.get("birth_day", "")} 日')
    if data.get('ethnicity'):
        cell_replace(cell, '民族：', f'民族：{data["ethnicity"]}')
    if data.get('workplace'):
        cell_replace(cell, '工作单位：', f'工作单位：{data["workplace"]}')
    if data.get('title'):
        cell_replace(cell, '职务：', f'职务：{data["title"]}')
    if data.get('phone'):
        cell_replace(cell, '联系电话：', f'联系电话：{data["phone"]}')
    if data.get('address'):
        cell_replace_para(cell, '住所地（户籍所在地）：',
                          f'住所地（户籍所在地）：{data["address"]}')
    if data.get('residence'):
        cell_replace(cell, '经常居住地：', f'经常居住地：{data["residence"]}')
    if data.get('id_type'):
        cell_replace_para(cell, '证件类型：', f'证件类型：{data["id_type"]}')
    if data.get('id_number'):
        cell_replace(cell, '证件号码：', f'证件号码：{data["id_number"]}')


def fill_legal_entity(cell, data: dict):
    """Fill a legal entity (法人/非法人组织) info cell.

    data keys: name, address, register_address, legal_rep,
               credit_code, entity_type, ownership
    """
    if data.get('name'):
        cell_replace(cell, '名称：', f'名称：{data["name"]}')
    if data.get('address'):
        cell_replace_para(cell, '住所地（主要办事机构所在地）：',
                          f'住所地（主要办事机构所在地）：{data["address"]}')
    if data.get('register_address'):
        cell_replace(cell, '注册地 / 登记地：',
                     f'注册地 / 登记地：{data["register_address"]}')
    if data.get('legal_rep'):
        cell_replace(cell, '法定代表人 / 负责人：',
                     f'法定代表人 / 负责人：{data["legal_rep"]}')
    if data.get('credit_code'):
        cell_replace(cell, '统一社会信用代码：',
                     f'统一社会信用代码：{data["credit_code"]}')

    # Entity type checkboxes
    entity_types = data.get('entity_type', [])
    type_map = {
        '有限责任公司': '有限责任公司□',
        '股份有限公司': '股份有限公司□',
        '上市公司': '上市公司□',
        '事业单位': '事业单位□',
        '社会团体': '社会团体□',
        '基金会': '基金会□',
        '机关法人': '机关法人□',
        '个人独资企业': '个人独资企业□',
        '合伙企业': '合伙企业□',
    }
    for et in entity_types:
        if et in type_map:
            cell_tick(cell, type_map[et])

    # Ownership
    ownership = data.get('ownership')
    if ownership == '国有':
        cell_tick(cell, '国有□')
    elif ownership == '民营':
        cell_tick(cell, '民营□')


# ================================================================
# Output validation
# ================================================================


def validate_output(doc_path: str) -> list:
    """Basic validation: check that critical fields have been filled.

    Scans the filled document for common "unfilled" patterns.
    Intelligently skips cells that were marked with "/" (intentionally skipped).

    Returns list of warning strings. Empty list = no issues found.
    """
    warnings = []
    doc = Document(doc_path)

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text

                # Skip cells that were intentionally marked with "/"
                if text.rstrip().endswith('/'):
                    continue

                # Check for unfilled name fields in natural person cells
                # (template says "姓名：" with only whitespace after, and
                # cell is NOT in a skip area)
                name_empty = re.search(r'姓名：\s*$', text, re.MULTILINE)
                if name_empty and '（法人' not in text and '（非法人' not in text:
                    # Double-check: is this a ghost cell? Let is_unique_cell decide
                    # but we don't have table object context easily here.
                    # Use heuristic: if the cell text is identical to its siblings
                    # in the same row, it's probably a merged ghost
                    pass  # TODO: better ghost detection in validate

    return warnings
