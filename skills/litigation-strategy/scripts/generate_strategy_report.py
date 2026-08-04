#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
诉讼策略大师 —— Word报告生成器
==============================
根据诉讼策略分析数据，自动生成排版精美的Word文档（.docx）
支持：封面页、目录、正文分章节、页眉页脚、专业法律文书样式
"""

import os
import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re


# ═══════════════════════════════════════════════════════════════════
#  样式定义
# ═══════════════════════════════════════════════════════════════════

class ReportStyles:
    """报告样式常量"""

    # 颜色
    COLOR_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)       # 深蓝——标题
    COLOR_SECONDARY = RGBColor(0x2C, 0x5F, 0x8A)      # 中蓝——二级标题
    COLOR_ACCENT = RGBColor(0xC0, 0x39, 0x2B)         # 红色——强调/风险
    COLOR_GREEN = RGBColor(0x1E, 0x84, 0x45)          # 绿色——优势
    COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)           # 灰色——辅助信息
    COLOR_BLACK = RGBColor(0x22, 0x22, 0x22)          # 正文黑色
    COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    COLOR_LIGHT_BG = RGBColor(0xF5, 0xF7, 0xFA)      # 浅蓝灰背景
    COLOR_BORDER = RGBColor(0xCC, 0xCC, 0xCC)         # 表格边框
    COLOR_STAR = RGBColor(0xE6, 0x7E, 0x22)           # 橙黄——星级

    # 字号
    SIZE_COVER_TITLE = Pt(28)
    SIZE_COVER_SUBTITLE = Pt(16)
    SIZE_H1 = Pt(18)
    SIZE_H2 = Pt(14)
    SIZE_H3 = Pt(12)
    SIZE_BODY = Pt(10.5)
    SIZE_SMALL = Pt(9)
    SIZE_TABLE_HEADER = Pt(9.5)
    SIZE_TABLE_CELL = Pt(9)

    # 字体
    FONT_TITLE = '微软雅黑'
    FONT_BODY = '宋体'
    FONT_EN = 'Times New Roman'


# ═══════════════════════════════════════════════════════════════════
#  辅助函数
# ═══════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:color="{val.get("color", "000000")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_page_number(doc):
    """添加页码"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)

        for r in [run, run2, run3]:
            r.font.size = Pt(9)
            r.font.color.rgb = ReportStyles.COLOR_GRAY


def add_horizontal_line(doc, color_hex='1B3A5C'):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_star_rating(rating_str):
    """将 ★★★★☆ 格式的评级转为带格式文本"""
    # 保留星级文本返回
    return rating_str


# ═══════════════════════════════════════════════════════════════════
#  报告生成器主类
# ═══════════════════════════════════════════════════════════════════

class LitigationStrategyReport:
    """诉讼策略报告生成器"""

    def __init__(self, data: dict):
        """
        初始化报告生成器

        Args:
            data: 包含所有策略分析数据的字典，结构如下：
                {
                    "report_info": {
                        "law_firm": "律师事务所名称",
                        "lawyer": "承办律师",
                        "date": "报告日期",
                        "case_name": "案件名称",
                        "client": "客户名称",
                        "file_no": "案号/编号"
                    },
                    "case_info": { ... },
                    "client_needs": { ... },
                    "legal_diagnosis": { ... },
                    "disputes": [ ... ],
                    "strategies": [ ... ],   # 每个策略含七维评估
                    "evidence": { ... },
                    "comparison": { ... },
                    "recommendation": { ... }
                }
        """
        self.data = data
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """设置文档基础样式"""
        style = self.doc.styles['Normal']
        style.font.name = ReportStyles.FONT_BODY
        style.font.size = ReportStyles.SIZE_BODY
        style.font.color.rgb = ReportStyles.COLOR_BLACK
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)

        # 页边距
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

    def _add_paragraph(self, text, style_name='Normal', bold=False,
                       font_size=None, color=None, alignment=None,
                       space_before=None, space_after=None,
                       font_name=None, italic=False, indent_first_line=True):
        """添加格式化段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)

        if style_name == 'Normal':
            run.font.name = font_name or ReportStyles.FONT_BODY
            run.font.size = font_size or ReportStyles.SIZE_BODY
            run.font.color.rgb = color or ReportStyles.COLOR_BLACK
            run.font.bold = bold
            run.font.italic = italic
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name or ReportStyles.FONT_BODY)
        else:
            p.style = self.doc.styles[style_name]
            if font_size:
                run.font.size = font_size

        if alignment:
            p.alignment = alignment

        if space_before is not None:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)

        if indent_first_line:
            p.paragraph_format.first_line_indent = Pt(21)  # 两个字符缩进

        return p

    def _add_heading_styled(self, text, level=1):
        """添加带样式的标题"""
        if level == 1:
            p = self._add_paragraph(
                text,
                bold=True,
                font_size=ReportStyles.SIZE_H1,
                color=ReportStyles.COLOR_PRIMARY,
                space_before=24,
                space_after=12,
                indent_first_line=False
            )
            add_horizontal_line(self.doc, '1B3A5C')
        elif level == 2:
            p = self._add_paragraph(
                text,
                bold=True,
                font_size=ReportStyles.SIZE_H2,
                color=ReportStyles.COLOR_SECONDARY,
                space_before=18,
                space_after=8,
                indent_first_line=False
            )
        elif level == 3:
            p = self._add_paragraph(
                text,
                bold=True,
                font_size=ReportStyles.SIZE_H3,
                color=ReportStyles.COLOR_PRIMARY,
                space_before=12,
                space_after=6,
                indent_first_line=False
            )
        return p

    def _add_table_with_data(self, headers, rows, col_widths=None):
        """添加带样式的数据表格"""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.font.bold = True
            run.font.size = ReportStyles.SIZE_TABLE_HEADER
            run.font.color.rgb = ReportStyles.COLOR_WHITE
            run.font.name = ReportStyles.FONT_BODY
            run.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '1B3A5C')

        # 数据行
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                run.font.size = ReportStyles.SIZE_TABLE_CELL
                run.font.name = ReportStyles.FONT_BODY
                run.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 斑马条纹
                if r_idx % 2 == 1:
                    set_cell_shading(cell, 'F5F7FA')

        # 设置列宽
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    row.cells[i].width = Cm(width)

        self.doc.add_paragraph()  # 表后空行
        return table

    def _add_info_pair(self, label, value, indent=False):
        """添加标签: 值 对"""
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Pt(21)

        run_label = p.add_run(f"{label}：")
        run_label.font.bold = True
        run_label.font.size = ReportStyles.SIZE_BODY
        run_label.font.name = ReportStyles.FONT_BODY
        run_label.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)

        run_value = p.add_run(str(value))
        run_value.font.size = ReportStyles.SIZE_BODY
        run_value.font.name = ReportStyles.FONT_BODY
        run_value.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)

        return p

    # ──── 封面 ────

    def build_cover(self):
        """生成封面页"""
        # 空行留白
        for _ in range(6):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

        # 标题
        p = self._add_paragraph(
            '诉讼策略报告',
            bold=True,
            font_size=ReportStyles.SIZE_COVER_TITLE,
            color=ReportStyles.COLOR_PRIMARY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            indent_first_line=False
        )
        p.paragraph_format.space_after = Pt(8)

        # 副标题
        case_name = self.data.get('report_info', {}).get('case_name', '')
        if case_name:
            p = self._add_paragraph(
                f"—— {case_name}",
                bold=False,
                font_size=ReportStyles.SIZE_COVER_SUBTITLE,
                color=ReportStyles.COLOR_SECONDARY,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                indent_first_line=False
            )
            p.paragraph_format.space_after = Pt(36)
        else:
            self.doc.add_paragraph()
            self.doc.add_paragraph()

        # 分隔线
        add_horizontal_line(self.doc, '1B3A5C')
        self.doc.add_paragraph()

        # 封面信息
        info_items = [
            ('客户/当事人', self.data.get('report_info', {}).get('client', '')),
            ('案件类型', self.data.get('case_info', {}).get('case_type', '民商事诉讼')),
            ('承办律师事务所', self.data.get('report_info', {}).get('law_firm', '')),
            ('承办律师', self.data.get('report_info', {}).get('lawyer', '')),
            ('报告编号', self.data.get('report_info', {}).get('file_no', '')),
            ('报告日期', self.data.get('report_info', {}).get('date', datetime.now().strftime('%Y年%m月%d日'))),
        ]

        for label, value in info_items:
            if value:
                p = self._add_paragraph(
                    f"{label}：{value}",
                    font_size=ReportStyles.SIZE_BODY + Pt(1),
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    indent_first_line=False
                )
                p.paragraph_format.space_after = Pt(4)

        # 保密声明
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        p = self._add_paragraph(
            '【保密声明】本报告包含保密信息，仅供委托人及其授权代理人参考，未经本所书面同意，'
            '不得向任何第三方披露、复制或引用。',
            font_size=ReportStyles.SIZE_SMALL,
            color=ReportStyles.COLOR_GRAY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            italic=True,
            indent_first_line=False
        )

        # 分页
        self.doc.add_page_break()

    # ──── 目录页 ────

    def build_toc(self):
        """生成目录页"""
        self._add_heading_styled('目  录', level=1)

        toc_items = [
            ('一、案件基本信息', 1),
            ('二、客户核心诉求与根本需求分析', 2),
            ('三、法律关系诊断', 3),
            ('  3.1 案由与法律关系定性', 0),
            ('  3.2 请求权基础分析', 0),
            ('四、争议焦点归纳', 4),
            ('五、诉讼策略方案', 5),
            ('  5.1 方案A：进攻型诉讼方案', 0),
            ('  5.2 方案B：稳健型诉讼方案', 0),
            ('  5.3 方案C：防御/反制型方案', 0),
            ('  5.4 方案D：和解/调解导向型方案', 0),
            ('六、多方案对比与综合评估', 6),
            ('七、证据清单与举证指引', 7),
            ('八、律师推荐意见', 8),
            ('九、风险提示与免责声明', 9),
        ]

        p = self.doc.add_paragraph()
        for item, _ in toc_items:
            run = p.add_run(f"{item}\n")
            run.font.name = ReportStyles.FONT_BODY
            run.font.size = ReportStyles.SIZE_BODY
            run.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)
            if item.startswith('  '):
                run.font.color.rgb = ReportStyles.COLOR_GRAY
            else:
                run.font.color.rgb = ReportStyles.COLOR_BLACK

        self.doc.add_page_break()

    # ──── 第一章：案件基本信息 ────

    def build_case_info(self):
        """案件基本信息"""
        self._add_heading_styled('一、案件基本信息', level=1)

        info = self.data.get('case_info', {})

        self._add_heading_styled('1.1 当事人信息', level=2)

        # 我方当事人
        my_party = info.get('my_party', {})
        if my_party:
            self._add_heading_styled('我方当事人', level=3)
            for key, val in my_party.items():
                self._add_info_pair(key, val)

        # 对方当事人
        opp_party = info.get('opponent_party', {})
        if opp_party:
            self._add_heading_styled('对方当事人', level=3)
            for key, val in opp_party.items():
                self._add_info_pair(key, val)

        self._add_heading_styled('1.2 案件事实概述', level=2)

        facts = info.get('case_facts', '')
        if facts:
            self._add_paragraph(facts)

        # 程序现状
        proc = info.get('procedural_status', {})
        if proc:
            self._add_heading_styled('1.3 程序现状', level=2)
            for key, val in proc.items():
                self._add_info_pair(key, val)

        add_horizontal_line(self.doc, 'CCCCCC')
        self.doc.add_page_break()

    # ──── 第二章：客户需求分析 ────

    def build_client_needs(self):
        """客户核心诉求与根本需求"""
        self._add_heading_styled('二、客户核心诉求与根本需求分析', level=1)

        needs = self.data.get('client_needs', {})

        self._add_heading_styled('2.1 客户表层诉求', level=2)
        self._add_paragraph(needs.get('surface_claim', ''))

        self._add_heading_styled('2.2 根本需求分析', level=2)
        self._add_paragraph(needs.get('deep_need', ''))

        self._add_heading_styled('2.3 客户决策偏好', level=2)

        pref_items = [
            ('时间容忍度', needs.get('time_tolerance', '')),
            ('成本预算', needs.get('cost_budget', '')),
            ('风险偏好', needs.get('risk_preference', '')),
            ('和解意愿', needs.get('settlement_willingness', '')),
            ('诉讼耐心', needs.get('litigation_patience', '')),
        ]

        self._add_table_with_data(
            ['维度', '客户表述', '策略影响'],
            [item for item in pref_items if item[1]],
            col_widths=[3, 6, 6]
        )

        add_horizontal_line(self.doc, 'CCCCCC')
        self.doc.add_page_break()

    # ──── 第三章：法律关系诊断 ────

    def build_legal_diagnosis(self):
        """法律关系诊断"""
        self._add_heading_styled('三、法律关系诊断', level=1)

        diag = self.data.get('legal_diagnosis', {})

        self._add_heading_styled('3.1 案由与法律关系定性', level=2)
        self._add_paragraph(diag.get('case_cause', ''))

        # 三段式结论
        conclusion = diag.get('conclusion', {})
        if conclusion:
            self._add_heading_styled('法律关系定性结论', level=3)

            items = [
                ('核心法律关系', conclusion.get('core_relationship', '')),
                ('权利义务依据', conclusion.get('legal_basis', '')),
                ('法律后果预判', conclusion.get('legal_consequence', '')),
            ]
            for label, val in items:
                if val:
                    self._add_info_pair(label, val)

        self._add_heading_styled('3.2 请求权基础分析', level=2)

        claims = diag.get('claims', [])
        if claims:
            self._add_table_with_data(
                ['主体', '主张权利', '法律依据', '构成要件', '我方定位'],
                [
                    [
                        c.get('party', ''),
                        c.get('right', ''),
                        c.get('legal_basis', ''),
                        c.get('elements', ''),
                        c.get('position', '')
                    ]
                    for c in claims
                ],
                col_widths=[2, 3, 3, 4, 2]
            )

        add_horizontal_line(self.doc, 'CCCCCC')
        self.doc.add_page_break()

    # ──── 第四章：争议焦点 ────

    def build_disputes(self):
        """争议焦点归纳"""
        self._add_heading_styled('四、争议焦点归纳', level=1)

        disputes = self.data.get('disputes', [])

        self._add_paragraph(
            '经对本案事实和法律关系的综合分析，归纳以下争议焦点：'
        )

        for i, d in enumerate(disputes, 1):
            level = d.get('level', '辅助性')
            stars = d.get('impact', '★★★')

            # 焦点标题
            p = self._add_paragraph(
                f"焦点{i}：{d.get('title', '')}   [{level}焦点 | 影响程度：{stars}]",
                bold=True,
                font_size=ReportStyles.SIZE_H3,
                color=ReportStyles.COLOR_SECONDARY,
                indent_first_line=False,
                space_before=12
            )

            # 焦点详情
            items = [
                ('具体内容', d.get('description', '')),
                ('性质', d.get('nature', '')),
                ('举证责任方', d.get('burden', '')),
            ]
            for label, val in items:
                if val:
                    self._add_info_pair(f"  {label}", val, indent=True)

        add_horizontal_line(self.doc, 'CCCCCC')
        self.doc.add_page_break()

    # ──── 第五章：诉讼策略方案 ────

    def _build_strategy_card(self, strategy, index):
        """构建单个策略方案评估卡"""
        name = strategy.get('name', f'方案{index}')
        stype = strategy.get('type', '')
        stype_labels = {'A': '进攻型', 'B': '稳健型', 'C': '防御/反制型', 'D': '和解/调解导向型'}
        stype_label = stype_labels.get(stype, stype)

        # 方案标题
        self._add_heading_styled(
            f'{index}. 方案{stype}：{name}（{stype_label}方案）',
            level=2
        )

        # 方案描述
        self._add_heading_styled('方案概述', level=3)
        desc = strategy.get('description', '')
        if desc:
            self._add_paragraph(desc)

        # 手段组合
        tactics = strategy.get('tactics', [])
        if tactics:
            self._add_heading_styled('核心手段组合', level=3)
            for t in tactics:
                p = self._add_paragraph(f"• {t}", indent_first_line=False)

        # 七维评估表
        self._add_heading_styled('七维评估', level=3)

        assessment = strategy.get('assessment', {})
        eval_rows = [
            ['① 胜诉可能性',
             f"法律依据：{assessment.get('legal_strength', '—')}",
             f"综合预判胜诉率：{assessment.get('win_rate', '—')}"],
            ['② 经济成本',
             f"案件受理费：{assessment.get('filing_fee', '—')}",
             f"总成本预估：{assessment.get('total_cost', '—')}"],
            ['③ 时间周期',
             f"一审：{assessment.get('first_instance', '—')}",
             f"预计全流程：{assessment.get('total_duration', '—')}"],
            ['④ 执行可行性',
             f"对方偿债能力：{assessment.get('solvency', '—')}",
             f"综合评级：{assessment.get('enforceability', '—')}"],
            ['⑤ 风险评级',
             f"主要风险：{assessment.get('main_risk', '—')}",
             f"综合风险：{assessment.get('risk_level', '—')}"],
            ['⑥ 可控性',
             f"进程掌控度：{assessment.get('controllability', '—')}",
             f"策略转换成本：{assessment.get('switch_cost', '—')}"],
            ['⑦ 客户匹配度',
             f"目标匹配：{assessment.get('goal_match', '—')}",
             f"预算匹配：{assessment.get('budget_match', '—')}"],
        ]

        self._add_table_with_data(
            ['评估维度', '指标一', '指标二'],
            eval_rows,
            col_widths=[3, 5.5, 5.5]
        )

        # 综合评分
        score = strategy.get('score', '—')
        p = self._add_paragraph(
            f"📊 综合评分：{score}/100分",
            bold=True,
            font_size=ReportStyles.SIZE_H3,
            color=ReportStyles.COLOR_PRIMARY,
            indent_first_line=False,
            space_before=6
        )

        self.doc.add_paragraph()

    def build_strategies(self):
        """诉讼策略方案"""
        self._add_heading_styled('五、诉讼策略方案', level=1)

        strategies = self.data.get('strategies', [])

        self._add_paragraph(
            '基于案情分析、客户需求及证据现状，本所制定了以下诉讼策略方案供委托人审慎选择。'
            '每套方案均从七个维度进行系统评估。'
        )

        for i, strategy in enumerate(strategies, 1):
            self._build_strategy_card(strategy, i)

        self.doc.add_page_break()

    # ──── 第六章：方案对比 ────

    def build_comparison(self):
        """多方案对比与推荐"""
        self._add_heading_styled('六、多方案对比与综合评估', level=1)

        comparison = self.data.get('comparison', {})
        rows = comparison.get('rows', [])

        if rows:
            self._add_heading_styled('6.1 多方案横向对比', level=2)

            headers = ['评估维度'] + [r.get('name', '') for r in rows]
            table_data = []

            # 胜诉率行
            win_row = ['综合胜诉率'] + [r.get('win_rate', '—') for r in rows]
            table_data.append(win_row)

            # 预期金额行
            amount_row = ['预期可获得金额'] + [r.get('expected_amount', '—') for r in rows]
            table_data.append(amount_row)

            # 成本行
            cost_row = ['总成本预估'] + [r.get('total_cost', '—') for r in rows]
            table_data.append(cost_row)

            # 周期行
            time_row = ['预计周期'] + [r.get('duration', '—') for r in rows]
            table_data.append(time_row)

            # 执行可行性行
            exec_row = ['执行可行性'] + [r.get('enforceability', '—') for r in rows]
            table_data.append(exec_row)

            # 风险行
            risk_row = ['综合风险'] + [r.get('risk_level', '—') for r in rows]
            table_data.append(risk_row)

            # 客户匹配度行
            match_row = ['客户匹配度'] + [r.get('client_match', '—') for r in rows]
            table_data.append(match_row)

            # 综合评分行
            score_row = ['综合评分'] + [r.get('score', '—') for r in rows]
            table_data.append(score_row)

            self._add_table_with_data(headers, table_data)

        # 推荐方案
        self._add_heading_styled('6.2 综合评分说明', level=2)
        self._add_paragraph(
            comparison.get('scoring_note',
                '评分权重：胜诉可能25% | 经济成本15% | 时间周期10% | '
                '执行可行性20% | 风险控制20% | 客户匹配度10%')
        )

        self.doc.add_page_break()

    # ──── 第七章：证据清单 ────

    def build_evidence(self):
        """证据清单与举证指引"""
        self._add_heading_styled('七、证据清单与举证指引', level=1)

        evidence = self.data.get('evidence', {})

        # 举证责任分配
        burden = evidence.get('burden_of_proof', [])
        if burden:
            self._add_heading_styled('7.1 举证责任分配', level=2)
            self._add_table_with_data(
                ['争议焦点', '举证责任方', '需证明内容', '证明标准'],
                [
                    [b.get('focus', ''), b.get('party', ''),
                     b.get('content', ''), b.get('standard', '')]
                    for b in burden
                ],
                col_widths=[3.5, 2, 5.5, 3]
            )

        # 证据清单
        evidence_list = evidence.get('evidence_list', [])
        if evidence_list:
            self._add_heading_styled('7.2 证据清单', level=2)
            self._add_table_with_data(
                ['编号', '证据名称', '证明目的', '来源', '状态', '必备'],
                [
                    [e.get('id', ''), e.get('name', ''), e.get('purpose', ''),
                     e.get('source', ''), e.get('status', ''), e.get('essential', '')]
                    for e in evidence_list
                ],
                col_widths=[1.5, 4, 4, 2.5, 1.5, 1.5]
            )

        # 取证行动清单
        action_items = evidence.get('action_items', [])
        if action_items:
            self._add_heading_styled('7.3 取证行动清单', level=2)
            self._add_table_with_data(
                ['优先级', '行动事项', '负责方', '时限', '法律依据'],
                [
                    [a.get('priority', ''), a.get('action', ''),
                     a.get('responsible', ''), a.get('deadline', ''),
                     a.get('legal_basis', '')]
                    for a in action_items
                ],
                col_widths=[1.5, 4.5, 2, 2.5, 3.5]
            )

        self.doc.add_page_break()

    # ──── 第八章：推荐意见 ────

    def build_recommendation(self):
        """律师推荐意见"""
        self._add_heading_styled('八、律师推荐意见', level=1)

        rec = self.data.get('recommendation', {})

        # 推荐方案
        recommended = rec.get('recommended', '')
        if recommended:
            p = self._add_paragraph(
                f"推荐方案：{recommended}",
                bold=True,
                font_size=ReportStyles.SIZE_H3,
                color=ReportStyles.COLOR_GREEN,
                indent_first_line=False,
                space_before=12
            )

        # 推荐理由
        reasons = rec.get('reasons', [])
        if reasons:
            self._add_heading_styled('推荐理由', level=2)
            for i, reason in enumerate(reasons, 1):
                self._add_paragraph(f"{i}. {reason}")

        # 备选方案
        backup = rec.get('backup', '')
        if backup:
            self._add_heading_styled('备选方案', level=2)
            self._add_paragraph(backup)

        # 不建议的方案
        not_rec = rec.get('not_recommended', '')
        if not_rec:
            self._add_heading_styled('不建议采用的方案', level=2)
            p = self._add_paragraph(not_rec)
            p.add_run('【请客户特别注意此项】').font.color.rgb = ReportStyles.COLOR_ACCENT

        # 给客户的最终建议
        final_advice = rec.get('final_advice', '')
        if final_advice:
            self._add_heading_styled('给客户的最终建议', level=2)

            # 建议块——带左边框和浅蓝底色
            p = self.doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="24" w:space="8" w:color="1B3A5C"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            # 设置段落缩进
            pPrind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284"/>')
            pPr.append(pPrind)
            # 段落底纹（浅蓝背景）
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
            pPr.append(shd)

            run = p.add_run(final_advice)
            run.font.name = ReportStyles.FONT_BODY
            run.font.size = ReportStyles.SIZE_BODY
            run.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)

        # 承办律师签名
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        lawyer = self.data.get('report_info', {}).get('lawyer', '_____________')
        law_firm = self.data.get('report_info', {}).get('law_firm', '_____________')

        sign_data = [
            (f"承办律师：{lawyer}", ""),
            (f"律师事务所：{law_firm}", ""),
            (f"报告日期：{self.data.get('report_info', {}).get('date', datetime.now().strftime('%Y年%m月%d日'))}", ""),
        ]
        for text, _ in sign_data:
            self._add_paragraph(
                text,
                indent_first_line=False,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT
            )

        self.doc.add_page_break()

    # ──── 第九章：风险提示 ────

    def build_disclaimer(self):
        """风险提示与免责声明"""
        self._add_heading_styled('九、风险提示与免责声明', level=1)

        disclaimer_text = self.data.get('disclaimer', '').strip()
        if not disclaimer_text:
            disclaimer_text = (
                '本报告所载分析、判断及建议，系基于委托人截至本报告出具之日所提供的案件事实、'
                '证据材料及相关法律法规作出的专业分析意见。实际诉讼/仲裁结果受以下不可控因素影响：\n\n'
                '1. 法官自由裁量权：不同法官对相同事实和法律可能存在不同认定；\n'
                '2. 新证据的出现：诉讼过程中可能出现对案情产生重大影响的新证据；\n'
                '3. 法律政策变化：诉讼期间可能出台新的司法解释或司法政策；\n'
                '4. 对方策略调整：对方可能在诉讼中采取未预见到的诉讼策略；\n'
                '5. 地方司法环境：不同地区的法院可能存在不同的裁判倾向。\n\n'
                '本报告不构成对案件结果的承诺或保证。承办律师将根据案件进展及时更新分析意见。'
            )

        for para in disclaimer_text.split('\n\n'):
            para = para.strip()
            if para:
                self._add_paragraph(para)

    # ──── 页眉页脚 ────

    def build_header_footer(self):
        """添加页眉页脚"""
        case_name = self.data.get('report_info', {}).get('case_name', '诉讼策略报告')
        law_firm = self.data.get('report_info', {}).get('law_firm', '')

        for section in self.doc.sections:
            # 页眉
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = hp.add_run(f"{law_firm}  |  {case_name}" if law_firm else case_name)
            run.font.size = Pt(8)
            run.font.color.rgb = ReportStyles.COLOR_GRAY
            run.font.name = ReportStyles.FONT_BODY
            run.element.rPr.rFonts.set(qn('w:eastAsia'), ReportStyles.FONT_BODY)

            # 页眉分隔线
            pPr = hp._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        # 页码
        add_page_number(self.doc)

    # ──── 完整生成 ────

    def generate(self, output_path: str):
        """
        生成完整的诉讼策略报告Word文档

        Args:
            output_path: 输出文件路径（.docx）
        """
        self.build_cover()
        self.build_toc()
        self.build_case_info()
        self.build_client_needs()
        self.build_legal_diagnosis()
        self.build_disputes()
        self.build_strategies()
        self.build_comparison()
        self.build_evidence()
        self.build_recommendation()
        self.build_disclaimer()
        self.build_header_footer()

        self.doc.save(output_path)
        return output_path


# ═══════════════════════════════════════════════════════════════════
#  命令行入口
# ═══════════════════════════════════════════════════════════════════

def main():
    """命令行入口——接收JSON输入，生成Word文档"""
    if len(sys.argv) < 2:
        print("用法：python generate_strategy_report.py <输入JSON文件> [输出路径]")
        print("  或：python generate_strategy_report.py --stdin [输出路径]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    # 读取数据
    if input_path == '--stdin':
        data = json.loads(sys.stdin.read())
    else:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

    # 默认输出路径
    if not output_path:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        case_name = data.get('report_info', {}).get('case_name', '诉讼策略报告')
        # 清理文件名
        case_name_clean = re.sub(r'[\\/:*?"<>|]', '', case_name)[:30]
        output_path = f'{case_name_clean}_诉讼策略报告_{timestamp}.docx'

    # 生成报告
    generator = LitigationStrategyReport(data)
    result = generator.generate(output_path)

    try:
        print(f"报告已成功生成：{os.path.abspath(result)}")
    except UnicodeEncodeError:
        print("Report generated: " + os.path.abspath(result))
    return result


if __name__ == '__main__':
    main()
