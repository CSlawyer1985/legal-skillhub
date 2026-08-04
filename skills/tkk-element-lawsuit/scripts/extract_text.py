#!/usr/bin/env python3
"""Extract text from DOCX or PDF file for BUDE 要素式起诉状转换工具."""

import sys
import os

def extract_docx(filepath):
    """Extract all text from a .docx file using python-docx."""
    from docx import Document
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return '\n'.join(lines)

def extract_pdf(filepath):
    """Extract all text from a .pdf file using pdfplumber."""
    import pdfplumber
    lines = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
    return '\n'.join(lines)

def main():
    if len(sys.argv) < 2:
        print("Usage: extract_text.py <filepath>", file=sys.stderr)
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"ERROR: File not found: {filepath}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        text = extract_docx(filepath)
    elif ext == '.pdf':
        text = extract_pdf(filepath)
    else:
        print(f"ERROR: Unsupported format: {ext}", file=sys.stderr)
        sys.exit(1)

    if not text.strip():
        print("ERROR: No text could be extracted from the file.", file=sys.stderr)
        sys.exit(1)

    print(text)

if __name__ == '__main__':
    main()
