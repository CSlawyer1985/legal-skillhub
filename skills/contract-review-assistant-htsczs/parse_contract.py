#!/usr/bin/env python3
"""
合同文档解析脚本
用于提取PDF/Word文档中的文本内容，辅助智能体进行合同审查

支持格式：PDF (.pdf)、Word (.docx, .doc)

使用方法：
python scripts/parse_contract.py <合同文件路径>
"""

import sys
import os

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None


def parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    if pdfplumber is None:
        return "错误：未安装 pdfplumber 库，请运行 pip install pdfplumber"

    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n\n".join(text)


def parse_docx(file_path: str) -> str:
    """解析Word文档"""
    if docx is None:
        return "错误：未安装 python-docx 库，请运行 pip install python-docx"

    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text.append(" | ".join(row_text))

    return "\n".join(text)


def parse_contract(file_path: str) -> str:
    """根据文件类型解析合同"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_path)
    else:
        return f"不支持的文件格式：{ext}，仅支持 PDF 和 Word 文档"


def main():
    if len(sys.argv) < 2:
        print("使用方法：python parse_contract.py <合同文件路径>")
        sys.exit(1)

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(f"错误：文件不存在 - {file_path}")
        sys.exit(1)

    print(f"正在解析文件：{file_path}")
    content = parse_contract(file_path)

    print("\n=== 解析结果 ===\n")
    print(content)

    # 输出到文件（方便后续处理）
    output_path = file_path + ".txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"\n文本内容已保存至：{output_path}")


if __name__ == "__main__":
    main()
