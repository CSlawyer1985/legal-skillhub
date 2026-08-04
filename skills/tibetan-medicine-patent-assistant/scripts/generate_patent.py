#!/usr/bin/env python3
"""
藏药专利初稿生成器

接受结构化专利数据输入（JSON格式），生成完整的 Word (.docx) 专利初稿。
使用 python-docx 1.2.0 API。

调用方式：
    python generate_patent.py --input input.json --output output.docx

输入JSON结构见 generate_patent_from_input() 函数。
"""

import json
import argparse
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 需要 python-docx 库。请运行: pip install python-docx")
    sys.exit(1)


def _set_cell_font(run, font_name_cn="宋体", font_name_en="Times New Roman", size=12):
    """设置 run 的中英文字体"""
    run.font.size = Pt(size)
    run.font.name = font_name_en
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name_cn)


def add_heading(doc, text, level=1):
    """添加标题，设置中文字体"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_cell_font(run, font_name_cn="黑体", size=16 if level == 1 else 14 if level == 2 else 12)
    return heading


def add_paragraph(doc, text, bold=False, font_size=12, alignment=None, first_line_indent=True, font_name_cn="仿宋"):
    """添加段落，自动处理缩进和字体"""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进
    run = para.add_run(text)
    run.bold = bold
    _set_cell_font(run, font_name_cn=font_name_cn, size=font_size)
    return para


def add_claim(doc, text, claim_num):
    """添加权利要求项"""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(f"{claim_num}. {text}")
    _set_cell_font(run, font_name_cn="宋体", size=12)
    return para


def generate_patent(data, output_path):
    """
    生成完整的专利说明书 Word 文档

    参数 data 结构：
    {
        "title": "发明名称",
        "formula": {
            "description": "处方描述",
            "ingredients": [
                {"name": "药材名", "weight_parts": "重量份数", "weight_range": "重量份数范围"}
            ]
        },
        "dosage_form": {
            "type": "剂型",
            "features": "剂型特征描述"
        },
        "process": {
            "steps": ["步骤1", "步骤2", ...],
            "key_parameters": {"参数名": "参数值", ...},
            "description": "工艺整体描述"
        },
        "quality_standard": {
            "items": [
                {"item": "鉴别/检查/含量测定等", "method": "方法", "criteria": "标准"}
            ],
            "description": "质量标准概述"
        },
        "technical_field": "技术领域描述",
        "background": "背景技术",
        "technical_problem": "要解决的技术问题",
        "beneficial_effects": {
            "description": "有益效果描述",
            "data": "实验数据"
        },
        "pharmacology_data": "药效学实验数据",
        "embodiments": [
            {"title": "实施例1", "content": "具体内容"}
        ],
        "abstract": "说明书摘要",
        "claims": [
            "权利要求1文本",
            "权利要求2文本"
        ]
    }
    """
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ========== 说明书摘要 ==========
    add_heading(doc, "说明书摘要", level=1)
    abstract = data.get("abstract", "")
    if abstract:
        add_paragraph(doc, abstract)
    else:
        add_paragraph(doc, "（摘要内容待补充）")

    doc.add_page_break()

    # ========== 权利要求书 ==========
    add_heading(doc, "权利要求书", level=1)
    claims = data.get("claims", [])
    if claims:
        for i, claim_text in enumerate(claims, 1):
            add_claim(doc, claim_text, i)
    else:
        add_paragraph(doc, "（权利要求内容待补充）")

    doc.add_page_break()

    # ========== 说明书 ==========
    add_heading(doc, "说明书", level=1)

    # 发明名称
    title = data.get("title", "（发明名称待补充）")
    add_heading(doc, "发明名称", level=2)
    add_paragraph(doc, title)

    # 技术领域
    add_heading(doc, "技术领域", level=2)
    tech_field = data.get("technical_field", "")
    if tech_field:
        add_paragraph(doc, tech_field)
    else:
        add_paragraph(doc, f"本发明属于藏药/中药技术领域，具体涉及一种{title}。")

    # 背景技术
    add_heading(doc, "背景技术", level=2)
    bg = data.get("background", "")
    if bg:
        add_paragraph(doc, bg)
    else:
        add_paragraph(doc, "（背景技术待补充，需包含藏医药理论依据、现有技术不足等）")

    # 发明内容
    add_heading(doc, "发明内容", level=2)

    # 要解决的技术问题
    add_heading(doc, "要解决的技术问题", level=3)
    problem = data.get("technical_problem", "")
    if problem:
        add_paragraph(doc, problem)
    else:
        add_paragraph(doc, "本发明旨在提供一种……，以解决现有技术中……的技术问题。")

    # 技术方案
    add_heading(doc, "技术方案", level=3)

    # 处方
    formula = data.get("formula", {})
    formula_desc = formula.get("description", "")
    if formula_desc:
        add_paragraph(doc, formula_desc)
    ingredients = formula.get("ingredients", [])
    if ingredients:
        add_paragraph(doc, "本发明所述藏药组合物由以下重量份的原料药组成：", bold=True)
        parts_text = ""
        for ing in ingredients:
            name = ing.get("name", "")
            weight = ing.get("weight_parts", "")
            w_range = ing.get("weight_range", "")
            if w_range:
                parts_text += f"{name} {w_range}份；"
            elif weight:
                parts_text += f"{name} {weight}份；"
            else:
                parts_text += f"{name} 适量；"
        add_paragraph(doc, parts_text.rstrip("；") + "。")

    # 剂型
    dosage = data.get("dosage_form", {})
    d_type = dosage.get("type", "")
    d_features = dosage.get("features", "")
    if d_type:
        add_paragraph(doc, f"本发明所述组合物的剂型为{d_type}。" + (d_features or ""))
        add_paragraph(doc, d_features)

    # 制备工艺
    process = data.get("process", {})
    process_desc = process.get("description", "")
    if process_desc:
        add_paragraph(doc, f"制备工艺：{process_desc}")
    steps = process.get("steps", [])
    if steps:
        add_paragraph(doc, "具体制备步骤如下：", bold=True)
        for idx, step in enumerate(steps, 1):
            add_paragraph(doc, f"（{idx}）{step}")
    params = process.get("key_parameters", {})
    if params:
        add_paragraph(doc, "其中关键工艺参数为：", bold=True)
        for k, v in params.items():
            add_paragraph(doc, f"{k}：{v}")

    # 质量标准
    qs = data.get("quality_standard", {})
    qs_desc = qs.get("description", "")
    if qs_desc:
        add_paragraph(doc, f"质量标准：{qs_desc}")
    qs_items = qs.get("items", [])
    if qs_items:
        add_paragraph(doc, "本发明的质量检测方法包括：", bold=True)
        for item in qs_items:
            i_name = item.get("item", "")
            i_method = item.get("method", "")
            i_criteria = item.get("criteria", "")
            item_text = f"• {i_name}"
            if i_method:
                item_text += f"：{i_method}"
            if i_criteria:
                item_text += f"（标准：{i_criteria}）"
            add_paragraph(doc, item_text)

    # 有益效果
    add_heading(doc, "有益效果", level=3)
    eff = data.get("beneficial_effects", {})
    eff_desc = eff.get("description", "")
    if eff_desc:
        add_paragraph(doc, eff_desc)
    eff_data = eff.get("data", "")
    if eff_data:
        add_paragraph(doc, eff_data)

    # 药效学数据
    pharma = data.get("pharmacology_data", "")
    if pharma:
        add_paragraph(doc, pharma)

    # 具体实施方式
    add_heading(doc, "具体实施方式", level=2)
    add_paragraph(doc, "下面结合具体实施例对本发明作进一步说明，但本发明并不限于以下实施例。")
    add_paragraph(doc, "")

    embodiments = data.get("embodiments", [])
    if embodiments:
        for emb in embodiments:
            e_title = emb.get("title", "实施例")
            e_content = emb.get("content", "")
            add_heading(doc, e_title, level=3)
            if e_content:
                add_paragraph(doc, e_content)
    else:
        add_paragraph(doc, "（实施例内容待补充，需包含至少一个具体实施配方和工艺）")

    # 保存文档
    doc.save(output_path)
    return output_path


def generate_patent_from_input(input_data, output_path):
    """
    从用户提供的结构数据快速构建专利初稿。

    用户提供的简化输入：
    {
        "title": "发明名称",
        "formula": {
            "description": "处方描述",
            "ingredients": [{"name": "A", "weight_parts": "10", "weight_range": "5-15"}]
        },
        "dosage_form": {"type": "贴剂", "features": "..."},
        "process": {"description": "...", "steps": ["步骤1..."]},
        "quality_standard": {...},
        "technical_field": "...",
        "background": "...",
        "technical_problem": "...",
        "technical_solution": "...",
        "beneficial_effects": {"description": "...", "data": "..."},
        "pharmacology_data": "...",
        "embodiments": [...],
        "abstract": "...",
        "claims": [...]
    }
    """
    return generate_patent(input_data, output_path)


def main():
    parser = argparse.ArgumentParser(description="藏药专利初稿生成器")
    parser.add_argument("--input", "-i", required=True, help="输入JSON文件路径")
    parser.add_argument("--output", "-o", default="patent_draft.docx", help="输出DOCX文件路径")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    output_path = generate_patent(data, args.output)
    print(f"专利初稿已生成: {os.path.abspath(output_path)}")


if __name__ == "__main__":
    main()
