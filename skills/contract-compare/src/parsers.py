"""
Contract file parsers for multiple formats.
Supports: TXT, DOCX, PDF, Image (JPG/PNG with OCR).
"""
import os
import chardet
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class EncodingError(Exception):
    """Raised when text file encoding cannot be detected."""
    pass


class ParseError(Exception):
    """Raised when a file cannot be parsed."""
    pass


def detect_encoding(file_path: str) -> str:
    """Detect text file encoding using chardet.

    Args:
        file_path: Path to the text file

    Returns:
        Detected encoding name (e.g., 'utf-8', 'gbk')

    Raises:
        EncodingError: If encoding cannot be detected
    """
    with open(file_path, 'rb') as f:
        raw_data = f.read()
    result = chardet.detect(raw_data)
    encoding = result.get('encoding', 'utf-8')
    confidence = result.get('confidence', 0)

    if not encoding or confidence < 0.6:
        # Fallback to utf-8
        encoding = 'utf-8'

    return encoding


def parse_txt(file_path: str) -> str:
    """Parse a plain text file with encoding detection.

    Args:
        file_path: Path to the TXT file

    Returns:
        Extracted text content

    Raises:
        ParseError: If file cannot be read
    """
    try:
        encoding = detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            text = f.read()
        return text.strip()
    except Exception as e:
        raise ParseError(f"Failed to read TXT file: {e}")


def parse_docx(file_path: str) -> str:
    """Parse a DOCX file and extract text.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content

    Raises:
        ParseError: If file cannot be read
    """
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(' | '.join(cells))
        return '\n'.join(paragraphs)
    except ImportError:
        raise ParseError("python-docx not installed")
    except Exception as e:
        raise ParseError(f"Failed to read DOCX file: {e}")


def parse_pdf(file_path: str) -> str:
    """Parse a PDF file and extract text.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content

    Raises:
        ParseError: If file cannot be read
    """
    text_parts = []

    # Try PyMuPDF first
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        doc.close()
        if text_parts:
            return '\n'.join(text_parts)
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}")

    # Fallback to pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
        if text_parts:
            return '\n'.join(text_parts)
    except ImportError:
        raise ParseError("No PDF library available (install PyMuPDF or pdfplumber)")
    except Exception as e:
        raise ParseError(f"Failed to read PDF file: {e}")

    if not text_parts:
        raise ParseError("PDF contains no extractable text (may be scanned image)")
    return '\n'.join(text_parts)


def parse_image(file_path: str) -> str:
    """Parse an image file using OCR.

    Args:
        file_path: Path to the image file (JPG/PNG)

    Returns:
        Extracted text content

    Raises:
        ParseError: If OCR fails or tesseract not available
    """
    try:
        from PIL import Image
        import pytesseract
    except ImportError as e:
        raise ParseError(f"OCR libraries not installed: {e}")

    try:
        image = Image.open(file_path)
        # Convert to RGB if necessary
        if image.mode not in ('L', 'RGB'):
            image = image.convert('RGB')

        text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        return text.strip()
    except Exception as e:
        raise ParseError(f"OCR failed: {e}")


def parse_contract(file_path: str) -> str:
    """Parse a contract file of any supported type.

    Automatically detects file type from extension.

    Args:
        file_path: Path to the contract file

    Returns:
        Extracted text content

    Raises:
        ParseError: If file type is unsupported or parsing fails
    """
    if not os.path.isfile(file_path):
        raise ParseError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower().lstrip('.')

    parsers = {
        'txt': parse_txt,
        'docx': parse_docx,
        'pdf': parse_pdf,
        'jpg': parse_image,
        'jpeg': parse_image,
        'png': parse_image,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ParseError(f"Unsupported file type: {ext}")

    return parser(file_path)


def sanitize_path(user_input: str) -> str:
    """Sanitize user input for use in file paths.

    Only allows alphanumeric characters, underscores, hyphens, and dots.

    Args:
        user_input: Raw user input

    Returns:
        Sanitized string safe for use in paths
    """
    import re
    sanitized = re.sub(r'[^a-zA-Z0-9_.-]', '_', user_input)
    return sanitized
