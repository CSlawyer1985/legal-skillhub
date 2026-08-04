# -*- coding: utf-8 -*-
import sys
import os
import subprocess
from pathlib import Path

# ── 自动依赖检查（首次自动安装） ──
_REQUIREMENTS = Path(__file__).resolve().parent.parent / "requirements.txt"
def _ensure_deps():
    try:
        import docx  # noqa: F401
    except ImportError:
        print("[chinese-contract-drafting] 检测到依赖缺失，正在自动安装...", file=sys.stderr)
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "-r", str(_REQUIREMENTS), "--quiet"]
        )
        print("[chinese-contract-drafting] ✅ 依赖安装完成", file=sys.stderr)

_ensure_deps()

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx_utils import (
    create_formatted_doc, set_page_layout, add_title, add_contract_info,
    add_body_text, add_heading_level1, add_heading_level2, add_heading_level3, add_clause,
    add_signature_block, save_as, append_clause, add_cover_page
)

doc = create_formatted_doc()
set_page_layout(doc)

add_cover_page(
    doc,
    title_cn='软件产品售卖合同',
    title_en='Software Product Sales Contract',
    contract_no='xxx',
    party_a_label='甲  方（出卖方）：',
    party_b_label='乙  方（买受方）：'
)

add_title(doc, '\u8f6f\u4ef6\u4ea7\u54c1\u552e\u5356\u5408\u540c')

add_body_text(doc, '\u7532\u65b9\uff08\u51fa\u5356\u65b9\uff09\uff1axxx')
add_body_text(doc, '\u4e59\u65b9\uff08\u4e70\u53d7\u65b9\uff09\uff1axxx')
add_body_text(doc, '\u5408\u540c\u7f16\u53f7\uff1axxx')

add_body_text(doc, '\u9274\u4e8e\u7532\u65b9\u662f\u201c\u65c5\u6e38\u5206\u9500\u5546\u667a\u80fd\u4f53\u7cfb\u7edf\u201d\u8f6f\u4ef6\u7684\u5408\u6cd5\u5f00\u53d1\u8005\u548c\u77e5\u8bc6\u4ea7\u6743\u6743\u5229\u4eba\uff0c\u4e59\u65b9\u6709\u610f\u8d2d\u4e70\u8be5\u8f6f\u4ef6\u4ea7\u54c1\uff0c\u53cc\u65b9\u6839\u636e\u300a\u4e2d\u534e\u4eba\u6c11\u5171\u548c\u56fd\u6c11\u6cd5\u5178\u300b\u53ca\u76f8\u5173\u6cd5\u5f8b\u6cd5\u89c4\u7684\u89c4\u5b9a\uff0c\u5728\u5e73\u7b49\u81ea\u613f\u7684\u57fa\u7840\u4e0a\uff0c\u7ecf\u53cb\u597d\u534f\u5546\uff0c\u5c31\u7532\u65b9\u51fa\u552e\u3001\u4e59\u65b9\u8d2d\u4e70\u4e0a\u8ff0\u8f6f\u4ef6\u4ea7\u54c1\u4e8b\u5b9c\u8fbe\u6210\u5982\u4e0b\u534f\u8bae\uff1a')

# 第一条 定义与解释
add_heading_level1(doc, '第一条  定义与解释')
add_body_text(doc, '1.1 \u201c软件产品\u201d：指甲方开发的\u201c旅游分销商智能体系统\u201d软件（以下简称\u201c本软件\u201d），具体功能模块以本合同附件一《软件功能清单》为准。')
add_body_text(doc, '1.2 \u201c交付\u201d：指甲方按照本合同约定的方式和期限，将软件产品及相关文档移交乙方的行为。')
add_body_text(doc, '1.3 \u201c验收\u201d：指乙方按照本合同约定的标准和程序，对软件产品进行检验并确认合格的行为。')
add_body_text(doc, '1.4 \u201c质保期\u201d：指自验收合格之日起，甲方对软件产品承担质量保证责任的期间。')
add_body_text(doc, '1.5 \u201c知识产权\u201d：指与软件产品相关的全部著作权、专利权、商标权、商业秘密及其他合法权益。')

# 第二条 合同标的
add_heading_level1(doc, '第二条  合同标的')
add_body_text(doc, '2.1 甲方向乙方出售的软件产品为\u201c旅游分销商智能体系统\u201d，具体功能模块、技术参数及性能指标详见附件一《软件功能清单》。')
add_body_text(doc, '2.2 甲方授予乙方对本软件的非独占、不可转让的使用许可，使用范围限于乙方自身业务经营，不得向任何第三方转售、转授权或提供托管服务。')
add_body_text(doc, '2.3 本软件的使用许可地域范围为中国境内。如需境外使用，双方另行协商并签订补充协议。')

# 第三条 合同金额与支付
add_heading_level1(doc, '第三条  合同金额与支付')
add_body_text(doc, '3.1 本合同软件产品总价款为人民币xxx元（大写：xxx元整），该价格为含税价格（增值税税率为xxx%）。')
add_body_text(doc, '3.2 付款方式：xxx')
add_body_text(doc, '3.3 甲方收款账户信息：xxx')
add_body_text(doc, '3.4 乙方逾期付款的，每逾期一日，应按逾期未付金额的万分之五向甲方支付违约金。逾期超过三十日的，甲方有权解除本合同，已收取的款项不予退还，并有权要求乙方赔偿由此造成的全部损失。')

# 第四条 交付
add_heading_level1(doc, '第四条  交付')
add_body_text(doc, '4.1 交付内容：甲方应向乙方交付以下内容：')
add_body_text(doc, '（1）旅游分销商智能体系统软件安装包；')
add_body_text(doc, '（2）软件使用手册（电子版）；')
add_body_text(doc, '（3）软件部署文档（电子版）；')
add_body_text(doc, '（4）附件一《软件功能清单》所列明的其他交付物。')
add_body_text(doc, '4.2 交付方式：甲方采用远程交付与现场部署相结合的方式。甲方先通过远程方式向乙方提供软件安装包及相关文档，随后派技术人员到乙方现场进行部署安装和调试。')
add_body_text(doc, '4.3 交付期限：甲方应在本合同签订且收到乙方首付款后xxx个工作日内完成全部交付。')
add_body_text(doc, '4.4 因乙方原因（包括但不限于未及时提供部署环境、网络条件不满足、现场配合不到位等）导致交付延迟的，交付期限相应顺延，甲方不承担违约责任。')
add_body_text(doc, '4.5 软件产品自交付之日起，毁损、灭失的风险由乙方承担。')

# 第五条 验收
add_heading_level1(doc, '第五条  验收')
add_body_text(doc, '5.1 验收标准：以附件一《软件功能清单》约定的功能和技术指标为验收依据。')
add_body_text(doc, '5.2 验收期限：乙方应在甲方完成交付后十（10）个工作日内组织验收。验收期限届满，乙方未提出书面异议的，视为验收合格。')
add_body_text(doc, '5.3 验收不合格的，乙方应在验收期限内以书面形式向甲方提出具体不合格事项。甲方在收到书面异议后十五（15）个工作日内进行修复或调整，并再次提交验收。')
add_body_text(doc, '5.4 经两次验收仍不合格的，双方应协商解决方案。协商不成的，任何一方有权解除合同，甲方退还乙方已付款项扣除已交付成果对应价值后的余额。')
add_body_text(doc, '5.5 乙方在验收期限内未组织验收，或已在实际业务中使用本软件的，均视为验收合格。')

# 第六条 知识产权
add_heading_level1(doc, '第六条  知识产权')
add_body_text(doc, '6.1 本软件的全部知识产权（包括但不限于著作权、专利权、商标权）归甲方所有。本合同的签订不构成任何知识产权的转让。')
add_body_text(doc, '6.2 甲方授予乙方对本软件的非独占使用许可，许可范围仅限于乙方自身业务使用，不得转授权。')
add_body_text(doc, '6.3 乙方不得对本软件进行反向工程、反编译、反汇编或以其他方式试图获取本软件的源代码。')
add_body_text(doc, '6.4 乙方不得删除、修改或遮盖本软件中的任何版权声明、商标或其他知识产权标识。')
add_body_text(doc, '6.5 甲方保证其为软件产品的合法开发者，拥有签订和履行本合同所需的完整知识产权。如因知识产权纠纷导致乙方无法正常使用本软件，甲方应承担相应责任，但赔偿上限不超过本合同总金额。')
add_body_text(doc, '6.6 乙方在使用本软件过程中产生的业务数据归乙方所有，甲方不得以任何方式获取、使用或泄露乙方的业务数据。')

# 第七条 技术服务与质量保证
add_heading_level1(doc, '第七条  技术服务与质量保证')
add_body_text(doc, '7.1 质保期为自验收合格之日起六（6）个月。')
add_body_text(doc, '7.2 质保期内，甲方负责对软件自身缺陷（Bug）进行修复，响应时间不超过四十八（48）小时。')
add_body_text(doc, '7.3 以下情形不属于免费质保范围：')
add_body_text(doc, '（1）乙方未按使用手册操作或擅自修改软件导致的故障；')
add_body_text(doc, '（2）乙方硬件环境、网络环境变化导致的软件运行异常；')
add_body_text(doc, '（3）第三方软件或系统与本软件兼容性问题；')
add_body_text(doc, '（4）因不可抗力导致的软件损坏；')
add_body_text(doc, '（5）乙方要求新增功能或修改现有功能的开发需求。')
add_body_text(doc, '7.4 质保期满后，如乙方需要甲方继续提供技术支持或软件升级服务，双方另行签订技术服务合同。')
add_body_text(doc, '7.5 甲方在质保期内的累计赔偿责任不超过本合同总金额的百分之三十（30%）。')

# 第八条 保密条款
add_heading_level1(doc, '第八条  保密条款')
add_body_text(doc, '8.1 双方对在履行本合同过程中知悉的对方商业秘密和技术秘密负有保密义务。')
add_body_text(doc, '8.2 乙方特别承诺对以下信息严格保密：')
add_body_text(doc, '（1）本软件的源代码、算法、架构设计等技术细节；')
add_body_text(doc, '（2）甲方提供的技术文档、接口文档、数据库结构；')
add_body_text(doc, '（3）甲方在部署和技术支持过程中接触到的乙方系统配置信息。')
add_body_text(doc, '8.3 保密期限自本合同签订之日起至合同终止后三（3）年止。')
add_body_text(doc, '8.4 违反保密义务的一方应赔偿对方因此遭受的全部损失。')

# 第九条 违约责任
add_heading_level1(doc, '第九条  违约责任')
add_body_text(doc, '9.1 甲方违约责任：')
add_body_text(doc, '（1）甲方逾期交付的，每逾期一日，应按合同总金额的万分之三向乙方支付违约金；')
add_body_text(doc, '（2）甲方累计违约金总额不超过合同总金额的百分之十（10%）；')
add_body_text(doc, '（3）因甲方原因导致合同无法继续履行的，乙方有权解除合同，甲方应退还乙方已付款项中对应未完成服务部分的费用。')
add_body_text(doc, '9.2 乙方违约责任：')
add_body_text(doc, '（1）乙方逾期付款的，适用本合同第3.4条约定；')
add_body_text(doc, '（2）乙方违反知识产权条款（第六条）的，应立即停止违约行为，并向甲方支付合同总金额百分之五十（50%）的违约金，违约金不足以弥补甲方损失的，乙方还应赔偿差额部分；')
add_body_text(doc, '（3）乙方违反保密条款（第八条）的，适用第8.4条约定。')
add_body_text(doc, '9.3 在任何情况下，甲方对本合同项下承担的累计赔偿责任（包括违约金、损害赔偿等）不超过本合同总金额。')

# 第十条 不可抗力
add_heading_level1(doc, '第十条  不可抗力')
add_body_text(doc, '10.1 因不可抗力（包括但不限于自然灾害、战争、政府行为、法律法规变更、网络攻击、电信故障、疫情等）导致合同无法履行的，受影响一方不承担违约责任。')
add_body_text(doc, '10.2 遭受不可抗力的一方应在事件发生后五（5）日内书面通知对方，并在十五（15）日内提供相关证明材料。')
add_body_text(doc, '10.3 不可抗力事件持续超过六十（60）日的，任何一方有权书面通知对方解除本合同，双方互不承担违约责任，甲方退还乙方已付款项中对应未履行部分的费用。')

# 第十一条 合同变更与解除
add_heading_level1(doc, '第十一条  合同变更与解除')
add_body_text(doc, '11.1 本合同的任何修改、补充均须经双方协商一致，并以书面形式签订补充协议。补充协议与本合同具有同等法律效力。')
add_body_text(doc, '11.2 有下列情形之一的，甲方有权书面通知乙方解除本合同：')
add_body_text(doc, '（1）乙方逾期付款超过三十日的；')
add_body_text(doc, '（2）乙方违反知识产权条款或保密条款的；')
add_body_text(doc, '（3）乙方将本软件用于违法活动的。')
add_body_text(doc, '11.3 合同解除后，乙方应立即停止使用本软件并删除全部软件副本，甲方已收取的款项不予退还。')

# 第十二条 争议解决
add_heading_level1(doc, '第十二条  争议解决')
add_body_text(doc, '12.1 因本合同引起的或与本合同有关的任何争议，双方应首先通过友好协商解决。')
add_body_text(doc, '12.2 协商不成的，任何一方有权向甲方所在地有管辖权的人民法院提起诉讼。')

# 第十三条 其他约定
add_heading_level1(doc, '第十三条  其他约定')
add_body_text(doc, '13.1 本合同自双方法定代表人或授权代表签字并加盖公章之日起生效。')
add_body_text(doc, '13.2 本合同一式肆份，甲乙双方各执贰份，具有同等法律效力。')
add_body_text(doc, '13.3 本合同附件为本合同不可分割的组成部分，与本合同正文具有同等法律效力。')
add_body_text(doc, '13.4 本合同未尽事宜，双方可另行签订补充协议。补充协议与本合同不一致的，以补充协议为准。')
add_body_text(doc, '13.5 本合同任何条款被认定为无效或不可执行的，不影响其他条款的效力。')

# 附件清单
add_heading_level1(doc, '附件清单')
add_body_text(doc, '附件一：《软件功能清单》')
add_body_text(doc, '附件二：《验收标准》')

# 签署区
add_body_text(doc, '（以下为签署页，无正文）')
add_signature_block(doc, party_a_name='甲方（签章）：', party_b_name='乙方（签章）：')

# ========== 附件一：软件功能清单 ==========
doc.add_page_break()
add_heading_level1(doc, '附件一：软件功能清单')

add_body_text(doc, '软件名称：旅游分销商智能体系统')
add_body_text(doc, '软件版本：V1.0')
add_body_text(doc, '交付形式：软件安装包 + 使用手册 + 部署文档')

add_heading_level2(doc, '一、核心功能模块')

add_heading_level3(doc, '1. 分销商管理模块')
add_body_text(doc, '（1）分销商注册与审核：支持分销商在线注册、资质审核、等级评定；')
add_body_text(doc, '（2）分销商信息管理：分销商基本信息维护、联系人管理、账户状态管理；')
add_body_text(doc, '（3）分销商业绩统计：实时统计分销商销售额、订单量、佣金收益。')

add_heading_level3(doc, '2. 产品管理模块')
add_body_text(doc, '（1）旅游产品上架：支持线路、酒店、门票、机票等多类型产品发布；')
add_body_text(doc, '（2）产品价格管理：支持多价格体系、动态定价、促销活动配置；')
add_body_text(doc, '（3）库存管理：实时库存同步、库存预警、超售防护。')

add_heading_level3(doc, '3. 订单管理模块')
add_body_text(doc, '（1）订单全流程管理：下单、支付、确认、出行、结算全链路跟踪；')
add_body_text(doc, '（2）退改签处理：支持在线退改签申请、审核、退款；')
add_body_text(doc, '（3）订单异常处理：自动识别异常订单、人工干预处理。')

add_heading_level3(doc, '4. 智能推荐模块')
add_body_text(doc, '（1）客户画像分析：基于历史数据的客户偏好分析；')
add_body_text(doc, '（2）产品智能推荐：根据客户特征自动推荐匹配产品；')
add_body_text(doc, '（3）营销策略建议：基于数据分析的营销方案生成。')

add_heading_level3(doc, '5. 财务管理模块')
add_body_text(doc, '（1）佣金结算：自动计算分销商佣金、支持多结算周期；')
add_body_text(doc, '（2）对账管理：与供应商、分销商自动对账；')
add_body_text(doc, '（3）财务报表：收入、支出、利润多维度报表生成。')

add_heading_level3(doc, '6. 数据分析模块')
add_body_text(doc, '（1）经营数据看板：核心指标实时展示；')
add_body_text(doc, '（2）销售趋势分析：多维度销售数据分析；')
add_body_text(doc, '（3）用户行为分析：访问、转化、留存数据分析。')

add_heading_level2(doc, '二、系统技术要求')
add_body_text(doc, '（1）部署方式：支持私有化部署/云服务器部署；')
add_body_text(doc, '（2）并发能力：支持不少于500并发用户；')
add_body_text(doc, '（3）数据安全：支持数据加密存储、定期备份；')
add_body_text(doc, '（4）接口规范：提供标准RESTful API接口。')

add_heading_level2(doc, '三、交付物清单')
add_body_text(doc, '（1）旅游分销商智能体系统软件安装包；')
add_body_text(doc, '（2）系统部署文档（PDF格式）；')
add_body_text(doc, '（3）用户操作手册（PDF格式）；')
add_body_text(doc, '（4）API接口文档（PDF格式）；')
add_body_text(doc, '（5）数据库设计文档（PDF格式）。')

# ========== 附件二：验收标准 ==========
doc.add_page_break()
add_heading_level1(doc, '附件二：验收标准')

add_heading_level2(doc, '一、验收依据')
add_body_text(doc, '（1）本合同正文及附件一《软件功能清单》；')
add_body_text(doc, '（2）甲方提供的产品需求规格说明书；')
add_body_text(doc, '（3）国家及行业相关标准。')

add_heading_level2(doc, '二、验收流程')
add_body_text(doc, '（1）甲方完成部署后，向乙方提交验收申请；')
add_body_text(doc, '（2）乙方在收到验收申请后十（10）个工作日内组织验收；')
add_body_text(doc, '（3）验收过程中发现的问题，乙方以书面形式提交甲方；')
add_body_text(doc, '（4）甲方在收到问题清单后十五（15）个工作日内完成整改；')
add_body_text(doc, '（5）整改完成后重新提交验收，直至验收通过。')

add_heading_level2(doc, '三、功能验收标准')
add_body_text(doc, '（1）附件一《软件功能清单》所列全部功能均已实现并可正常使用；')
add_body_text(doc, '（2）各功能模块运行稳定，无严重缺陷（导致系统崩溃或数据丢失的缺陷）；')
add_body_text(doc, '（3）一般缺陷（不影响主要功能使用的缺陷）数量不超过五个，且有明确修复计划。')

add_heading_level2(doc, '四、性能验收标准')
add_body_text(doc, '（1）系统响应时间：页面加载时间不超过3秒（正常网络环境下）；')
add_body_text(doc, '（2）并发处理能力：支持不少于500用户同时在线操作；')
add_body_text(doc, '（3）数据处理能力：单次查询响应时间不超过5秒（百万级数据量）。')

add_heading_level2(doc, '五、安全验收标准')
add_body_text(doc, '（1）用户认证：支持用户名密码登录、短信验证码登录；')
add_body_text(doc, '（2）权限控制：支持角色权限管理，敏感操作有权限校验；')
add_body_text(doc, '（3）数据安全：用户密码加密存储，敏感信息脱敏展示；')
add_body_text(doc, '（4）日志审计：关键操作有日志记录，支持日志查询。')

add_heading_level2(doc, '六、文档验收标准')
add_body_text(doc, '（1）交付物清单所列文档齐全；')
add_body_text(doc, '（2）文档内容完整、表述清晰，能够指导实际使用；')
add_body_text(doc, '（3）文档版本与软件版本一致。')

add_heading_level2(doc, '七、验收通过条件')
add_body_text(doc, '同时满足以下条件视为验收通过：')
add_body_text(doc, '（1）功能验收、性能验收、安全验收、文档验收均符合标准；')
add_body_text(doc, '（2）无严重缺陷，一般缺陷数量在允许范围内；')
add_body_text(doc, '（3）乙方授权代表在验收报告上签字确认。')

add_heading_level2(doc, '八、验收不合格处理')
add_body_text(doc, '（1）验收不合格的，乙方应在验收期限内出具书面整改通知；')
add_body_text(doc, '（2）甲方应在收到通知后十五（15）个工作日内完成整改并重新提交验收；')
add_body_text(doc, '（3）经两次整改仍不合格的，双方协商解决；协商不成的，按本合同第5.4条处理。')

# 法律依据
add_heading_level1(doc, '法律依据')
add_body_text(doc, '1.《中华人民共和国民法典》第三编 合同，第四百六十九条至第四百七十八条（合同的订立）')
add_body_text(doc, '2.《中华人民共和国民法典》第五百零九条至第五百五十八条（合同的履行）')
add_body_text(doc, '3.《中华人民共和国民法典》第五百六十二条至第五百六十七条（合同的权利义务终止）')
add_body_text(doc, '4.《中华人民共和国民法典》第五百七十七条至第五百九十四条（违约责任）')
add_body_text(doc, '5.《中华人民共和国民法典》第六百一十条至第六百二十三条（买卖合同）')
add_body_text(doc, '6.《中华人民共和国著作权法》')
add_body_text(doc, '7.《计算机软件保护条例》')

output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '\u8f93\u51fa', '\u8f6f\u4ef6\u4ea7\u54c1\u552e\u5356\u5408\u540c_v3.docx')
save_as(doc, output_path)
print('done: ' + output_path)
