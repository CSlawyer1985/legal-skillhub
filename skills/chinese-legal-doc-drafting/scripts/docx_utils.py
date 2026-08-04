from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

FONT_FANGSONG = "仿宋"
FONT_HEITI = "黑体"
FONT_KAITI = "楷体"

PT_TITLE = Pt(22)
PT_BODY = Pt(16)
PT_PAGE_NUM = Pt(14)
PT_TABLE = Pt(12)
PT_FOOTNOTE = Pt(10.5)
PT_HEADER_FOOTER = Pt(9)

LINE_SPACING = Pt(28.95)
FIRST_LINE_INDENT = Pt(32)

MARGIN_TOP = Mm(37)
MARGIN_BOTTOM = Mm(35)
MARGIN_LEFT = Mm(28)
MARGIN_RIGHT = Mm(26)


def create_formatted_doc() -> Document:
    doc = Document()
    set_page_layout(doc)
    return doc


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def _set_run_font(run, font_name: str, font_size, bold: bool = False):
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_format(para, alignment=None, line_spacing=None, first_line_indent=None,
                          space_before=Pt(0), space_after=Pt(0)):
    pf = para.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = 4
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.space_before = space_before
    pf.space_after = space_after


def load_template(path: str) -> Document:
    return Document(path)


def extract_all_text(doc: Document) -> str:
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return "\n".join(lines)


def replace_text(doc: Document, old: str, new: str) -> bool:
    found = False
    for para in doc.paragraphs:
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    found = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            found = True
    return found


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          line_spacing=LINE_SPACING, space_before=Pt(22), space_after=Pt(22))
    run = para.add_run(text)
    _set_run_font(run, FONT_HEITI, PT_TITLE, bold=True)


def add_contract_info(doc: Document, contract_no: str = "", date: str = "",
                      party_a: str = "", party_b: str = "") -> None:
    if contract_no:
        para = doc.add_paragraph()
        _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=LINE_SPACING, space_after=Pt(16))
        run = para.add_run(contract_no)
        _set_run_font(run, FONT_FANGSONG, PT_BODY)

    if date:
        para = doc.add_paragraph()
        _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=LINE_SPACING, space_after=Pt(16))
        run = para.add_run(date)
        _set_run_font(run, FONT_FANGSONG, PT_BODY)

    for party_info in [party_a, party_b]:
        if party_info:
            para = doc.add_paragraph()
            _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                  line_spacing=LINE_SPACING, first_line_indent=FIRST_LINE_INDENT)
            run = para.add_run(party_info)
            _set_run_font(run, FONT_FANGSONG, PT_BODY)


def add_heading_level1(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing=LINE_SPACING, first_line_indent=FIRST_LINE_INDENT,
                          space_before=Pt(16), space_after=Pt(0))
    run = para.add_run(text)
    _set_run_font(run, FONT_HEITI, PT_BODY, bold=True)


def add_heading_level2(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing=LINE_SPACING, first_line_indent=FIRST_LINE_INDENT)
    run = para.add_run(text)
    _set_run_font(run, FONT_KAITI, PT_BODY, bold=True)


def add_heading_level3(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing=LINE_SPACING, first_line_indent=FIRST_LINE_INDENT)
    run = para.add_run(text)
    _set_run_font(run, FONT_FANGSONG, PT_BODY, bold=True)


def add_body_text(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing=LINE_SPACING, first_line_indent=FIRST_LINE_INDENT)
    run = para.add_run(text)
    _set_run_font(run, FONT_FANGSONG, PT_BODY, bold=bold)


def add_clause(doc: Document, clause_no: str = "", text: str = "") -> None:
    if not text and clause_no:
        text = clause_no
        clause_no = ""
    para = doc.add_paragraph()
    _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing=LINE_SPACING, first_line_indent=FIRST_LINE_INDENT)
    if clause_no:
        run_no = para.add_run(clause_no)
        _set_run_font(run_no, FONT_HEITI, PT_BODY, bold=True)
    run_text = para.add_run(text)
    _set_run_font(run_text, FONT_FANGSONG, PT_BODY)


def add_signature_block(doc: Document, party_a_name: str = "甲方（签章）：",
                        party_b_name: str = "乙方（签章）：",
                        date_placeholder: str = "日期：    年   月   日") -> None:
    for _ in range(2):
        empty = doc.add_paragraph()
        _set_paragraph_format(empty, line_spacing=LINE_SPACING)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cells_data = [
        (party_a_name, party_b_name),
        ("", ""),
        ("法定代表人/授权代表：", "法定代表人/授权代表："),
        (date_placeholder, date_placeholder),
    ]

    for row_idx, (left_text, right_text) in enumerate(cells_data):
        left_cell = table.cell(row_idx, 0)
        right_cell = table.cell(row_idx, 1)
        _clear_cell_and_write(left_cell, left_text)
        _clear_cell_and_write(right_cell, right_text)

    for _ in range(2):
        empty = doc.add_paragraph()
        _set_paragraph_format(empty, line_spacing=LINE_SPACING)


def _clear_cell_and_write(cell, text: str) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(text)
        _set_run_font(run, FONT_FANGSONG, PT_BODY)
        _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              line_spacing=LINE_SPACING)


def append_clause(doc: Document, text: str, style: str = "Normal") -> None:
    para = doc.add_paragraph(text, style=style)
    _set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing=LINE_SPACING, first_line_indent=FIRST_LINE_INDENT)
    for run in para.runs:
        _set_run_font(run, FONT_FANGSONG, PT_BODY)


def save_as(doc: Document, path: str) -> None:
    doc.save(path)


def add_cover_page(doc: Document, title_cn: str, title_en: str = "",
                   contract_no: str = "xxx", party_a_label: str = "甲  方：",
                   party_b_label: str = "乙  方：",
                   location: str = "xxx", date: str = "xxx年xxx月xxx日",
                   copies: str = "【本合同共计壹份，具有同等法律效力】") -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        _set_paragraph_format(p, line_spacing=LINE_SPACING)

    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          line_spacing=Pt(70), space_after=Pt(20))
    run = p.add_run(title_cn)
    _set_run_font(run, FONT_HEITI, Pt(42), bold=True)

    if title_en:
        p = doc.add_paragraph()
        _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=Pt(50), space_before=Pt(30))
        run = p.add_run(title_en)
        _set_run_font(run, FONT_KAITI, Pt(18))

    for _ in range(3):
        p = doc.add_paragraph()
        _set_paragraph_format(p, line_spacing=Pt(40))

    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(50))
    run = p.add_run(f"合同编号：{contract_no}")
    _set_run_font(run, FONT_FANGSONG, Pt(20))

    for _ in range(6):
        p = doc.add_paragraph()
        _set_paragraph_format(p, line_spacing=Pt(40))

    cover_items = [
        (party_a_label, "xxx"),
        (party_b_label, "xxx"),
        ("签订地点：", location),
        ("签订日期：", date),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(45))
        run1 = p.add_run(label)
        _set_run_font(run1, FONT_FANGSONG, Pt(18), bold=True)
        run2 = p.add_run(value)
        _set_run_font(run2, FONT_FANGSONG, Pt(18))

    for _ in range(3):
        p = doc.add_paragraph()
        _set_paragraph_format(p, line_spacing=Pt(40))

    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(40))
    run = p.add_run(copies)
    _set_run_font(run, FONT_FANGSONG, Pt(14))

    doc.add_page_break()
