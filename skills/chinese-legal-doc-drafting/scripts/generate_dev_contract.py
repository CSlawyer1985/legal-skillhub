# -*- coding: utf-8 -*-
import sys
import os
import subprocess
from pathlib import Path

# ── 自动依赖检查（首次自动安装） ──
_REQUIREMENTS = Path(__file__).resolve().parent.parent / "requirements.txt"
def _ensure_deps():
    try:
        import docx  # noqa: F401
    except ImportError:
        print("[chinese-contract-drafting] 检测到依赖缺失，正在自动安装...", file=sys.stderr)
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "-r", str(_REQUIREMENTS), "--quiet"]
        )
        print("[chinese-contract-drafting] ✅ 依赖安装完成", file=sys.stderr)

_ensure_deps()

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx_utils import (
    create_formatted_doc, set_page_layout, add_title, add_contract_info,
    add_body_text, add_heading_level1, add_heading_level2, add_heading_level3, add_clause,
    add_signature_block, save_as, append_clause, add_cover_page
)

doc = create_formatted_doc()
set_page_layout(doc)

add_cover_page(
    doc,
    title_cn='软件开发及技术服务合同',
    title_en='Software Development and Technical Service Contract',
    contract_no='xxx',
    party_a_label='甲  方（委托方）：',
    party_b_label='乙  方（开发方）：'
)

add_title(doc, '软件开发及技术服务合同')

add_body_text(doc, '甲方（委托方）：xxx')
add_body_text(doc, '乙方（开发方）：xxx')
add_body_text(doc, '合同编号：xxx')

add_body_text(doc, '鉴于甲方有软件开发及技术服务需求，乙方具备相应的技术能力和开发经验，双方根据《中华人民共和国民法典》及相关法律法规的规定，在平等自愿的基础上，经友好协商，就甲方委托乙方开发"旅游分销商智能体系统"并提供相关技术服务事宜达成如下协议：')

# 第一条 定义与解释
add_heading_level1(doc, '第一条  定义与解释')
add_body_text(doc, '1.1 "本项目"：指甲方委托乙方开发的"旅游分销商智能体系统"软件项目。')
add_body_text(doc, '1.2 "需求规格说明书"：指双方确认的软件功能需求文档，作为开发和验收的依据。')
add_body_text(doc, '1.3 "交付物"：指乙方按照本合同约定应向甲方交付的软件成果及相关文档。')
add_body_text(doc, '1.4 "验收"：指甲方按照本合同约定的标准和程序，对乙方交付的成果进行检验并确认合格的行为。')
add_body_text(doc, '1.5 "技术服务"：指乙方为本项目提供的部署、培训、运维支持等技术保障工作。')
add_body_text(doc, '1.6 "知识产权"：指与本项目相关的全部著作权、专利权、商标权、商业秘密及其他合法权益。')

# 第二条 项目内容与技术要求
add_heading_level1(doc, '第二条  项目内容与技术要求')
add_body_text(doc, '2.1 项目名称：旅游分销商智能体系统定制开发。')
add_body_text(doc, '2.2 项目内容：乙方按照甲方确认的需求规格说明书（附件一）进行软件开发，具体功能模块详见附件一《需求规格说明书》。')
add_body_text(doc, '2.3 技术要求：')
add_body_text(doc, '（1）系统架构：采用成熟的B/S架构，支持主流浏览器访问；')
add_body_text(doc, '（2）性能要求：支持不少于500并发用户，页面响应时间不超过3秒；')
add_body_text(doc, '（3）安全要求：支持数据加密存储、权限控制、操作日志审计；')
add_body_text(doc, '（4）接口规范：提供标准RESTful API接口，支持与第三方系统对接。')
add_body_text(doc, '2.4 需求变更：甲方如需变更需求，应以书面形式提出。乙方评估变更影响后，双方协商调整开发周期和费用。因需求变更导致的工期延误，乙方不承担违约责任。')

# 第三条 开发周期与交付
add_heading_level1(doc, '第三条  开发周期与交付')
add_body_text(doc, '3.1 开发周期：自本合同生效且乙方收到甲方首付款之日起xxx个工作日内完成开发并交付。')
add_body_text(doc, '3.2 交付内容：')
add_body_text(doc, '（1）旅游分销商智能体系统软件安装包；')
add_body_text(doc, '（2）系统部署文档；')
add_body_text(doc, '（3）用户操作手册；')
add_body_text(doc, '（4）API接口文档；')
add_body_text(doc, '（5）数据库设计文档。')
add_body_text(doc, '3.3 交付方式：乙方通过远程方式向甲方交付软件安装包及相关文档，并派技术人员到甲方现场进行部署安装。')
add_body_text(doc, '3.4 交付延迟：因甲方原因（包括但不限于需求确认延迟、提供资料不完整、环境准备不到位等）导致交付延迟的，交付期限相应顺延，乙方不承担违约责任。')

# 第四条 合同金额与支付
add_heading_level1(doc, '第四条  合同金额与支付')
add_body_text(doc, '4.1 合同总金额：人民币xxx元（大写：xxx元整），其中：')
add_body_text(doc, '（1）软件开发费：xxx元；')
add_body_text(doc, '（2）技术服务费：xxx元。')
add_body_text(doc, '4.2 付款方式：')
add_body_text(doc, '（1）合同签订后xxx个工作日内，甲方向乙方支付合同总金额的xxx%作为预付款；')
add_body_text(doc, '（2）开发完成并通过验收后xxx个工作日内，甲方向乙方支付合同总金额的xxx%作为尾款。')
add_body_text(doc, '4.3 乙方收款账户：xxx')
add_body_text(doc, '4.4 甲方逾期付款的，每逾期一日，应按逾期未付金额的万分之五向乙方支付违约金。逾期超过三十日的，乙方有权暂停开发工作或解除合同，已收取的款项不予退还。')

# 第五条 验收
add_heading_level1(doc, '第五条  验收')
add_body_text(doc, '5.1 验收标准：以附件一《需求规格说明书》约定的功能和技术指标为验收依据。')
add_body_text(doc, '5.2 验收期限：甲方应在乙方提交交付物后十（10）个工作日内组织验收。验收期限届满，甲方未提出书面异议的，视为验收合格。')
add_body_text(doc, '5.3 验收流程：')
add_body_text(doc, '（1）甲方在验收期限内完成验收，并出具书面验收意见；')
add_body_text(doc, '（2）验收不合格的，甲方应一次性提出全部修改意见；')
add_body_text(doc, '（3）乙方在收到修改意见后十五（15）个工作日内完成修改并重新提交验收；')
add_body_text(doc, '（4）经两次验收仍不合格的，双方协商解决。')
add_body_text(doc, '5.4 视为验收合格：甲方在验收期限内未组织验收，或已在实际业务中使用本系统的，均视为验收合格。')
add_body_text(doc, '5.5 超出需求范围的功能或修改，不纳入验收范围，双方另行协商费用。')

# 第六条 知识产权
add_heading_level1(doc, '第六条  知识产权')
add_body_text(doc, '6.1 本项目开发成果（包括但不限于软件程序、源代码、文档）的全部知识产权归乙方所有。')
add_body_text(doc, '6.2 乙方授予甲方对本项目的非独占、不可转让的使用许可，使用范围限于甲方自身业务经营。')
add_body_text(doc, '6.3 甲方不得对本项目成果进行反向工程、反编译、反汇编或试图获取源代码。')
add_body_text(doc, '6.4 甲方不得将本项目成果转售、转授权或提供给第三方使用。')
add_body_text(doc, '6.5 乙方在开发过程中使用的自有技术、工具、组件的知识产权仍归乙方所有。')
add_body_text(doc, '6.6 甲方提供的资料、数据的知识产权仍归甲方所有。')
add_body_text(doc, '6.7 乙方保证本项目成果不侵犯第三方知识产权。如因知识产权纠纷导致甲方无法正常使用，乙方应承担相应责任，但赔偿上限不超过本合同总金额。')

# 第七条 技术服务
add_heading_level1(doc, '第七条  技术服务')
add_heading_level2(doc, '7.1 系统部署')
add_body_text(doc, '（1）乙方负责将系统部署到甲方指定的服务器环境；')
add_body_text(doc, '（2）甲方应提供符合要求的服务器环境及必要的技术配合；')
add_body_text(doc, '（3）因甲方环境不符合要求导致的部署延迟，乙方不承担责任。')

add_heading_level2(doc, '7.2 用户培训')
add_body_text(doc, '（1）乙方为甲方提供系统使用培训，培训时长不少于xxx小时；')
add_body_text(doc, '（2）培训方式：现场培训或远程培训；')
add_body_text(doc, '（3）培训对象：甲方系统管理员及业务操作人员。')

add_heading_level2(doc, '7.3 运维支持')
add_body_text(doc, '（1）免费运维期：自验收合格之日起六（6）个月；')
add_body_text(doc, '（2）运维内容：系统故障排除、Bug修复、技术咨询；')
add_body_text(doc, '（3）响应时间：工作日xxx小时内响应，非工作日xxx小时内响应；')
add_body_text(doc, '（4）以下情形不属于免费运维范围：')
add_body_text(doc, '    a. 甲方操作不当或擅自修改系统导致的故障；')
add_body_text(doc, '    b. 甲方服务器环境变化导致的系统异常；')
add_body_text(doc, '    c. 甲方要求新增功能或修改现有功能；')
add_body_text(doc, '    d. 第三方软件或系统兼容性问题；')
add_body_text(doc, '    e. 不可抗力导致的系统损坏。')
add_body_text(doc, '（5）免费运维期满后，如甲方需要继续获得运维支持，双方另行签订运维服务合同。')

# 第八条 双方权利义务
add_heading_level1(doc, '第八条  双方权利义务')
add_heading_level2(doc, '8.1 甲方权利义务')
add_body_text(doc, '（1）按时支付合同款项；')
add_body_text(doc, '（2）及时确认需求、提供开发所需资料；')
add_body_text(doc, '（3）提供符合要求的服务器环境及网络条件；')
add_body_text(doc, '（4）指定专人配合乙方工作，及时响应乙方沟通需求；')
add_body_text(doc, '（5）按照约定使用软件，不得侵犯乙方知识产权。')

add_heading_level2(doc, '8.2 乙方权利义务')
add_body_text(doc, '（1）按照约定时间和质量完成开发工作；')
add_body_text(doc, '（2）提供符合约定的技术服务；')
add_body_text(doc, '（3）对甲方提供的资料负有保密义务；')
add_body_text(doc, '（4）有权要求甲方按时支付款项、提供必要配合；')
add_body_text(doc, '（5）有权在甲方违约时暂停工作或解除合同。')

# 第九条 保密条款
add_heading_level1(doc, '第九条  保密条款')
add_body_text(doc, '9.1 双方对在履行本合同过程中知悉的对方商业秘密和技术秘密负有保密义务。')
add_body_text(doc, '9.2 甲方特别承诺对以下信息严格保密：')
add_body_text(doc, '（1）乙方的技术方案、开发工具、组件库；')
add_body_text(doc, '（2）乙方的报价信息、商业模式。')
add_body_text(doc, '9.3 乙方特别承诺对以下信息严格保密：')
add_body_text(doc, '（1）甲方的业务数据、客户信息；')
add_body_text(doc, '（2）甲方的系统配置、网络架构。')
add_body_text(doc, '9.4 保密期限自本合同签订之日起至合同终止后三（3）年止。')
add_body_text(doc, '9.5 违反保密义务的一方应赔偿对方因此遭受的全部损失。')

# 第十条 违约责任
add_heading_level1(doc, '第十条  违约责任')
add_body_text(doc, '10.1 甲方违约责任：')
add_body_text(doc, '（1）甲方逾期付款的，适用第4.4条约定；')
add_body_text(doc, '（2）甲方违反知识产权条款（第六条）的，应立即停止违约行为，并向乙方支付合同总金额百分之五十（50%）的违约金；')
add_body_text(doc, '（3）甲方违反保密条款（第九条）的，适用第9.5条约定。')

add_body_text(doc, '10.2 乙方违约责任：')
add_body_text(doc, '（1）乙方逾期交付的，每逾期一日，应按合同总金额的万分之三向甲方支付违约金；')
add_body_text(doc, '（2）乙方累计违约金总额不超过合同总金额的百分之十（10%）；')
add_body_text(doc, '（3）因乙方原因导致合同无法继续履行的，甲方有权解除合同，乙方应退还甲方已付款项中对应未完成部分的费用。')

add_body_text(doc, '10.3 责任上限：在任何情况下，乙方对本合同项下承担的累计赔偿责任（包括违约金、损害赔偿等）不超过本合同总金额。')

# 第十一条 不可抗力
add_heading_level1(doc, '第十一条  不可抗力')
add_body_text(doc, '11.1 因不可抗力（包括但不限于自然灾害、战争、政府行为、法律法规变更、网络攻击、疫情等）导致合同无法履行的，受影响一方不承担违约责任。')
add_body_text(doc, '11.2 遭受不可抗力的一方应在事件发生后五（5）日内书面通知对方，并在十五（15）日内提供相关证明材料。')
add_body_text(doc, '11.3 不可抗力事件持续超过六十（60）日的，任何一方有权书面通知对方解除本合同，双方互不承担违约责任。')

# 第十二条 合同变更与解除
add_heading_level1(doc, '第十二条  合同变更与解除')
add_body_text(doc, '12.1 本合同的任何修改、补充均须经双方协商一致，并以书面形式签订补充协议。')
add_body_text(doc, '12.2 有下列情形之一的，乙方有权书面通知甲方解除合同：')
add_body_text(doc, '（1）甲方逾期付款超过三十日的；')
add_body_text(doc, '（2）甲方违反知识产权条款或保密条款的；')
add_body_text(doc, '（3）甲方长期不配合导致项目无法推进的。')
add_body_text(doc, '12.3 合同解除后，甲方应支付乙方已完成工作对应的费用，已收取的款项不予退还。')

# 第十三条 争议解决
add_heading_level1(doc, '第十三条  争议解决')
add_body_text(doc, '13.1 因本合同引起的或与本合同有关的任何争议，双方应首先通过友好协商解决。')
add_body_text(doc, '13.2 协商不成的，任何一方有权向乙方所在地有管辖权的人民法院提起诉讼。')

# 第十四条 其他约定
add_heading_level1(doc, '第十四条  其他约定')
add_body_text(doc, '14.1 本合同自双方法定代表人或授权代表签字并加盖公章之日起生效。')
add_body_text(doc, '14.2 本合同一式肆份，甲乙双方各执贰份，具有同等法律效力。')
add_body_text(doc, '14.3 本合同附件为本合同不可分割的组成部分，与本合同正文具有同等法律效力。')
add_body_text(doc, '14.4 本合同未尽事宜，双方可另行签订补充协议。补充协议与本合同不一致的，以补充协议为准。')
add_body_text(doc, '14.5 本合同任何条款被认定为无效或不可执行的，不影响其他条款的效力。')

# 附件清单
add_heading_level1(doc, '附件清单')
add_body_text(doc, '附件一：《需求规格说明书》')
add_body_text(doc, '附件二：《验收标准》')
add_body_text(doc, '附件三：《技术服务内容清单》')

# 签署区
add_body_text(doc, '（以下为签署页，无正文）')
add_signature_block(doc, party_a_name='甲方（签章）：', party_b_name='乙方（签章）：')

# 法律依据
add_heading_level1(doc, '法律依据')
add_body_text(doc, '1.《中华人民共和国民法典》第三编 合同，第八百五十一条至第八百五十五条（委托开发合同）')
add_body_text(doc, '2.《中华人民共和国民法典》第五百零九条至第五百五十八条（合同的履行）')
add_body_text(doc, '3.《中华人民共和国民法典》第五百七十七条至第五百九十四条（违约责任）')
add_body_text(doc, '4.《中华人民共和国著作权法》')
add_body_text(doc, '5.《计算机软件保护条例》')

# ========== 附件一：需求规格说明书 ==========
doc.add_page_break()
add_heading_level1(doc, '附件一：需求规格说明书')

add_body_text(doc, '项目名称：旅游分销商智能体系统')
add_body_text(doc, '版本号：V1.0')

add_heading_level2(doc, '一、项目概述')
add_body_text(doc, '本系统为旅游行业分销商提供智能化的业务管理平台，实现分销商管理、产品管理、订单管理、智能推荐、财务管理、数据分析等核心功能。')

add_heading_level2(doc, '二、功能模块')
add_heading_level3(doc, '2.1 分销商管理模块')
add_body_text(doc, '（1）分销商注册与审核：支持在线注册、资质审核、等级评定；')
add_body_text(doc, '（2）分销商信息管理：基本信息维护、联系人管理、账户状态管理；')
add_body_text(doc, '（3）分销商业绩统计：实时统计销售额、订单量、佣金收益。')

add_heading_level3(doc, '2.2 产品管理模块')
add_body_text(doc, '（1）旅游产品上架：支持线路、酒店、门票、机票等多类型产品；')
add_body_text(doc, '（2）产品价格管理：多价格体系、动态定价、促销活动配置；')
add_body_text(doc, '（3）库存管理：实时库存同步、库存预警、超售防护。')

add_heading_level3(doc, '2.3 订单管理模块')
add_body_text(doc, '（1）订单全流程管理：下单、支付、确认、出行、结算全链路跟踪；')
add_body_text(doc, '（2）退改签处理：在线退改签申请、审核、退款；')
add_body_text(doc, '（3）订单异常处理：自动识别异常订单、人工干预处理。')

add_heading_level3(doc, '2.4 智能推荐模块')
add_body_text(doc, '（1）客户画像分析：基于历史数据的客户偏好分析；')
add_body_text(doc, '（2）产品智能推荐：根据客户特征自动推荐匹配产品；')
add_body_text(doc, '（3）营销策略建议：基于数据分析的营销方案生成。')

add_heading_level3(doc, '2.5 财务管理模块')
add_body_text(doc, '（1）佣金结算：自动计算分销商佣金、支持多结算周期；')
add_body_text(doc, '（2）对账管理：与供应商、分销商自动对账；')
add_body_text(doc, '（3）财务报表：收入、支出、利润多维度报表。')

add_heading_level3(doc, '2.6 数据分析模块')
add_body_text(doc, '（1）经营数据看板：核心指标实时展示；')
add_body_text(doc, '（2）销售趋势分析：多维度销售数据分析；')
add_body_text(doc, '（3）用户行为分析：访问、转化、留存数据分析。')

add_heading_level2(doc, '三、非功能性需求')
add_body_text(doc, '（1）性能要求：支持500并发用户，页面响应≤3秒；')
add_body_text(doc, '（2）安全要求：数据加密存储、权限控制、操作日志；')
add_body_text(doc, '（3）可用性要求：系统可用性≥99.5%；')
add_body_text(doc, '（4）兼容性要求：支持Chrome、Firefox、Edge等主流浏览器。')

# ========== 附件二：验收标准 ==========
doc.add_page_break()
add_heading_level1(doc, '附件二：验收标准')

add_heading_level2(doc, '一、验收依据')
add_body_text(doc, '（1）本合同正文；')
add_body_text(doc, '（2）附件一《需求规格说明书》；')
add_body_text(doc, '（3）国家及行业相关标准。')

add_heading_level2(doc, '二、功能验收标准')
add_body_text(doc, '（1）附件一所列全部功能均已实现并可正常使用；')
add_body_text(doc, '（2）各功能模块运行稳定，无严重缺陷；')
add_body_text(doc, '（3）一般缺陷数量不超过五个，且有明确修复计划。')

add_heading_level2(doc, '三、性能验收标准')
add_body_text(doc, '（1）页面加载时间不超过3秒；')
add_body_text(doc, '（2）支持不少于500用户同时在线；')
add_body_text(doc, '（3）单次查询响应时间不超过5秒。')

add_heading_level2(doc, '四、安全验收标准')
add_body_text(doc, '（1）用户认证：支持用户名密码登录、短信验证码登录；')
add_body_text(doc, '（2）权限控制：支持角色权限管理；')
add_body_text(doc, '（3）数据安全：密码加密存储，敏感信息脱敏展示；')
add_body_text(doc, '（4）日志审计：关键操作有日志记录。')

add_heading_level2(doc, '五、文档验收标准')
add_body_text(doc, '（1）交付物清单所列文档齐全；')
add_body_text(doc, '（2）文档内容完整、表述清晰；')
add_body_text(doc, '（3）文档版本与软件版本一致。')

add_heading_level2(doc, '六、验收通过条件')
add_body_text(doc, '（1）功能、性能、安全、文档验收均符合标准；')
add_body_text(doc, '（2）无严重缺陷，一般缺陷在允许范围内；')
add_body_text(doc, '（3）甲方授权代表在验收报告上签字确认。')

# ========== 附件三：技术服务内容清单 ==========
doc.add_page_break()
add_heading_level1(doc, '附件三：技术服务内容清单')

add_heading_level2(doc, '一、系统部署服务')
add_body_text(doc, '（1）部署环境检查与确认；')
add_body_text(doc, '（2）系统安装与配置；')
add_body_text(doc, '（3）系统联调与测试；')
add_body_text(doc, '（4）上线支持与保障。')

add_heading_level2(doc, '二、用户培训服务')
add_body_text(doc, '（1）培训对象：系统管理员、业务操作人员；')
add_body_text(doc, '（2）培训内容：系统功能操作、常见问题处理；')
add_body_text(doc, '（3）培训方式：现场培训或远程培训；')
add_body_text(doc, '（4）培训时长：不少于xxx小时；')
add_body_text(doc, '（5）培训资料：提供培训手册（电子版）。')

add_heading_level2(doc, '三、运维支持服务')
add_body_text(doc, '（1）服务期限：验收合格后六（6）个月；')
add_body_text(doc, '（2）服务内容：')
add_body_text(doc, '    a. 系统故障排除；')
add_body_text(doc, '    b. Bug修复；')
add_body_text(doc, '    c. 技术咨询；')
add_body_text(doc, '    d. 系统巡检（每月一次）。')
add_body_text(doc, '（3）响应时间：')
add_body_text(doc, '    a. 严重故障：xxx小时内响应，xxx小时内解决；')
add_body_text(doc, '    b. 一般故障：xxx小时内响应，xxx小时内解决；')
add_body_text(doc, '    c. 技术咨询：xxx小时内响应。')
add_body_text(doc, '（4）服务方式：电话、邮件、远程协助。')

add_heading_level2(doc, '四、增值服务（可选）')
add_body_text(doc, '（1）系统升级服务：按次收费，费用另议；')
add_body_text(doc, '（2）功能扩展服务：按工作量评估，费用另议；')
add_body_text(doc, '（3）数据迁移服务：按数据量评估，费用另议。')

output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '输出', '软件开发及技术服务合同.docx')
save_as(doc, output_path)
print('done: ' + output_path)
