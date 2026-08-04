#!/usr/bin/env python3
"""
markdown_to_docx.py — 将技能产出的 Markdown 文件批量转换为 DOCX 并打包 ZIP。

用法：
  python3 scripts/markdown_to_docx.py <md_file_or_dir> [output_path]

参数：
  md_file_or_dir    单个 .md 文件，或包含多个 .md 文件的目录
  output_path       可选。输出路径，默认为用户主目录下的 legal-skills-export/（跨平台，自动适配 Windows/macOS/Linux）

输出：
  - 单个文件：生成同名的 .docx 文件
  - 目录：生成所有 .docx 文件并打包为 ZIP

示例：
  python3 scripts/markdown_to_docx.py /path/to/output.md
  python3 scripts/markdown_to_docx.py /path/to/analysis_dir/
"""

import os
import sys
import re
import zipfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import sys
try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
except (AttributeError, ValueError):
    pass


# ── 排版工具函数 ──

def set_cjk_font(run_or_style, cn_font='宋体', en_font='Times New Roman'):
    rPr = run_or_style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def clean_heading_style(style):
    style.font.color.rgb = RGBColor(0, 0, 0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
            if rFonts.get(qn(attr)):
                del rFonts.attrib[qn(attr)]
    pPr = style.element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)


def setup_document(doc, title_text=None):
    """政府标准行文格式（GB/T 9704-2012 党政机关公文格式）"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)      # 上3.7cm
    section.bottom_margin = Cm(3.5)   # 下3.5cm
    section.left_margin = Cm(2.8)     # 左2.8cm
    section.right_margin = Cm(2.6)    # 右2.6cm

    # 正文：仿宋_GB2312 三号(16pt)，1.5倍行距，首行缩进2字符
    normal = doc.styles['Normal']
    normal.font.size = Pt(16)         # 三号=16pt
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(32)  # 2字符缩进
    set_cjk_font(normal, '仿宋_GB2312', 'Times New Roman')

    # 标题样式
    for i in range(0, 4):
        h_style_name = 'Title' if i == 0 else f'Heading {i}'
        h_style = doc.styles[h_style_name]
        clean_heading_style(h_style)
        # 一级标题(##)：黑体 二号(22pt) 居中
        if i == 1:
            h_style.font.size = Pt(22)
            h_style.font.bold = True
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cjk_font(h_style, '黑体', 'Arial')
        # 二级标题(###)：黑体 三号(16pt) 加粗
        elif i == 2:
            h_style.font.size = Pt(16)
            h_style.font.bold = True
            set_cjk_font(h_style, '黑体', 'Arial')
        # 三级标题(####)：楷体 三号(16pt)
        elif i == 3:
            h_style.font.size = Pt(16)
            h_style.font.bold = False
            set_cjk_font(h_style, '楷体_GB2312', 'Arial')
        # 文档主标题（#）：黑体 小标宋 22pt 居中
        else:
            h_style.font.size = Pt(22)
            h_style.font.bold = True
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cjk_font(h_style, '黑体', 'Arial')
        h_style.paragraph_format.first_line_indent = Pt(0)

    # 页脚页码
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    run.font.size = Pt(14)  # 表格字号四号=14pt
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run = footer_para.add_run()
    run.font.size = Pt(14)  # 表格字号四号=14pt
    run._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    run = footer_para.add_run()
    run.font.size = Pt(14)  # 表格字号四号=14pt
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    if title_text:
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def _add_rich_text(paragraph, text):
    """添加富文本（处理 **加粗** 标记）"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_table_from_lines(doc, lines):
    rows_data = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|---') or line.startswith('|--'):
            continue
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows_data.append(cells)
    if not rows_data:
        return
    max_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=max_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_cells in enumerate(rows_data):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]
            cell.text = row_cells[j] if j < len(row_cells) else ''
            for para in cell.paragraphs:
                para.style = doc.styles['Normal']
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(14)  # 表格字号四号=14pt
                if i == 0:
                    for run in para.runs:
                        run.bold = True
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def markdown_to_docx(md_path, docx_path, doc_title=None):
    """将单个 .md 文件转换为 .docx"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    # 去掉 YAML frontmatter
    content = re.sub(r'^---\n.*?\n---\n', '', content, flags=re.DOTALL)

    doc = Document()
    setup_document(doc, doc_title)

    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    in_code = False
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue

        s = line.strip()
        # 表格
        if s.startswith('|') and s.endswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                add_table_from_lines(doc, table_lines)
                in_table = False
                table_lines = []

        # 标题
        if s.startswith('# ') and not s.startswith('## '):
            doc.add_heading(s[2:], level=1)
        elif s.startswith('## ') and not s.startswith('### '):
            doc.add_heading(s[3:], level=2)
        elif s.startswith('### '):
            doc.add_heading(s[4:], level=3)
        elif s.startswith('- [') or s.startswith('* ['):
            prefix = '☑ ' if '[x]' in s.lower() else '☐ '
            text = re.sub(r'^[-*]\s*\[\s*[xX]?\s*\]\s*', '', s)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(prefix + text)
        elif s.startswith('- ') or s.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, s[2:])
        elif re.match(r'^\d+[\.\)]\s', s):
            p = doc.add_paragraph(style='List Number')
            _add_rich_text(p, re.sub(r'^\d+[\.\)]\s*', '', s))
        elif s.startswith('> '):
            text = s[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)  # 2字符缩进
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Pt(0)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '6')
            left.set(qn('w:color'), '4472C4')
            pBdr.append(left)
            pPr.append(pBdr)
            _add_rich_text(p, text)
        elif s.startswith('---') and len(s) >= 3:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif s:
            p = doc.add_paragraph()
            _add_rich_text(p, s)
        else:
            doc.add_paragraph()
        i += 1

    if in_table and table_lines:
        add_table_from_lines(doc, table_lines)

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    return docx_path


def convert_and_zip(file_or_dir, output_base=None):
    """
    主入口：将文件或目录中的 .md 转为 .docx 并可选打包 ZIP。

    参数:
        file_or_dir: 单个 .md 文件路径，或包含 .md 的目录路径
        output_base: 输出基础路径。为 None 时使用跨平台默认路径（用户主目录下的 legal-skills-export/）
    返回:
        (zip_path, file_count) 或 (docx_path, 1)
    """
    if output_base is None:
        output_base = os.path.join(os.path.expanduser("~"), "legal-skills-export")

    if os.path.isfile(file_or_dir) and file_or_dir.endswith('.md'):
        # 单个文件
        base = os.path.splitext(os.path.basename(file_or_dir))[0]
        docx_path = os.path.join(output_base, f'{base}.docx')
        markdown_to_docx(file_or_dir, docx_path, base)
        return docx_path, 1

    elif os.path.isdir(file_or_dir):
        # 目录 → 批量转换并打包 ZIP
        md_files = sorted([
            f for f in os.listdir(file_or_dir)
            if f.endswith('.md')
        ])
        if not md_files:
            print("未找到 .md 文件")
            return None, 0

        for fname in md_files:
            md_path = os.path.join(file_or_dir, fname)
            base = os.path.splitext(fname)[0]
            docx_path = os.path.join(output_base, 'docx', f'{base}.docx')
            title = base.replace('-', ' ').replace('_', ' ')
            markdown_to_docx(md_path, docx_path, title)
            print(f"  ✓ {fname} → {base}.docx")

        # 打包 ZIP
        zip_name = os.path.basename(os.path.normpath(file_or_dir))
        zip_path = os.path.join(output_base, f'{zip_name}.zip')
        os.makedirs(os.path.dirname(zip_path), exist_ok=True)
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            docx_dir = os.path.join(output_base, 'docx')
            for fname in os.listdir(docx_dir):
                if fname.endswith('.docx'):
                    fpath = os.path.join(docx_dir, fname)
                    zf.write(fpath, fname)

        print(f"\n已完成！共 {len(md_files)} 个文件")
        print(f"ZIP 包: {zip_path}")
        return zip_path, len(md_files)

    else:
        raise ValueError(f"路径无效或不是 .md 文件: {file_or_dir}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    source = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else None
    result, count = convert_and_zip(source, output)
    if result:
        print(f"输出: {result}")
    else:
        sys.exit(1)
