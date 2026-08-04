#!/usr/bin/env python3
"""执行案件助手：docx 文书生成脚本。

用法：
    python3 gen_docx.py --template <模板.docx> --map <填空.json> --out <输出.docx>

填空.json 格式：{"占位符原文": "替换值", ...}
原理：复制模板文件，在段落级对占位符做文本替换，保留原排版样式。
对 .doc 旧格式模板，先用 textutil（macOS）或 soffice 转成 .docx 再处理。
"""
import argparse
import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


def convert_doc_to_docx(doc_path: str) -> str:
    """把 .doc 转成 .docx（优先 soffice，退化 textutil）。"""
    tmp = Path(tempfile.mkdtemp())
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(tmp), doc_path],
            check=True, capture_output=True, timeout=120,
        )
        out = tmp / (Path(doc_path).stem + ".docx")
        if out.exists():
            return str(out)
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        pass
    # textutil 兜底（macOS 自带，但可能丢部分格式）
    try:
        subprocess.run(["textutil", "-convert", "docx", "-output", str(tmp / (Path(doc_path).stem + ".docx")), doc_path],
                       check=True, capture_output=True, timeout=60)
        out = tmp / (Path(doc_path).stem + ".docx")
        if out.exists():
            return str(out)
    except Exception:
        pass
    sys.exit(f"错误：无法将 {doc_path} 转换为 docx，请手动用 WPS/Word 另存为 .docx 后重试。")


def replace_in_paragraph(paragraph, mapping: dict) -> int:
    """段落级占位符替换。占位符可能跨 run，先做整段文本替换再重写第一个 run。"""
    hits = 0
    full = paragraph.text
    new = full
    for k, v in mapping.items():
        if k in new:
            new = new.replace(k, v)
            hits += 1
    if new != full:
        # 保留第一个 run 的样式，清空其余 run
        runs = paragraph.runs
        if runs:
            runs[0].text = new
            for r in runs[1:]:
                r.text = ""
        else:
            paragraph.text = new
    return hits


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True, help="模板文件（.docx 或 .doc）")
    ap.add_argument("--map", required=True, help="填空 JSON 文件")
    ap.add_argument("--out", required=True, help="输出 docx 路径")
    args = ap.parse_args()

    tpl = args.template
    if tpl.lower().endswith(".doc") and not tpl.lower().endswith(".docx"):
        tpl = convert_doc_to_docx(tpl)

    with open(args.map, encoding="utf-8") as f:
        mapping = json.load(f)
    if not isinstance(mapping, dict) or not mapping:
        sys.exit("错误：填空 JSON 必须是非空对象。")

    from docx import Document
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(tpl, out_path)

    doc = Document(out_path)
    total = 0
    for p in doc.paragraphs:
        total += replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += replace_in_paragraph(p, mapping)
    # 页眉页脚
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                total += replace_in_paragraph(p, mapping)

    doc.save(out_path)

    # 残留占位符检查
    doc2 = Document(out_path)
    leftovers = []
    for p in doc2.paragraphs:
        if "【" in p.text:
            leftovers.append(p.text.strip()[:60])
    print(f"[ok] 已生成: {out_path}")
    print(f"[ok] 替换命中: {total} 处")
    if leftovers:
        print(f"[warn] 仍有 {len(leftovers)} 处【】占位符未替换，请检查：")
        for t in leftovers[:10]:
            print(f"  - {t}")


if __name__ == "__main__":
    main()
