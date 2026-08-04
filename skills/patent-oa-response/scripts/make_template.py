# -*- coding: utf-8 -*-
"""生成 assets/意见陈述书_模板.docx：标准章节骨架 + 对比表占位。"""
import os
from docx import Document
from docx.shared import Pt

OUT = os.path.join(os.path.dirname(__file__), '..', 'assets', '意见陈述书_模板.docx')

SECTIONS = [
    ('一、', '收到审查意见通知书的基本情况'),
    ('二、', '对比文件1（最接近现有技术）不适合作为评述本申请创造性的技术原点'),
    ('    2.1', '本申请与对比文件1的技术特征对比表'),
    ('    2.2', '对比文件1不适合作为最接近现有技术的理由'),
    ('三、', '关于审查员归纳的区别技术特征是否准确的核查（协同性回击）'),
    ('四、', '关于其余对比文件技术特征的实质与作用的核查（作用不同则无启示）'),
    ('    4.1', '对比文件2的核查'),
    ('    4.2', '对比文件3的核查'),
    ('五、', '关于"常规技术手段"认定的驳斥'),
    ('六、', '本申请技术方案的整体性与"事后诸葛亮"分析'),
    ('七、', '本申请技术的产业价值（辅助论点）'),
    ('八、', '创造性审查原则的重申与结论'),
]


def main():
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc = Document()
    # 标题
    title = doc.add_heading('意见陈述书', level=0)
    title.alignment = 1  # center
    sub = doc.add_paragraph('（答复第____次审查意见通知书）')
    sub.alignment = 1

    for pre, name in SECTIONS:
        if pre.endswith('、'):
            h = doc.add_heading(f'{pre}{name}', level=1)
        else:
            h = doc.add_heading(f'{pre} {name}', level=2)

    # 在 2.1 后插入对比表占位
    for p in doc.paragraphs:
        if '2.1' in p.text and '技术特征对比表' in p.text:
            tbl = doc.add_table(rows=2, cols=4)
            tbl.style = 'Table Grid'
            hdr = tbl.rows[0].cells
            for i, t in enumerate(['技术特征', '本申请（权利要求/说明书）', '对比文件1（标号/段落）', '差异与作用']):
                hdr[i].paragraphs[0].add_run(t).bold = True
            # 示例行
            row = tbl.rows[1].cells
            for i, t in enumerate(['（示例特征）', '（填入本申请对应内容）', '（填入对比文件内容）', '（差异说明）']):
                row[i].paragraphs[0].add_run(t)
            # 提示
            note = doc.add_paragraph('〔提示：本表逐特征比对；嵌入附图时用 extract_figs.py 生成左右并列对比图，置于本表或对应论证之后。〕')
            note.runs[0].font.size = Pt(9)
            note.runs[0].font.italic = True
            break

    # 落款
    doc.add_paragraph()
    doc.add_paragraph('申请人：________________________')
    doc.add_paragraph('代理机构：______________________')
    doc.add_paragraph('日期：______年______月______日')

    doc.save(OUT)
    print('Template saved:', os.path.abspath(OUT))


if __name__ == '__main__':
    main()
