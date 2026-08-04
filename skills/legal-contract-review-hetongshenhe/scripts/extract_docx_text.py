#!/usr/bin/env python3
"""
Extract text content from a .docx file.

Extracts text from:
- Body paragraphs
- Tables (including nested tables)
- Headers and footers

Replaces Dify's document-extractor node.

Usage:
    python extract_docx_text.py <input.docx> [--output <output.txt>]

If --output is not specified, text is printed to stdout.
"""

import sys
import os
import argparse

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def extract_table_text(table):
    """Extract text from a table, including nested tables."""
    lines = []
    for row in table.rows:
        row_texts = []
        for cell in row.cells:
            cell_text = []
            for para in cell.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    cell_text.append(para_text)
            # Check for nested tables
            for nested_table in cell.tables:
                nested_text = extract_table_text(nested_table)
                if nested_text:
                    cell_text.append(nested_text)
            row_texts.append(" ".join(cell_text))
        lines.append("\t".join(row_texts))
    return "\n".join(lines)


def extract_docx_text(docx_path):
    """Extract all text content from a .docx file."""
    doc = Document(docx_path)
    parts = []

    # Extract body paragraphs and tables in document order
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Find the corresponding paragraph object
            for para in doc.paragraphs:
                if para._element is child:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
                    break
        elif tag == 'tbl':
            # Find the corresponding table object
            for table in doc.tables:
                if table._element is child:
                    table_text = extract_table_text(table)
                    if table_text:
                        parts.append(table_text)
                    break

    # Extract header text
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            try:
                if header and not header.is_linked_to_previous:
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
            except Exception:
                pass

    # Extract footer text
    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            try:
                if footer and not footer.is_linked_to_previous:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
            except Exception:
                pass

    return "\n".join(parts)


def main():
    parser = argparse.ArgumentParser(description="Extract text from a .docx file")
    parser.add_argument("input", help="Path to the input .docx file")
    parser.add_argument("--output", "-o", help="Path to the output text file (default: stdout)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"ERROR: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    if not args.input.lower().endswith(".docx"):
        print(f"WARNING: Input file does not have .docx extension: {args.input}", file=sys.stderr)

    text = extract_docx_text(args.input)

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"Extracted text written to: {args.output}", file=sys.stderr)
    else:
        print(text)


if __name__ == "__main__":
    main()
