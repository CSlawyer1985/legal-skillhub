#!/usr/bin/env python3
"""
律所委托材料 .docx 占位符填充脚本

功能：
1. fill: 读取模板，将 {{FIELD_NAME}} 替换为实际值，输出新文件
2. list: 列出模板中所有占位符

用法：
    # 列出模板中所有占位符
    python3 fill_docx.py list --template 授权委托书.docx

    # 单字段替换
    python3 fill_docx.py fill --template 授权委托书.docx --output 输出.docx \
        --fields '{"委托人姓名":"张三","案由":"民间借贷纠纷"}'

    # 从 JSON 文件读取字段
    python3 fill_docx.py fill --template 授权委托书.docx --output 输出.docx \
        --fields-file fields.json
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

# 匹配 {{中文或英文描述}} 占位符
PLACEHOLDER_PATTERN = re.compile(r'\{\{(.+?)\}\}')


def extract_text_from_cell(cell):
    """提取表格单元格中的完整文本"""
    paragraphs_text = []
    for paragraph in cell.paragraphs:
        paragraphs_text.append(paragraph.text)
    return '\n'.join(paragraphs_text)


def find_placeholders_in_doc(doc):
    """扫描文档，返回所有占位符及其位置信息

    返回结构：
    {
        "paragraphs": [  # 正文段落
            {
                "paragraph_index": 0,
                "placeholders": ["委托人姓名", "案由"],
                "full_text": "..."
            }
        ],
        "tables": [  # 表格
            {
                "table_index": 0,
                "cells": [
                    {
                        "row": 0, "col": 0,
                        "placeholders": ["委托人姓名"],
                        "full_text": "..."
                    }
                ]
            }
        ]
    }
    """
    result = {"paragraphs": [], "tables": []}

    # 扫描正文段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        matches = PLACEHOLDER_PATTERN.findall(text)
        if matches:
            result["paragraphs"].append({
                "paragraph_index": i,
                "placeholders": matches,
                "full_text": text
            })

    # 扫描表格
    for ti, table in enumerate(doc.tables):
        table_info = {"table_index": ti, "cells": []}
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = extract_text_from_cell(cell)
                matches = PLACEHOLDER_PATTERN.findall(text)
                if matches:
                    table_info["cells"].append({
                        "row": ri, "col": ci,
                        "placeholders": matches,
                        "full_text": text
                    })
        if table_info["cells"]:
            result["tables"].append(table_info)

    return result


def get_all_placeholder_names(doc):
    """获取文档中所有唯一的占位符名称"""
    info = find_placeholders_in_doc(doc)
    names = set()

    for p in info["paragraphs"]:
        names.update(p["placeholders"])

    for t in info["tables"]:
        for c in t["cells"]:
            names.update(c["placeholders"])

    return sorted(names)


def split_lawyer_names(value):
    """将 '朱彬、陈志宇' 拆成 ['朱彬', '陈志宇']

    支持中英文逗号分隔。
    """
    if not value:
        return []
    # 用中英文逗号、顿号、斜杠分隔
    parts = re.split(r'[,，、/；;]', str(value))
    return [p.strip() for p in parts if p.strip()]


def expand_block_placeholders(field_values):
    """展开块占位符（如 {{代理律师_受托人声明}}）

    这些占位符会根据 代理律师 字段的多个姓名展开为多行内容（使用 \n 分隔）。
    """
    if '代理律师' not in field_values:
        return field_values

    lawyer_value = str(field_values['代理律师'])
    lawyer_names = split_lawyer_names(lawyer_value)

    if not lawyer_names:
        return field_values

    expanded = dict(field_values)

    # 受托人声明块：每个律师完整的受托人声明行
    if '代理律师_受托人声明' not in expanded:
        merchant_text = "\n".join(
            f"受托人：{name}，北京天驰君泰（扬州）律师事务所律师"
            for name in lawyer_names
        )
        expanded['代理律师_受托人声明'] = merchant_text

    # 代理权限声明块：每个律师完整的代理权限行
    if '代理律师_代理权限声明' not in expanded:
        auth_text = "\n".join(
            f"受托人{name}律师，代理权限为：特别授权"
            f"（有权代为承认、变更、放弃诉讼请求，有权进行和解、调解，有权代收法律文书等）。"
            for name in lawyer_names
        )
        expanded['代理律师_代理权限声明'] = auth_text

    # 联系方式声明块：每个律师完整的联系方式行
    if '代理律师_联系方式声明' not in expanded:
        phone = str(field_values.get('律师电话', '________________'))
        contact_text = "\n".join(
            f"受托人：{name}，联系方式：{phone}"
            for name in lawyer_names
        )
        expanded['代理律师_联系方式声明'] = contact_text

    return expanded


def expand_letter_case_description(field_values):
    """构建所函案件描述占位符

    {{所函案件描述}} 根据委托人地位自动拼接：
    - 原告：{{委托人姓名}}诉{{案件对方及案由}}
    - 被告：{{案件对方及案由}}诉{{委托人姓名}}
    """
    if '所函案件描述' not in field_values:
        client = field_values.get('委托人姓名', '')
        other = field_values.get('案件对方及案由', '')
        role = field_values.get('当事人地位', '原告')
        if role == '被告':
            return f"{other}诉{client}"
        else:
            return f"{client}诉{other}"
    return field_values['所函案件描述']


def replace_in_paragraph(paragraph, field_values):
    """替换段落中的占位符，保留格式

    处理跨 run 的占位符：先将所有 run 文本拼接，找到占位符位置，
    然后分解到各 run 中进行替换。
    """
    # 获取所有 run
    runs = paragraph.runs
    if not runs:
        return

    # 构建 run 文本的字符级映射
    # run_map[i] = (run_index, char_position_in_that_run)
    run_map = []
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            run_map.append((ri, ci))

    # 拼接完整文本
    full_text = paragraph.text

    # 查找占位符
    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return

    # 反向处理（从后往前），避免位置偏移
    for match in reversed(matches):
        key = match.group(1)
        if key not in field_values:
            continue

        replacement = str(field_values[key])
        start, end = match.start(), match.end()

        # 找到起止位置对应的 run 和字符位置
        start_run_idx, start_char_pos = run_map[start]
        end_run_idx, end_char_pos = run_map[end - 1]

        if start_run_idx == end_run_idx:
            # 占位符在同一个 run 内，简单替换
            run = runs[start_run_idx]
            if '\n' in replacement:
                _set_text_with_line_breaks(paragraph, start_run_idx, start_char_pos, end_char_pos, replacement)
            else:
                run.text = run.text[:start_char_pos] + replacement + run.text[end_char_pos + 1:]
        else:
            # 占位符跨多个 run
            # 第一步：处理中间 run（全部清空其文本中属于占位符的部分）
            for ri in range(end_run_idx, start_run_idx, -1):
                run = runs[ri]
                r_start = 0
                r_end = len(run.text) - 1
                if ri == end_run_idx:
                    r_end = end_char_pos
                if ri == start_run_idx:
                    r_start = start_char_pos

                # 删除占位符对应部分的字符
                before = run.text[:r_start]
                after = run.text[r_end + 1:]
                run.text = before + after

            # 第二步：在起始 run 中插入替换文本
            run = runs[start_run_idx]
            if '\n' in replacement:
                _set_text_with_line_breaks(paragraph, start_run_idx, start_char_pos, len(run.text), replacement)
            else:
                run.text = run.text[:start_char_pos] + replacement + run.text[start_char_pos:]

            # 第三步：删除中间被清空的 run（可选，保持文档整洁）
            # 这里保留空 run 以避免破坏 XML 结构


def _set_text_with_line_breaks(paragraph, run_idx, char_start, char_end, text_with_breaks):
    """在指定 run 中插入包含换行符的文本

    将 text_with_breaks 中的 \\n 转换为 <w:br/> 元素（行内换行）。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    run = paragraph.runs[run_idx]
    # run 原本的完整文本
    original = run.text
    prefix = original[:char_start]  # 占位符前的部分
    suffix = original[char_end + 1:]  # 占位符后的部分

    # 清空 run 文本
    run.text = ""

    # 插入前缀
    if prefix:
        t = etree.SubElement(run._element, qn('w:t'))
        t.text = prefix
        t.set(qn('xml:space'), 'preserve')

    # 分割多行内容，用 <w:br/> 分隔
    lines = text_with_breaks.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            # 插入换行符
            etree.SubElement(run._element, qn('w:br'))
        if line:
            t = etree.SubElement(run._element, qn('w:t'))
            t.text = line
            t.set(qn('xml:space'), 'preserve')

    # 插入后缀
    if suffix:
        t = etree.SubElement(run._element, qn('w:t'))
        t.text = suffix
        t.set(qn('xml:space'), 'preserve')


def replace_in_cell(cell, field_values):
    """替换表格单元格中的占位符"""
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, field_values)


def fill_document(template_path, output_path, field_values):
    """填充文档模板"""
    # 展开块占位符
    field_values = expand_block_placeholders(field_values)

    # 自动填入默认值
    from datetime import datetime, date
    if '年份' not in field_values:
        field_values['年份'] = str(datetime.now().year)
    if '签署日期' not in field_values:
        today = date.today()
        field_values['签署日期'] = f"{today.year}年{today.month}月{today.day}日"

    doc = Document(template_path)

    # 替换正文段落
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, field_values)

    # 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, field_values)

    # 同时检查页眉页脚
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, field_values)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, field_values)

    doc.save(output_path)


def cmd_list(args):
    """列出模板中的所有占位符"""
    template = Path(args.template)
    if not template.exists():
        print(f"错误: 模板文件不存在: {template}")
        sys.exit(1)

    doc = Document(str(template))
    names = get_all_placeholder_names(doc)

    if not names:
        print("未发现任何占位符（{{...}}）。")
    else:
        print(f"发现 {len(names)} 个占位符：")
        print("-" * 40)
        for name in names:
            print(f"  {{{{ {name} }}}}")

    # 详细位置信息
    if args.verbose:
        info = find_placeholders_in_doc(doc)
        print("\n详细位置:")
        if info["paragraphs"]:
            print(f"\n正文段落 ({len(info['paragraphs'])} 处):")
            for p in info["paragraphs"]:
                print(f"  段落 {p['paragraph_index']}: {p['placeholders']}")

        if info["tables"]:
            print(f"\n表格 ({len(info['tables'])} 个表):")
            for t in info["tables"]:
                print(f"  表 {t['table_index']}:")
                for c in t["cells"]:
                    print(f"    单元格({c['row']},{c['col']}): {c['placeholders']}")


def cmd_fill(args):
    """填充模板"""
    template = Path(args.template)
    if not template.exists():
        print(f"错误: 模板文件不存在: {template}")
        sys.exit(1)

    # 读取字段值
    field_values = {}

    if args.fields_file:
        with open(args.fields_file, 'r', encoding='utf-8') as f:
            field_values = json.load(f)
    elif args.fields:
        field_values = json.loads(args.fields)

    if not field_values:
        print("错误: 未提供任何字段值")
        print("使用 --fields 或 --fields-file 指定")
        sys.exit(1)

    # 自动填入默认值
    from datetime import datetime, date
    if '年份' not in field_values:
        field_values['年份'] = str(datetime.now().year)
    if '签署日期' not in field_values:
        today = date.today()
        field_values['签署日期'] = f"{today.year}年{today.month}月{today.day}日"

    # 展开块占位符
    field_values = expand_block_placeholders(field_values)

    # 构建所函案件描述
    if '所函案件描述' not in field_values:
        field_values['所函案件描述'] = expand_letter_case_description(field_values)

    # 检查模板中的占位符
    doc = Document(str(template))
    template_fields = set(get_all_placeholder_names(doc))
    provided_fields = set(field_values.keys())

    # 提示未填充的字段
    missing = template_fields - provided_fields
    if missing:
        print(f"警告: 以下 {len(missing)} 个占位符未提供值，将保留原样:")
        for f in sorted(missing):
            print(f"  {{{{ {f} }}}}")

    unknown = provided_fields - template_fields
    if unknown:
        print(f"提示: 以下字段在模板中未找到:")
        for f in sorted(unknown):
            print(f"  {f}")

    # 执行填充
    output = Path(args.output)
    fill_document(str(template), str(output), field_values)

    print(f"\n填充完成，已保存至: {output}")
    print(f"模板: {template.name}")
    print(f"已填充 {len(provided_fields & template_fields)} / {len(template_fields)} 个占位符")


def main():
    parser = argparse.ArgumentParser(description='律所委托材料 .docx 占位符填充工具')
    subparsers = parser.add_subparsers(dest='command', help='操作命令')

    # list 子命令
    list_parser = subparsers.add_parser('list', help='列出模板中的所有占位符')
    list_parser.add_argument('--template', '-t', required=True, help='模板文件路径')
    list_parser.add_argument('--verbose', '-v', action='store_true', help='显示详细位置信息')

    # fill 子命令
    fill_parser = subparsers.add_parser('fill', help='填充模板')
    fill_parser.add_argument('--template', '-t', required=True, help='模板文件路径')
    fill_parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    fill_parser.add_argument('--fields', '-f', help='JSON 格式的字段值，如 \'{"委托人姓名":"张三"}\'')
    fill_parser.add_argument('--fields-file', '-F', help='从 JSON 文件读取字段值')

    args = parser.parse_args()

    if args.command == 'list':
        cmd_list(args)
    elif args.command == 'fill':
        cmd_fill(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == '__main__':
    main()
