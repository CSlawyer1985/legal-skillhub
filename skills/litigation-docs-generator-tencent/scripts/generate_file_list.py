#!/usr/bin/env python3
"""
文件清单及盖章指引生成脚本
读取 references/文件清单及盖章指引模板.md 中的表格，生成 Word (.docx) 版文件清单。
格式与诉讼文书一致：宋体标题、仿宋正文、A4 页面。

用法：
  python3 generate_file_list.py <模板.md> <输出.docx> [--case "案件：XX诉XX XX纠纷案"]
"""
import re, sys, argparse
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_table(md_path):
    """解析 markdown 表格为二维数组"""
    rows = []
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    in_table = False
    for line in lines:
        s = line.strip()
        if s.startswith("|"):
            if not in_table:
                in_table = True
            cells = [c.strip() for c in s.strip("|").split("|")]
            # 跳过分隔行 |---|
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                continue
            rows.append(cells)
        elif in_table:
            break
    return rows


def set_font(run, name, size, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_solid_borders(table):
    """设置表格为实线边框（single），替代 python-docx 默认的虚线/无边框"""
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    # 移除已有 borders 定义
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")      # 1pt
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def render(md_path, out_path, case_line=""):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(3), Cm(2.5)
    section.left_margin, section.right_margin = Cm(3), Cm(2.5)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("文件清单"), "宋体", 18, bold=True)

    # 案件信息行（如有）
    if case_line:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        set_font(p.add_run(case_line), "仿宋", 14)

    # 说明
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    set_font(p.add_run("打印时按本清单份数准备，备注列标「公司盖章」的位置需加盖公章。"), "仿宋", 14)

    # 表格
    rows = parse_table(md_path)
    if rows:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_solid_borders(table)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.cell(ri, ci)
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                is_header = (ri == 0)
                set_font(para.add_run(val), "宋体" if is_header else "仿宋", 12, bold=is_header)

    doc.save(out_path)
    print(f"✅ {out_path}")
    return out_path


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("template", help="模板 .md 路径")
    ap.add_argument("output", help="输出 .docx 路径")
    ap.add_argument("--case", default="", help="案件信息行（可选）")
    args = ap.parse_args()
    render(args.template, args.output, args.case)
