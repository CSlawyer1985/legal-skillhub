import io
import sys
import re
import json
import argparse
from pathlib import Path

try:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except (AttributeError, io.UnsupportedOperation):
    pass

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from scripts.doc_converter import ensure_docx


def _extract_text_from_textboxes(doc):
    textbox_texts = []
    for txbx_content in doc.element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
        for para in txbx_content.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            text_parts = []
            for t in para.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    text_parts.append(t.text)
            if text_parts:
                textbox_texts.append(''.join(text_parts))
    for txbx in doc.element.iter('{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'):
        for para in txbx.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            text_parts = []
            for t in para.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    text_parts.append(t.text)
            if text_parts:
                textbox_texts.append(''.join(text_parts))
    return textbox_texts


def extract_text_from_docx(file_path):
    input_path = Path(file_path)
    converted_docx = None
    actual_input = str(input_path)

    if input_path.suffix.lower() == ".doc":
        print(f"检测到 .doc 格式文件，正在转换为 .docx 以提取文本 ...")
        try:
            docx_path, was_converted = ensure_docx(str(input_path))
            if was_converted:
                converted_docx = docx_path
                actual_input = docx_path
                print(f"转换完成，使用临时文件: {Path(docx_path).name}")
        except Exception as e:
            raise RuntimeError(f"无法转换 .doc 文件: {e}")

    try:
        doc = Document(actual_input)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append('\t'.join(row_texts))
        textbox_texts = _extract_text_from_textboxes(doc)
        for idx, text in enumerate(textbox_texts):
            text = text.strip()
            if text:
                parts.append(f'[文本框{idx+1}] {text}')
        return '\n'.join(parts)
    finally:
        if converted_docx:
            try:
                Path(converted_docx).unlink()
            except OSError:
                pass


def extract_abstract_text(text):
    pattern = re.compile(r'说\s*明\s*书\s*摘\s*要')
    match = pattern.search(text)
    if match:
        start_pos = match.end()
        content_pattern = re.compile(r'(本实用新型|本发明)')
        content_match = content_pattern.search(text[start_pos:])
        if content_match:
            abstract_text = text[start_pos + content_match.start():]
        else:
            abstract_text = text[start_pos:]
    else:
        content_pattern = re.compile(r'(本实用新型|本发明)')
        content_match = content_pattern.search(text)
        if not content_match:
            return ""
        abstract_text = text[content_match.start():]

    end_patterns = [
        r'\n1\s*[\.\、]\s*一种',
        r'\n技术领域',
        r'\n背景技术',
        r'\n附图说明',
        r'\n具体实施方式',
        r'\n权利要求',
        r'\n说明书附图',
        r'\n摘\s*要\s*附\s*图',
    ]

    min_end_pos = len(abstract_text)
    for ep in end_patterns:
        m = re.search(ep, abstract_text)
        if m:
            min_end_pos = min(min_end_pos, m.start())

    abstract_text = abstract_text[:min_end_pos]
    abstract_text = ' '.join(abstract_text.split())

    figure_match = re.search(r'\s+(图\d+)\s*$', abstract_text)
    if figure_match:
        abstract_text = abstract_text[:figure_match.start()].strip()

    return abstract_text


def extract_abstract_fig(text):
    patterns = [
        r'摘\s*要\s*附\s*图',
        r'摘要附图',
    ]

    start_pos = -1
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            start_pos = match.end()
            break

    if start_pos == -1:
        return ""

    remaining_text = text[start_pos:]

    end_patterns = [
        r'\n\s*权\s*利\s*要\s*求\s*书',
        r'\n\s*说\s*明\s*书',
        r'\n权利要求书',
        r'\n说明书',
    ]

    end_pos = len(remaining_text)
    for pattern in end_patterns:
        match = re.search(pattern, remaining_text)
        if match:
            end_pos = min(end_pos, match.start())

    figure_text = remaining_text[:end_pos].strip()
    figure_text = re.sub(r'^图\d+\s*$', '', figure_text, flags=re.MULTILINE)
    figure_text = figure_text.strip()

    return figure_text


def extract_claims(text):
    claim_pattern = r'(权\s*利\s*要\s*求\s*书)'

    header_text = text[:5000]
    match = re.search(claim_pattern, header_text)

    if not match:
        match = re.search(claim_pattern, text)

    if match:
        start_pos = match.end()
        next_section_pattern = r'(\n\s*说\s*明\s*书|\n\s*技术领域|^\s*说\s*明\s*书|^\s*技术领域)'
        next_match = re.search(next_section_pattern, text[start_pos:])

        if next_match:
            end_pos = start_pos + next_match.start()
            return text[start_pos:end_pos].strip()
        else:
            return text[start_pos:].strip()
    else:
        claim_item_pattern = r'(\d+)\s*[\.\、]'
        matches = list(re.finditer(claim_item_pattern, text))

        if matches:
            start_pos = matches[0].start()
            next_section_pattern = r'(\n\s*说\s*明\s*书|\n\s*技术领域|^\s*说\s*明\s*书|^\s*技术领域)'
            next_match = re.search(next_section_pattern, text[start_pos:])

            if next_match:
                end_pos = start_pos + next_match.start()
                return text[start_pos:end_pos].strip()
            else:
                return text[start_pos:].strip()
        else:
            return ""


def extract_description(text):
    lines = text.split('\n')

    start_idx = -1
    section_markers = ['技术领域', '背景技术', '发明内容', '实用新型内容', '附图说明', '具体实施方式']

    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped in section_markers:
            for j in range(i - 1, -1, -1):
                prev_line = lines[j].strip()
                if prev_line and not prev_line.startswith('图') and len(prev_line) <= 60 and '。' not in prev_line:
                    start_idx = j
                    break
                elif prev_line and prev_line not in section_markers:
                    break
            if start_idx == -1:
                start_idx = i - 1 if i > 0 else 0
            break

    if start_idx == -1:
        return _extract_description_fallback(text)

    end_idx = len(lines)

    for i, line in enumerate(lines[start_idx:], start=start_idx):
        stripped = line.strip()

        if re.match(r'^说\s*明\s*书\s*附\s*图$', stripped):
            end_idx = i
            break

        if re.match(r'^权\s*利\s*要\s*求\s*书$', stripped):
            end_idx = i
            break

        if re.match(r'^摘\s*要$', stripped):
            end_idx = i
            break

    result_lines = lines[start_idx:end_idx]
    cleaned_lines = []
    for line in result_lines:
        stripped = line.strip()
        if re.match(r'^说\s*明\s*书$', stripped):
            continue
        if re.match(r'^CN\s*\d+\s*\w*\s*\d+/\d+\s*页$', stripped):
            continue
        if re.match(r'^\d+$', stripped) and len(stripped) <= 4:
            continue
        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines)


def _extract_description_fallback(text):
    pattern = r'^[\s]*(说明书)[\s]*$'
    lines = text.split('\n')

    start_idx = -1
    for i, line in enumerate(lines):
        if re.match(pattern, line.strip()):
            start_idx = i
            break

    if start_idx == -1:
        pattern2 = r'说明书'
        for i, line in enumerate(lines):
            if re.search(pattern2, line):
                start_idx = i
                break

    if start_idx >= 0:
        remaining_lines = lines[start_idx + 1:]
        end_markers = ['权利要求书', '摘要', '说明书附图']

        end_idx = len(remaining_lines)
        for i, line in enumerate(remaining_lines):
            stripped = line.strip()
            for marker in end_markers:
                if stripped == marker or (marker in stripped and len(stripped) <= len(marker) + 5):
                    end_idx = i
                    break
            if end_idx != len(remaining_lines):
                break

        return '\n'.join(remaining_lines[:end_idx])

    return text


def extract_description_figs(text):
    patterns = [
        r'说\s*明\s*书\s*附\s*图',
        r'说明书附图',
    ]

    start_pos = -1
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            start_pos = match.end()
            break

    if start_pos == -1:
        figure_matches = list(re.finditer(r'\n\s*图\s*(\d+)\s*$', text, re.MULTILINE))
        if figure_matches:
            for i, m in enumerate(figure_matches):
                if m.group(1) == '1':
                    if i + 1 < len(figure_matches):
                        next_m = figure_matches[i + 1]
                        if next_m.group(1) == '2':
                            start_pos = m.start()
                            break

    if start_pos == -1:
        return ""

    remaining_text = text[start_pos:]
    figure_text = remaining_text.strip()

    lines = figure_text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^CN\s*\d+\s*\w*\s*\d+/\d+\s*页$', stripped):
            continue
        if re.match(r'^\d+$', stripped) and len(stripped) <= 4:
            continue
        if re.match(r'^说\s*明\s*书\s*附\s*图$', stripped):
            continue
        if stripped:
            cleaned_lines.append(stripped)

    return '\n'.join(cleaned_lines)


def split_patent_text(text):
    if not text or not text.strip():
        return {
            "abstract_text": "",
            "abstract_fig": "",
            "claims": "",
            "description": "",
            "description_figs": ""
        }

    text = text.replace('\r\n', '\n').replace('\r', '\n')

    abstract_text = extract_abstract_text(text)
    abstract_fig = extract_abstract_fig(text)
    claims = extract_claims(text)
    description = extract_description(text)
    description_figs = extract_description_figs(text)

    return {
        "abstract_text": abstract_text,
        "abstract_fig": abstract_fig,
        "claims": claims,
        "description": description,
        "description_figs": description_figs
    }


def main():
    parser = argparse.ArgumentParser(description='Extract and split Chinese patent .docx')
    parser.add_argument('input', type=str, help='Input .docx file path')
    parser.add_argument('--extract-only', action='store_true', help='Extract full text only, print to stdout')
    parser.add_argument('--output-json', type=str, metavar='FILE', help='Output split sections as JSON to file')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if not input_path.suffix.lower() in ('.docx', '.doc'):
        print(f"Error: Only .docx and .doc files are supported, got: {input_path.suffix}", file=sys.stderr)
        sys.exit(1)

    text = extract_text_from_docx(input_path)

    if args.extract_only:
        print(text)
        return

    result = split_patent_text(text)

    if args.output_json:
        output_path = Path(args.output_json)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"Output written to: {output_path}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
