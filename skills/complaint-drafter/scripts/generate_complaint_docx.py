#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
要素式 / 一般 起诉状 Word 生成器（自包含，仅需 python-docx）

═══════════════════════════════════════════════════════════════════════
核心原则（与最高法 67 类要素式示范文本要求一致）：
  · 忠实渲染官方模板：保留模板中【全部表格、行列与文字】，不得删减任何
    表格或单元格内容；未知项留空即可，但不得删除表格结构或提示文字。
  · 不改变字体：正文统一宋体、标题与表头统一黑体（与官方版式一致），
    不对单元格内文字做字体替换或样式破坏。
  · 不臆造格式：字号、列宽等均按官方 67 类要素式示范文本标准设置
    （标题黑体二号、节标题黑体小四、表内正文宋体小四），不随意变更。
═══════════════════════════════════════════════════════════════════════

两种模式：
  A) 模板忠实渲染（要素式 · 推荐）：
       python generate_complaint_docx.py <官方模板.txt> <output.docx>
     脚本读取官方 67 类要素式模板（references/templates/*.txt），在保留全部
     表格结构与文字的前提下渲染为 Word。适用于“复制官方模板 → 在原模板逐项
     填写 → 渲染”的流程。未知字段保持空白，所有原表格/提示文字均保留。

  B) 通用数据模式（一般起诉状 / 兜底）：
       python generate_complaint_docx.py data.json <output.docx>
     通过 DATA（JSON）构造 原告/被告/诉讼请求/事实与理由/证据 等通用结构。
"""

import sys
import json
import os
import re
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 与官方 67 类要素式示范文本一致的字体：正文宋体、标题/表头黑体
CJK = '宋体'
HEI = '黑体'

# A4 版心宽度（21cm - 左右各 2.5cm 页边距）
CONTENT_WIDTH_CM = 16.0

# 字号（与官方 67 类要素式示范文本标准一致）
SZ_TITLE = Pt(22)     # 标题“民事起诉状”：黑体二号
SZ_REASON = Pt(16)    # 案由（民间借贷纠纷）：黑体三号
SZ_SECTION = Pt(12)   # 节标题（表内首行合并单元格）：黑体小四
SZ_BODY = Pt(12)      # 表格内正文：宋体小四
SZ_NOTE = Pt(10.5)    # 说明框 / 提示性小字：宋体五号
SZ_SIGN = Pt(12)      # 落款：宋体小四

# 占位符：标记模式下写入可填写单元格，填空时按 id 精确替换（格式不变）
MARK_PRE = '⟦F'
MARK_POST = '⟧'

# 已知“节标题”行（官方模板中以表格首行合并单元格形式出现）
SECTION_HEADERS = {
    '当事人信息', '诉讼请求', '约定管辖和诉前保全', '诉前保全',
    '事实与理由', '对纠纷解决方式的意愿',
}

# 一般起诉状模式使用的“说明”提示框（要素式模板自带该文字，无需此常量）
NOTICE_TEXT = (
    "说明：\n"
    "为了方便您更好地参加诉讼，保护您的合法权利，请填写本表。\n"
    "1. 起诉时需向人民法院提交证明您身份的材料，如身份证复印件、营业执照复印件等。\n"
    "2. 本表所列内容是您提起诉讼以及人民法院查明案件事实所需，请务必如实填写。\n"
    "3. 本表有些内容可能与您的案件无关，您认为与案件无关的项目可以填“无”或不填；"
    "对于本表中勾选项可以在对应项打“√”；您认为另有重要内容需要列明的，可以另附页填写。\n"
    "4. 本表 word 电子版填写时，相关栏目可复制粘贴或扩容，但不得改变要素内容、格式设置。"
    "例如，多原告、多被告或多委托诉讼代理人等情况，可根据实际情况复制粘贴；需填写文字较多时，可根据实际对栏目进行扩容等。\n"
    "★特别提示★\n"
    "诉讼参加人应遵守诚信原则如实认真填写表格。如果诉讼参加人违反有关规定，"
    "虚假诉讼、恶意诉讼、滥用诉权，人民法院将视违法情形依法追究责任。"
)


# ───────────────────────── 字体 / 单元格工具 ─────────────────────────

def _set_cjk(run, name=CJK):
    """设置中英文字体（含东亚字体），不破坏既有文字内容。"""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)


def set_cell_shading(cell, fill='D9D9D9'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_vmerge(cell, val):
    """val='restart' 表示纵向合并起点；val=None 表示接续（continue）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for el in tcPr.findall(qn('w:vMerge')):
        tcPr.remove(el)
    vm = OxmlElement('w:vMerge')
    if val:
        vm.set(qn('w:val'), val)
    tcPr.append(vm)


def set_gspan(cell, n):
    """横向合并：当前单元格向右合并 n 列（n>1 才写 gridSpan）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for el in tcPr.findall(qn('w:gridSpan')):
        tcPr.remove(el)
    if n and n > 1:
        gs = OxmlElement('w:gridSpan')
        gs.set(qn('w:val'), str(n))
        tcPr.append(gs)


def _raw_cell(table, ri, ci):
    """直接取第 ri 行第 ci 列的 <w:tc> 元素（绕过 python-docx 的纵向合并栅格映射）。

    Cell 类在不同 python-docx 版本中的导出名不同（Cell / _Cell），故运行时动态获取。
    """
    CellCls = type(table.cell(0, 0))
    tr = table._tbl.findall(qn('w:tr'))[ri]
    tc = tr.findall(qn('w:tc'))[ci]
    return CellCls(tc, table)


# ───────────────────────── 表格连续性 ─────────────────────────
#
# 官方 67 类要素式母版（及我们截取的标记模板）的表格行默认没有
# w:cantSplit / w:tblHeader —— 一旦单元格内容变长导致表格跨页，行会被分页拆断、
# 节标题行也不重复，破坏“表格连续性”。这里在生成/填空时统一补齐这两个属性。

# w:trPr 子元素正确顺序（ECMA-376 CT_TrPr）；插入新元素时必须遵循，否则 Word
# 会因 schema 顺序错误而报错或自动修复。
_TRPR_ORDER = ['cnfStyle', 'divId', 'gridBefore', 'gridAfter', 'widthBefore',
               'widthAfter', 'cantSplit', 'trHeight', 'tblHeader',
               'tableCellSpacing', 'jc', 'hidden']


def _add_trpr_child(row, tag):
    """在表格行的 w:trPr 中按 schema 顺序插入（或重置）指定子元素（cantSplit/tblHeader）。"""
    trpr = row._tr.get_or_add_trPr()
    for el in trpr.findall(qn('w:' + tag)):
        trpr.remove(el)
    my_idx = _TRPR_ORDER.index(tag)
    pos = None
    for i, child in enumerate(trpr):
        ctag = child.tag.split('}')[-1]
        if ctag in _TRPR_ORDER and _TRPR_ORDER.index(ctag) > my_idx:
            pos = i
            break
    el = OxmlElement('w:' + tag)
    if pos is None:
        trpr.append(el)
    else:
        trpr.insert(pos, el)


def ensure_table_continuity(doc):
    """保证输出文档中每张表格的连续性：
      · 每个表格的每一行加 w:cantSplit —— 禁止单行被分页拆断（避免行跨页断裂）；
      · 每个表格的首行（节标题 / 表头合并行）加 w:tblHeader —— 表格跨页时表头自动重复。
    这样无论单元格填多长，表格都保持视觉连续，与官方 67 类要素式版式一致。"""
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            _add_trpr_child(row, 'cantSplit')
            if ri == 0:
                _add_trpr_child(row, 'tblHeader')


# ───────────────────────── 表格线补全 ─────────────────────────
#
# 官方 67 类母版（经 LibreOffice 由 .doc 转 .docx）部分单元格的边框是“缺省缺失”
# 而非显式 nil —— 例如节标题行与首内容行之间、合并单元格邻接处，常出现某条内边
# 线在相邻两格都未声明，渲染时该格线断开（用户反馈“个别地方表格线不完整”）。
# 该步骤在不改动任何已有边框、字体、合并的前提下，仅为“两侧都未声明”的内/外边线
# 按官方同款样式补全一条可见边框；纵向/横向合并（vMerge/gridSpan）区域的内部线
# 不补（保持合并外观）。边框样式严格复刻母版：
#   · 横向边（上/下）：single / sz=2 / color=231F20
#   · 纵向左边：      single / sz=4 / color=000000
#   · 纵向右边：      single / sz=2 / color=231F20

_TC_BORDERS_ORDER = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
_H_SPEC = ('single', 2, '231F20')     # 横向边（上/下）样式
_VL_SPEC = ('single', 4, '000000')    # 纵向左边样式
_VR_SPEC = ('single', 2, '231F20')    # 纵向右边样式


def _has_edge(tc, edge):
    """该单元格是否已声明指定边线（且非 nil/none）。"""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return False
    tb = tcPr.find(qn('w:tcBorders'))
    if tb is None:
        return False
    for e in tb:
        if e.tag.split('}')[-1] == edge:
            return e.get(qn('w:val')) not in ('nil', 'none')
    return False


def _ensure_cell_edge(tc, edge, spec):
    """若该单元格未声明指定边线（或显式 nil），按 spec=(val,sz,color) 补全之；
    已存在非 nil 边线则保持不动。按 schema 顺序插入 w:tcBorders 子元素。"""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tb = tcPr.find(qn('w:tcBorders'))
    if tb is None:
        tb = OxmlElement('w:tcBorders')
        tcPr.append(tb)
    # 先移除同边已有元素（含显式 nil），再按序插入真实边线
    for e in list(tb):
        if e.tag.split('}')[-1] == edge:
            tb.remove(e)
    el = OxmlElement('w:' + edge)
    el.set(qn('w:val'), spec[0])
    el.set(qn('w:sz'), str(spec[1]))
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), spec[2])
    pos = None
    for i, child in enumerate(tb):
        ct = child.tag.split('}')[-1]
        if ct in _TC_BORDERS_ORDER and _TC_BORDERS_ORDER.index(ct) > _TC_BORDERS_ORDER.index(edge):
            pos = i
            break
    if pos is None:
        tb.append(el)
    else:
        tb.insert(pos, el)


def _cell_at(rowmap, cv):
    """在某一行的 (tc, cv, gs, vm) 列表中，找覆盖视觉列 cv 的单元格。"""
    for (tc, cv2, gs, vm) in rowmap:
        if cv2 <= cv < cv2 + gs:
            return tc, cv2, gs, vm
    return None


def ensure_table_borders(doc):
    """补全文档中所有表格缺失的边框线（详见模块说明）。幂等、可重复执行；
    不触碰已有边框、字体、合并结构与文字内容。"""
    for table in doc.tables:
        tbl = table._tbl
        rows = tbl.findall(qn('w:tr'))
        if not rows:
            continue
        # 1) 构造视觉网格：每行 (tc, 视觉列cv, 网格跨度gs, 纵向合并vm)
        grid = []
        for tr in rows:
            tcs = tr.findall(qn('w:tc'))
            rowmap = []
            cv = 0
            for tc in tcs:
                gs, vm = 1, None
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    g = tcPr.find(qn('w:gridSpan'))
                    if g is not None:
                        try:
                            gs = int(g.get(qn('w:val')))
                        except (TypeError, ValueError):
                            gs = 1
                    v = tcPr.find(qn('w:vMerge'))
                    if v is not None:
                        vm = v.get(qn('w:val')) or 'continue'
                rowmap.append((tc, cv, gs, vm))
                cv += gs
            grid.append(rowmap)
        nrows = len(grid)
        ncols = max((cv + gs for rowmap in grid for (_, cv, gs, _) in rowmap), default=1)
        last_ri = nrows - 1

        # 2) 预计算每列纵向合并区域的下界 region_last[(ri,cv)]
        region_last = {}
        for cv in range(ncols):
            ri = 0
            while ri < nrows:
                cell = _cell_at(grid[ri], cv)
                if cell is None:
                    ri += 1
                    continue
                vm = cell[3]
                if vm == 'restart':
                    k = ri + 1
                    while k < nrows:
                        c2 = _cell_at(grid[k], cv)
                        if c2 is not None and c2[3] == 'continue':
                            k += 1
                        else:
                            break
                    for rr in range(ri, k):
                        region_last[(rr, cv)] = k - 1
                    ri = k
                else:
                    region_last[(ri, cv)] = ri
                    ri += 1

        # 3) 逐单元格补全四条边
        for ri, rowmap in enumerate(grid):
            for (tc, cv, gs, vm) in rowmap:
                is_merged = vm is not None
                rlast = region_last.get((ri, cv), ri)
                is_last_of_region = (rlast == ri)

                # —— 上边 ——
                if vm != 'continue':   # 合并续接格的上边在合并区内，不补
                    if ri == 0:
                        _ensure_cell_edge(tc, 'top', _H_SPEC)
                    else:
                        above = _cell_at(grid[ri - 1], cv)
                        if above is not None and (not _has_edge(above[0], 'bottom')) and (not _has_edge(tc, 'top')):
                            _ensure_cell_edge(tc, 'top', _H_SPEC)
                # —— 下边 ——
                if ri == last_ri:
                    _ensure_cell_edge(tc, 'bottom', _H_SPEC)
                elif is_merged and not is_last_of_region:
                    pass  # 合并区内，不补
                else:
                    below = _cell_at(grid[ri + 1], cv)
                    if below is not None and (not _has_edge(below[0], 'top')) and (not _has_edge(tc, 'bottom')):
                        _ensure_cell_edge(tc, 'bottom', _H_SPEC)
                # —— 左边 ——
                if cv == 0:
                    _ensure_cell_edge(tc, 'left', _VL_SPEC)
                else:
                    left = _cell_at(rowmap, cv - 1)
                    if left is not None and (not _has_edge(left[0], 'right')) and (not _has_edge(tc, 'left')):
                        _ensure_cell_edge(tc, 'left', _VL_SPEC)
                # —— 右边（跨度的右边界）——
                right_cv = cv + gs - 1
                if right_cv == ncols - 1:
                    _ensure_cell_edge(tc, 'right', _VR_SPEC)
                else:
                    nxt = _cell_at(rowmap, right_cv + 1)
                    if nxt is not None and (not _has_edge(nxt[0], 'left')) and (not _has_edge(tc, 'right')):
                        _ensure_cell_edge(tc, 'right', _VR_SPEC)
    return doc


def _para_is_empty(p):
    """判断 body 级段落是否为"空白段"：无文字、无图片、不含分页符。
    （分页符 w:br type=page 视为应剔除的对象，故也算"空"。）"""
    if p.find(qn('w:drawing')) is not None or p.find(qn('w:pict')) is not None:
        return False
    for t in p.findall('.//' + qn('w:t')):
        if (t.text or '').strip():
            return False
    return True


def normalize_layout(doc):
    """消除要素式起诉状中的空白页 / 大段空白，保证内容连续：
      1. 删除 body 级段落里内嵌的"下一页"分节符（w:sectPr type=nextPage 等）。
         这些分节符来自 67 类母版（每个案由/表格各自成节），会导致表格被强制
         推到下一页、形成空白页。仅移除"段落内的分节符"，保留 body 末尾的
         全局版心 sectPr（页边距/纸张设置）。
      2. 压缩连续空白段：连续多个空段只保留 1 个，并清掉其多余的段前/段后间距，
         避免堆叠出大段空白。标题与首表之间、各节之间仅留 1 个适度空段。
    幂等、可重复执行；不触碰任何表格结构、单元格内容与字体。"""
    body = doc.element.body

    # 1) 移除段落内嵌的"下一页"分节符（保留 body 级 sectPr）
    for p in body.findall(qn('w:p')):
        ppr = p.find(qn('w:pPr'))
        if ppr is None:
            continue
        sp = ppr.find(qn('w:sectPr'))
        if sp is not None:
            # 段落内的分节符 = 分节（分页）符号，删除之
            ppr.remove(sp)
            # 清掉可能因分节产生的多余属性（保留段内字体/对齐）
            # （无需额外处理，删除 sectPr 即可恢复为普通段落）

    # 2) 压缩连续空白段 + 归一化间距
    children = list(body)
    keep = []          # 保留的 body 子元素
    prev_was_empty = False
    for el in children:
        tag = el.tag.split('}')[-1]
        if tag != 'p':
            keep.append(el)
            prev_was_empty = False
            continue
        if _para_is_empty(el):
            if prev_was_empty:
                # 连续第二个及以上空段 -> 丢弃
                continue
            # 第一个空段：保留，但清掉多余段前/段后间距，避免大段空白
            ppr = el.find(qn('w:pPr'))
            if ppr is not None:
                sp = ppr.find(qn('w:spacing'))
                if sp is not None:
                    ppr.remove(sp)
            prev_was_empty = True
        else:
            keep.append(el)
            prev_was_empty = False
    # 重建 body 子元素顺序
    for el in list(body):
        body.remove(el)
    for el in keep:
        body.append(el)


def _add_text(cell, text, size=SZ_BODY, bold=False, font=CJK,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    """把文本写入单元格（保留多行），按指定字体/字号/对齐写入。"""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text if text else ' ')
    run.font.size = size
    run.bold = bold
    _set_cjk(run, font)


# ───────────────────────── 模板解析（模式 A） ─────────────────────────

def _parse_template_blocks(lines):
    """把官方模板文本解析为结构化块列表。

    返回块：
      {'t':'title', 'text':...}                 民事起诉状
      {'t':'reason', 'text':...}                （案由）
      {'t':'notice', 'rows':[...]}              说明框（1 列多行）
      {'t':'section', 'title':..., 'rows':[...]} 节（首行为节标题，rows 含提示行/要素行）
      {'t':'signature', 'text':...}             具状人…
      {'t':'text', 'text':...}                  其它自由文本
    """
    blocks = []
    cur_section = None
    notice_rows = []
    in_notice = False

    def flush_section():
        nonlocal cur_section
        if cur_section is not None:
            blocks.append(cur_section)
            cur_section = None

    def flush_notice():
        nonlocal notice_rows, in_notice
        if notice_rows:
            blocks.append({'t': 'notice', 'rows': list(notice_rows)})
            notice_rows = []
            in_notice = False

    pre = []
    for raw in lines:
        line = raw.rstrip('\n')
        stripped = line.strip()
        if stripped.startswith('|'):
            # 表格行：用 | 拆分，去掉首尾空段
            parts = line.split('|')
            cells = [p.rstrip() for p in parts[1:-1]] if len(parts) >= 2 else []
            if not cells and parts:
                cells = [p.rstrip() for p in parts if p != '']
            if not cells:
                continue
            c0 = cells[0].strip()
            if c0 in SECTION_HEADERS:
                flush_notice()
                flush_section()
                cur_section = {'t': 'section', 'title': c0, 'rows': []}
                continue
            # 内容 / 提示行
            if cur_section is not None:
                cur_section['rows'].append(cells)
            else:
                # 首个节之前：说明框
                if c0.startswith('说明：') or in_notice:
                    in_notice = True
                    notice_rows.append(cells[0])
                else:
                    blocks.append({'t': 'text', 'text': '\n'.join(cells)})
        else:
            flush_notice()
            if stripped == '':
                continue
            if '具状人' in stripped:
                flush_section()
                blocks.append({'t': 'signature', 'text': stripped})
            elif not pre:
                pre.append(stripped)
            else:
                if stripped.startswith('（') and stripped.endswith('）'):
                    blocks.append({'t': 'reason', 'text': stripped})
                else:
                    blocks.append({'t': 'title', 'text': stripped})
    flush_notice()
    flush_section()

    # 前置纯文本：第一个为标题，其余作标题补充
    for i, t in enumerate(pre):
        blocks.insert(i, {'t': 'title', 'text': t})
    return blocks


# ───────────────────────── 列宽配置 ─────────────────────────

def _widths_for(section_title, ncols):
    """按官方 67 类要素式版式比例返回每列宽度（cm），总长度≈版心 16cm。"""
    if section_title == '当事人信息':
        if ncols >= 3:
            return [1.9, 4.3, 9.8][:ncols]
        return [1.9, CONTENT_WIDTH_CM - 1.9]
    if ncols == 2:
        return [4.8, CONTENT_WIDTH_CM - 4.8]
    if ncols >= 3:
        return [5.2, 4.8, CONTENT_WIDTH_CM - 5.2 - 4.8][:ncols]
    return [CONTENT_WIDTH_CM]


# ───────────────────────── 渲染：说明框 / 节表格 ─────────────────────────

def _render_notice(doc, rows):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.width = Cm(CONTENT_WIDTH_CM)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F2F2F2')
    _add_text(cell, '\n'.join(rows), size=SZ_NOTE, font=CJK,
              align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()


ROLE_SET = {'原告', '被告', '第三人', '委托诉讼代理人'}


def _render_section(doc, block, fields=None):
    """渲染一个节：首行为节标题合并单元格（黑体小四），其下为提示行/要素行。

    保留官方模板的全部行列与文字；首列（诉讼地位 / 要素名）做纵向合并；
    提示性单行（如“（可完整表述…）”）横向合并整行。

    fields 不为 None 时进入“标记模式”：对可填写单元格（首列之外的列）写入唯一
    占位符 ⟦F{n}⟧，并把该字段（位置/原文/角色）记入 fields 列表，供后续“原封不动
    填空”使用——填空只是替换占位符文本，因此表格结构、字体、列宽完全不变。
    """
    title = block['title']
    rows = block['rows']
    ncols = max((len(r) for r in rows), default=1)
    total = len(rows) + 1  # + 节标题行
    table = doc.add_table(rows=total, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.width = Cm(CONTENT_WIDTH_CM)
    ws = _widths_for(title, ncols)
    total_w = sum(ws[:ncols])

    # 节标题行（首行合并）
    tcell = table.rows[0].cells[0]
    set_gspan(tcell, ncols)
    tcell.width = Cm(total_w)
    tcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _add_text(tcell, title, size=SZ_SECTION, bold=True, font=HEI,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # 首列纵向合并区间（仅“当事人信息”等含诉讼地位标签的表需要）
    groups = []
    i = 0
    while i < len(rows):
        if rows[i][0].strip():
            j = i + 1
            while j < len(rows) and not rows[j][0].strip():
                j += 1
            if j - 1 > i:
                groups.append((i, j - 1))
            i = j
        else:
            i += 1
    cont_rows = {k for g in groups for k in range(g[0] + 1, g[1] + 1)}

    cur_role = ''
    for di, r in enumerate(rows):
        ri = di + 1
        k = len(r)
        if k >= 1 and r[0].strip() and r[0].strip() in ROLE_SET:
            cur_role = r[0].strip()
        for ci in range(ncols):
            cell = _raw_cell(table, ri, ci)
            cell.width = Cm(ws[ci]) if ci < len(ws) else Cm(total_w / ncols)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if k == 1:
                # 提示性单行：横向合并整行（静态说明，不标记）
                if ci == 0:
                    set_gspan(cell, ncols)
                    _add_text(cell, r[0], size=SZ_NOTE, font=CJK,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                continue
            if ci >= k:
                continue  # 该行未使用的列：留空
            if ci == 0 and di in cont_rows:
                continue  # 纵向合并接续格：不写文字
            if ci == 0:
                # 首列：诉讼地位 / 要素名（静态标签，不标记）
                _add_text(cell, r[ci], size=SZ_BODY, bold=True, font=HEI,
                          align=WD_ALIGN_PARAGRAPH.LEFT)
                continue
            # ci >= 1：可填写单元格
            if fields is not None:
                fid = len(fields) + 1
                tok = MARK_PRE + str(fid) + MARK_POST
                fields.append({
                    'id': fid,
                    'section': title,
                    'role': cur_role,
                    'label': r[ci].strip(),
                    'original': r[ci],
                })
                _add_text(cell, tok, size=SZ_BODY, bold=False, font=CJK,
                          align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                _add_text(cell, r[ci], size=SZ_BODY, bold=False, font=CJK,
                          align=WD_ALIGN_PARAGRAPH.LEFT)

    # 写纵向合并标记
    for (s, e) in groups:
        set_vmerge(_raw_cell(table, s + 1, 0), 'restart')
        for kk in range(s + 1, e + 1):
            set_vmerge(_raw_cell(table, kk + 1, 0), None)
    doc.add_paragraph()


def render_template(txt_path, out_path):
    with open(txt_path, encoding='utf-8') as f:
        lines = f.readlines()
    blocks = _parse_template_blocks(lines)

    doc = _new_doc()

    for b in blocks:
        t = b['t']
        if t == 'title':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b['text'])
            r.bold = True
            r.font.size = SZ_TITLE
            _set_cjk(r, HEI)
        elif t == 'reason':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b['text'])
            r.font.size = SZ_REASON
            r.bold = True
            _set_cjk(r, HEI)
        elif t == 'notice':
            _render_notice(doc, b['rows'])
        elif t == 'section':
            _render_section(doc, b)
        elif t == 'signature':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_text_p(p, b['text'], size=SZ_SIGN, font=CJK)
        elif t == 'text':
            p = doc.add_paragraph()
            _add_text_p(p, b['text'], size=SZ_BODY, font=CJK)
    ensure_table_continuity(doc)
    normalize_layout(doc)
    ensure_table_borders(doc)
    doc.save(out_path)
    print('OK saved ->', out_path)


def _add_text_p(paragraph, text, size=SZ_BODY, bold=False, font=CJK):
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    for line in (text or '').split('\n'):
        if paragraph.runs:
            paragraph.add_run('\n')
        run = paragraph.add_run(line if line else ' ')
        run.font.size = size
        run.bold = bold
        _set_cjk(run, font)


# ───────────────────────── 标记模式（模式 C · 推荐） ─────────────────────────
#
# 思路：先把官方模板渲染成带唯一占位符 ⟦F{n}⟧ 的 .docx（每个可填写单元格一个），
# 并导出字段映射 JSON。用户上传内容后，专家识别并抽取各字段值，调用 fill_marked
# 按 id 替换占位符即可。因为只改单元格内的“文本”，表格结构 / 字体 / 列宽 / 合并
# 完全不变，从根本上杜绝“重新解析填写稿导致表格错乱”的问题。

def _new_doc():
    """创建一篇 A4、页边距 2.5cm、正文宋体小四的空白文档（与官方版式一致）。"""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    normal = doc.styles['Normal']
    normal.font.name = CJK
    normal.font.size = SZ_BODY
    _npr = normal.element.get_or_add_rPr()
    _nf = _npr.find(qn('w:rFonts'))
    if _nf is None:
        _nf = OxmlElement('w:rFonts')
        _npr.append(_nf)
    _nf.set(qn('w:eastAsia'), CJK)
    _nf.set(qn('w:ascii'), CJK)
    _nf.set(qn('w:hAnsi'), CJK)
    return doc


def render_marked(txt_path, out_docx, out_json):
    """标记模式入口：渲染官方模板为带占位符的 .docx，并导出字段映射 JSON。"""
    with open(txt_path, encoding='utf-8') as f:
        lines = f.readlines()
    blocks = _parse_template_blocks(lines)
    fields = []
    doc = _new_doc()
    for b in blocks:
        t = b['t']
        if t == 'title':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b['text'])
            r.bold = True
            r.font.size = SZ_TITLE
            _set_cjk(r, HEI)
        elif t == 'reason':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b['text'])
            r.font.size = SZ_REASON
            r.bold = True
            _set_cjk(r, HEI)
        elif t == 'notice':
            _render_notice(doc, b['rows'])
        elif t == 'section':
            _render_section(doc, b, fields=fields)
        elif t == 'signature':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_text_p(p, b['text'], size=SZ_SIGN, font=CJK)
        elif t == 'text':
            p = doc.add_paragraph()
            _add_text_p(p, b['text'], size=SZ_BODY, font=CJK)
    ensure_table_continuity(doc)
    normalize_layout(doc)
    ensure_table_borders(doc)
    doc.save(out_docx)
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump(fields, f, ensure_ascii=False, indent=2)
    print('OK marked ->', out_docx, '| fields =', len(fields))


def _maybe_break(text):
    """结构化填充值（含多个“：”且非整段叙述）按“：”自动断行，避免整段黏连（串行）；
    含“。”的整段叙述（如诉讼请求/事实与理由全文）不自动断行，交由原文段落控制。"""
    if '\n' in text or '。' in text:
        return text
    if text.count('：') >= 2:
        text = re.sub(r'：(?![　 ]*?[□▢])', '：\n', text)
    return text


def _replace_token_in_run(r, token, text):
    """在保持 run 的 rPr（字体）前提下，把 run 文本中的 token 替换为 text；
    text 含 \\n 时转为 <w:br/> 换行。返回是否命中 token。"""
    full = r.text
    if token not in full:
        return False
    new = full.replace(token, text)
    for t in r._element.findall(qn('w:t')):
        r._element.remove(t)
    for br in r._element.findall(qn('w:br')):
        r._element.remove(br)
    lines = new.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            r._element.append(OxmlElement('w:br'))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r._element.append(t)
    return True


def _set_cell_text(cell, text, token=None):
    """把单元格中的占位符 token 原地替换为 text，保留周围文字 / 标签 / 勾选框与字体；
    text 含 \\n 时自动换行。

    - 提供 token 时：在单元格内定位该 token（嵌入标记场景，token 可能只是单元格
      文本的一部分），仅替换占位符本身，标签、提示语、勾选框原样保留 → 不再“串行”。
    - 不提供 token（兼容）时：清空整格写入 text（旧行为）。
    """
    text = text if text else ''
    text = _maybe_break(text)
    if token:
        for p in cell.paragraphs:
            for r in p.runs:
                if _replace_token_in_run(r, token, text):
                    return
        # token 不在任何 run（极少见，如跨 run 拆分）→ 退回整格写入，保留原字体
    ps = cell.paragraphs
    p0 = ps[0]
    # 保留首段首个有字 run 的 rPr（字体/字号/颜色），避免整格重写时丢失官方字体
    src_rpr = None
    for r in p0.runs:
        if r.text.strip():
            src_rpr = r._element.find(qn('w:rPr'))
            break
    for p in ps[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            br = OxmlElement('w:br')
            p0.add_run('')._element.append(br)
        run = p0.add_run(line)
        if src_rpr is not None:
            run._element.append(copy.deepcopy(src_rpr))
        else:
            run.font.size = SZ_BODY
            _set_cjk(run, CJK)


def separate_party_fields(doc):
    """要素式「当事人信息」自然人行里，「住所地（户籍所在地）：____ 经常居住地：」
    与「证件类型：____ 证件号码：」在官方 67 类母版中本就分行显示；经 LibreOffice
    由 .doc 转 .docx 后，子字段之间的 <w:br/> 换行丢失，填空时空占位符塌陷，导致
    住址 / 证件号码 的填充内容直接黏在相邻空字段标签前，出现「串行」观感。

    此处在每个当事人行的「经常居住地：」「证件号码：」run 之前补一个 <w:br/>，
    使 住址 / 证件类型 / 证件号码 的填充内容各占一行、与其后的相邻空字段清晰分隔，
    即用户要求的「填充内容不换行（不黏连）」。

    仅插入换行符：不改动任何表格结构 / 字体 / 列宽 / 合并 / 既有文字；幂等可重复执行。
    """
    for table in doc.tables:
        for row in table.rows:
            for tc in row._tr.findall(qn('w:tc')):
                txt = ''.join(t.text or '' for t in tc.iter(qn('w:t')))
                for marker in ('经常居住地：', '证件号码：'):
                    if marker not in txt:
                        continue
                    for p in tc.findall(qn('w:p')):
                        for r in p.findall(qn('w:r')):
                            rt = ''.join(t.text or '' for t in r.findall(qn('w:t')))
                            if marker in rt:
                                prev = r.getprevious()
                                if prev is None or prev.tag != qn('w:br'):
                                    br = OxmlElement('w:br')
                                    r.addprevious(br)
                                break
    return doc


def fill_marked(docx_path, fields_json, values_json, out_docx):
    """按字段映射把占位符替换为用户内容；未提供值的字段恢复模板原文，
    从而“原封不动”保留表格与提示文字，格式绝不紊乱。

    关键：用原始 w:tc 元素遍历来定位占位符，规避 python-docx 对纵向合并
    （vMerge）续接单元格的“折叠”行为——续接格的 cell.text 会返回合并锚点
    的文字而非自身内容，导致续接格内的占位符（如空占位的 ⟦F15⟧/⟦F18⟧）被漏检、
    用户填值被静默丢弃。直接遍历 w:tc 即可稳定命中。
    """
    from docx.table import _Cell
    with open(fields_json, encoding='utf-8') as f:
        fields = json.load(f)
    with open(values_json, encoding='utf-8') as f:
        values = json.load(f)
    doc = Document(docx_path)
    for fld in fields:
        fid = fld['id']
        tok = MARK_PRE + str(fid) + MARK_POST
        v = values.get(str(fid))
        if v is None:
            v = values.get(fid)
        new_text = v if v else fld.get('original', '')
        if new_text is None:
            new_text = ' '
        found = False
        for table in doc.tables:
            for row in table.rows:
                for tc in row._tr.findall(qn('w:tc')):
                    ctext = ''.join(t.text or '' for t in tc.iter(qn('w:t')))
                    if tok in ctext:
                        cell = _Cell(tc, table)
                        _set_cell_text(cell, new_text, token=tok)
                        found = True
                        break
                if found:
                    break
            if found:
                break
        if not found:
            print('  [warn] 未找到占位符', tok, '(field original=%r)' % (fld.get('original', ''),))
    separate_party_fields(doc)
    ensure_table_continuity(doc)
    normalize_layout(doc)
    ensure_table_borders(doc)
    doc.save(out_docx)
    print('OK filled ->', out_docx)


# ───────────────────────── 通用数据模式（模式 B，兜底） ─────────────────────────

def _render_data_table(doc, rows):
    """通用 2 列表格（标签 | 内容），兜底用。"""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    n = len(rows)
    table = doc.add_table(rows=n, cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.width = Cm(CONTENT_WIDTH_CM)
    ws = [4.8, CONTENT_WIDTH_CM - 4.8] if max_cols == 2 else [CONTENT_WIDTH_CM / max_cols] * max_cols
    for ri, r in enumerate(rows):
        for ci in range(max_cols):
            cell = _raw_cell(table, ri, ci)
            cell.width = Cm(ws[ci]) if ci < len(ws) else Cm(CONTENT_WIDTH_CM / max_cols)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ci >= len(r):
                continue
            is_label = (ci == 0)
            _add_text(cell, r[ci], size=SZ_BODY, bold=is_label,
                      font=HEI if is_label else CJK)
    doc.add_paragraph()


def build(data, out_path):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    normal = doc.styles['Normal']
    normal.font.name = CJK
    normal.font.size = SZ_BODY
    _npr = normal.element.get_or_add_rPr()
    _nf = _npr.find(qn('w:rFonts'))
    if _nf is None:
        _nf = OxmlElement('w:rFonts')
        _npr.append(_nf)
    _nf.set(qn('w:eastAsia'), CJK)
    _nf.set(qn('w:ascii'), CJK)
    _nf.set(qn('w:hAnsi'), CJK)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(data.get('doc_title', '民事起诉状'))
    run.bold = True
    run.font.size = SZ_TITLE
    _set_cjk(run, HEI)

    if data.get('reason'):
        r = doc.add_paragraph()
        r.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = r.add_run('（' + data['reason'] + '）')
        rr.font.size = SZ_REASON
        rr.bold = True
        _set_cjk(rr, HEI)

    if data.get('include_notice', False):
        nt = doc.add_table(rows=1, cols=1)
        nt.style = 'Table Grid'
        nt.allow_autofit = False
        nt.width = Cm(CONTENT_WIDTH_CM)
        c = nt.rows[0].cells[0]
        set_cell_shading(c, 'F2F2F2')
        _add_text(c, NOTICE_TEXT, size=SZ_NOTE, font=CJK)
        doc.add_paragraph()

    if data.get('parties'):
        h = doc.add_paragraph()
        hr = h.add_run('当事人信息')
        hr.bold = True
        hr.font.size = SZ_SECTION
        _set_cjk(hr, HEI)
        rows = [[p.get('role', ''), p.get('fields', '')] for p in data['parties']]
        _render_data_table(doc, rows)

    for sec in data.get('sections', []):
        h = doc.add_paragraph()
        hr = h.add_run(sec.get('title', ''))
        hr.bold = True
        hr.font.size = SZ_SECTION
        _set_cjk(hr, HEI)
        if sec.get('free'):
            fp = doc.add_paragraph()
            _add_text_p(fp, sec['free'], size=SZ_BODY, font=CJK)
        if sec.get('elements'):
            _render_data_table(doc, sec['elements'])
        if sec.get('rows'):
            _render_data_table(doc, sec['rows'])
        doc.add_paragraph()

    if data.get('mediation'):
        h = doc.add_paragraph()
        hr = h.add_run('对纠纷解决方式的意愿')
        hr.bold = True
        hr.font.size = SZ_SECTION
        _set_cjk(hr, HEI)
        mp = doc.add_paragraph()
        _add_text_p(mp, data['mediation'], size=SZ_BODY, font=CJK)

    if data.get('signer_label'):
        sp = doc.add_paragraph()
        _add_text_p(sp, data['signer_label'], size=SZ_SIGN, font=CJK)

    separate_party_fields(doc)
    ensure_table_continuity(doc)
    normalize_layout(doc)
    ensure_table_borders(doc)
    doc.save(out_path)
    print('OK saved ->', out_path)


# ───────────────────────── 入口 ─────────────────────────

if __name__ == '__main__':
    args = sys.argv[1:]
    if len(args) < 1:
        print(__doc__)
        sys.exit(1)
    if args[0] == '--mark':
        # 模式 C：标记模式（推荐）
        #   python generate_complaint_docx.py --mark <官方模板.txt> <输出.docx> <输出字段.json>
        if len(args) < 4:
            print('用法: --mark <tpl.txt> <out.docx> <out.json>')
            sys.exit(1)
        render_marked(args[1], args[2], args[3])
    elif args[0] == '--fill':
        # 模式 C：填空（原封不动替换占位符，格式不变）
        #   python generate_complaint_docx.py --fill <标记.docx> <字段.json> <值.json> <out.docx>
        if len(args) < 5:
            print('用法: --fill <marked.docx> <fields.json> <values.json> <out.docx>')
            sys.exit(1)
        fill_marked(args[1], args[2], args[3], args[4])
    else:
        # 模式 A / B（兼容旧流程）
        if len(args) < 2:
            print(__doc__)
            sys.exit(1)
        inp, out = args[0], args[1]
        if inp.endswith('.txt'):
            # 模式 A：官方模板忠实渲染
            if not os.path.exists(inp):
                print('模板文件不存在：', inp)
                sys.exit(1)
            render_template(inp, out)
        elif inp.endswith('.json'):
            # 模式 B：通用数据
            with open(inp, encoding='utf-8') as f:
                data = json.load(f)
            build(data, out)
        else:
            print('不支持的输入类型，请使用 .txt 模板或 .json 数据。')
            sys.exit(1)
