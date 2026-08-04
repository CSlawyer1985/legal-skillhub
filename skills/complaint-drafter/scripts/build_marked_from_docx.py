#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从官方原文件（已转成 .docx 的母版）为每个案由生成"标记模板"：

  1. 在母版 body 中按"表单标题 + （案由名）"定位每个案由的【起诉状】表单
     （跳过 答辩状 / 实例），把该段 body 子元素原样截取到独立 per-case .docx。
     因为直接裁剪自母版，字体 / 列宽 / 合并 / 底纹等官方版式 100% 保留。
  2. 对每个 per-case 文档的"可填写单元格"写入唯一占位符 ⟦F{n}⟧，并导出
     fields.json（id / section / role / label / original）。填空走现有
     --fill 模式（按 id 替换文本），只改单元格内文字，表格结构永不紊乱。

用法：
  python build_marked_from_docx.py <母版.docx> <templates_dir> <out_marked_dir>
"""

import sys
import os
import re
import json
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MARK_PRE = '⟦F'
MARK_POST = '⟧'
CJK = '宋体'

# 表单标题集合（用于识别内容区的表单起止）
FORM_TITLES = {
    '民事起诉状', '行政起诉状', '刑事（附带民事）自诉状',
    '民事答辩状', '行政答辩状', '刑事（附带民事）自诉答辩状',
}
COMPLAINT_TITLES = {'民事起诉状', '行政起诉状', '刑事（附带民事）自诉状'}

ROLE_SET = {'原告', '被告', '第三人', '委托诉讼代理人', '法定代表人', '监护人'}
# 官方模板的角色标题常带限定说明，如“原告（自然人）”“第三人（法人、非法人组织）”
# “法定代表人（基层组织、单位工作人员等）”，需按前缀匹配，而非精确相等。
ROLE_PREFIX = ['原告', '被告', '第三人', '委托诉讼代理人', '法定代表人', '监护人']


def _match_role(t):
    # 角色标题可能带内部空格（如“第三人     （自然人）”）或限定说明，先去空白再匹配
    t = re.sub(r'\s+', '', t.strip())
    for r in ROLE_PREFIX:
        if t == r or t.startswith(r + '（') or t.startswith(r + '('):
            return r
    return None


SECTION_HEADERS = {
    '当事人信息', '诉讼请求', '约定管辖和诉前保全', '诉前保全',
    '事实与理由', '对纠纷解决方式的意愿', '证据清单', '请求依据',
    '答辩事项', '答辩依据', '答辩理由',
}


def _cell_text(cell):
    return ''.join(p.text or '' for p in cell.paragraphs)


# 官方模板用空格占位的填空格通常含这些关键字（年/月/日/元/编号/金额…）
FILL_KEYS = set('年月日元编号金额大写％%')


def _classify(t):
    """判断单元格是否可填写。"""
    t = t.strip()
    if t == '':
        return 'blank'
    if re.fullmatch(r'[_　\s]+', t):
        return 'blank'
    if re.search(r'_{2,}', t) or '□' in t or '▢' in t:
        return 'inline'  # 含下划线占位或勾选框 → 整格标记
    # 官方模板用连续空格占位（年  月  日 / 元 / 编号  等）：含填空关键字才标记为可填。
    # 排除“说明/提示”等说明性长文本：这些通常含句号“。”且不含上述填空关键字。
    if (re.search(r'[　 ]{2,}', t) and any(k in t for k in FILL_KEYS)
            and '。' not in t and not t.startswith('说明')):
        return 'inline'
    if t.endswith('：') or t.endswith(':'):
        return 'label'   # 标签列，值在其右侧空白格
    return 'static'


def _has_checkbox(t):
    """单元格含勾选框（□/▢）→ 视为“选项”单元格，整格标记由 agent 打钩。"""
    return ('□' in t) or ('▢' in t)


def _merge_spans(spans):
    if not spans:
        return []
    spans = sorted(spans)
    out = [list(spans[0])]
    for s, e in spans[1:]:
        if s <= out[-1][1]:
            out[-1][1] = max(out[-1][1], e)
        else:
            out.append([s, e])
    return [tuple(x) for x in out]


def _find_blanks(text):
    """返回 text 中需被占位符替换的“填空段”[(start,end)...]。

    仅识别连续下划线（官方模板的填空线）以及整格空白。空格占位（年 月 日）
    因上下文依赖强、易碎片化，不在此逐个标记，整格交给 agent 还原原句。
    """
    spans = []
    for m in re.finditer(r'_{2,}', text):
        spans.append((m.start(), m.end()))
    stripped = text.strip()
    if not spans and (stripped == '' or re.fullmatch(r'[_　\s]+', stripped)):
        spans = [(0, len(text))]
    return _merge_spans(spans)


def _embed_cell(cell, new_text):
    """把单元格文字原地改写为 new_text（含嵌入占位符），保留原段落首个有字 run 的字体。"""
    from docx.shared import Pt
    p0 = cell.paragraphs[0]
    src_rpr = None
    for r in p0.runs:
        if r.text.strip():
            src_rpr = r._element.find(qn('w:rPr'))
            break
    # 去掉多余段落（仅保留首个）
    for q in list(cell.paragraphs[1:]):
        q._element.getparent().remove(q._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    run = p0.add_run(new_text)
    if src_rpr is not None:
        run._element.append(copy.deepcopy(src_rpr))
    else:
        run.font.name = CJK
        run.font.size = Pt(12)
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.append(rf)
        rf.set(qn('w:eastAsia'), CJK)
        rf.set(qn('w:ascii'), CJK)
        rf.set(qn('w:hAnsi'), CJK)


def _mark_cell(cell, token):
    """整格标记：保留单元格【全部段落与 run 结构】（w:pPr / run rPr / 字符间距 /
    字体 / 颜色 / 段落数 一律不动），仅把首个有字 run 的文字替换为占位符，其余 run
    清空文字（保留其 rPr）。这样标记模板的【格式签名】与官方母版逐单元格一致；填空时
    token 被就地替换，保留该 run 的字体，输出与 Word 模板版式一模一样。

    适用：自由表述框 / 空格占位整短语 / 纯整格空白等整格填充字段。"""
    from docx.shared import Pt
    # 首段首个有字 run 的 rPr（仅当全空单元格时作为回退字体）
    src_rpr = None
    for p in cell.paragraphs:
        for r in p.runs:
            if r.text.strip():
                src_rpr = r._element.find(qn('w:rPr'))
                break
        if src_rpr is not None:
            break
    placed = False
    for p in cell.paragraphs:
        for r in p.runs:
            if not placed:
                r.text = token          # 复用首个 run（空或非空），保留其 rPr
                placed = True
            else:
                r.text = ''             # 其余 run 清空文字但保留 rPr / 结构
    if not placed:
        # 真正无任何 run 的单元格：在首个段落补一个 run（沿用首段 rPr，无则回退宋体小四）
        p0 = cell.paragraphs[0]
        run = p0.add_run(token)
        if src_rpr is not None:
            run._element.append(copy.deepcopy(src_rpr))
        else:
            run.font.name = CJK
            run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rf = rpr.find(qn('w:rFonts'))
            if rf is None:
                rf = OxmlElement('w:rFonts')
                rpr.append(rf)
            rf.set(qn('w:eastAsia'), CJK)
            rf.set(qn('w:ascii'), CJK)
            rf.set(qn('w:hAnsi'), CJK)


def _embed_scaffold_cell(cell, fields, cur_section, cur_role):
    """在【完整保留】单元格所有段落与 run 的既有格式（pStyle、w:spacing、w:ind、
    字符间距 w:spacing val=-N、字体、颜色、加粗）前提下，仅在以下位置嵌入占位符：

      · 勾选框（□ / ▢）→ 保留“选项文字”，把方框替换为占位符，agent 填 √ 或 □；
      · 连续下划线（官方填空线 ____）→ 替换为占位符；
      · 冒号后的空格填空位（如“担保债权的确定时间： 担保额度：”中冒号后的空格）
        → 该空格 run 替换为占位符。

    其余模板文字（如“担保债权的确定时间：”）原样保留。这样填空后右侧单元格的
    段落结构、样式与官方 Word 模板逐像素一致，彻底解决“勾选框/最高额担保/先行调解
    等右侧表格内容格式段落不对”的问题。

    返回新增字段数（0 表示本格无可填项，原样保留）。
    """
    added = 0
    ctx = ''
    for p in cell.paragraphs:
        prev_text = ''
        for r in list(p.runs):
            t = r.text or ''
            # 1) 纯空白 run：若上一 run 以全角冒号结尾，视为冒号后填空位
            if t.strip() == '':
                if prev_text.endswith('：'):
                    fid = len(fields) + 1
                    tok = MARK_PRE + str(fid) + MARK_POST
                    r.text = tok
                    fields.append({
                        'id': fid, 'section': cur_section, 'role': cur_role,
                        'label': (ctx + '（填空）').strip(), 'original': ' ',
                    })
                    added += 1
                prev_text = t
                continue
            # 2) 含勾选框（同一 run 可能含多个选项，逐个赋唯一占位符）
            if '□' in t or '▢' in t:
                matches = list(re.finditer(r'([^\s□▢]+)([□▢])', t))
                out = ''
                last = 0
                for m in matches:
                    out += t[last:m.start()]
                    opt = m.group(1)
                    fid = len(fields) + 1
                    tok = MARK_PRE + str(fid) + MARK_POST
                    # 注意：单元格 run 内已保留选项文字（如“是”），占位符只替换了方框，
                    # 因此 original 必须【仅】为方框字符（m.group(2)），不能是 m.group(0)
                    # （选项+方框）。否则填空未提供值时还原 original 会变成
                    # “是” + “是□” = “是是□”（方框重复），这正是用户反馈的“右侧格式
                    # 段落不对 / 勾选框重复”的根因。改为 m.group(2) 后：
                    #   未填 → 还原“□” → “是□”（与官方母版一致）；
                    #   填 √  → “是√”； 填 □ → “是□”。全部正确。
                    fields.append({
                        'id': fid, 'section': cur_section, 'role': cur_role,
                        'label': (opt or '选项').strip(), 'original': m.group(2),
                    })
                    out += opt + tok
                    last = m.end()
                    added += 1
                out += t[last:]
                r.text = out
                ctx = matches[-1].group(1) if matches else ''
                prev_text = out
                continue
            # 3) 含下划线填空
            if re.search(r'_{2,}', t):
                new = ''
                last = 0
                for um in re.finditer(r'_{2,}', t):
                    new += t[last:um.start()]
                    fid = len(fields) + 1
                    tok = MARK_PRE + str(fid) + MARK_POST
                    pre = t[max(0, um.start() - 12):um.start()].strip()
                    fields.append({
                        'id': fid, 'section': cur_section, 'role': cur_role,
                        'label': pre, 'original': t[um.start():um.end()],
                    })
                    new += tok
                    last = um.end()
                    added += 1
                new += t[last:]
                r.text = new
                prev_text = new
                continue
            # 4) 模板标签文字（含“：”的 label 或选项文字）→ 原样保留，更新上下文
            if '：' in t:
                ctx = t.rstrip('：').strip()
            else:
                ctx = t.strip()
            prev_text = t
    return added


# ───────────────────────── 表格连续性 ─────────────────────────
#
# 官方母版表格行默认没有 w:cantSplit / w:tblHeader，长内容跨页时行会被拆断、
# 节标题不重复，破坏连续性。在截取标记模板时一并补齐这两个属性。

_TRPR_ORDER = ['cnfStyle', 'divId', 'gridBefore', 'gridAfter', 'widthBefore',
               'widthAfter', 'cantSplit', 'trHeight', 'tblHeader',
               'tableCellSpacing', 'jc', 'hidden']


def _add_trpr_child(row, tag):
    """在表格行的 w:trPr 中按 schema 顺序插入（或重置）指定子元素。"""
    trpr = row._tr.get_or_add_trPr()
    for el in trpr.findall(qn('w:' + tag)):
        trpr.remove(el)
    my_idx = _TRPR_ORDER.index(tag)
    pos = None
    for i, child in enumerate(trpr):
        ctag = child.tag.split('}')[-1]
        if ctag in _TRPR_ORDER and _TRPR_ORDER.index(ctag) > my_idx:
            pos = i
            break
    el = OxmlElement('w:' + tag)
    if pos is None:
        trpr.append(el)
    else:
        trpr.insert(pos, el)


def ensure_table_continuity(doc):
    """每张表每行加 w:cantSplit（禁止行跨页拆断）；首行加 w:tblHeader（跨页重复表头）。"""
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            _add_trpr_child(row, 'cantSplit')
            if ri == 0:
                _add_trpr_child(row, 'tblHeader')


def _para_is_empty(p):
    if p.find(qn('w:drawing')) is not None or p.find(qn('w:pict')) is not None:
        return False
    for t in p.findall('.//' + qn('w:t')):
        if (t.text or '').strip():
            return False
    return True


def normalize_layout(doc):
    """消除空白页 / 大段空白：删除段落内嵌的"下一页"分节符（保留 body 级版心
    sectPr），压缩连续空白段。使单个案由的要素式模板连续成页，不留空白页。"""
    body = doc.element.body
    # 1) 移除段落内嵌的"下一页"分节符（不碰 body 末尾全局 sectPr）
    for p in body.findall(qn('w:p')):
        ppr = p.find(qn('w:pPr'))
        if ppr is None:
            continue
        sp = ppr.find(qn('w:sectPr'))
        if sp is not None:
            ppr.remove(sp)
    # 2) 压缩连续空白段 + 归一化间距
    children = list(body)
    keep = []
    prev_empty = False
    for el in children:
        tag = el.tag.split('}')[-1]
        if tag != 'p':
            keep.append(el)
            prev_empty = False
            continue
        if _para_is_empty(el):
            if prev_empty:
                continue
            ppr = el.find(qn('w:pPr'))
            if ppr is not None:
                sp = ppr.find(qn('w:spacing'))
                if sp is not None:
                    ppr.remove(sp)
            prev_empty = True
        else:
            keep.append(el)
            prev_empty = False
    for el in list(body):
        body.remove(el)
    for el in keep:
        body.append(el)


def _find_form_blocks(body, ch):
    """返回内容区表单块列表：{idx, title, casename, is_example}。

    母版两种排版都支持：
      (a) 两段落：标题段(民事起诉状) + 案由段（案由名）
      (b) 单段落：民事起诉状    （案由名）
    """
    blocks = []
    n = len(ch)
    title_re = re.compile(
        r'^(民事起诉状|行政起诉状|刑事（附带民事）自诉状|'
        r'民事答辩状|行政答辩状|刑事（附带民事）自诉答辩状)[\s　]*（(.+?)）\s*$')
    for i, c in enumerate(ch):
        if c.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(x.text or '' for x in c.iter(qn('w:t'))).strip()
        title = None
        casename = None
        if txt in FORM_TITLES:
            # 两段落式：下一段应为（案由名）
            if i + 1 < n and ch[i + 1].tag.split('}')[-1] == 'p':
                cn = ''.join(x.text or '' for x in ch[i + 1].iter(qn('w:t'))).strip()
                if cn.startswith('（') and cn.endswith('）'):
                    title, casename = txt, cn
        else:
            m = title_re.match(txt)
            if m:
                title = m.group(1)
                casename = '（' + m.group(2) + '）'
        if title is None:
            continue
        # 是否为实例：向前查 1~3 段是否有"实例"
        is_example = False
        for k in range(max(0, i - 3), i):
            pt = ''.join(x.text or '' for x in ch[k].iter(qn('w:t'))).strip()
            if pt == '实例':
                is_example = True
                break
        blocks.append({'idx': i, 'title': title, 'casename': casename, 'is_example': is_example})
    return blocks


def _extract_case(out_path, master_el, s, e, master_sectPr):
    """从 master 的 w:document 元素深拷贝 [s:e) 段落到独立 .docx（仅一次解析母版）。"""
    out = Document()  # 空白包（含内置 Table Grid / Normal 样式）
    out_body = out.element.body
    for c in list(out_body.iterchildren()):
        out_body.remove(c)
    src_body = master_el.body
    ch = list(src_body.iterchildren())
    for i in range(s, e):
        out_body.append(copy.deepcopy(ch[i]))
    # 清理可能混入的"实例"标记段（部分案由示例紧跟模板之后）
    for c in list(out_body.iterchildren()):
        if c.tag.split('}')[-1] == 'p':
            t = ''.join(x.text or '' for x in c.iter(qn('w:t'))).strip()
            if t == '实例':
                out_body.remove(c)
    if out_body.find(qn('w:sectPr')) is None and master_sectPr is not None:
        out_body.append(copy.deepcopy(master_sectPr))
    out.save(out_path)


def _row_tcs(row):
    """返回本行真实 w:tc 列表（按逻辑列展开，含 gridSpan 跨度）。

    python-docx 的 row.cells 会复用 cell 对象 / _tc 元素（合并、续接单元格），
    导致用 id() 去重时漏标。这里直接遍历原始 w:tc，每个 w:tc 是独立 XML 节点，
    并以 (table,row,逻辑列) 作为稳定去重键。
    """
    out = []
    col = 0
    for tc in row._tr.findall(qn('w:tc')):
        gs = tc.find(qn('w:tcPr') + '/' + qn('w:gridSpan'))
        span = int(gs.get(qn('w:val'))) if gs is not None else 1
        out.append((tc, col, span))
        col += span
    return out


def _mark_doc(docx_path):
    """在 per-case 文档上标记可填写单元格，返回 fields 列表。

    直接用原始 w:tc 元素遍历，避免 python-docx 复用 cell 对象导致的漏标。
    """
    from docx.table import _Cell
    doc = Document(docx_path)
    fields = []
    cur_role = ''
    cur_section = ''
    marked = set()  # 去重键 = (ti, ri, 逻辑列)
    for ti, table in enumerate(doc.tables):
        # 预解析每行的真实 tc（含逻辑列）
        per_row = []
        for row in table.rows:
            per_row.append([(_Cell(tc, table), ci, span)
                            for tc, ci, span in _row_tcs(row)])
        # 逐行：先识别本节 / 本行角色，再标记可填写单元格。
        # 角色按“行”判定——官方模板中角色标题与取值单元格在同一行（左侧），
        # 纵向合并的续接行继承其锚点行的角色；故必须逐行判定，不能用“全表最后
        # 一个角色标题”覆盖，否则前面的角色块会被后面的角色标题误标。
        for ri, cells in enumerate(per_row):
            n = len(cells)
            txts = [_cell_text(c) for c, _, _ in cells]
            if n == 1 or (len(set(txts)) == 1):
                head = txts[0].strip()
                if head in SECTION_HEADERS:
                    cur_section = head
                    cur_role = ''   # 进入新节，角色重置；角色仅在本节内由角色标题设置
            for c, _, _ in cells:
                r = _match_role(_cell_text(c).strip())
                if r:
                    cur_role = r
            kinds = [_classify(_cell_text(c)) for c, _, _ in cells]
            for idx in range(n):
                cell, ci, _ = cells[idx]
                key = (ti, ri, ci)
                if key in marked:
                    continue
                if MARK_PRE in _cell_text(cell):
                    continue  # 已标记，幂等跳过
                txt = _cell_text(cell)
                # 自由表述方框（诉讼请求 / 事实与理由 下方提示框）：整格标记，
                # 由 agent 填入用户上传的全部诉讼请求 / 事实与理由。
                if txt.startswith('（可完整表述'):
                    fid = len(fields) + 1
                    tok = MARK_PRE + str(fid) + MARK_POST
                    sec = ('诉讼请求' if '诉讼请求' in txt
                           else '事实与理由' if '事实与理由' in txt else cur_section)
                    fields.append({
                        'id': fid, 'section': sec, 'role': cur_role,
                        'label': sec + '（完整表述）', 'original': txt,
                    })
                    _mark_cell(cell, tok)
                    marked.add(key)
                    continue
                kind = kinds[idx]
                if kind == 'label':
                    if idx + 1 < n and kinds[idx + 1] == 'blank' \
                            and (ti, ri, cells[idx + 1][1]) not in marked:
                        rc = cells[idx + 1][0]
                        fid = len(fields) + 1
                        tok = MARK_PRE + str(fid) + MARK_POST
                        fields.append({
                            'id': fid,
                            'section': cur_section,
                            'role': cur_role,
                            'label': _cell_text(cell).strip(),
                            'original': _cell_text(rc),
                        })
                        _mark_cell(rc, tok)
                        marked.add((ti, ri, cells[idx + 1][1]))
                    continue
                if kind in ('blank', 'inline'):
                    blanks = _find_blanks(txt)
                    # 整格空白（整格下划线或纯空格占位）→ 整格标记
                    whole_blank = bool(blanks) and blanks == [(0, len(txt))]
                    # 含勾选框 / 下划线 / 冒号后填空位 → 原地嵌入标记，
                    # 完整保留段落结构、pStyle、间距、字符间距、字体（修复“右侧格式段落不对”）
                    if (not whole_blank) and (
                            _has_checkbox(txt)
                            or bool(blanks)
                            or ('：' in txt and re.search(r'：\s', txt))):
                        added = _embed_scaffold_cell(cell, fields, cur_section, cur_role)
                        if added:
                            marked.add(key)
                        continue
                    # 其余（整格空白 / 纯空格占位整短语等）→ 整格标记
                    fid = len(fields) + 1
                    tok = MARK_PRE + str(fid) + MARK_POST
                    fields.append({
                        'id': fid, 'section': cur_section, 'role': cur_role,
                        'label': txt.strip()[:40], 'original': txt,
                    })
                    _mark_cell(cell, tok)
                    marked.add(key)
        ensure_table_continuity(doc)
        normalize_layout(doc)
        doc.save(docx_path)
    return fields


def main():
    if len(sys.argv) < 4:
        print('用法: build_marked_from_docx.py <母版.docx> <templates_dir> <out_marked_dir>')
        sys.exit(1)
    master_path, tpl_dir, out_dir = sys.argv[1], sys.argv[2], sys.argv[3]
    os.makedirs(out_dir, exist_ok=True)

    master = Document(master_path)
    master_el = master.element
    mbody = master_el.body
    mch = list(mbody.iterchildren())
    master_sectPr = mbody.find(qn('w:sectPr'))
    blocks = _find_form_blocks(mbody, mch)
    print(f'母版中识别到表单块 {len(blocks)} 个')

    # 模板文件 -> 目标案由名 / 期望标题
    tpl_files = [f for f in os.listdir(tpl_dir) if f.endswith('.txt')]
    ok, skip = 0, 0
    for tf in sorted(tpl_files):
        name = tf[:-4]  # 去掉 .txt
        if name.endswith('_行政起诉状'):
            want_title = '行政起诉状'
            base = name[:-len('_行政起诉状')]
        else:
            want_title = None
            base = name
        target_casename = '（' + base + '）'

        cands = [b for b in blocks if b['casename'] == target_casename
                 and not b['is_example'] and b['title'] in COMPLAINT_TITLES]
        chosen = None
        if want_title:
            for b in cands:
                if b['title'] == want_title:
                    chosen = b
                    break
        else:
            for pref in ('民事起诉状', '行政起诉状', '刑事（附带民事）自诉状'):
                for b in cands:
                    if b['title'] == pref:
                        chosen = b
                        break
                if chosen:
                    break
        if chosen is None:
            print(f'  [跳过] 未匹配到 {name} (casename={target_casename})')
            skip += 1
            continue

        # 下一个表单块作为结束
        e = len(mch)
        for b in blocks:
            if b['idx'] > chosen['idx']:
                e = b['idx']
                break
        out_docx = os.path.join(out_dir, name + '.docx')
        out_json = os.path.join(out_dir, name + '.fields.json')
        _extract_case(out_docx, master_el, chosen['idx'], e, master_sectPr)
        fields = _mark_doc(out_docx)
        with open(out_json, 'w', encoding='utf-8') as f:
            json.dump(fields, f, ensure_ascii=False, indent=2)
        ok += 1
        print(f'  [OK] {name:24} title={chosen["title"]} fields={len(fields)}')
    print(f'\n完成：成功 {ok}，跳过 {skip}')


if __name__ == '__main__':
    main()
