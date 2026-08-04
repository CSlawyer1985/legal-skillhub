"""Element-based（要素式起诉状）模板填充引擎。

负责 token 替换、当事人区块裁剪与勾选渲染。

Token 约定：
  {{T:key}}            文本占位
  {{C:group=option}}   勾选占位 → □ / ☑

值表（values）形如：
  {
    "text": {key: 值},
    "checkboxes": {group: 选项或选项列表},
    "party_types": {"原告": "自然人", "被告": "法人", "第三人": "无"},
    "checked_mark": "☑" (可选)
  }

party_types 是可选的向后兼容字段。提供后，会删除模板中不适用的当事人
信息表格行：自然人只保留自然人行，法人/非法人组织只保留法人行；第三人
为“无”时，删除第三人的两类信息行。

缺失文本渲染为空、未选勾选渲染为 □，支持分步填充；调用方负责完整性核对。
替换会合并段内 run，模板占位符应置于格式中性的填空区。
"""
import re

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover - 仅在依赖缺失时执行
    raise SystemExit("缺少 python-docx，请运行：pip install python-docx") from exc

CHECKED_MARK = "☑"
UNCHECKED_MARK = "□"

_TOKEN_RE = re.compile(r"\{\{(T:[^{}]+|C:[^{}]+)\}\}")
_PARTY_TYPE_ALIASES = {
    "自然人": "自然人",
    "法人": "法人",
    "非法人组织": "法人",
    "法人、非法人组织": "法人",
    "法人/非法人组织": "法人",
    "无": "无",
    "不存在": "无",
    "没有": "无",
    "none": "无",
    "null": "无",
}


def _normalise_inline_value(value):
    """防止来源资料中的换行把模板同一逻辑行拆开。"""
    return re.sub(r"\s*[\r\n]+\s*", " ", str(value))


def render_token(body, values, inline=False):
    """body 形如 'T:原告.姓名' 或 'C:原告.性别=女'。"""
    if body.startswith("T:"):
        key = body[2:]
        value = values.get("text", {}).get(key, "")
        return _normalise_inline_value(value) if inline else str(value)
    # 'C:group=option'
    group, _, option = body[2:].partition("=")
    sel = values.get("checkboxes", {}).get(group)
    if isinstance(sel, list):
        checked = option in sel
    else:
        checked = option == sel
    return values.get("checked_mark", CHECKED_MARK) if checked else UNCHECKED_MARK


def _use_full_table_cell_width(paragraph):
    """清除表格段落缩进以使用完整单元格宽度，不改变原有字号。"""
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)


def _sync_party_addresses(values):
    """在未提供不同地址时，同步自然人与法人当事人的对应地址字段。"""
    result = dict(values)
    text = dict(values.get("text", {}))
    result["text"] = text
    for party in ("原告", "被告", "第三人"):
        # 自然人：同步住所地与经常居住地；两个地址均有值时绝不覆盖。
        for prefix in (f"{party}.自然人.", f"{party}."):
            residence = f"{prefix}住所地"
            habitual = f"{prefix}经常居住地"
            residence_value = str(text.get(residence, "")).strip()
            habitual_value = str(text.get(habitual, "")).strip()
            if residence_value and not habitual_value:
                text[habitual] = residence_value
            elif habitual_value and not residence_value:
                text[residence] = habitual_value
        # 法人/非法人组织：同步主要办事机构所在地与注册地/登记地；两个
        # 地址均有值时绝不覆盖，保留已经收集到的不同信息。
        prefix = f"{party}.法人."
        residence = f"{prefix}住所地"
        registration_keys = (
            f"{prefix}注册地 / 登记地",
            f"{prefix}注册地/登记地",
            f"{prefix}注册地",
        )
        residence_value = str(text.get(residence, "")).strip()
        registration_value = next(
            (str(text.get(key, "")).strip() for key in registration_keys
             if str(text.get(key, "")).strip()),
            "",
        )
        if residence_value and not registration_value:
            for key in registration_keys:
                text[key] = residence_value
        elif registration_value and not residence_value:
            text[residence] = registration_value
    return result


def replace_in_paragraph(paragraph, values):
    """对单个段落做 token 替换；通过合并 run 文本处理 token 跨 run 拆分。

    返回 True 表示发生了替换。替换后把结果写入首个 run、清空其余 run。
    """
    full = "".join(run.text for run in paragraph.runs)
    if "{{" not in full:
        return False
    def replace(match):
        # 仅当 token 所在的模板逻辑行本身还含有标签/其他文字时，才压平输入值。
        # 这样“证件类型：… 证件号码：…”等同一行不会被换行拆开，而单独的大
        # 文本填写区仍可保留用户有意提供的分段。
        line_start = full.rfind("\n", 0, match.start()) + 1
        line_end = full.find("\n", match.end())
        if line_end == -1:
            line_end = len(full)
        line_without_tokens = _TOKEN_RE.sub("", full[line_start:line_end])
        return render_token(match.group(1), values, inline=bool(line_without_tokens.strip()))

    new = _TOKEN_RE.sub(replace, full)
    if new == full:
        return False
    # 到这里 full 非空 ⇒ 至少有一个 run：把替换结果写入首个 run，清空其余
    paragraph.runs[0].text = new
    for r in paragraph.runs[1:]:
        r.text = ""
    return True


def iter_all_paragraphs(container):
    """遍历正文与表格（含嵌套表格）的所有段落。container 为 Document 或 _Cell。"""
    for p in container.paragraphs:
        yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)


def iter_table_paragraphs(container):
    """遍历 container 中全部表格单元格的段落，包含嵌套表格且避免合并单元格重复。"""
    seen_cells = set()

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                table_cell = cell._tc
                if table_cell in seen_cells:
                    continue
                # 保留 XML 单元格对象本身，而非其 id；临时 wrapper 被回收后 id
                # 可能被复用，导致后续单元格被误判为已遍历。
                seen_cells.add(table_cell)
                yield from cell.paragraphs
                for nested_table in cell.tables:
                    yield from walk_table(nested_table)

    for table in container.tables:
        yield from walk_table(table)


def _is_large_table_title(paragraph):
    """识别模板中以较大左缩进模拟居中的表格大标题。"""
    text = re.sub(r"\s+", "", paragraph.text)
    left_indent = paragraph.paragraph_format.left_indent
    return bool(text) and len(text) <= 20 and left_indent and left_indent.pt >= 100


def _format_signature_and_date(doc):
    """将具状人（签字、签章）与日期分成两行，并清除落款段落缩进。"""
    for paragraph in iter_all_paragraphs(doc):
        full = "".join(run.text for run in paragraph.runs)
        if "具状人" not in full or "日期" not in full:
            continue
        formatted = re.sub(r"[ \t　]*(日期[：:])", r"\n\1", full, count=1)
        if formatted != full:
            paragraph.runs[0].text = formatted
            for run in paragraph.runs[1:]:
                run.text = ""
        fmt = paragraph.paragraph_format
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _format_party_information_lines(doc):
    """将易拥挤的当事人信息字段固定为独立行。"""
    for paragraph in iter_all_paragraphs(doc):
        full = "".join(run.text for run in paragraph.runs)
        if "经常居住地" not in full and "证件号码" not in full:
            continue
        # 仅在标签前插入换行，保留同一字段中的长地址或号码完整显示。
        formatted = re.sub(r"[ \t　]*(经常居住地[：:])", r"\n\1", full)
        formatted = re.sub(r"[ \t　]*(证件号码[：:])", r"\n\1", formatted)
        if formatted == full:
            continue
        paragraph.runs[0].text = formatted
        for run in paragraph.runs[1:]:
            run.text = ""


def _normalize_all_table_paragraph_indents(doc):
    """让所有表格文字使用完整单元格宽度，同时保持大标题居中。"""
    for paragraph in iter_table_paragraphs(doc):
        if _is_large_table_title(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _use_full_table_cell_width(paragraph)


def _party_type(value):
    """将调用方的主体类型写法归一化；未知或缺失值不裁剪模板。"""
    if value is None:
        return None
    return _PARTY_TYPE_ALIASES.get(str(value).strip().lower())


def _row_text(row):
    """读取一行的唯一单元格文本，避免合并单元格被重复计入。"""
    seen, parts = set(), []
    for cell in row.cells:
        table_cell = cell._tc
        if table_cell not in seen:
            seen.add(table_cell)
            parts.append(cell.text)
    return "\n".join(parts)


def _remove_unused_party_rows(doc, values):
    """删除主体类型不适用的当事人信息行，返回删除的行数。"""
    raw_types = values.get("party_types", {})
    if not isinstance(raw_types, dict):
        return 0

    unwanted_prefixes = []
    for party in ("原告", "被告", "第三人"):
        kind = _party_type(raw_types.get(party))
        if kind == "自然人":
            unwanted_prefixes.append(f"{party}.法人.")
        elif kind == "法人":
            unwanted_prefixes.append(f"{party}.自然人.")
        elif party == "第三人" and kind == "无":
            unwanted_prefixes.extend(("第三人.自然人.", "第三人.法人."))

    if not unwanted_prefixes:
        return 0

    removed = 0
    for table in doc.tables:
        for row in list(table.rows):
            if any(prefix in _row_text(row) for prefix in unwanted_prefixes):
                row._tr.getparent().remove(row._tr)
                removed += 1
    return removed


def _is_blank_paragraph(element):
    return element.tag.endswith("}p") and not "".join(element.itertext()).strip()


def _remove_blank_paragraphs_between_tables(doc):
    """移除仅夹在两个表格之间的空段落，令相邻表格保持连续。"""
    body = doc.element.body
    children = list(body)
    removed = 0
    index = 0
    while index < len(children):
        if not _is_blank_paragraph(children[index]):
            index += 1
            continue
        start = index
        while index < len(children) and _is_blank_paragraph(children[index]):
            index += 1
        if (start > 0 and index < len(children)
                and children[start - 1].tag.endswith("}tbl")
                and children[index].tag.endswith("}tbl")):
            for element in children[start:index]:
                body.remove(element)
                removed += 1
    return removed


def fill_document(template_path, values, out_path):
    """加载模板 → 地址同步、裁剪/填充 → 规范版式 → 另存。"""
    values = _sync_party_addresses(values)
    doc = Document(template_path)
    _remove_unused_party_rows(doc, values)
    count = 0
    for p in iter_all_paragraphs(doc):
        if replace_in_paragraph(p, values):
            count += 1
    _remove_blank_paragraphs_between_tables(doc)
    _normalize_all_table_paragraph_indents(doc)
    _format_party_information_lines(doc)
    _format_signature_and_date(doc)
    doc.save(out_path)
    return count
