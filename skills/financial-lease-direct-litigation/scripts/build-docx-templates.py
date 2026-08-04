#!/usr/bin/env python3
"""Build the three editable A4 DOCX templates shipped with this package."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


AUTHOR = "李时瑀律师"
BODY_FONT = "Arial Unicode MS"
HEADING_FONT = "Arial Unicode MS"
BLUE = "2E5F86"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F3F4F6"
MUTED = RGBColor(90, 96, 104)
USABLE_WIDTH_DXA = 9072


def set_run_font(run, *, font: str = BODY_FONT, size: float = 10.5, bold: bool = False, color=None) -> None:
    run.font.name = font
    r_fonts = run._element.get_or_add_rPr().rFonts
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{script}"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_a4(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{script}"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 22, "000000", 0, 12),
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12, BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = HEADING_FONT
        for script in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{script}"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_header_footer(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run(f"{short_title} | 空白模板"), size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run(f"{AUTHOR}  |  第 "), size=8.5, color=MUTED)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), size=8.5, color=MUTED)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_in: list[float], *, header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        if row_index == 0 and header:
            set_repeat_header(row)
        for index, cell in enumerate(row.cells):
            width = Inches(widths_in[index])
            cell.width = width
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(text), size=size, bold=bold)


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title_paragraph.add_run(title), font=HEADING_FONT, size=22, bold=True)
    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_after = Pt(12)
    set_run_font(subtitle_paragraph.add_run(subtitle), size=10.5, color=MUTED)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    set_run_font(
        note.add_run("使用前先完成材料可读性、普通直租路由、权利基础和证据门禁；未知事实保持待核。"),
        size=9.5,
        color=MUTED,
    )


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        set_cell_text(table.cell(index, 0), label, bold=True, size=9.5)
        shade_cell(table.cell(index, 0), LIGHT_GRAY)
        set_cell_text(table.cell(index, 1), value, size=9.5)
    set_table_geometry(table, [1.25, 5.05], header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def setup_document(title: str, subject: str) -> Document:
    doc = Document()
    set_a4(doc.sections[0])
    configure_styles(doc)
    configure_header_footer(doc, title)
    props = doc.core_properties
    props.author = AUTHOR
    props.last_modified_by = AUTHOR
    props.title = title
    props.subject = subject
    props.keywords = "融资租赁, 普通直租, 空白模板, 可编辑DOCX"
    props.category = "公开方法模板"
    props.comments = "仅含通用字段，不含个案信息。"
    return doc


def build_complaint(path: Path) -> None:
    doc = setup_document("民事起诉状", "普通直租型融资租赁合同纠纷空白模板")
    add_title_block(doc, "民事起诉状", "普通直租型融资租赁合同纠纷 | 可编辑空白模板")
    add_metadata_table(
        doc,
        [
            ("路由状态", "【仅普通直租；其他结构填写 HOLD-ROUTE】"),
            ("材料状态", "【material_readable / HOLD-MATERIAL-READABILITY】"),
            ("证据状态", "【READY / HOLD-EVIDENCE】"),
            ("法源核验日", "【待填】"),
        ],
    )

    doc.add_heading("一、当事人及权利基础", level=1)
    party_table = doc.add_table(rows=5, cols=2)
    party_table.style = "Table Grid"
    party_rows = [
        ("原告", "【现行名称、住所、统一代码、法定代表人、联系方式】"),
        ("原告权利基础", "【合同主体／债权转让／回购承接／其他，附证据编号】"),
        ("被告一", "【主体、住所、送达地址、责任基础】"),
        ("被告二", "【如有：保证、抵押、配偶承诺或债务加入基础】"),
        ("其他主体", "【如有：逐一说明适格性和责任层级】"),
    ]
    for index, (label, value) in enumerate(party_rows):
        set_cell_text(party_table.cell(index, 0), label, bold=True)
        shade_cell(party_table.cell(index, 0), LIGHT_GRAY)
        set_cell_text(party_table.cell(index, 1), value)
    set_table_geometry(party_table, [1.25, 5.05], header=False)

    doc.add_heading("二、诉讼请求", level=1)
    requests = [
        "【选择并固定：支付全部未付租金路径，或解除合同并收回租赁物路径；不得机械并列】",
        "【违约责任、损失差额或租赁物价值处理，写明计算依据和截止日】",
        "【保证、抵押、配偶承诺或其他责任主体的具体责任请求】",
        "【诉讼费用、保全费用及其他有证据的程序性请求】",
    ]
    for text in requests:
        paragraph = doc.add_paragraph(style="List Number")
        set_run_font(paragraph.add_run(text), size=10.5)

    doc.add_heading("三、事实与理由", level=1)
    for heading, prompt in [
        ("1. 交易结构与合同订立", "【主体链、合同链、物权流、资金流和普通直租判断；附证据编号】"),
        ("2. 出租人履行", "【购买、价款支付、交付、验收和租赁物识别；无付款凭证不得写已付款】"),
        ("3. 承租人履行与违约", "【付款流水、冲抵顺序、起诉时现余额、催告和送达】"),
        ("4. 当前状态", "【当前权利人、时效证据、管辖连接点、租赁物位置及控制状态】"),
        ("5. 各被告责任基础", "【逐一对应合同、承诺或担保文件及责任期间】"),
        ("6. 法律依据状态", "【已核验官方法源及待核命题；不得以模板替代现行法回读】"),
    ]:
        paragraph = doc.add_paragraph()
        set_run_font(paragraph.add_run(heading), font=HEADING_FONT, size=11, bold=True, color=RGBColor.from_string(BLUE))
        paragraph = doc.add_paragraph(prompt)
        paragraph.paragraph_format.first_line_indent = Mm(7.4)

    doc.add_heading("四、附件与一致性检查", level=1)
    checks = [
        "【起诉状、计算表、证据目录和缺失材料清单编号一致】",
        "【所有当前事实均有当前证据，未知事项保持待核】",
        "【诉请路径选择、金额口径、时效和管辖状态前后一致】",
        "【唯一署名为李时瑀律师，立案前完成现行法与当地要求复核】",
    ]
    for text in checks:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(text), size=10.5)

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(18)
    set_run_font(closing.add_run("此致\n【有管辖权的人民法院】\n\n具状人：【待填】\n日期：【待填】"), size=10.5)
    doc.save(path)


def build_evidence_index(path: Path) -> None:
    doc = setup_document("证据目录", "普通直租型融资租赁合同纠纷证据目录空白模板")
    add_title_block(doc, "证据目录", "普通直租型融资租赁合同纠纷 | 可编辑空白模板")
    add_metadata_table(
        doc,
        [
            ("事项名称", "【待填】"),
            ("材料盘点状态", "【完整 / OCR待复核 / HOLD-MATERIAL-READABILITY】"),
            ("请求路径", "【支付全部未付租金 / 解除并收回租赁物 / 待选择】"),
        ],
    )

    doc.add_heading("一、证据编排表", level=1)
    headers = ["序号", "证据名称", "要件事实与证明目的", "原文件及页码", "读取状态", "对应诉请", "补证缺口"]
    table = doc.add_table(rows=8, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.cell(0, index), header, bold=True, center=True, size=8.5)
        shade_cell(table.cell(0, index), LIGHT_BLUE)
    for row_index in range(1, 8):
        values = [str(row_index), "【待填】", "【待填】", "【待填】", "【已读/OCR/人工复核】", "【待填】", "【无/待补】"]
        for col_index, value in enumerate(values):
            set_cell_text(table.cell(row_index, col_index), value, center=col_index in {0, 4}, size=8.5)
    set_table_geometry(table, [0.4, 1.0, 1.35, 0.9, 0.75, 0.8, 1.1])

    doc.add_heading("二、分组检查", level=1)
    groups = [
        "主体资格、授权及当前权利人",
        "融资租赁关系、购买、价款支付、交付和验收",
        "付款流水、冲抵顺序、现余额、违约和催告",
        "保证、抵押、配偶承诺、债务加入及责任期间",
        "债权转让、回购、承接、重组和第三方付款",
        "时效、管辖、送达、租赁物现状和保全线索",
    ]
    for text in groups:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(text), size=10.5)

    doc.add_heading("三、交叉核对", level=1)
    paragraph = doc.add_paragraph(
        "【逐项核对起诉状、计算表和缺失材料清单；不得出现有主张无证据、已取得材料被误列缺失、历史台账替代现余额或 OCR 未读内容被当成缺失。】"
    )
    paragraph.paragraph_format.line_spacing = 1.25
    doc.save(path)


def build_missing_materials(path: Path) -> None:
    doc = setup_document("缺失材料清单", "普通直租型融资租赁合同纠纷用户补件空白模板")
    add_title_block(doc, "缺失材料清单", "普通直租型融资租赁合同纠纷 | 用户补件可编辑模板")
    add_metadata_table(
        doc,
        [
            ("事项名称", "【待填】"),
            ("清单日期", "【待填】"),
            ("读取说明", "【已完成全量盘点；无法读取材料另列，不重复索要】"),
        ],
    )

    doc.add_heading("一、补件清单", level=1)
    headers = ["序号", "材料类别", "当前核验", "需补内容与可接受材料", "优先级", "缺失影响"]
    categories = [
        "原告主体、授权与当前权利人",
        "被告主体、身份与送达信息",
        "主合同、附件、补充及变更文件",
        "买卖合同、价款支付与发票",
        "租赁物识别、交付、验收与登记",
        "付款流水、冲抵顺序与现余额",
        "违约、催告、送达与时效证据",
        "保证、抵押、配偶承诺与债务加入",
        "债权转让、回购、承接或重组",
        "管辖条款、履行地与住所连接点",
        "租赁物现状、价值、控制与保全线索",
        "费用、既有诉讼、执行与回收材料",
    ]
    table = doc.add_table(rows=len(categories) + 1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.cell(0, index), header, bold=True, center=True, size=8.5)
        shade_cell(table.cell(0, index), LIGHT_BLUE)
    for row_index, category in enumerate(categories, start=1):
        values = [str(row_index), category, "【待确认】", "【填写缺口、用途及可接受材料】", "【待定】", "【填写直接影响】"]
        for col_index, value in enumerate(values):
            set_cell_text(table.cell(row_index, col_index), value, center=col_index in {0, 2, 4}, size=8.5)
    set_table_geometry(table, [0.4, 1.1, 0.75, 2.05, 0.75, 1.25])

    doc.add_heading("二、填写纪律", level=1)
    rules = [
        "已存在但 OCR 或人工回看尚未完成的材料，不列为缺失。",
        "每一项写明为什么需要、什么材料可以接受以及不补充的直接影响。",
        "未知事实保持待确认，不写内部谈判底线、胜诉概率或确定性结论。",
        "清单与证据目录、起诉状和计算表状态保持一致。",
        "唯一署名为李时瑀律师。",
    ]
    for text in rules:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(text), size=10.5)
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=Path("templates"))
    args = parser.parse_args()
    output_dir = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    build_complaint(output_dir / "民事起诉状-空白模板.docx")
    build_evidence_index(output_dir / "证据目录-空白模板.docx")
    build_missing_materials(output_dir / "缺失材料清单-空白模板.docx")
    print("built=3 author=李时瑀律师 page=A4 editable=true")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
