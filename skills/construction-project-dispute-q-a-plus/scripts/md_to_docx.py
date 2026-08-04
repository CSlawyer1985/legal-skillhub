#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 文本转换为 docx 文件。
简单解析规则：
- # 标题 → 标题 (Heading 1)
- ## 标题 → 标题 2 (Heading 2)
- ### 标题 → 标题 3 (Heading 3)
- 以数字开头、后跟 . 或 、的列表项/段落 → 普通段落
- - 或 * 开头的行 → 列表项（List Bullet）
- 空行分隔段落
- 粗体 **text** / __text__ → 粗体
- 引用块 > 文本 → 普通斜体段落
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_inline_formatting(paragraph, text):
    """解析行内粗体 **text** 或 __text__，追加到段落。"""
    # 统一处理 **...** 和 __...__
    parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
    for part in parts:
        run = paragraph.add_run()
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('__') and part.endswith('__') and len(part) >= 4:
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part


def markdown_to_docx(md_text, output_path, title=None):
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题样式
    for level in (1, 2, 3):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = '宋体'
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            heading_font.size = Pt(18)
        elif level == 2:
            heading_font.size = Pt(16)
        else:
            heading_font.size = Pt(14)

    # 文档标题
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = '宋体'
        run.font.size = Pt(20)
        run.font.bold = True
        doc.add_paragraph()

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过 YAML frontmatter
        if line.startswith('---') and i == 0:
            j = i + 1
            while j < len(lines) and not lines[j].startswith('---'):
                j += 1
            i = j + 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        if line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                run.font.name = '宋体'
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                run.font.name = '宋体'
            i += 1
            continue
        elif line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            for run in p.runs:
                run.font.name = '宋体'
            i += 1
            continue

        # 引用块
        if line.startswith('>'):
            quote_text = line[1:].strip()
            # 合并连续引用行
            j = i + 1
            while j < len(lines) and lines[j].startswith('>'):
                quote_text += '\n' + lines[j][1:].strip()
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(quote_text)
            run.font.name = '宋体'
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
            i = j
            continue

        # 列表项
        if re.match(r'^[\*\-\+]\s', line):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, re.sub(r'^[\*\-\+]\s+', '', line))
            i += 1
            continue

        # 表格：简单解析 Markdown 表格
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                rows.append(lines[i])
                i += 1
            # 过滤表头分隔行（如 |---|---|）
            data_rows = [r for r in rows if not re.match(r'^\|\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?$', r.strip())]
            if data_rows:
                cells_per_row = [r.strip('|').split('|') for r in data_rows]
                max_cols = max(len(cells) for cells in cells_per_row) if cells_per_row else 0
                if max_cols > 0:
                    table = doc.add_table(rows=len(data_rows), cols=max_cols)
                    table.style = 'Table Grid'
                    for r_idx, cells in enumerate(cells_per_row):
                        for c_idx in range(max_cols):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = cells[c_idx].strip() if c_idx < len(cells) else ''
                    doc.add_paragraph()
            continue

        # 普通段落
        p = doc.add_paragraph()
        parse_inline_formatting(p, line)
        i += 1

    doc.save(output_path)
    print(f"已保存：{output_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法：python md_to_docx.py <input.md> <output.docx> [标题]")
        sys.exit(1)
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else None

    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    markdown_to_docx(md_text, output_path, title)
