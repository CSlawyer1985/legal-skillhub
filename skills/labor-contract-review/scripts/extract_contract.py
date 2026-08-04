#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从劳动合同 PDF / DOCX 抽取纯文本，便于逐条合规检查。

用法:
    python extract_contract.py <合同路径> [-o 输出txt路径]

说明:
    - PDF 使用 pdfplumber 抽取；DOCX 使用 python-docx 抽取。
    - 首次运行前在托管 venv 安装依赖: pip install pdfplumber python-docx
    - .doc 旧格式不支持，请另存为 .docx 或以图片方式上传。
"""
import sys
import os
import argparse


def extract_pdf(path):
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("缺少 pdfplumber，请运行: pip install pdfplumber")
    chunks = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            chunks.append(f"--- 第{i}页 ---\n{text}")
    return "\n\n".join(chunks)


def extract_docx(path):
    try:
        import docx
    except ImportError:
        raise RuntimeError("缺少 python-docx，请运行: pip install python-docx")
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            line = " | ".join(c for c in cells if c.strip())
            if line:
                parts.append(line)
    return "\n".join(parts)


def main():
    ap = argparse.ArgumentParser(description="抽取劳动合同文本")
    ap.add_argument("path", help="合同文件路径 (.pdf / .docx)")
    ap.add_argument("-o", "--out", help="输出 txt 路径")
    args = ap.parse_args()

    ext = os.path.splitext(args.path)[1].lower()
    if ext == ".pdf":
        text = extract_pdf(args.path)
    elif ext == ".docx":
        text = extract_docx(args.path)
    elif ext == ".doc":
        raise RuntimeError(".doc 为旧格式，请另存为 .docx 或以图片方式上传")
    else:
        raise RuntimeError(f"不支持的格式: {ext}，仅支持 .pdf / .docx")

    out = args.out or (os.path.splitext(args.path)[0] + "_extracted.txt")
    with open(out, "w", encoding="utf-8") as f:
        f.write(text)
    print(out)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
