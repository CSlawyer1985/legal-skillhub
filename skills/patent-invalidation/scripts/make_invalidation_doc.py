#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_invalidation_doc.py —— 无效宣告请求书生成器（v1.0.8 增强）
=============================================================

两种模式:
  1. **模板模式**（无 --content）：生成 10 章空骨架 + 特征比对表占位，供智能体/人工填充。
  2. **填实模式**（--content content.json）：从 JSON 读取完整内容，生成可直接使用的请求书。

JSON 结构（content.json 必填字段）:
    {
      "target_patent": {
        "name": "机电式后拨链器",
        "number": "ZL201410571813.6",
        "applicant": "什拉姆有限责任公司",
        "app_date": "2014-12-29",
        "grant_date": "2017-06-13",
        "ipc": "B62M 9/124, B62M 25/08"
      },
      "requester": {
        "name": "兰溪轮峰车料有限公司",
        "address": "浙江省金华市兰溪市...",
        "agent": "×××专利代理事务所"
      },
      "request_items": "请求宣告 ZL201410571813.6 号发明专利权全部无效",
      "reasons": [
        {
          "name": "创造性(法22.3)",
          "main_argument": "本专利权利要求1相对于证据1+公知常识的结合不具备创造性",
          "evidence_refs": ["证据1", "证据2"],
          "analysis": "三步法分析：...（详细论证）...",
          "conclusion": "本专利权利要求1-5均不具备创造性，应宣告全部无效"
        }
      ],
      "evidence_list": [
        {"no":"1","name":"...","source":"...","pubdate":"...","type":"...","reasons":"...","form_legality":"...","remark":"..."}
      ],
      "feature_mapping": [
        {"claim_feature":"...","compare_feature":"...","type":"公开/实质相同/上位/未公开","evidence_ref":"...","remark":"..."}
      ]
    }

依赖:
    pip install python-docx
"""
import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: 需要安装 python-docx。请运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── 共用工具 ──────────────────────────────────────────────

def set_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def fill_placeholder(doc, label, hint, value=None):
    """填占位符：value 为 None 时显示 hint；否则显示 value。"""
    p = doc.add_paragraph()
    r = p.add_run(f'【{label}】')
    r.bold = True
    if value is not None and value:
        doc.add_paragraph(str(value))
    else:
        doc.add_paragraph(hint)


def add_table_from_rows(doc, header, rows, col_widths=None):
    """根据 header + rows 添加表格。"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = 'Table Grid'
    # 表头
    hdr = table.rows[0].cells
    for i, t in enumerate(header):
        cell = hdr[i]
        cell.paragraphs[0].text = ''
        run = cell.paragraphs[0].add_run(t)
        run.bold = True
    # 数据
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val) if val is not None else ''
    return table


# ── 模式 1: 模板模式（原行为）────────────────────────────

def generate_template(out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    title = doc.add_heading('无效宣告请求书', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_heading(doc, '一、首部', 1)
    fill_placeholder(doc, '涉案专利信息', '名称、专利号、专利权人、申请日/优先权日、授权公告日')
    fill_placeholder(doc, '请求人信息', '名称/姓名、地址、联系人；代理机构及代理人（附委托书）')

    set_heading(doc, '二、请求事项', 1)
    fill_placeholder(doc, '请求事项', '请求宣告涉案专利权全部无效 / 部分无效（具体到权利要求项）')

    set_heading(doc, '三、无效宣告范围', 1)
    fill_placeholder(doc, '无效范围', '列明请求宣告无效的权利要求编号')

    set_heading(doc, '四、无效理由与证据概述', 1)
    fill_placeholder(doc, '理由-证据对应', '逐条：无效理由（法/条款）+ 所依据证据编号；表明均具体说明并结合证据')

    set_heading(doc, '五、对比文件与证据说明', 1)
    fill_placeholder(doc, '证据清单', '编号、名称、来源、公开日、类型（文献/使用公开/档案）')

    set_heading(doc, '六、技术方案对比（特征映射）', 1)
    doc.add_paragraph('下表逐特征比对涉案专利权利要求与对比文件：')
    add_table_from_rows(
        doc,
        ['涉案专利权利要求特征', '对比文件对应特征', '是否公开/实质相同/上位/未公开', '备注'],
        [['（待填）'] * 4] * 4,
    )

    set_heading(doc, '七、各无效理由的具体论证', 1)
    reasons = [
        ('7.1 新颖性（法22.2）', '单篇实质相同论证 + 特征映射 + 结论'),
        ('7.2 创造性（法22.3）', '三步法：最接近现有技术→区别特征→实际解决的技术问题→技术启示击破 + 公知常识证据'),
        ('7.3 充分公开（法26.3）', '技术领域/问题/方案/效果四要素缺失论证'),
        ('7.4 清楚性（法26.4/细则22，旧细则20.1）', '含糊术语/自造词论证'),
        ('7.5 支持（法26.4）', '权利要求概括与说明书实施例落差'),
        ('7.6 修改超范围（法33）', '原始申请 vs 授权文本比对（可设两难陷阱）'),
        ('7.7 缺必特/重复授权/客体等', '按需展开（细则23.2（旧20.2） / 法9.1 / 法5·25）'),
    ]
    for name, hint in reasons:
        set_heading(doc, name, 2)
        fill_placeholder(doc, '论证', hint)

    set_heading(doc, '八、证据链与证明标准说明', 1)
    fill_placeholder(doc, '证明标准', '使用公开类：证据链闭合性、高度盖然性、域外公证认证、译文规范')

    set_heading(doc, '九、结论', 1)
    fill_placeholder(doc, '结论', '涉案专利不符合授权条件，请求宣告无效')

    set_heading(doc, '十、附件目录', 1)
    fill_placeholder(doc, '附件', '证据复印件、译文、公证认证、特征比对表、委托书')

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or '.', exist_ok=True)
    doc.save(out_path)
    return out_path


# ── 模式 2: 填实模式（v1.0.8 新增）────────────────────────

def generate_from_content(content: dict, out_path: str) -> str:
    """从 content dict 生成填好内容的请求书。"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    title = doc.add_heading('无效宣告请求书', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    target = content.get('target_patent', {})
    requester = content.get('requester', {})

    # 一、首部
    set_heading(doc, '一、首部', 1)
    fill_placeholder(doc, '涉案专利信息',
        '名称、专利号、专利权人、申请日/优先权日、授权公告日',
        value=(f"名称：{target.get('name', '')}；"
               f"专利号：{target.get('number', '')}；"
               f"专利权人：{target.get('applicant', '')}；"
               f"申请日：{target.get('app_date', '')}；"
               f"授权公告日：{target.get('grant_date', '')}；"
               f"IPC：{target.get('ipc', '')}"))
    fill_placeholder(doc, '请求人信息',
        '名称/姓名、地址、联系人；代理机构及代理人（附委托书）',
        value=(f"名称：{requester.get('name', '')}；"
               f"地址：{requester.get('address', '')}；"
               f"代理机构：{requester.get('agent', '')}"))

    # 二、请求事项
    set_heading(doc, '二、请求事项', 1)
    fill_placeholder(doc, '请求事项',
        '请求宣告涉案专利权全部无效 / 部分无效（具体到权利要求项）',
        value=content.get('request_items', ''))

    # 三、无效宣告范围
    set_heading(doc, '三、无效宣告范围', 1)
    fill_placeholder(doc, '无效范围',
        '列明请求宣告无效的权利要求编号',
        value=content.get('request_scope', '权利要求 1-5'))

    # 四、无效理由与证据概述
    set_heading(doc, '四、无效理由与证据概述', 1)
    reasons = content.get('reasons', [])
    overview_lines = []
    for i, r in enumerate(reasons, 1):
        ref_str = '、'.join(r.get('evidence_refs', []))
        overview_lines.append(f"{i}. **{r.get('name', '')}** —— 主攻论据：{r.get('main_argument', '')}（依据证据：{ref_str}）")
    fill_placeholder(doc, '理由-证据对应',
        '逐条：无效理由（法/条款）+ 所依据证据编号；表明均具体说明并结合证据',
        value='\n'.join(overview_lines) if overview_lines else None)

    # 五、对比文件与证据说明
    set_heading(doc, '五、对比文件与证据说明', 1)
    evidence_list = content.get('evidence_list', [])
    if evidence_list:
        rows = [[e.get('no', ''), e.get('name', ''), e.get('source', ''),
                 e.get('pubdate', ''), e.get('type', ''), e.get('reasons', ''),
                 e.get('form_legality', ''), e.get('remark', '')] for e in evidence_list]
        add_table_from_rows(
            doc,
            ['编号', '证据名称', '来源', '公开日', '类型', '与无效理由对应', '形式合法性', '备注'],
            rows,
        )
    else:
        fill_placeholder(doc, '证据清单', '编号、名称、来源、公开日、类型')

    # 六、技术方案对比（特征映射）
    set_heading(doc, '六、技术方案对比（特征映射）', 1)
    doc.add_paragraph('下表逐特征比对涉案专利权利要求与对比文件：')
    feature_mapping = content.get('feature_mapping', [])
    if feature_mapping:
        rows = [[fm.get('claim_feature', ''), fm.get('compare_feature', ''),
                 fm.get('type', ''), fm.get('evidence_ref', ''), fm.get('remark', '')]
                for fm in feature_mapping]
        add_table_from_rows(
            doc,
            ['涉案专利权利要求特征', '对比文件对应特征', '是否公开/实质相同/上位/未公开', '证据具体出处', '备注'],
            rows,
        )
    else:
        add_table_from_rows(
            doc,
            ['涉案专利权利要求特征', '对比文件对应特征', '是否公开/实质相同/上位/未公开', '备注'],
            [['（待填）'] * 4] * 4,
        )

    # 七、各无效理由的具体论证
    set_heading(doc, '七、各无效理由的具体论证', 1)
    if reasons:
        for i, r in enumerate(reasons, 1):
            set_heading(doc, f"7.{i} {r.get('name', '')}", 2)
            # 主攻论据
            fill_placeholder(doc, '主攻论据',
                '本理由的核心论据',
                value=r.get('main_argument', ''))
            # 依据证据
            ref_str = '、'.join(r.get('evidence_refs', []))
            fill_placeholder(doc, '依据证据',
                '本理由所依据的证据编号',
                value=ref_str or None)
            # 具体分析
            fill_placeholder(doc, '具体分析',
                '详细论证（含法条 + 事实 + 证据 + 结论）',
                value=r.get('analysis', ''))
            # 结论
            fill_placeholder(doc, '结论',
                '本理由的小结',
                value=r.get('conclusion', ''))
    else:
        # 退化为模板
        default_reasons = [
            ('7.1 新颖性（法22.2）', '单篇实质相同论证 + 特征映射 + 结论'),
            ('7.2 创造性（法22.3）', '三步法：最接近现有技术→区别特征→实际解决的技术问题→技术启示击破 + 公知常识证据'),
            ('7.3 充分公开（法26.3）', '技术领域/问题/方案/效果四要素缺失论证'),
            ('7.4 清楚性（法26.4/细则22，旧细则20.1）', '含糊术语/自造词论证'),
            ('7.5 支持（法26.4）', '权利要求概括与说明书实施例落差'),
            ('7.6 修改超范围（法33）', '原始申请 vs 授权文本比对（可设两难陷阱）'),
            ('7.7 缺必特/重复授权/客体等', '按需展开（细则23.2（旧20.2） / 法9.1 / 法5·25）'),
        ]
        for name, hint in default_reasons:
            set_heading(doc, name, 2)
            fill_placeholder(doc, '论证', hint)

    # 八、证据链与证明标准说明
    set_heading(doc, '八、证据链与证明标准说明', 1)
    fill_placeholder(doc, '证明标准',
        '使用公开类：证据链闭合性、高度盖然性、域外公证认证、译文规范',
        value=content.get('proof_standard',
            '本请求中各证据形式合法、内容真实，证据链相互印证达"高度盖然性"标准。'
            '对于中国专利文献证据 1-3，经国家知识产权局中国专利公布公告系统（epub.cnipa.gov.cn）'
            '下载官方专利单行本 PDF 即可采信；其他文献证据 4-5 经查证具有公开性与真实性。'))

    # 九、结论
    set_heading(doc, '九、结论', 1)
    fill_placeholder(doc, '结论',
        '涉案专利不符合授权条件，请求宣告无效',
        value=content.get('conclusion',
            f"综上所述，涉案 ZL{target.get('number', '××××××××.×')} 号发明专利的独立权利要求 1 "
            "相对于现有技术与公知常识的结合不具备突出的实质性特点和显著的进步，不符合专利法第 22 条第 3 款"
            "关于创造性的规定，从属权利要求 2-5 也不具备创造性。"
            f"依据《中华人民共和国专利法》第 22 条第 3 款、第 46 条的规定，"
            f"请求宣告 ZL{target.get('number', '××××××××.×')} 号发明专利权全部无效。"))

    # 十、附件目录
    set_heading(doc, '十、附件目录', 1)
    fill_placeholder(doc, '附件',
        '证据复印件、译文、公证认证、特征比对表、委托书',
        value=('附件 1：证据 1-5 的复印件\n'
               '附件 2：证据 4-5 的外文中文译文（如适用）\n'
               '附件 3：特征对照表（涉案专利 vs 对比文件同色标注）\n'
               '附件 4：无效宣告程序授权委托书'))

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or '.', exist_ok=True)
    doc.save(out_path)
    return out_path


# ── CLI 入口 ─────────────────────────────────────────────

def main() -> int:
    ap = argparse.ArgumentParser(
        description='无效宣告请求书生成器（模板模式 + 填实模式）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 模板模式（生成空骨架）
  python make_invalidation_doc.py --out template.docx

  # 填实模式（从 JSON 读取完整内容）
  python make_invalidation_doc.py --out request.docx --content content.json

content.json 格式见脚本顶部 docstring。
        """,
    )
    ap.add_argument('--out', required=True, help='输出 .docx 路径')
    ap.add_argument('--content', default='', help='填实模式：JSON 内容文件路径（可选）')
    args = ap.parse_args()

    if args.content:
        # 填实模式
        if not os.path.isfile(args.content):
            print(f"错误: content 文件不存在: {args.content}", file=sys.stderr)
            return 1
        with open(args.content, 'r', encoding='utf-8') as f:
            content = json.load(f)
        out = generate_from_content(content, args.out)
    else:
        # 模板模式
        out = generate_template(args.out)

    print(f"[OK] 已生成: {out}")
    return 0


if __name__ == '__main__':
    sys.exit(main())
