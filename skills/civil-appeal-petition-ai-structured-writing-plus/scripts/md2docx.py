#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
民事上诉状（AI结构化书写）Plus —— Markdown → docx 导出脚本

用途：
    将「第六步最终呈现、并经第八步附加 AI 免责提醒后的完整 Markdown 上诉状」
    转换为 .docx 文件，从工程层面固化 SKILL.md 第七步「一致性强制要求」：
    docx 必须与源 Markdown 在格式、内容、结构上一模一样（逐段、逐句、逐列表项对应）。

层级映射规则（与 SKILL.md 第七步严格一致）：
    - 一级标题 (# )            → Heading 1，居中
    - 二级标题 (## )           → Heading 2
    - 三级标题 (### )          → Heading 3
    - 有序列表 (1. 2. 3. )      → 有序列表（numbered list），文本/序号与 md 一致
    - 无序列表 (- / * )         → 无序列表（bulleted list）
    - 加粗 (**xxx**)            → 对应 run 加粗（保留）
    - 引用块 (> ...)            → 独立段落（用于底部 AI 免责提醒，逐行保留）
    - 普通段落 / 空行           → Normal 样式段落，文本逐字一致

依赖：python-docx（Skill 依赖项之一）
用法：
    python3 md2docx.py <input.md> [output.docx]
    若不指定 output.docx，默认保存到用户桌面：~/Desktop/民事上诉状.docx
    （若 ~/Desktop 不存在则尝试 ~/桌面）

注意（重要）：
    1. 调用本脚本前，必须确保传入的 Markdown 已是「含第八步免责提醒」的最终版本，
       因为脚本对源文件不做任何增删，仅忠实渲染。
    2. 脚本不联网、不调用任何 MCP，仅做格式转换。
"""

import sys
import os
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write(
        "错误：未安装 python-docx。请先安装：pip install python-docx\n"
    )
    sys.exit(2)


# --------------------------------------------------------------------------- #
# 行内格式解析：支持 **加粗**，其余原样保留
# --------------------------------------------------------------------------- #
def add_runs_with_bold(paragraph, text):
    """将含 **加粗** 的文本按片段写入 paragraph，保留加粗。"""
    # 用正则切分 **...** 片段
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def strip_inline(text):
    """去除行内标记但不改可见文本（用于纯文本兜底场景，本脚本主要用带格式版本）。"""
    return text.replace('**', '')


# --------------------------------------------------------------------------- #
# 判定行类型
# --------------------------------------------------------------------------- #
def detect_heading(line):
    m = re.match(r'^(#{1,6})\s+(.*)$', line)
    if m:
        return len(m.group(1)), m.group(2).strip()
    return None


def detect_ordered(line):
    m = re.match(r'^(\d+)[.)]\s+(.*)$', line)
    if m:
        return m.group(2).strip()
    return None


def detect_unordered(line):
    m = re.match(r'^[-*]\s+(.*)$', line)
    if m:
        return m.group(1).strip()
    return None


def detect_blockquote(line):
    m = re.match(r'^>\s?(.*)$', line)
    if m:
        return m.group(1)
    return None


# --------------------------------------------------------------------------- #
# 主转换逻辑
# --------------------------------------------------------------------------- #
def convert(md_text):
    doc = Document()

    # 基础正文样式（Normal 中文字体友好）
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(
        __import__('docx').oxml.ns.qn('w:eastAsia'), '宋体'
    )

    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # 跳过完全空行（不生成空段落，保持与 md 视觉一致；
        # 注：SKILL.md 要求"空行段落分隔尽量一致"，此处以不插入空段为准，
        # 因 md 空行本身多用于分隔块，word 段落间距已能体现）
        if line.strip() == '':
            i += 1
            continue

        # 标题
        h = detect_heading(line)
        if h:
            level, content = h
            # 显式指定样式名，避免 add_heading 内部映射偏移（实测 level=3 会写成 Heading 4）
            # 与 SKILL.md 第七步约定严格一致：#→Heading1, ##→Heading2, ###→Heading3
            style_name = "Heading %d" % min(max(level, 1), 9)
            p = doc.add_paragraph(content, style=style_name)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 有序列表（连续若干行合并为一个 numbered list）
        if detect_ordered(line) is not None:
            while i < n and detect_ordered(lines[i].rstrip()) is not None:
                item = detect_ordered(lines[i].rstrip())
                p = doc.add_paragraph(style='List Number')
                add_runs_with_bold(p, item)
                i += 1
            continue

        # 无序列表
        if detect_unordered(line) is not None:
            while i < n and detect_unordered(lines[i].rstrip()) is not None:
                item = detect_unordered(lines[i].rstrip())
                p = doc.add_paragraph(style='List Bullet')
                add_runs_with_bold(p, item)
                i += 1
            continue

        # 引用块（可能连续多行，如第八步免责提醒）
        if detect_blockquote(line) is not None:
            # 收集连续引用行
            quote_lines = []
            while i < n and detect_blockquote(lines[i].rstrip()) is not None:
                quote_lines.append(detect_blockquote(lines[i].rstrip()))
                i += 1
            # 引用块整体作为若干段落（逐行还原，保留换行）
            for q in quote_lines:
                p = doc.add_paragraph()
                # 引用文本保留加粗标记
                add_runs_with_bold(p, q)
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_runs_with_bold(p, line.strip())
        i += 1

    return doc


def resolve_output_path(output_arg):
    if output_arg:
        return output_arg
    desktop = os.path.expanduser("~/Desktop")
    if not os.path.isdir(desktop):
        desktop = os.path.expanduser("~/桌面")
    return os.path.join(desktop, "民事上诉状.docx")


def main():
    if len(sys.argv) < 2:
        sys.stderr.write(
            "用法：python3 md2docx.py <input.md> [output.docx]\n"
        )
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.isfile(input_path):
        sys.stderr.write("错误：找不到输入文件：%s\n" % input_path)
        sys.exit(1)

    output_path = resolve_output_path(sys.argv[2] if len(sys.argv) > 2 else None)

    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = convert(md_text)

    out_dir = os.path.dirname(os.path.abspath(output_path))
    if not os.path.isdir(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    doc.save(output_path)
    sys.stdout.write("✓ 已导出 docx：%s\n" % output_path)


if __name__ == '__main__':
    main()
