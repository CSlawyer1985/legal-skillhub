# -*- coding: utf-8 -*-
"""生成 assets/复审请求书_模板.docx：标准结构（请求事项/事实与理由/修改对照/证据清单）+ 对比表占位。"""
import os
from docx import Document
from docx.shared import Pt

OUT = os.path.join(os.path.dirname(__file__), '..', 'assets', '复审请求书_模板.docx')

# (前缀, 名称, 级别)  level: 0=title 1=h1 2=h2 3=h3  'table'=修改对照表占位
ITEMS = [
    ('', '复审请求书', 0),
    ('', '请求事项', 1),
    ('一、', '对本申请权利要求做进一步修改（若采用修改策略）', 2),
    ('', '__修改对照表__', 'table'),
    ('二、', '对本申请技术方案的进一步阐释说明', 2),
    ('三、', '关于对驳回决定的抗辩说明', 2),
    ('3.1', '针对创造性（/新颖性）驳回的逐条抗辩', 3),
    ('3.2', '驳回决定的违法/违规/事实错误点', 3),
    ('', '证据清单', 1),
    ('', '附件目录', 1),
]


def main():
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc = Document()

    for pre, name, level in ITEMS:
        if level == 0:
            h = doc.add_heading(name, level=0); h.alignment = 1
        elif level == 1:
            doc.add_heading(name, level=1)
        elif level == 2:
            doc.add_heading(f'{pre}{name}', level=2)
        elif level == 3:
            doc.add_heading(f'{pre} {name}', level=3)
        elif level == 'table':
            tbl = doc.add_table(rows=2, cols=4)
            tbl.style = 'Table Grid'
            hdr = tbl.rows[0].cells
            for i, t in enumerate(['权利要求项', '修改前（被驳回文本）', '修改后', '修改依据（法33）']):
                hdr[i].paragraphs[0].add_run(t).bold = True
            row = tbl.rows[1].cells
            for i, t in enumerate(['（示例）权1', '（填入原文）', '（填入修改后）', '原说明书__段/图__']):
                row[i].paragraphs[0].add_run(t)
            note = doc.add_paragraph('〔提示：修改仅消除驳回缺陷、不得扩大保护范围（细则66）；每处标注法33依据。〕')
            note.runs[0].font.size = Pt(9)
            note.runs[0].font.italic = True

    # 落款
    doc.add_paragraph()
    doc.add_paragraph('请求人（申请人）：________________________')
    doc.add_paragraph('专利代理机构：__________________________')
    doc.add_paragraph('日期：______年______月______日')

    doc.save(OUT)
    print('Template saved:', os.path.abspath(OUT))


if __name__ == '__main__':
    main()
