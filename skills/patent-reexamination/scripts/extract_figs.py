# -*- coding: utf-8 -*-
"""
extract_figs.py —— 专利审查意见答复附图提取与并列比对工具

功能：
  1. 从申请文件(docx)按 document.xml 真实嵌入顺序提取附图（避免 media/imageN 序号≠图号）。
  2. 从对比文件(PDF)定位并渲染关键附图页（自动寻找含"图1"的页；扫描件按指定页号）。
  3. 可选：将"本申请图 vs 对比文件图"左右并列嵌入到意见陈述书对应段落之后。

依赖：pip install python-docx PyMuPDF

用法示例：
  # 仅提取附图到输出目录
  python extract_figs.py --app 申请文件.docx --out figs/

  # 提取并在意见陈述书锚点段落后插入并列对比图
  python extract_figs.py --app 申请文件.docx \
      --refs "对比文件1.pdf:图1:6" "对比文件2.pdf:系统图:8" \
      --doc 意见陈述书.docx --anchor "图示清楚显示：制冷机"
"""

import argparse
import os
import re
import shutil
import zipfile

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None


# ---------------------------------------------------------------------------
# 1) 从 docx 按真实嵌入顺序提取附图
# ---------------------------------------------------------------------------
def extract_app_figures(docx_path, out_dir, max_n=6):
    """返回 [(图号, 文件路径), ...] 按文档出现顺序。"""
    os.makedirs(out_dir, exist_ok=True)
    results = []
    with zipfile.ZipFile(docx_path) as z:
        rels = z.read('word/_rels/document.xml.rels').decode('utf-8')
        docxml = z.read('word/document.xml').decode('utf-8')
        rid2target = {}
        for m in re.finditer(r'Id="(rId\d+)"[^>]*Target="(media/([^"]+))"', rels):
            rid2target[m.group(1)] = 'word/' + m.group(2)
        order = [m.group(1) for m in re.finditer(r'<a:blip r:embed="(rId\d+)"', docxml)]
        seen = []
        for rid in order:
            if rid in rid2target and rid not in seen:
                seen.append(rid)
        for i, rid in enumerate(seen[:max_n], 1):
            data = z.read(rid2target[rid])
            out = os.path.join(out_dir, f'app_fig{i}.png')
            with open(out, 'wb') as f:
                f.write(data)
            results.append((i, out))
    return results


# ---------------------------------------------------------------------------
# 2) 从 PDF 定位并渲染附图页
# ---------------------------------------------------------------------------
def render_pdf_figure(pdf_path, out_path, page_hint=None, figure_marker='图1', dpi=2.2):
    """page_hint: 0-based 页号（扫描件用）；否则自动按 figure_marker 查找。"""
    if fitz is None:
        raise RuntimeError('需安装 PyMuPDF: pip install PyMuPDF')
    doc = fitz.open(pdf_path)
    target = None
    if page_hint is not None:
        target = int(page_hint)
    else:
        for i in range(doc.page_count):
            t = doc[i].get_text()
            if re.search(rf'{re.escape(figure_marker)}[\s\b]', t) or figure_marker in t:
                target = i
                break
        if target is None:  # 扫描件无文本：回退到中部页（由调用方指定更稳妥）
            target = min(6, doc.page_count - 1)
    doc[target].get_pixmap(matrix=fitz.Matrix(dpi, dpi)).save(out_path)
    doc.close()
    return target + 1  # 返回 1-based 页号


# ---------------------------------------------------------------------------
# 3) 可选：将并列对比图嵌入 docx
# ---------------------------------------------------------------------------
def embed_side_by_side(docx_path, anchor_text, left_img, right_img,
                       left_cap, right_cap, note, out_path=None):
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree

    if out_path is None:
        out_path = docx_path
    if os.path.abspath(docx_path) != os.path.abspath(out_path):
        shutil.copy2(docx_path, out_path)
    doc = Document(out_path)

    anchor = None
    for p in doc.paragraphs:
        if anchor_text in p.text:
            anchor = p
            break
    if anchor is None:
        raise ValueError(f'未找到锚点段落: {anchor_text[:30]}')

    IMG_W = Cm(7.0)
    table = doc.add_table(rows=2, cols=2)
    for col, img, cap in [(0, left_img, left_cap), (1, right_img, right_cap)]:
        cell = table.cell(0, col)
        cell.width = IMG_W
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(img, width=IMG_W)
        c2 = table.cell(1, col)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = cap.split('\n')
        for j, line in enumerate(lines):
            r2 = p2.add_run(line)
            r2.font.size = Pt(9)
            r2.font.name = '宋体'
            r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            if j != len(lines) - 1:
                p2.add_run('\n')
        # 浅灰边框
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        b = tcPr.find(qn('w:tcBorders'))
        if b is None:
            b = etree.SubElement(tcPr, qn('w:tcBorders'))
        for edge in ('top', 'left', 'bottom', 'right'):
            be = b.find(qn(f'w:{edge}'))
            if be is None:
                be = etree.SubElement(b, qn(f'w:{edge}'))
            be.set(qn('w:val'), 'single')
            be.set(qn('w:sz'), '4')
            be.set(qn('w:color'), 'CCCCCC')

    anchor._p.addnext(table._tbl)

    # 追加对比小结段
    new_p = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    r = etree.SubElement(new_p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    t = etree.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    t.text = note
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    table._tbl.addnext(new_p)

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description='专利附图提取与并列比对')
    ap.add_argument('--app', required=True, help='申请文件 docx 路径')
    ap.add_argument('--out', default='figs', help='附图输出目录')
    ap.add_argument('--refs', nargs='*', default=[],
                    help='对比文件规格: "pdf路径:图标记或系统图:页号(可选)"，如 "D1.pdf:图1:6"')
    ap.add_argument('--doc', help='意见陈述书 docx（用于嵌入对比图）')
    ap.add_argument('--anchor', help='嵌入锚点段落文本')
    ap.add_argument('--left-cap', default='本申请图')
    ap.add_argument('--right-cap', default='对比文件图')
    ap.add_argument('--note', default='')
    args = ap.parse_args()

    figs = extract_app_figures(args.app, args.out)
    print('申请文件附图提取:', [(n, os.path.basename(p)) for n, p in figs])

    for spec in args.refs:
        parts = spec.split(':')
        pdf = parts[0]
        marker = parts[1] if len(parts) > 1 else '图1'
        hint = parts[2] if len(parts) > 2 and parts[2].isdigit() else None
        outp = os.path.join(args.out, f'ref_{os.path.splitext(os.path.basename(pdf))[0]}.png')
        try:
            page = render_pdf_figure(pdf, outp, page_hint=hint, figure_marker=marker)
            print(f'对比文件 {os.path.basename(pdf)} 附图 -> {outp} (页{page})')
        except Exception as e:
            print(f'渲染失败 {pdf}: {e}')

    if args.doc and args.anchor and figs:
        embed_side_by_side(
            args.doc, args.anchor,
            figs[0][1],
            os.path.join(args.out, f'ref_{os.path.splitext(os.path.basename(args.refs[0].split(":")[0]))[0]}.png'),
            args.left_cap, args.right_cap, args.note,
            out_path=args.doc.replace('.docx', '_附图比对版.docx')
        )
        print('已生成附图比对版意见陈述书')


if __name__ == '__main__':
    main()
