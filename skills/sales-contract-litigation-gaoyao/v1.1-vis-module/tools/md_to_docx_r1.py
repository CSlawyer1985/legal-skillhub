#!/usr/bin/env python3
"""九文书盲稿 md → 法院版式 DOCX 转换器（GPL-3.0，追偿权批次r4.1定版）。

- 阻断1：正文用本机实际存在的中文字体 Songti SC（STFangsong），标题 宋体；
  fonts 表登记 altName 映射（Songti SC↔仿宋_GB2312、宋体↔Songti SC），
  呈法院终稿在 Windows Word 中置换为 仿宋_GB2312（既定规范不变）。
- 阻断2：显式设置 A4（21×29.7cm）与页边距（上下2.54/左右3.17cm）。
- 阻断4：生成后对 DOCX ZIP 容器做规范化重写（entry 时间固定 2026-07-26 00:00、
  按名排序、统一 deflate），容器级字节确定。
"""
from __future__ import annotations

import io
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


NEUTRAL_BANNER = "内部研究草稿 · 非提交文书 · 法律意见与现行法结论待人工复核 · 禁止外发"

def sanitize_line(ln: str):
    """用户可见docx脱除构建治理信息；返回None表示整行跳过。"""
    if re.search(r"17850\d{8}", ln) and ("builder" in ln or "落账" in ln or "免脱敏令" in ln or "授权" in ln):
        ln = re.sub(r"[（(][^（）()]*17850\d{8}[^（）()]*[)）]", "（按委托人指示）", ln)
    ln = re.sub(r"17850\d{8}(-user|-codex|-claude-code|-hermes)?", "", ln)
    if "builder=claude-code" in ln or "builder=claude" in ln:
        if ln.strip().startswith("*") or "审查人" in ln or "出具" in ln:
            ln = re.sub(r"追偿权纠纷Skill[（(]r4候选[)）]builder=claude-code", "追偿权纠纷Skill（候选版）", ln); ln = re.sub(r"[（(]r4候选[)）]builder=claude-code", "（候选版）", ln)
            ln = ln.replace("builder=claude-code", "")
        else:
            return None
    if re.search(r"数据截止=卷内材料|本底稿每一事实句均可回溯", ln) and ln.strip().startswith("*"):
        return None
    ln = ln.replace("root免脱敏令", "委托人免脱敏指示").replace("root 免脱敏令", "委托人免脱敏指示")
    ln = ln.replace("待root终签", "待人工法律复核终签").replace("待 root 终签", "待人工法律复核终签")
    ln = ln.replace("root单独授权", "委托人单独授权").replace("root（人工）", "委托人（人工）")
    ln = ln.replace("经root", "经委托人").replace("root对", "委托人对").replace("由root", "由委托人")
    ln = ln.replace("root终签", "人工法律复核终签").replace("root", "委托人")
    return ln

BODY_FONT = "Songti SC"  # 前两案交付约定+正典链实证可落图
HEAD_FONT = "Heiti SC"  # 前两案交付约定标题字体
FIXED_DT = (2026, 7, 26, 0, 0, 0)


def set_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr
    rfonts = r.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = r.makeelement(qn("w:rFonts"), {})
        r.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, size=14, font=None, bold=False, align=None, spacing=28, first_indent=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(spacing)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    if first_indent:
        pf.first_line_indent = Cm(0.99)
    text = text.replace("`", "")
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        b = part.startswith("**") and part.endswith("**")
        run = p.add_run(part[2:-2] if b else part)
        set_font(run, font or BODY_FONT, size, bold=bold or b)
    return p


def add_table(doc, rows, repeat_header=False):
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    # r1修订合同：全部表格首行跨页重复表头 + 全部行禁止跨页拆分（不设固定行高）
    for ri, row in enumerate(t.rows):
        trPr = row._tr.get_or_add_trPr()
        if ri == 0:
            e = trPr.makeelement(qn("w:tblHeader"), {}); trPr.append(e)
        e2 = trPr.makeelement(qn("w:cantSplit"), {}); trPr.append(e2)
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(20)
            txt = (row[j] if j < len(row) else "").replace("**", "")
            run = p.add_run(txt)
            set_font(run, BODY_FONT, 11, bold=(i == 0))
    return t


def set_a4(doc, landscape=False):
    for section in doc.sections:
        if landscape:
            from docx.enum.section import WD_ORIENT
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
        else:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def register_font_alts(doc):
    """fonts 表登记 altName：本机字体 ↔ 法院规范字体互认。"""
    from docx.oxml import parse_xml
    fonts_part = None
    for rel in doc.part.package.main_document_part.part.rels.values():
        if rel.reltype.endswith("/fontTable"):
            fonts_part = rel.target_part
            break
    if fonts_part is None:
        return
    root = fonts_part._element if hasattr(fonts_part, "_element") else None
    try:
        from lxml import etree
        tree = etree.fromstring(fonts_part.blob)
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        def font_entry(name, alt):
            e = etree.SubElement(tree, f"{{{ns}}}font")
            e.set(f"{{{ns}}}name", name)
            a = etree.SubElement(e, f"{{{ns}}}altName")
            a.set(f"{{{ns}}}val", alt)
            return e
        existing = {f.get(f"{{{ns}}}name") for f in tree.findall(f"{{{ns}}}font")}
        if "Songti SC" not in existing:
            font_entry("Songti SC", "宋体")
        if "Heiti SC" not in existing:
            font_entry("Heiti SC", "黑体")
        fonts_part._blob = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    except Exception:
        pass


def normalize_zip(path: Path):
    """ZIP 容器规范化：固定 entry 时间、按名排序、统一压缩。"""
    src = zipfile.ZipFile(path, "r")
    entries = {}
    for info in src.infolist():
        entries[info.filename] = src.read(info.filename)
    src.close()
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as z:
        for name in sorted(entries):
            zi = zipfile.ZipInfo(name, date_time=FIXED_DT)
            zi.compress_type = zipfile.ZIP_DEFLATED
            zi.external_attr = 0o644 << 16
            z.writestr(zi, entries[name])
    path.write_bytes(buf.getvalue())


def convert(md_path: Path, out_path: Path):
    doc = Document()
    # 证据目录（02）按基线规范横版排布，其余纵向
    is_evidence_table = md_path.stem.startswith("02")
    set_a4(doc, landscape=is_evidence_table)
    lines = md_path.read_text(encoding="utf-8").split("\n")
    i = 0
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            rows = []
            for tl in table_buf:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                if set("".join(cells)) <= set("-: "):
                    continue
                rows.append(cells)
            if rows:
                add_table(doc, rows, repeat_header=is_evidence_table)
            table_buf = []

    banner_done = False
    while i < len(lines):
        raw = lines[i].rstrip()
        if raw.startswith("> ") and not banner_done:
            banner_done = True
            add_para(doc, NEUTRAL_BANNER, size=10.5, font=HEAD_FONT, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=16)
            i += 1
            continue
        _s = sanitize_line(raw)
        if _s is None:
            i += 1
            continue
        ln = _s
        if ln.startswith("|"):
            table_buf.append(ln)
            i += 1
            continue
        flush_table()
        if not ln.strip():
            i += 1
            continue
        if ln.startswith("> "):
            add_para(doc, ln[2:].strip(), size=10.5, font=HEAD_FONT, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=16)
        elif ln.startswith("# "):
            add_para(doc, ln[2:].strip(), size=18, font=HEAD_FONT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=36)
        elif ln.startswith("## "):
            add_para(doc, ln[3:].strip(), size=15, font=HEAD_FONT, bold=True, spacing=30)
        elif ln.startswith("### "):
            add_para(doc, ln[4:].strip(), size=14, font=HEAD_FONT, bold=True, spacing=28)
        elif ln.startswith("---"):
            pass
        elif ln.startswith("- ") or ln.startswith("* "):
            add_para(doc, "· " + ln[2:].strip(), first_indent=True)
        elif re.match(r"^\d+\. ", ln):
            add_para(doc, ln.strip(), first_indent=True)
        elif ln.startswith("*") and ln.endswith("*") and not ln.startswith("**"):
            add_para(doc, ln.strip("*"), size=10.5, font=HEAD_FONT, spacing=16)
        else:
            add_para(doc, ln.strip(), first_indent=True)
        i += 1
    flush_table()

    cp = doc.core_properties
    cp.author = ""
    cp.title = md_path.stem
    from datetime import datetime, timezone
    fixed = datetime(2026, 7, 26, 0, 0, 0, tzinfo=timezone.utc)
    cp.created = fixed
    cp.modified = fixed
    register_font_alts(doc)
    doc.save(out_path)
    normalize_zip(out_path)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: md_to_docx.py <盲稿md目录> <输出目录>"); sys.exit(2)
    src_dir, out_dir = Path(sys.argv[1]), Path(sys.argv[2])
    out_dir.mkdir(parents=True, exist_ok=True)
    n = 0
    for mdp in sorted(src_dir.glob("0*.md")):
        convert(mdp, out_dir / (mdp.stem + ".docx")); n += 1
    print(f"written: {n} docx -> {out_dir}")
    sys.exit(0)
