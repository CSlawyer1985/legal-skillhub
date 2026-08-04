#!/usr/bin/env python3
"""
招标文件文本提取脚本
支持 PDF（含OCR）、DOCX、TXT 格式
"""
import argparse
import json
import os
import sys


def extract_from_pdf(file_path, enable_ocr=True):
    """从PDF提取文本，支持OCR"""
    try:
        import pdfplumber
    except ImportError:
        print("ERROR: pdfplumber not installed. Run: pip install pdfplumber", file=sys.stderr)
        sys.exit(1)

    pages_text = []
    ocr_used = False
    page_count = 0

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text()
            if (not text or text.strip() == "") and enable_ocr:
                # OCR fallback
                try:
                    import pytesseract
                    from PIL import Image
                    img = page.to_image(resolution=300)
                    pil_img = img.original
                    text = pytesseract.image_to_string(pil_img, lang='chi_sim+eng')
                    ocr_used = True
                except ImportError:
                    text = ""
                    print(f"WARNING: OCR not available for page {page.page_number}", file=sys.stderr)
            pages_text.append(text or "")

    full_text = "\n\n".join(pages_text)
    return full_text, page_count, ocr_used


def extract_from_docx(file_path):
    """从DOCX提取文本"""
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 提取表格
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            tables_text.append(row_text)

    full_text = "\n".join(paragraphs)
    if tables_text:
        full_text += "\n\n--- 表格内容 ---\n" + "\n".join(tables_text)

    return full_text, len(paragraphs), False


def extract_from_txt(file_path):
    """从TXT提取文本"""
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return text, text.count('\n') + 1, False


def main():
    parser = argparse.ArgumentParser(description='招标文件文本提取')
    parser.add_argument('--file_path', required=True, help='招标文件路径')
    parser.add_argument('--project_id', required=True, help='项目ID')
    parser.add_argument('--enable_ocr', default='true', help='是否启用OCR')
    parser.add_argument('--output_dir', required=True, help='输出目录')
    args = parser.parse_args()

    enable_ocr = args.enable_ocr.lower() == 'true'
    file_path = args.file_path
    ext = os.path.splitext(file_path)[1].lower()

    result = {
        "project_id": args.project_id,
        "file_name": os.path.basename(file_path),
        "file_type": ext.lstrip('.'),
        "page_count": 0,
        "ocr_used": False,
        "text_length": 0,
        "output_path": "",
        "status": "success",
        "error_message": ""
    }

    try:
        if ext == '.pdf':
            text, page_count, ocr_used = extract_from_pdf(file_path, enable_ocr)
            result["page_count"] = page_count
            result["ocr_used"] = ocr_used
        elif ext in ('.docx', '.doc'):
            text, line_count, _ = extract_from_docx(file_path)
            result["page_count"] = line_count
        elif ext == '.txt':
            text, line_count, _ = extract_from_txt(file_path)
            result["page_count"] = line_count
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

        # 保存文本
        os.makedirs(args.output_dir, exist_ok=True)
        output_path = os.path.join(args.output_dir, "raw_text.txt")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)

        result["text_length"] = len(text)
        result["output_path"] = output_path

    except Exception as e:
        result["status"] = "failed"
        result["error_message"] = str(e)

    # 输出元信息
    meta_path = os.path.join(args.output_dir, "extraction_meta.json")
    with open(meta_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["status"] == "success" else 1


if __name__ == "__main__":
    sys.exit(main())
